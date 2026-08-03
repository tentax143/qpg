"""Regression tests for compound-paper marks handling.

These lock in two production bugs so they can't silently return:

  1. ExamPattern.get_total_marks / get_total_questions double-counted compound
     papers (section marks + subsection marks) — an 80-mark paper reported 160.

  2. The section generator crashed on compound sections because their
     marks_per_question is the literal string "varies" (real per-type marks live
     in `subsections`):  TypeError: '>=' not supported between 'str' and 'int'.

Run with:  python manage.py test core
"""

import collections
import copy
import json
from unittest import mock

from django.test import TestCase, TransactionTestCase, SimpleTestCase

from django.contrib.auth.models import User

from core.models import ExamPattern, School, ExamBlueprint
from core import section_generator as sg
from core import mantle_client as mc
from core import generator as sg_gen
from core import paper_edit as pe
from api.views import _scoped_blueprints


def _sub(name, qtype, marks, q, mpq):
    return {"name": name, "question_types": [qtype], "marks": marks,
            "questions_count": q, "marks_per_question": mpq}


def _compound_sections():
    """A compound Science board paper: each section's `marks` equals the sum of
    its subsections, marks_per_question == "varies", question_types is a flat list
    of type-name strings. Mirrors the real DB pattern that triggered both bugs."""
    types = ["MCQ", "Assertion-Reason", "VSA", "SA", "Source-Based/CBQ", "LA"]
    bio = [_sub("MCQ", "MCQ", 7, 7, 1), _sub("AR", "Assertion-Reason", 2, 2, 1),
           _sub("VSA", "VSA", 6, 3, 2), _sub("SA", "SA", 6, 2, 3),
           _sub("CBQ", "Source-Based/CBQ", 4, 1, 4), _sub("LA", "LA", 5, 1, 5)]   # = 30 / 16q
    chem = [_sub("MCQ", "MCQ", 7, 7, 1), _sub("AR", "Assertion-Reason", 1, 1, 1),
            _sub("VSA", "VSA", 2, 1, 2), _sub("SA", "SA", 6, 2, 3),
            _sub("CBQ", "Source-Based/CBQ", 4, 1, 4), _sub("LA", "LA", 5, 1, 5)]   # = 25 / 13q
    return [
        {"name": "Biology", "marks": 30, "questions_count": 16,
         "marks_per_question": "varies", "question_types": types, "subsections": bio},
        {"name": "Chemistry", "marks": 25, "questions_count": 13,
         "marks_per_question": "varies", "question_types": types, "subsections": chem},
    ]


class TotalMarksTest(TestCase):
    def test_compound_total_not_doubled(self):
        p = ExamPattern(name="t", sections=_compound_sections())
        self.assertEqual(p.get_total_marks(), 55)        # 30 + 25, NOT 110
        self.assertEqual(p.get_total_questions(), 29)    # 16 + 13, NOT 58

    def test_save_stores_correct_totals(self):
        p = ExamPattern.objects.create(name="t2", sections=_compound_sections())
        self.assertEqual(p.total_marks, 55)
        self.assertEqual(p.total_questions, 29)

    def test_plain_section_unchanged(self):
        # A normal (non-compound) section with no subsections must still total normally.
        p = ExamPattern(name="t3", sections=[{"name": "A", "marks": 20, "questions_count": 10}])
        self.assertEqual(p.get_total_marks(), 20)
        self.assertEqual(p.get_total_questions(), 10)


class VariesWorkOrderTest(TestCase):
    def setUp(self):
        self.pattern = ExamPattern(name="wo", sections=_compound_sections())
        self.blueprint = {s["name"]: s for s in self.pattern.sections}

    def _build(self):
        return sg.build_work_orders(self.blueprint, self.pattern, {}, "Hard",
                                    "10", "Science", ["Life Processes"])

    def test_per_q_tokens_tolerates_non_numeric(self):
        self.assertIsInstance(sg._per_q_tokens("varies"), int)
        self.assertIsInstance(sg._per_q_tokens("3"), int)
        self.assertIsInstance(sg._per_q_tokens(None), int)

    def test_build_work_orders_numeric_and_mixed(self):
        wos = self._build()
        self.assertEqual(len(wos), 2)
        for wo in wos:
            self.assertIsInstance(wo.marks_per_question, (int, float))   # never "varies"
            self.assertTrue(wo.mixed_marks)
            self.assertTrue(wo.question_types and all(isinstance(qt, dict) for qt in wo.question_types))
            self.assertTrue(all("marks_each" in qt for qt in wo.question_types))

    def test_estimate_token_budget_no_crash(self):
        for wo in self._build():
            budget = sg.estimate_token_budget(wo)        # this is the call that used to crash
            self.assertIsInstance(budget, int)
            self.assertGreaterEqual(budget, 3000)
            self.assertLessEqual(budget, 8192)

    def test_prompt_emits_per_type_marks(self):
        prompt = sg.build_section_prompt(self._build()[0])
        self.assertIn("VARIES BY TYPE", prompt)
        self.assertIn("assertion_reason", prompt)


class CompoundChapterRoutingTest(TestCase):
    """A Biology section of a Science paper must only see Biology chapters — otherwise
    it generates Chemistry/Physics questions (the cross-subject leak bug)."""

    CHAPTERS = ["Acids, Bases and Salts", "Carbon and its Compounds",
                "Chemical Reactions and Equations", "Control and Coordination",
                "Electricity", "Heredity", "How do Organisms Reproduce", "Life Processes",
                "Light – Reflection and Refraction", "Magnetic Effects of Electric Current",
                "Metals and Non-metals", "Our Environment",
                "The Human Eye and the Colourful World"]

    def test_section_subject_inferred_from_name(self):
        # Pattern omits section_subject — it must be inferred from the section name.
        self.assertEqual(sg._resolve_section_subject("Science", "Biology", ""), "Biology")
        self.assertEqual(sg._resolve_section_subject("Science", "Chemistry", ""), "Chemistry")
        # Explicit value wins; ordinary papers stay unscoped.
        self.assertEqual(sg._resolve_section_subject("Science", "Biology", "Physics"), "Physics")
        self.assertEqual(sg._resolve_section_subject("Physics", "Section A", ""), "")

    def test_biology_section_excludes_chem_and_physics(self):
        bio = sg._chapters_for_subject("Biology", "Science", self.CHAPTERS)
        self.assertIn("Life Processes", bio)
        self.assertIn("Heredity", bio)
        for leaked in ("Acids, Bases and Salts", "Metals and Non-metals",
                       "Electricity", "Magnetic Effects of Electric Current"):
            self.assertNotIn(leaked, bio)

    def test_each_subject_partitions_chapters(self):
        bio = sg._chapters_for_subject("Biology", "Science", self.CHAPTERS)
        chem = sg._chapters_for_subject("Chemistry", "Science", self.CHAPTERS)
        phys = sg._chapters_for_subject("Physics", "Science", self.CHAPTERS)
        self.assertEqual(len(bio) + len(chem) + len(phys), len(self.CHAPTERS))
        self.assertEqual(set(bio) | set(chem) | set(phys), set(self.CHAPTERS))

    def test_single_subject_paper_keeps_all_chapters(self):
        # A real Biology paper (parent == section) must not be scoped down.
        self.assertEqual(sg._chapters_for_subject("Biology", "Biology", self.CHAPTERS),
                         list(self.CHAPTERS))


class RenderGroupingTest(TestCase):
    """Render layer must group questions by type and restore per-type marks, so an SA
    never lands in the MCQ block and a 1-mark MCQ never shows the section average."""

    SEC_INFO = {
        "name": "Biology", "marks": 30,
        "subsections": [
            {"name": "MCQ", "question_types": ["MCQ"], "marks": 7, "questions_count": 7, "marks_per_question": 1},
            {"name": "Assertion-Reason", "question_types": ["Assertion-Reason"], "marks": 2, "questions_count": 2, "marks_per_question": 1},
            {"name": "VSA", "question_types": ["VSA"], "marks": 6, "questions_count": 3, "marks_per_question": 2},
            {"name": "SA", "question_types": ["SA"], "marks": 6, "questions_count": 2, "marks_per_question": 3},
            {"name": "CBQ", "question_types": ["Source-Based/CBQ"], "marks": 4, "questions_count": 1, "marks_per_question": 4},
            {"name": "LA", "question_types": ["LA"], "marks": 5, "questions_count": 1, "marks_per_question": 5},
        ],
    }

    def _questions(self):
        # Deliberately jumbled order + wrong average marks (mimics fallback output).
        return [
            {"qnum": 1, "type": "MCQ", "subtype": "assertion_reason", "marks": 1.9},
            {"qnum": 2, "type": "SA", "subtype": "standard", "marks": 1.9},
            {"qnum": 3, "type": "MCQ", "subtype": "standard", "marks": 1.9},
            {"qnum": 4, "type": "LA", "subtype": "standard", "marks": 1.9},
            {"qnum": 5, "type": "VSA", "subtype": "standard", "marks": 1.9},
            {"qnum": 6, "type": "CBQ", "subtype": "source_based", "marks": 1.9},
        ]

    def test_marks_map_from_subsections(self):
        m = sg_gen._section_type_marks(self.SEC_INFO)
        self.assertEqual(m, {"mcq": 1.0, "ar": 1.0, "vsa": 2.0, "sa": 3.0, "cbq": 4.0, "la": 5.0})

    def test_category_assignment(self):
        self.assertEqual(sg_gen._question_category({"type": "MCQ", "subtype": "assertion_reason"}), "ar")
        self.assertEqual(sg_gen._question_category({"type": "MCQ", "subtype": "standard"}), "mcq")
        self.assertEqual(sg_gen._question_category({"type": "CBQ", "subtype": "source_based"}), "cbq")
        self.assertEqual(sg_gen._question_category({"type": "SA"}), "sa")
        self.assertEqual(sg_gen._question_category({"type": "LA"}), "la")

    def test_regroup_orders_and_fixes_marks(self):
        groups = sg_gen._regroup_section(self._questions(), self.SEC_INFO)
        labels = [lbl for lbl, _ in groups if lbl]
        # Canonical order, MCQ group first, LA last.
        self.assertTrue(labels[0].startswith("I.") and "Multiple Choice" in labels[0])
        self.assertIn("Long Answer", labels[-1])
        # Per-type marks restored — no 1.9 anywhere.
        flat = [q for _, qs in groups for q in qs]
        by_cat = {sg_gen._question_category(q): q["marks"] for q in flat}
        self.assertEqual(by_cat["mcq"], 1)
        self.assertEqual(by_cat["sa"], 3)
        self.assertEqual(by_cat["cbq"], 4)
        self.assertEqual(by_cat["la"], 5)
        self.assertFalse(any(q["marks"] == 1.9 for q in flat))

    def test_uniform_section_gets_no_group_labels(self):
        # All-MCQ section → single group, no roman-numeral headers, no behaviour change.
        qs = [{"type": "MCQ", "subtype": "standard", "marks": 1} for _ in range(5)]
        groups = sg_gen._regroup_section(qs, {"name": "A", "subsections": []})
        self.assertEqual([lbl for lbl, _ in groups if lbl], [])

    def test_render_stamps_display_qnums_on_shared_dicts(self):
        # _render_paper_from_stored_data persists paper_data after rendering BECAUSE the
        # renderer writes each question's PRINTED number back into the shared question dicts
        # (blueprint section order + type regrouping), regardless of paper_data dict order or
        # stale stored qnums. paper_edit.renumber() numbers in storage order instead, so
        # without this write-back+save the ai_edit planner targets numbers the teacher
        # doesn't see. This test pins the write-back contract the fix relies on.
        blueprint = {
            "Section A": {"title": "Objective", "marks": 2},
            "Section B": {"title": "Short Answer", "marks": 3},
        }
        a_mcq = {"qnum": 99, "type": "MCQ", "subtype": "standard", "marks": 1, "text": "pick",
                 "options": {"a": "1", "b": "2", "c": "3", "d": "4"}, "answer": "a"}
        a_sa = {"qnum": 98, "type": "SA", "subtype": "standard", "marks": 1, "text": "why"}
        b_sa = {"qnum": 97, "type": "SA", "subtype": "standard", "marks": 3, "text": "explain"}
        # dict order (B first) deliberately differs from blueprint order (A first); Section A
        # stores its SA before its MCQ while the print regroups MCQ first.
        data = {
            "Section B": {"questions": [b_sa]},
            "Section A": {"questions": [a_sa, a_mcq]},
        }
        sg_gen.render_section_questions([], data, blueprint)
        self.assertEqual({a_mcq["qnum"], a_sa["qnum"]}, {1, 2})   # Section A printed first
        self.assertLess(a_mcq["qnum"], a_sa["qnum"])              # regrouped: MCQ before SA
        self.assertEqual(b_sa["qnum"], 3)                         # B last despite dict order


class SectionHeaderLabelTest(TestCase):
    """AI/teacher-authored patterns often name a section 'Section A — Objective Type'
    (the label already spelled out) rather than a bare letter. The renderer used to
    always prefix 'SECTION – ', producing a doubled word: 'SECTION – Section A —
    Objective Type'. Reported production bug on a Social Science paper."""

    def _headers(self, blueprint, data):
        all_q = []
        sg_gen.render_section_questions(all_q, data, blueprint)
        return [text for typ, text in all_q if typ == "header"]

    def test_no_duplicate_section_word_when_name_already_labelled(self):
        blueprint = {"Section A — Objective Type": {"title": "", "marks": 10}}
        data = {"Section A — Objective Type": {"questions": []}}
        self.assertEqual(self._headers(blueprint, data),
                         ["Section A — Objective Type (10 MARKS)"])

    def test_bare_letter_keeps_section_prefix(self):
        blueprint = {"A": {"title": "", "marks": 5}}
        data = {"A": {"questions": []}}
        self.assertEqual(self._headers(blueprint, data), ["SECTION – A (5 MARKS)"])

    def test_title_suffix_still_appended_for_bare_letter(self):
        blueprint = {"A": {"title": "Objective Type", "marks": 5}}
        data = {"A": {"questions": []}}
        self.assertEqual(self._headers(blueprint, data),
                         ["SECTION – A: OBJECTIVE TYPE (5 MARKS)"])

    def test_compound_sub_subject_no_duplicate_when_already_labelled(self):
        blueprint = {"Section A": {"title": "", "marks": 26, "section_subject": "Biology"}}
        data = {"Section A": {"questions": []}}
        self.assertEqual(self._headers(blueprint, data), ["Section A — BIOLOGY (26 MARKS)"])

    def test_compound_sub_subject_bare_letter_unchanged(self):
        blueprint = {"A": {"title": "", "marks": 26, "section_subject": "Biology"}}
        data = {"A": {"questions": []}}
        self.assertEqual(self._headers(blueprint, data), ["SECTION A — BIOLOGY (26 MARKS)"])


class CountBasedValidationTest(TestCase):
    """Section validation must check the right NUMBER of each type (renderer handles order),
    not that each type sits at an exact position."""

    def _wo(self):
        pattern = ExamPattern(name="cv", subject="Science", class_name="10", sections=[{
            "name": "Biology", "marks": 8, "questions_count": 5, "marks_per_question": "varies",
            "question_types": ["MCQ", "Assertion-Reason", "VSA", "SA"],
            "subsections": [
                {"name": "MCQ", "question_types": ["MCQ"], "marks": 2, "questions_count": 2, "marks_per_question": 1},
                {"name": "AR", "question_types": ["Assertion-Reason"], "marks": 1, "questions_count": 1, "marks_per_question": 1},
                {"name": "VSA", "question_types": ["VSA"], "marks": 2, "questions_count": 1, "marks_per_question": 2},
                {"name": "SA", "question_types": ["SA"], "marks": 3, "questions_count": 1, "marks_per_question": 3},
            ],
        }])
        bp = {s["name"]: s for s in pattern.sections}
        return sg.build_work_orders(bp, pattern, {}, "Medium", "10", "Science", ["Life"])[0]

    def _mcq(self, m=1, ar=False):
        q = {"type": "MCQ", "subtype": "assertion_reason" if ar else "standard", "marks": m,
             "text": "Assertion (A): x\nReason (R): y" if ar else "q", "answer": "a",
             "options": {"a": "1", "b": "2", "c": "3", "d": "4"}}
        return q

    def test_right_counts_wrong_order_passes_type_check(self):
        # Jumbled order, but exactly 2 MCQ + 1 AR + 1 VSA + 1 SA.
        qs = [
            {"type": "SA", "marks": 3, "text": "q", "answer_explanation": "a"},
            {"type": "VSA", "marks": 2, "text": "q", "answer_explanation": "a"},
            self._mcq(1),
            self._mcq(1, ar=True),
            self._mcq(1),
        ]
        errors = sg.validate_section_output({"questions": qs}, self._wo())
        self.assertFalse(any("distribution wrong" in e for e in errors), errors)
        self.assertFalse(any("type mismatch" in e for e in errors), errors)

    def test_wrong_counts_flagged(self):
        # 3 MCQ, 0 AR, 1 VSA, 1 SA — AR missing, MCQ over.
        qs = [self._mcq(1), self._mcq(1), self._mcq(1),
              {"type": "VSA", "marks": 2, "text": "q", "answer_explanation": "a"},
              {"type": "SA", "marks": 3, "text": "q", "answer_explanation": "a"}]
        errors = sg.validate_section_output({"questions": qs}, self._wo())
        dist = [e for e in errors if "distribution wrong" in e]
        self.assertTrue(dist, errors)
        self.assertIn("AR", dist[0])


class SectionTypeEnforcementTest(TestCase):
    """Uniform sections must reject foreign question types and the prompt must forbid them."""

    def _wo(self, types, count=5, mpq=2.0):
        from core import section_generator as sg
        return sg.SectionWorkOrder(
            section_name="Section B", section_id="B", title="Short Answer I", marks=int(count * mpq),
            questions_count=count, marks_per_question=mpq, question_types=types, instructions=[],
            constraints={}, context_text="ctx", difficulty="medium", subject="Chemistry",
            class_name="12", chapters=["Equilibrium"])

    def _sa_q(self, n):
        return {"qnum": n, "type": "SA", "subtype": "standard", "text": f"Explain concept {n}.",
                "marks": 2, "answer_explanation": "…", "competency_type": "constructed"}

    def _mcq_q(self, n):
        return {"qnum": n, "type": "MCQ", "subtype": "standard", "text": f"Pick {n}",
                "options": {"a": "1", "b": "2", "c": "3", "d": "4"}, "answer": "a",
                "marks": 2, "answer_explanation": "…", "competency_type": "recall"}

    def test_mcq_in_short_answer_section_is_rejected(self):
        from core import section_generator as sg
        wo = self._wo(["Short Answer I"])
        data = {"questions": [self._sa_q(1), self._sa_q(2), self._sa_q(3), self._mcq_q(4), self._mcq_q(5)]}
        errs = sg.validate_section_output(data, wo)
        self.assertTrue(any("Wrong question type" in e and "MCQ" in e for e in errs), errs)

    def test_all_short_answer_section_passes_type_check(self):
        from core import section_generator as sg
        wo = self._wo(["Short Answer I"])
        data = {"questions": [self._sa_q(i) for i in range(1, 6)]}
        errs = sg.validate_section_output(data, wo)
        self.assertFalse(any("Wrong question type" in e for e in errs), errs)

    def test_mcq_and_ar_both_allowed_in_objective_section(self):
        from core import section_generator as sg
        wo = self._wo(["MCQ", "Assertion-Reason"], count=2, mpq=1.0)
        ar = {"qnum": 2, "type": "MCQ", "subtype": "assertion_reason",
              "text": "Assertion (A): x.\nReason (R): y.",
              "options": {"a": "Both A and R are true and R is the correct explanation of A",
                          "b": "Both true, R not correct explanation", "c": "A true R false", "d": "A false R true"},
              "answer": "a", "marks": 1, "answer_explanation": "…", "competency_type": "application"}
        mcq = self._mcq_q(1); mcq["marks"] = 1
        errs = sg.validate_section_output({"questions": [mcq, ar]}, wo)
        self.assertFalse(any("Wrong question type" in e for e in errs), errs)

    def test_prompt_forbids_other_types_for_short_answer(self):
        from core import section_generator as sg
        wo = self._wo(["Short Answer I"])
        sg.plan_chapter_allocation([wo])
        p = sg.build_section_prompt(wo)
        self.assertIn("QUESTION TYPE — MANDATORY", p)
        self.assertIn("WRITTEN-ANSWER", p)
        self.assertIn('"type":"SA"', p)

    def test_mcq_token_budget_fits_sixteen(self):
        from core import section_generator as sg
        wo = self._wo(["MCQ", "Assertion-Reason"], count=16, mpq=1.0)
        # 16 MCQs must budget well above the old 3380 truncation point.
        self.assertGreater(sg.estimate_token_budget(wo), 4000)

    def test_enforce_pass_strips_mcq_from_sa_section(self):
        from core import section_generator as sg
        wo = self._wo(["Short Answer I"])
        paper_data = {"Section B": {"section_name": "Section B", "questions": [
            self._sa_q(1), self._mcq_q(2), self._sa_q(3)]}}
        out = sg.enforce_section_question_types(paper_data, [wo])
        kinds = [q["type"] for q in out["Section B"]["questions"]]
        self.assertEqual(kinds, ["SA", "SA"])                         # MCQ removed
        self.assertEqual(len(out["Section B"]["_dropped_wrong_type"]), 1)

    def test_enforce_pass_keeps_assertion_reason_in_objective_section(self):
        from core import section_generator as sg
        wo = self._wo(["MCQ", "Assertion-Reason"], count=2, mpq=1.0)
        ar = {"qnum": 2, "type": "MCQ", "subtype": "assertion_reason", "text": "A…R…",
              "options": {"a": "1", "b": "2", "c": "3", "d": "4"}, "answer": "a", "marks": 1}
        paper_data = {"Section B": {"section_name": "Section B", "questions": [self._mcq_q(1), ar]}}
        out = sg.enforce_section_question_types(paper_data, [wo])
        self.assertEqual(len(out["Section B"]["questions"]), 2)        # nothing dropped
        self.assertNotIn("_dropped_wrong_type", out["Section B"])

    def test_enforce_pass_keeps_cbq_with_subtype(self):
        from core import section_generator as sg
        wo = self._wo(["Case-Based Questions"], count=2, mpq=4.0)
        cbq = {"qnum": 1, "type": "CBQ", "subtype": "image_based", "text": "Observe…",
               "marks": 4, "sub_questions": [{"text": "a?", "marks": 2}, {"text": "b?", "marks": 2}]}
        paper_data = {"Section B": {"section_name": "Section B", "questions": [cbq]}}
        out = sg.enforce_section_question_types(paper_data, [wo])
        self.assertEqual(len(out["Section B"]["questions"]), 1)        # CBQ kept despite subtype

    def _map_q(self, n, marks=1):
        # Map questions are emitted as type "SA" + subtype "map_based" per the schema.
        return {"qnum": n, "type": "SA", "subtype": "map_based",
                "text": f"On the given outline map of India, locate and label: (a) place {n} (b) place {n + 1}",
                "marks": marks, "map_note": "[Attach outline map of India — examiner to supply]",
                "chapter_tag": "Ch", "competency_type": "application"}

    def test_enforce_pass_keeps_map_based_in_pure_map_section(self):
        # SS Section F ("Map location" / "Diagram-based"): every valid map question was
        # stripped as a foreign SA, shipping the section 0/3.
        from core import section_generator as sg
        wo = self._wo(["Map location", "Diagram-based"], count=3, mpq=1.0)
        paper_data = {"Section B": {"section_name": "Section B",
                                    "questions": [self._map_q(1), self._map_q(2), self._map_q(3)]}}
        out = sg.enforce_section_question_types(paper_data, [wo])
        self.assertEqual(len(out["Section B"]["questions"]), 3)        # nothing dropped
        self.assertNotIn("_dropped_wrong_type", out["Section B"])

    def test_enforce_pass_keeps_map_and_la_in_mixed_section(self):
        # SS Section D ("Long Answer" / "Map Work"): the map half was dropped, shipping 1/2.
        from core import section_generator as sg
        wo = self._wo(["Long Answer", "Map Work"], count=2, mpq=5.0)
        la = {"qnum": 1, "type": "LA", "subtype": "standard", "text": "Explain in detail…",
              "marks": 5, "or_alternative": {"text": "Or explain…"}}
        paper_data = {"Section B": {"section_name": "Section B",
                                    "questions": [la, self._map_q(2, marks=5)]}}
        out = sg.enforce_section_question_types(paper_data, [wo])
        self.assertEqual(len(out["Section B"]["questions"]), 2)        # LA and map both kept
        self.assertNotIn("_dropped_wrong_type", out["Section B"])

    def test_enforce_pass_keeps_map_by_map_note_when_subtype_missing(self):
        from core import section_generator as sg
        wo = self._wo(["Map Work"], count=1, mpq=2.0)
        q = self._map_q(1, marks=2)
        q["subtype"] = "standard"                                      # model forgot the subtype
        paper_data = {"Section B": {"section_name": "Section B", "questions": [q]}}
        out = sg.enforce_section_question_types(paper_data, [wo])
        self.assertEqual(len(out["Section B"]["questions"]), 1)        # map_note is enough
        self.assertNotIn("_dropped_wrong_type", out["Section B"])

    def test_enforce_pass_still_strips_mcq_from_map_section(self):
        from core import section_generator as sg
        wo = self._wo(["Map location"], count=2, mpq=1.0)
        paper_data = {"Section B": {"section_name": "Section B",
                                    "questions": [self._map_q(1), self._mcq_q(2)]}}
        out = sg.enforce_section_question_types(paper_data, [wo])
        self.assertEqual(len(out["Section B"]["questions"]), 1)        # MCQ still removed
        self.assertEqual(len(out["Section B"]["_dropped_wrong_type"]), 1)

    def test_top_up_applies_to_map_section(self):
        # A short map section used to be refused by the top-up (no recovery path at all).
        from unittest import mock
        from core import section_generator as sg
        wo = self._wo(["Map location"], count=3, mpq=1.0)
        wo.is_map_work = True
        sec = {"section_name": "Section B", "questions": [self._map_q(1), self._map_q(3)]}
        reply = ('{"questions": [{"qnum": 3, "type": "SA", "subtype": "map_based", '
                 '"text": "On the given outline map of India, locate and label: (a) Mumbai (b) Goa", '
                 '"marks": 1, "map_note": "[Attach outline map of India — examiner to supply]", '
                 '"chapter_tag": "Ch", "competency_type": "application"}]}')
        with mock.patch.object(sg.mantle_client, "converse", return_value=(reply, 10, 20)):
            in_tok, out_tok = sg._top_up_short_section(sec, wo)
        self.assertEqual(len(sec["questions"]), 3)                     # topped up to count
        self.assertEqual(sec.get("_topped_up"), 1)
        self.assertGreater(in_tok + out_tok, 0)

    def test_top_up_still_refuses_cbq_section(self):
        from unittest import mock
        from core import section_generator as sg
        wo = self._wo(["Case-Based Questions"], count=2, mpq=4.0)
        sec = {"section_name": "Section B", "questions": []}
        with mock.patch.object(sg.mantle_client, "converse") as conv:
            in_tok, out_tok = sg._top_up_short_section(sec, wo)
        self.assertEqual((in_tok, out_tok), (0, 0))
        conv.assert_not_called()


class MarksReconcileAndTopUpLoopTest(TestCase):
    """Uniform sections must end up at the right marks AND the right count — the two shapes
    behind the recurring 'Marks check' audit warnings (9/10 marks; 18/20 questions)."""

    def _wo(self, mixed=False, mpq=2.0, count=5, types=None, name="Section B"):
        from core import section_generator as sg
        return sg.SectionWorkOrder(
            section_name=name, section_id="B", title="VSA", marks=int(count * mpq),
            questions_count=count, marks_per_question=mpq, question_types=types or ["VSA"],
            instructions=[], constraints={}, context_text="ctx", difficulty="medium",
            subject="Science", class_name="10", chapters=["Light"], mixed_marks=mixed)

    def test_reconcile_clamps_drifted_uniform_marks(self):
        from core import section_generator as sg
        wo = self._wo(mpq=2.0, count=5)
        qs = [{"type": "VSA", "text": f"q{i}", "marks": 2} for i in range(4)]
        qs.append({"type": "VSA", "text": "drift", "marks": 1})          # one drifted to 1m
        pd = {"Section B": {"section_name": "Section B", "questions": qs}}
        sg.reconcile_uniform_marks(pd, [wo])
        self.assertEqual(sum(q["marks"] for q in pd["Section B"]["questions"]), 10)

    def test_reconcile_fixes_non_numeric_marks(self):
        from core import section_generator as sg
        wo = self._wo(mpq=3.0, count=2, types=["SA"])
        pd = {"Section B": {"section_name": "Section B", "questions": [
            {"type": "SA", "text": "a", "marks": 3},
            {"type": "SA", "text": "b", "marks": "varies"}]}}
        sg.reconcile_uniform_marks(pd, [wo])
        self.assertEqual([q["marks"] for q in pd["Section B"]["questions"]], [3.0, 3.0])

    def test_reconcile_leaves_mixed_marks_untouched(self):
        from core import section_generator as sg
        wo = self._wo(mixed=True, mpq=2.0, count=3)
        qs = [{"type": "MCQ", "text": "a", "marks": 1},
              {"type": "SA", "text": "b", "marks": 3},
              {"type": "VSA", "text": "c", "marks": 2}]
        pd = {"Section B": {"section_name": "Section B", "questions": qs}}
        sg.reconcile_uniform_marks(pd, [wo])
        self.assertEqual([q["marks"] for q in pd["Section B"]["questions"]], [1, 3, 2])

    def test_reconcile_skips_cbq_subquestions(self):
        from core import section_generator as sg
        wo = self._wo(mpq=4.0, count=1, types=["CBQ"])
        pd = {"Section B": {"section_name": "Section B", "questions": [
            {"type": "CBQ", "text": "x", "marks": 4,
             "sub_questions": [{"text": "a", "marks": 2}, {"text": "b", "marks": 2}]}]}}
        sg.reconcile_uniform_marks(pd, [wo])
        self.assertEqual(pd["Section B"]["questions"][0]["marks"], 4)

    def test_distributes_marks_when_count_times_mpq_below_total(self):
        # The exact log case: Section B declared 10m but 3 SA × 3m = 9. Distribute to 4,3,3.
        from core import section_generator as sg
        wo = self._wo(mpq=3.0, count=3, types=["Short Answer I (SA-I)"])
        wo.marks = 10                                          # pattern says 10m, 3×3 can't reach it
        qs = [{"type": "SA", "text": f"q{i}", "marks": 3} for i in range(3)]
        pd = {"Section B": {"section_name": "Section B", "questions": qs}}
        sg.reconcile_uniform_marks(pd, [wo])
        marks = sorted(q["marks"] for q in pd["Section B"]["questions"])
        self.assertEqual(sum(marks), 10)                       # section now totals exactly 10
        self.assertEqual(marks, [3, 3, 4])

    def test_objective_section_not_inflated_when_short(self):
        # A short MCQ section must NOT have its marks inflated to hide the count shortfall.
        from core import section_generator as sg
        wo = self._wo(mpq=1.0, count=20, types=["MCQ", "Assertion-Reason"])
        wo.marks = 20
        qs = [{"type": "MCQ", "text": f"q{i}", "marks": 1} for i in range(18)]  # 18/20 short
        pd = {"Section A": {"section_name": "Section A", "questions": qs}}
        # work order name must match the dict key
        wo.section_name = "Section A"
        sg.reconcile_uniform_marks(pd, [wo])
        self.assertTrue(all(q["marks"] == 1 for q in pd["Section A"]["questions"]))
        self.assertEqual(sum(q["marks"] for q in pd["Section A"]["questions"]), 18)  # stays 18, not 20

    def test_fill_loop_recovers_full_count_over_rounds(self):
        from unittest import mock
        from core import section_generator as sg
        wo = self._wo(mpq=1.0, count=20, types=["MCQ"], name="Section A")
        sec = {"section_name": "Section A",
               "questions": [{"type": "MCQ", "text": f"q{i}", "marks": 1} for i in range(18)]}
        seq = {"n": 0}

        def fake_topup(section_data, w):
            # Simulate the model recovering only ONE of the missing questions per call —
            # the exact behaviour that left single-shot top-up stuck at 19/20.
            cur = len(section_data["questions"])
            if cur < w.questions_count:
                section_data["questions"].append(
                    {"type": "MCQ", "text": f"new{seq['n']}", "marks": 1})
                seq["n"] += 1
                return 1, 1
            return 0, 0

        with mock.patch.object(sg, "_top_up_short_section", side_effect=fake_topup):
            sg._fill_short_section(sec, wo, max_rounds=5)
        self.assertEqual(len(sec["questions"]), 20)        # filled, not stuck at 19
        self.assertEqual(sec.get("_topped_up"), 2)         # cumulative across rounds

    def test_fill_loop_stops_when_a_round_adds_nothing(self):
        from unittest import mock
        from core import section_generator as sg
        wo = self._wo(mpq=1.0, count=20, types=["MCQ"], name="Section A")
        sec = {"section_name": "Section A",
               "questions": [{"type": "MCQ", "text": f"q{i}", "marks": 1} for i in range(18)]}
        calls = {"n": 0}

        def stuck_topup(section_data, w):
            calls["n"] += 1                                 # spends a call but adds nothing
            return 1, 1

        with mock.patch.object(sg, "_top_up_short_section", side_effect=stuck_topup):
            sg._fill_short_section(sec, wo, max_rounds=5)
        self.assertEqual(calls["n"], 1)                     # bails after the first no-progress round
        self.assertEqual(len(sec["questions"]), 18)


class V5L2DuplicateReplacementTypeTest(TestCase):
    """A duplicate-replacement must keep the original question's TYPE. A type-drifted
    replacement (an SA where the original was an MCQ) is stripped by the type-enforcer and
    drops the section below count — the proven cause of 'Section A: 18/20 questions'."""

    def _wo(self, types, count=20, mpq=1.0):
        from core import section_generator as sg
        return sg.SectionWorkOrder(
            section_name="Section A — Objective Type", section_id="A", title="Objective",
            marks=int(count * mpq), questions_count=count, marks_per_question=mpq,
            question_types=types, instructions=[], constraints={}, context_text="ctx",
            difficulty="medium", subject="Computer Science", class_name="11",
            chapters=["Encoding Schemes and Number System"])

    def _mcq(self, n, text):
        return {"qnum": n, "type": "MCQ", "subtype": "standard", "text": text,
                "options": {"a": "1", "b": "2", "c": "3", "d": "4"}, "answer": "a",
                "marks": 1, "answer_explanation": "because a", "chapter_tag": "X",
                "competency_type": "application"}

    def test_skeleton_preserves_mcq_type(self):
        from core import section_generator as sg
        instr, skel = sg._regen_question_skeleton({"type": "MCQ", "subtype": "standard"}, 5, 1)
        self.assertIn('"type": "MCQ"', skel)
        self.assertIn('"options"', skel)
        self.assertIn("MCQ", instr)

    def test_skeleton_preserves_la_type(self):
        from core import section_generator as sg
        instr, skel = sg._regen_question_skeleton({"type": "LA", "subtype": "standard"}, 31, 5)
        self.assertIn('"type": "LA"', skel)
        self.assertIn('"or_alternative"', skel)
        self.assertNotIn('"options"', skel)

    def test_drifted_sa_replacement_is_rejected_for_mcq(self):
        from unittest import mock
        from core import section_generator as sg
        wo = self._wo(["MCQ", "True/False", "Fill in the Blanks"])
        questions = [self._mcq(1, "What is binary grouping?"),
                     self._mcq(2, "What is binary grouping for octal?")]
        warnings = ["Q1 and Q2 overlap 70% — likely duplicate concept"]

        def fake_converse(**kw):
            if kw.get("model_id") == sg.mantle_client.VAL_MODEL:
                return ('{"same_concept": true, "reason": "dup"}', 0, 0)   # confirm dup
            # regen returns an SA (the old hardcoded-SA behaviour) — must be rejected
            return ('{"qnum": 2, "type": "SA", "text": "Explain binary.", "marks": 1, '
                    '"answer_explanation": "pts", "chapter_tag": "X", '
                    '"competency_type": "constructed"}', 0, 0)

        with mock.patch.object(sg.mantle_client, "converse", side_effect=fake_converse):
            out, remaining = sg.verify_and_fix_semantic_duplicates(questions, warnings, wo, [])
        self.assertEqual(out[1]["type"], "MCQ")              # original MCQ kept, not SA
        self.assertIn("options", out[1])
        self.assertEqual(len(out), 2)                         # count preserved

    def test_valid_mcq_replacement_is_applied(self):
        from unittest import mock
        from core import section_generator as sg
        wo = self._wo(["MCQ"])
        questions = [self._mcq(1, "What is binary grouping?"),
                     self._mcq(2, "What is binary grouping for octal?")]
        warnings = ["Q1 and Q2 overlap 70% — likely duplicate concept"]

        def fake_converse(**kw):
            if kw.get("model_id") == sg.mantle_client.VAL_MODEL:
                return ('{"same_concept": true, "reason": "dup"}', 0, 0)
            return ('{"qnum": 2, "type": "MCQ", "subtype": "standard", '
                    '"text": "Which radix does hexadecimal use?", '
                    '"options": {"a": "8", "b": "10", "c": "16", "d": "2"}, "answer": "c", '
                    '"answer_explanation": "hex is base 16", "chapter_tag": "X", '
                    '"competency_type": "application"}', 0, 0)

        with mock.patch.object(sg.mantle_client, "converse", side_effect=fake_converse):
            out, remaining = sg.verify_and_fix_semantic_duplicates(questions, warnings, wo, [])
        self.assertEqual(out[1]["type"], "MCQ")
        self.assertIn("radix", out[1]["text"])               # replacement applied
        self.assertEqual(remaining, [])                       # warning resolved


class SingularQuestionTypeTest(TestCase):
    """Patterns saved by the generate-page form store the type as singular 'question_type'
    (not the plural 'question_types' list). It must still reach the work order — otherwise the
    section carries no type and the model mixes MCQs into a Short-Answer section, etc."""

    SECTIONS = [
        {"id": "A", "name": "Section A", "marks": 10, "question_type": "MCQ",
         "questions_count": 10, "marks_per_question": 1},
        {"id": "B", "name": "Section B", "marks": 8, "question_type": "Short Answer",
         "questions_count": 4, "marks_per_question": 2},
        {"id": "C", "name": "Section C", "marks": 4, "question_type": "Long Answer",
         "questions_count": 1, "marks_per_question": 4},
    ]

    def test_blueprint_maps_singular_question_type(self):
        pattern = ExamPattern(name="singular", sections=self.SECTIONS)
        bp = sg_gen.pattern_sections_to_blueprint_dict(pattern)
        self.assertEqual(bp["Section A"]["question_types"], ["MCQ"])
        self.assertEqual(bp["Section B"]["question_types"], ["Short Answer"])
        self.assertEqual(bp["Section C"]["question_types"], ["Long Answer"])

    def test_work_orders_carry_type_and_enforce_it(self):
        pattern = ExamPattern(name="singular2", sections=self.SECTIONS)
        bp = sg_gen.pattern_sections_to_blueprint_dict(pattern)
        wos = sg.build_work_orders(bp, pattern, {}, "medium", "5", "Mathematics", ["Patterns"])
        by_name = {w.section_name: w for w in wos}
        self.assertEqual(by_name["Section A"].question_types, ["MCQ"])
        self.assertEqual(by_name["Section B"].question_types, ["Short Answer"])
        # The type constraint is now actually enforced: an MCQ in the SA section is rejected.
        mcq = {"qnum": 1, "type": "MCQ", "subtype": "standard", "text": "Pick", "marks": 2,
               "options": {"a": "1", "b": "2", "c": "3", "d": "4"}, "answer": "a"}
        errs = sg.validate_section_output({"questions": [mcq]}, by_name["Section B"])
        self.assertTrue(any("Wrong question type" in e for e in errs), errs)

    def test_build_work_orders_reads_singular_when_blueprint_lacks_plural(self):
        # A blueprint dict carrying only the singular field must still yield a typed work order.
        bp = {"Section B": {"id": "B", "name": "Section B", "marks": 8,
                            "question_type": "Short Answer", "questions_count": 4,
                            "marks_per_question": 2}}
        wos = sg.build_work_orders(bp, None, {}, "medium", "5", "Mathematics", ["Patterns"])
        self.assertEqual(wos[0].question_types, ["Short Answer"])


class PaperEditOperationsTest(TestCase):
    """The operation-based AI editor: move/delete/swap/set/add/edit + renumber. These are the
    deterministic guts of 'full control' — verified without any model call via a fake generator."""

    def _paper(self):
        return {
            "Section A": {"section_name": "Section A", "section_id": "A", "questions": [
                {"qnum": 1, "type": "MCQ", "marks": 1, "text": "A-one",
                 "options": {"a": "1", "b": "2", "c": "3", "d": "4"}, "answer": "a"},
                {"qnum": 2, "type": "MCQ", "marks": 1, "text": "A-two",
                 "options": {"a": "1", "b": "2", "c": "3", "d": "4"}, "answer": "b"},
            ]},
            "Section B": {"section_name": "Section B", "section_id": "B", "questions": [
                {"qnum": 3, "type": "SA", "marks": 3, "text": "B-one", "answer_explanation": "…"},
                {"qnum": 4, "type": "SA", "marks": 3, "text": "B-two", "answer_explanation": "…"},
            ]},
        }

    def _qtexts(self, pd, section):
        return [q["text"] for q in pd[section]["questions"]]

    def test_move_question_across_sections(self):
        pd, applied, notes = pe.apply_operations(
            self._paper(), [{"action": "move", "qnum": 1, "to_section": "Section B"}])
        self.assertNotIn("A-one", self._qtexts(pd, "Section A"))
        self.assertIn("A-one", self._qtexts(pd, "Section B"))
        self.assertTrue(any("moved" in a for a in applied))

    def test_move_resolves_section_by_letter(self):
        # "Section B" target given simply as "B" must still resolve via section_id.
        pd, applied, _ = pe.apply_operations(
            self._paper(), [{"action": "move", "qnum": 1, "to_section": "B"}])
        self.assertIn("A-one", self._qtexts(pd, "Section B"))

    def test_delete_question(self):
        pd, applied, _ = pe.apply_operations(self._paper(), [{"action": "delete", "qnum": 2}])
        self.assertNotIn("A-two", self._qtexts(pd, "Section A"))
        self.assertEqual(len(pd["Section A"]["questions"]), 1)

    def test_set_marks_field(self):
        pd, applied, _ = pe.apply_operations(
            self._paper(), [{"action": "set", "qnum": 3, "fields": {"marks": 5}}])
        self.assertEqual(pd["Section B"]["questions"][0]["marks"], 5)

    def test_set_flat_form(self):
        # Tolerate the flat shape {"action":"set","qnum":3,"marks":5} (no nested "fields").
        pd, _a, _n = pe.apply_operations(self._paper(), [{"action": "set", "qnum": 3, "marks": 5}])
        self.assertEqual(pd["Section B"]["questions"][0]["marks"], 5)

    def test_swap_questions(self):
        pd, applied, _ = pe.apply_operations(
            self._paper(), [{"action": "swap", "qnum_a": 1, "qnum_b": 2}])
        self.assertEqual(self._qtexts(pd, "Section A"), ["A-two", "A-one"])

    def test_renumber_is_sequential_after_structural_change(self):
        pd, _a, _n = pe.apply_operations(
            self._paper(), [{"action": "move", "qnum": 1, "to_section": "Section B"}])
        nums = [q["qnum"] for _n, s in pe.iter_sections(pd) for q in s["questions"]]
        self.assertEqual(nums, [1, 2, 3, 4])

    def test_add_uses_generator_and_inserts(self):
        def fake_gen(kind, ctx):
            self.assertEqual(kind, "add")
            return {"type": ctx["type"], "marks": ctx["marks"], "text": "fresh added Q",
                    "answer_explanation": "…"}
        pd, applied, _ = pe.apply_operations(
            self._paper(),
            [{"action": "add", "section": "Section B", "type": "SA", "marks": 3,
              "instruction": "on fractions"}],
            generate_fn=fake_gen)
        self.assertIn("fresh added Q", self._qtexts(pd, "Section B"))
        self.assertEqual(len(pd["Section B"]["questions"]), 3)

    def test_edit_replaces_text_and_preserves_identity(self):
        def fake_gen(kind, ctx):
            return {"text": "reworded", "type": ctx["question"]["type"]}
        pd, applied, _ = pe.apply_operations(
            self._paper(), [{"action": "edit", "qnum": 1, "instruction": "reword"}],
            generate_fn=fake_gen)
        q1 = pd["Section A"]["questions"][0]
        self.assertEqual(q1["text"], "reworded")
        self.assertEqual(q1["marks"], 1)          # marks preserved on a plain edit

    def test_unknown_and_missing_ops_are_noted_not_fatal(self):
        pd, applied, notes = pe.apply_operations(self._paper(), [
            {"action": "frobnicate", "qnum": 1},        # unknown
            {"action": "delete", "qnum": 999},           # missing target
            {"action": "delete", "qnum": 1},             # valid — must still apply
        ])
        self.assertEqual(applied, ["deleted Q1"])
        self.assertEqual(len(notes), 2)
        self.assertNotIn("A-one", self._qtexts(pd, "Section A"))

    def test_multiple_ops_apply_in_order(self):
        pd, applied, _ = pe.apply_operations(self._paper(), [
            {"action": "delete", "qnum": 2},
            {"action": "move", "qnum": 3, "to_section": "Section A"},
            {"action": "set", "qnum": 4, "fields": {"marks": 4}},
        ])
        self.assertIn("B-one", self._qtexts(pd, "Section A"))
        self.assertEqual(len(applied), 3)
        # renumbered 1..N with no gaps
        nums = [q["qnum"] for _n, s in pe.iter_sections(pd) for q in s["questions"]]
        self.assertEqual(nums, list(range(1, len(nums) + 1)))


class MarksRightAlignTest(TestCase):
    """Per-question marks must snap to a right-aligned tab at the supplied text-area width, so
    they sit flush at the right margin (not the old fixed 5.75\")."""

    def _render(self, text, tab_twips, left_indent=None):
        import re as _re
        from docx import Document
        from docx.shared import Inches
        from lxml import etree
        from core.generator import _add_question_with_marks
        marks_pattern = _re.compile(r"\s*\[(\d+)\s*marks?\]", _re.IGNORECASE)
        li = Inches(left_indent) if left_indent is not None else None
        p = _add_question_with_marks(Document(), text, marks_pattern,
                                     left_indent=li, right_tab_twips=tab_twips)
        return p, etree.tostring(p._element, encoding="unicode")

    def test_marks_use_supplied_right_tab(self):
        p, xml = self._render("1. What is 2+2? [2 marks]", 10068)
        self.assertIn('val="right"', xml)            # right-aligned tab stop
        self.assertIn('pos="10068"', xml)            # at the supplied text-area width
        self.assertIn("\t[2]", p.text)               # marks rendered number-only after a tab

    def test_marks_tab_tracks_position(self):
        _p, xml = self._render("3. Define osmosis. [3 marks]", 9750)
        self.assertIn('pos="9750"', xml)
        self.assertNotIn('pos="8280"', xml)          # no longer the hardcoded fallback

    def test_no_right_tab_when_no_marks(self):
        _p, xml = self._render("1. A plain question with no marks tag.", 10068)
        self.assertNotIn('val="right"', xml)

    def test_numbered_question_gets_hanging_indent(self):
        # The number is isolated in a hanging-indent column; body text snaps to it via a tab,
        # so wrapped lines align under the text — not under the number.
        p, xml = self._render("8. A long question that wraps onto a second line. [1 marks]", 10068)
        self.assertIn('hanging', xml)                # hanging indent applied
        self.assertIn('val="left"', xml)             # left tab stop at the hang column
        self.assertIn('pos="504"', xml)              # 0.35in * 1440 = 504 twips
        self.assertIn("8.\t", p.text)                # number + tab before the body text

    def test_subquestion_hang_nests_under_its_indent(self):
        # A sub-question indented 0.25" hangs relative to that → body column at 0.25+0.35=0.6".
        _p, xml = self._render("(ii) A wrapped sub-question. [2 marks]", 10068, left_indent=0.25)
        self.assertIn('hanging', xml)
        self.assertIn('pos="864"', xml)              # (0.25+0.35)*1440 = 864 twips

    def test_unnumbered_line_has_no_hanging_indent(self):
        _p, xml = self._render("A passage lead-in with no leading number.", 10068)
        self.assertNotIn('hanging', xml)

    def test_tabs_precede_ind_in_pPr(self):
        # Regression: <w:tabs> MUST come before <w:ind> in <w:pPr> (OOXML CT_PPr order). When it
        # was appended last, Word dropped the left tab and the first body line shot far right.
        p, _xml = self._render("8. A wrapped question. [1 marks]", 10068)
        order = [c.tag.split('}')[-1] for c in p._element.get_or_add_pPr()]
        self.assertIn('tabs', order)
        self.assertIn('ind', order)
        self.assertLess(order.index('tabs'), order.index('ind'))


class ChapterAllocationTest(TestCase):
    """Deterministic, CBSE-weighted, coverage-balanced chapter allocation."""

    def test_plan_length_equals_question_count(self):
        from core import section_generator as sg
        plan = sg._allocate_chapters_to_slots(["A", "B", "C", "D", "E"], 3, "X", {})
        self.assertEqual(len(plan), 3)

    def test_uniform_weights_spread_distinct_when_fewer_slots(self):
        from unittest import mock
        from core import section_generator as sg
        with mock.patch.object(sg, "_chapter_weight", return_value=1):
            plan = sg._allocate_chapters_to_slots(["A", "B", "C", "D", "E"], 3, "X", {})
        self.assertEqual(len(set(plan)), 3)        # 3 distinct chapters, no clustering

    def test_more_slots_than_chapters_covers_all_then_repeats(self):
        from collections import Counter
        from unittest import mock
        from core import section_generator as sg
        with mock.patch.object(sg, "_chapter_weight", return_value=1):
            plan = sg._allocate_chapters_to_slots(["A", "B", "C"], 7, "X", {})
        c = Counter(plan)
        self.assertEqual(set(c), {"A", "B", "C"})  # every chapter covered
        self.assertEqual(sum(c.values()), 7)
        self.assertGreaterEqual(min(c.values()), 2)

    def test_heavier_chapter_gets_more_questions(self):
        from collections import Counter
        from unittest import mock
        from core import section_generator as sg
        w = {"Heavy": 10, "Light": 1}
        # *a absorbs the class_name / default args _chapter_weights passes through.
        with mock.patch.object(sg, "_chapter_weight", side_effect=lambda subj, ch, *a: w.get(ch, 1)):
            plan = sg._allocate_chapters_to_slots(["Heavy", "Light"], 6, "X", {})
        c = Counter(plan)
        self.assertGreater(c["Heavy"], c.get("Light", 0))

    def test_sections_coordinate_to_fill_coverage_gaps(self):
        from unittest import mock
        from core import section_generator as sg
        with mock.patch.object(sg, "_chapter_weight", return_value=1):
            covered = {}
            p1 = sg._allocate_chapters_to_slots(["A", "B", "C", "D"], 2, "X", covered)
            p2 = sg._allocate_chapters_to_slots(["A", "B", "C", "D"], 2, "X", covered)
        self.assertEqual(len(set(p1) | set(p2)), 4)  # 2nd section fills what the 1st left out


class ChapterCoverageAuditTest(TestCase):
    """audit_chapter_coverage: every planned chapter should receive at least one question."""

    def test_no_plan_is_ok(self):
        from core.paper_audit import audit_chapter_coverage
        res = audit_chapter_coverage({"Section A": {"questions": [{"chapter_tag": "X"}]}})
        self.assertFalse(res["has_plan"])
        self.assertTrue(res["ok"])

    def test_all_planned_chapters_covered(self):
        from core.paper_audit import audit_chapter_coverage
        pd = {"Section A": {"_chapter_plan": ["Carbon", "Electricity"],
                            "questions": [{"chapter_tag": "Carbon and its Compounds"},
                                          {"chapter_tag": "Electricity"}]}}
        res = audit_chapter_coverage(pd)
        self.assertTrue(res["ok"])
        self.assertEqual(res["covered"], 2)
        self.assertEqual(res["missed"], [])

    def test_missed_chapter_is_flagged(self):
        from core.paper_audit import audit_chapter_coverage
        pd = {"Section A": {"_chapter_plan": ["Carbon", "Electricity"],
                            "questions": [{"chapter_tag": "Carbon"}, {"chapter_tag": "Carbon"}]}}
        res = audit_chapter_coverage(pd)
        self.assertFalse(res["ok"])
        self.assertEqual((res["planned"], res["covered"]), (2, 1))
        self.assertTrue(any("Electricity" in m for m in res["missed"]))

    def test_chapter_covered_in_another_section_is_not_a_gap(self):
        """A chapter planned for Section B but answered in Section A counts as covered — the
        old per-section audit wrongly reported it as 'Section B: <chapter> missing'."""
        from core.paper_audit import audit_chapter_coverage
        pd = {
            "Section A": {"_chapter_plan": ["Carbon"],
                          "questions": [{"chapter_tag": "Carbon"}]},
            "Section B": {"_chapter_plan": ["Carbon", "Acids"],
                          "questions": [{"chapter_tag": "Acids"}]},
        }
        res = audit_chapter_coverage(pd)
        self.assertTrue(res["ok"])
        self.assertEqual(res["missed"], [])
        self.assertEqual((res["planned"], res["covered"]), (2, 2))      # distinct chapters
        self.assertEqual((res["slot_planned"], res["slot_covered"]), (3, 2))  # section-slots


class SectionTopUpTest(TestCase):
    """A uniform-marks section that comes back short on count is filled by ONE focused top-up
    call, strictly additively (only valid, correctly-typed, non-duplicate questions, never
    more than the shortfall)."""

    # Genuinely distinct question stems so concept-overlap dedup doesn't falsely reject them.
    EXISTING = ["Atomic radius decreases across a period left to right.",
                "Noble gases have completely filled valence shells."]
    FRESH = ["Methane exhibits a tetrahedral molecular geometry.",
             "Sodium chloride forms via ionic bonding between Na and Cl.",
             "Benzene shows delocalised pi electron resonance stability.",
             "Hydrogen bonding raises the boiling point of water."]

    def _wo(self, n, types=("MCQ",)):
        from core.section_generator import SectionWorkOrder
        return SectionWorkOrder(
            section_name="Section A", section_id="A", title="MCQ", marks=n,
            questions_count=n, marks_per_question=1, question_types=list(types),
            instructions=[], constraints={}, context_text="", difficulty="medium",
            subject="Chemistry", class_name="11", chapters=["Structure of Atom", "Chemical Bonding"],
        )

    @staticmethod
    def _mcq(text, chapter="Chemical Bonding"):
        return {"type": "MCQ", "text": text, "options": ["a) 1", "b) 2", "c) 3", "d) 4"],
                "answer": "a", "marks": 1, "chapter_tag": chapter}

    def _existing(self, n):
        return [self._mcq(t, "Structure of Atom") for t in self.EXISTING[:n]]

    def test_topup_fills_missing_mcqs(self):
        import json
        from unittest import mock
        from core import section_generator as sg
        wo = self._wo(4)
        section_data = {"questions": self._existing(2)}
        fake = json.dumps({"questions": [self._mcq(t) for t in self.FRESH[:2]]})
        with mock.patch.object(sg.mantle_client, "converse", return_value=(fake, 10, 20)):
            tokens = sg._top_up_short_section(section_data, wo)
        self.assertEqual(len(section_data["questions"]), 4)
        self.assertEqual(section_data.get("_topped_up"), 2)
        self.assertEqual(tokens, (10, 20))

    def test_topup_fills_mcq_plus_assertion_reason_section(self):
        """The reported case: a 16-q 'MCQ + Assertion-Reason' Section A short by 1 — the old
        single-category guard skipped it. An AR question (type MCQ, subtype assertion_reason)
        is a valid fill for the allowed set."""
        import json
        from unittest import mock
        from core import section_generator as sg
        wo = self._wo(3, types=("MCQ", "Assertion-Reason"))
        section_data = {"questions": self._existing(2)}
        ar = {"type": "MCQ", "subtype": "assertion_reason",
              "text": "Assertion (A): Oxygen is paramagnetic.\nReason (R): It has two unpaired electrons.",
              "options": ["a) ...", "b) ...", "c) ...", "d) ..."], "answer": "a",
              "marks": 1, "chapter_tag": "Chemical Bonding"}
        with mock.patch.object(sg.mantle_client, "converse", return_value=(json.dumps({"questions": [ar]}), 5, 5)):
            sg._top_up_short_section(section_data, wo)
        self.assertEqual(len(section_data["questions"]), 3)
        self.assertEqual(section_data.get("_topped_up"), 1)

    def test_topup_rejects_near_duplicate(self):
        """A topped-up question that echoes an existing one is dropped (no V5 chain runs on it)."""
        import json
        from unittest import mock
        from core import section_generator as sg
        wo = self._wo(3)
        section_data = {"questions": self._existing(2)}
        dup = self._mcq(self.EXISTING[0])      # verbatim echo of an existing question
        with mock.patch.object(sg.mantle_client, "converse", return_value=(json.dumps({"questions": [dup]}), 1, 1)):
            sg._top_up_short_section(section_data, wo)
        self.assertEqual(len(section_data["questions"]), 2)        # duplicate rejected
        self.assertNotIn("_topped_up", section_data)

    def test_topup_never_exceeds_shortfall(self):
        import json
        from unittest import mock
        from core import section_generator as sg
        wo = self._wo(3)                                   # need 3, have 2 → only 1 missing
        section_data = {"questions": self._existing(2)}
        fake = json.dumps({"questions": [self._mcq(t) for t in self.FRESH[:4]]})  # over-delivers
        with mock.patch.object(sg.mantle_client, "converse", return_value=(fake, 1, 1)):
            sg._top_up_short_section(section_data, wo)
        self.assertEqual(len(section_data["questions"]), 3)

    def test_topup_noop_when_section_complete(self):
        from core import section_generator as sg
        wo = self._wo(2)
        section_data = {"questions": self._existing(2)}
        self.assertEqual(sg._top_up_short_section(section_data, wo), (0, 0))
        self.assertNotIn("_topped_up", section_data)

    def test_topup_fills_free_form_type_section(self):
        """AI-generated patterns declare descriptive types ('3-5 Mark Questions', 'Essay',
        'Letter Writing') that classify to 'other', so the old guard skipped the top-up and the
        section shipped short (the reported Tamil paper: SA 27/30, LA 15/20, etc.). A plain
        uniform-marks section with free-form types must still be filled (loose mode)."""
        import json
        from unittest import mock
        from core import section_generator as sg
        wo = self._wo(4, types=("3–5 Mark Questions",))      # descriptive label → 'other'
        sa = lambda t: {"type": "SA", "text": t, "marks": 1, "chapter_tag": "Chemical Bonding"}
        section_data = {"questions": [sa(t) for t in self.EXISTING[:2]]}
        fresh = json.dumps({"questions": [sa(t) for t in self.FRESH[:2]]})
        with mock.patch.object(sg.mantle_client, "converse", return_value=(fresh, 7, 9)):
            tokens = sg._top_up_short_section(section_data, wo)
        self.assertEqual(len(section_data["questions"]), 4)        # filled 2/2 missing
        self.assertEqual(section_data.get("_topped_up"), 2)
        self.assertEqual(tokens, (7, 9))
        self.assertTrue(all(q["marks"] == 1 for q in section_data["questions"]))

    def test_topup_freeform_still_rejects_structural_types(self):
        """Loose mode accepts any written-answer question but must NOT pull a CBQ/map question
        into a plain section it can't structurally host."""
        import json
        from unittest import mock
        from core import section_generator as sg
        wo = self._wo(3, types=("Essay",))
        sa = lambda t: {"type": "SA", "text": t, "marks": 1, "chapter_tag": "Chemical Bonding"}
        section_data = {"questions": [sa(t) for t in self.EXISTING[:2]]}
        cbq = {"type": "CBQ", "text": "Read the passage and answer.", "marks": 1,
               "sub_questions": [{"text": "x", "marks": 1}], "chapter_tag": "Chemical Bonding"}
        with mock.patch.object(sg.mantle_client, "converse", return_value=(json.dumps({"questions": [cbq]}), 1, 1)):
            sg._top_up_short_section(section_data, wo)
        self.assertEqual(len(section_data["questions"]), 2)        # CBQ rejected, nothing added
        self.assertNotIn("_topped_up", section_data)


class PatternAIEnhanceTest(TestCase):
    """validate_and_enhance_pattern: a section's own marks already cover its subsections, so the
    paper total must NOT add both (a subsectioned Hindi board paper summed to 146 instead of 80)."""

    def test_subsection_marks_not_double_counted(self):
        from core import pattern_ai_generator as pag
        raw = {"sections": [
            {"id": "A", "name": "Unseen", "marks": 14, "questions_count": 2, "marks_per_question": 7,
             "question_types": ["Comprehension"], "subsections": []},
            {"id": "B", "name": "Grammar", "marks": 16, "questions_count": 4, "marks_per_question": 4,
             "question_types": ["Short Answer"],
             "subsections": [{"marks": 4, "questions_count": 4} for _ in range(4)]},
            {"id": "D", "name": "Writing", "marks": 22, "questions_count": 5, "marks_per_question": None,
             "question_types": ["Paragraph", "Letter"],
             "subsections": [{"marks": m, "questions_count": 1} for m in (5, 5, 4, 3, 5)]},
        ]}
        out = pag.validate_and_enhance_pattern(raw, "10", "Hindi", "Board")
        self.assertEqual(out["total_marks"], 52)          # 14 + 16 + 22, subsections NOT re-added
        self.assertEqual(out["total_questions"], 11)      # 2 + 4 + 5
        # None marks_per_question is coerced to the section average, never left None.
        self.assertIsNotNone(out["sections"][2]["marks_per_question"])

    def test_section_marks_fallback_to_subsection_sum(self):
        from core import pattern_ai_generator as pag
        raw = {"sections": [
            {"id": "C", "name": "Textbook", "marks": 0, "questions_count": 0,
             "question_types": ["MCQ"],
             "subsections": [{"marks": 5, "questions_count": 5}, {"marks": 6, "questions_count": 3}]},
        ]}
        out = pag.validate_and_enhance_pattern(raw, "10", "Hindi", "Board")
        # Section declared 0 → fall back to the subsection sum (11), don't leave it empty.
        self.assertEqual(out["total_marks"], 11)
        self.assertEqual(out["sections"][0]["marks"], 11)
        self.assertEqual(out["sections"][0]["questions_count"], 8)


class LanguageSupportTest(TestCase):
    """Tamil & Hindi: the right complex-script font is applied at render, and the generation
    prompt instructs the model to write in the target language/script."""

    def test_script_font_picker(self):
        from core.generator import _pick_script_font
        self.assertEqual(_pick_script_font("Hindi Core", []), "Nirmala UI")
        self.assertEqual(_pick_script_font("Sanskrit", []), "Nirmala UI")
        self.assertEqual(_pick_script_font("Tamil", []), "Latha")
        self.assertIsNone(_pick_script_font("Science", [("q", "What is photosynthesis?")]))
        # Detected from the text even when the subject name is non-obvious.
        self.assertEqual(_pick_script_font("X", [("q", "जल का सूत्र?")]), "Nirmala UI")
        self.assertEqual(_pick_script_font("X", [("q", "இது ஒரு வினா")]), "Latha")

    def test_devanagari_run_gets_complex_script_font(self):
        import re
        from docx import Document
        from docx.oxml.ns import qn
        from core.generator import _add_question_with_marks
        doc = Document()
        _add_question_with_marks(doc, "1. जल का रासायनिक सूत्र क्या है? [1 marks]",
                                 re.compile(r"\s*\[(\d+)\s*marks?\]", re.I), None, "Nirmala UI")
        run = doc.paragraphs[-1].runs[-1]
        rFonts = run._element.get_or_add_rPr().find(qn("w:rFonts"))
        self.assertIsNotNone(rFonts)
        self.assertEqual(rFonts.get(qn("w:cs")), "Nirmala UI")

    def test_language_directive_for_language_subjects(self):
        from core.section_generator import _language_directive
        self.assertIn("Devanagari", _language_directive("Hindi Course B"))
        self.assertIn("Tamil", _language_directive("Tamil"))
        self.assertEqual(_language_directive("Science"), "")
        self.assertEqual(_language_directive("Mathematics"), "")

    def test_tamil_conventions_scoped_to_tamil_only(self):
        from core.section_generator import _language_directive
        tamil = _language_directive("Tamil")
        # Tamil papers carry the Samacheer-Kalvi question-type conventions...
        self.assertIn("TAMIL EXAM CONVENTIONS", tamil)
        self.assertIn("இலக்கணக் குறிப்பு தருக", tamil)
        self.assertIn("சீர்களால் நிரப்புக", tamil)
        # ...but Hindi/Sanskrit language papers must NOT inherit them.
        self.assertNotIn("TAMIL EXAM CONVENTIONS", _language_directive("Hindi Course B"))
        self.assertNotIn("TAMIL EXAM CONVENTIONS", _language_directive("Sanskrit"))

    def test_language_directive_injected_into_section_prompt(self):
        from core.section_generator import build_section_prompt, SectionWorkOrder
        wo = SectionWorkOrder(
            section_name="Section A", section_id="A", title="MCQ", marks=5, questions_count=5,
            marks_per_question=1, question_types=["MCQ"], instructions=[], constraints={},
            context_text="", difficulty="medium", subject="Hindi", class_name="10",
            chapters=["कविता"],
        )
        prompt = build_section_prompt(wo)
        self.assertIn("LANGUAGE — MANDATORY", prompt)
        self.assertIn("Devanagari", prompt)


class MatchingQuestionTest(SimpleTestCase):
    """Match-the-following: correct subtype, side-by-side (table) rendering, and the four
    a/b/c/d pairing choices the question is answered with."""

    # The shape the base model returns for a "matching" slot when it ignores the table
    # instruction: VSA / standard with both columns stacked into one text field by newlines.
    # Current label convention — Column I "(A)…", Column II "(1)…" — 4 pairs, scrambled.
    STACKED = (
        "Match the following and choose the correct option:\n"
        "(A) Chacha\n(B) Bua\n(C) Mausi\n(D) Mama\n"
        "(1) Mother's sister\n(2) Father's younger brother\n"
        "(3) Mother's brother\n(4) Father's sister"
    )
    KEY = "A-2, B-4, C-1, D-3"

    # Pre-2026 papers stacked roman "(i)" against lettered "(A)" — still parsed.
    LEGACY_STACKED = (
        "Match the following family values with their meanings:\n"
        "(i) Sevā\n(ii) Dāna\n(iii) Tyāga\n"
        "(A) Sacrifice\n(B) Selfless Service\n(C) Giving"
    )

    def _table(self):
        from core.section_generator import _matching_to_markdown
        return _matching_to_markdown(self.STACKED)

    def test_split_columns(self):
        from core.section_generator import _split_match_columns
        left, right = _split_match_columns(self.STACKED)
        self.assertEqual([l[0] for l in left], ["(A)", "(B)", "(C)", "(D)"])
        self.assertEqual([r[0] for r in right], ["(1)", "(2)", "(3)", "(4)"])
        self.assertEqual(left[0][1], "Chacha")
        self.assertEqual(right[1][1], "Father's younger brother")

    def test_split_columns_legacy_roman_convention(self):
        from core.section_generator import _split_match_columns
        left, right = _split_match_columns(self.LEGACY_STACKED)
        self.assertEqual([l[0] for l in left], ["(i)", "(ii)", "(iii)"])
        self.assertEqual([r[0] for r in right], ["(A)", "(B)", "(C)"])

    def test_to_markdown_builds_two_column_table(self):
        md = self._table()
        self.assertIn("| Column I | Column II |", md)
        self.assertIn("| --- | --- |", md)
        self.assertIn("| (A) Chacha | (1) Mother's sister |", md)
        # The stem is preserved above the table; the columns are no longer newline-stacked.
        self.assertTrue(md.startswith("Match the following and choose the correct option:"))
        self.assertNotIn("(A) Chacha\n(B)", md)

    def test_idempotent_on_existing_table(self):
        from core.section_generator import _matching_to_markdown
        once = self._table()
        self.assertEqual(_matching_to_markdown(once), once)

    def test_table_labels_read_back(self):
        from core.section_generator import _match_table_labels
        left, right = _match_table_labels(self._table())
        self.assertEqual(left, ["A", "B", "C", "D"])
        self.assertEqual(right, ["1", "2", "3", "4"])

    def test_table_labels_ignore_header_and_separator(self):
        from core.section_generator import _match_table_labels
        left, _ = _match_table_labels(
            "| Column I | Column II |\n| --- | --- |\n| (A) x | (1) y |\n| (B) p | (2) q |")
        self.assertEqual(left, ["A", "B"])

    def test_table_labels_dont_mistake_an_abbreviation_for_a_label(self):
        # An unlabelled cell must stay unlabelled — "Dr." / "The" are not column labels.
        from core.section_generator import _match_cell_label
        self.assertIsNone(_match_cell_label("Dr. B. R. Ambedkar"))
        self.assertIsNone(_match_cell_label("The capital of India"))
        self.assertIsNone(_match_cell_label("a small village in Bengal"))
        self.assertEqual(_match_cell_label("(A) Chacha"), "A")
        self.assertEqual(_match_cell_label("2. Father's sister"), "2")
        self.assertEqual(_match_cell_label("B  Statue of Liberty"), "B")

    def test_repair_retags_a_table_the_model_left_as_standard(self):
        # The model laid out the table correctly but forgot subtype="matching"; without the
        # retag the question would be validated as a plain VSA and ship with no options.
        from core.section_generator import _repair_section_data
        data = {"questions": [{
            "qnum": 6, "type": "VSA", "subtype": "standard", "text": self._table(),
            "answer_explanation": self.KEY, "marks": 1.0, "options": {},
        }]}
        _repair_section_data(data)
        q = data["questions"][0]
        self.assertEqual(q["subtype"], "matching")
        self.assertEqual(sorted(q["options"]), ["a", "b", "c", "d"])

    def test_parse_match_key_accepts_the_usual_notations(self):
        from core.section_generator import _parse_match_key
        self.assertEqual(_parse_match_key("A-2, B-4, C-1, D-3"),
                         {"A": "2", "B": "4", "C": "1", "D": "3"})
        self.assertEqual(_parse_match_key("A → 2; B: 4"), {"A": "2", "B": "4"})
        self.assertEqual(_parse_match_key("(i)-B, (ii)-C"), {"I": "B", "II": "C"})

    def test_repair_retags_reformats_and_builds_four_options(self):
        from core.section_generator import _repair_section_data, _parse_match_key
        data = {"questions": [{
            "qnum": 6, "type": "VSA", "subtype": "standard", "text": self.STACKED,
            "answer_explanation": self.KEY, "marks": 1.0, "options": {},
        }]}
        _repair_section_data(data)
        q = data["questions"][0]
        self.assertEqual(q["subtype"], "matching")           # type now identifies a match question
        self.assertIn("| Column I | Column II |", q["text"])  # columns are a table → side by side
        # Four distinct complete pairings, and 'answer' points at the one holding the key.
        self.assertEqual(sorted(q["options"]), ["a", "b", "c", "d"])
        self.assertEqual(len({v for v in q["options"].values()}), 4)
        for v in q["options"].values():
            self.assertEqual(sorted(_parse_match_key(v)), ["A", "B", "C", "D"])
        self.assertIn(q["answer"], ("a", "b", "c", "d"))
        self.assertEqual(_parse_match_key(q["options"][q["answer"]]),
                         _parse_match_key(self.KEY))

    def test_repair_is_deterministic(self):
        from core.section_generator import _repair_section_data
        def _built():
            data = {"questions": [{
                "qnum": 6, "type": "VSA", "subtype": "matching", "text": self.STACKED,
                "answer_explanation": self.KEY, "marks": 1.0, "options": {},
            }]}
            _repair_section_data(data)
            return data["questions"][0]["options"], data["questions"][0]["answer"]
        self.assertEqual(_built(), _built())

    def test_repair_fixes_a_mispointed_answer_letter(self):
        from core.section_generator import _repair_section_data
        data = {"questions": [{
            "qnum": 3, "type": "VSA", "subtype": "matching", "text": self.STACKED,
            "answer_explanation": self.KEY, "marks": 1.0, "answer": "a",
            "options": {"a": "A-1, B-2, C-3, D-4", "b": self.KEY,
                        "c": "A-3, B-1, C-4, D-2", "d": "A-4, B-3, C-2, D-1"},
        }]}
        _repair_section_data(data)
        q = data["questions"][0]
        self.assertEqual(q["answer"], "b")                  # the option that holds the key
        self.assertEqual(q["options"]["a"], "A-1, B-2, C-3, D-4")   # model's own options kept

    def test_repair_leaves_options_alone_without_a_usable_key(self):
        # Prose instead of a pairing → nothing to build from; validation asks for the retry.
        from core.section_generator import _repair_section_data
        data = {"questions": [{
            "qnum": 6, "type": "VSA", "subtype": "matching", "text": self.STACKED,
            "answer_explanation": "Chacha is the father's younger brother.",
            "marks": 1.0, "options": {},
        }]}
        _repair_section_data(data)
        self.assertEqual(data["questions"][0]["options"], {})

    def test_render_detects_the_table(self):
        # The rewritten body must be recognised by the DOCX table splitter so it renders
        # as a real two-column Word table rather than stacked pipe text.
        from core.generator import _md_table_segments
        segs = _md_table_segments(self._table())
        self.assertIsNotNone(segs)
        self.assertTrue(any(kind == "table" for kind, _ in segs))

    def test_options_survive_the_renderer(self):
        # process_question must print the pairing choices as an option block — a match
        # question with no printed options is unanswerable.
        from core.generator import process_question
        out = []
        process_question(out, {
            "qnum": 6, "type": "VSA", "subtype": "matching", "text": self._table(),
            "options": {"a": "A-1, B-2, C-3, D-4", "b": self.KEY,
                        "c": "A-3, B-1, C-4, D-2", "d": "A-4, B-3, C-2, D-1"},
            "answer": "b", "answer_explanation": self.KEY, "marks": 1.0,
        }, 6)
        blocks = [payload for kind, payload in out if kind == "opts_block"]
        self.assertTrue(blocks, out)
        self.assertEqual(len(blocks[0]), 4)
        self.assertIn("(b) A-2, B-4, C-1, D-3", blocks[0])

    def test_non_matching_vsa_untouched(self):
        from core.section_generator import _repair_section_data
        data = {"questions": [{
            "qnum": 1, "type": "VSA", "subtype": "standard",
            "text": "Define secularism in one sentence.",
            "answer_explanation": "…", "marks": 1.0, "options": {},
        }]}
        _repair_section_data(data)
        q = data["questions"][0]
        self.assertEqual(q["subtype"], "standard")
        self.assertNotIn("|", q["text"])


class MatchingValidationTest(SimpleTestCase):
    """A match question must ship 4+ pairs and 4 complete pairing choices — a 3-pair match
    or a bare table with no options is rejected so the section regenerates."""

    def _wo(self):
        from core import section_generator as sg
        return sg.SectionWorkOrder(
            section_name="Section A", section_id="A", title="Objective", marks=1,
            questions_count=1, marks_per_question=1.0, question_types=[
                {"type": "VSA", "count": 1, "marks_each": 1.0}],
            instructions=[], constraints={}, context_text="ctx", difficulty="medium",
            subject="Social Science", class_name="10", chapters=["Nationalism in Europe"],
            slots=[{"qnum": 1, "type": "matching", "marks": 1.0}])

    def _q(self, pairs=4, **over):
        rows = [f"| ({chr(64 + i)}) Column-I entry {i} | ({i}) match {i} |"
                for i in range(1, pairs + 1)]
        labels = [chr(64 + i) for i in range(1, pairs + 1)]
        def _key(vals):
            return ", ".join(f"{l}-{v}" for l, v in zip(labels, vals))
        vals = list(range(1, pairs + 1))
        q = {
            "qnum": 1, "type": "VSA", "subtype": "matching", "marks": 1.0,
            "text": "Match the following and choose the correct option:\n"
                    "| Column I | Column II |\n| --- | --- |\n" + "\n".join(rows),
            "options": {"a": _key(vals),
                        "b": _key(vals[1:] + vals[:1]),
                        "c": _key(vals[::-1]),
                        "d": _key(vals[2:] + vals[:2])},
            "answer": "a", "answer_explanation": _key(vals),
            "competency_type": "recall",
        }
        q.update(over)
        return q

    def _errs(self, q):
        from core import section_generator as sg
        return [e for e in sg.validate_section_output({"questions": [q]}, self._wo())
                if "VSA/matching" in e]

    def test_well_formed_four_pair_match_passes(self):
        self.assertEqual(self._errs(self._q()), [])

    def test_three_pairs_rejected(self):
        errs = self._errs(self._q(pairs=3))
        self.assertTrue(any("AT LEAST 4 pairs" in e for e in errs), errs)

    def test_missing_options_rejected(self):
        errs = self._errs(self._q(options={}))
        self.assertTrue(any("EXACTLY 4 options" in e for e in errs), errs)

    def test_partial_pairing_option_rejected(self):
        q = self._q()
        q["options"]["c"] = "A-2, B-1"          # only two of the four Column I entries
        errs = self._errs(q)
        self.assertTrue(any("do not pair EVERY Column I entry" in e for e in errs), errs)

    def test_duplicate_pairings_rejected(self):
        q = self._q()
        q["options"]["d"] = q["options"]["b"]
        errs = self._errs(q)
        self.assertTrue(any("4 DIFFERENT pairings" in e for e in errs), errs)

    def test_missing_answer_letter_rejected(self):
        errs = self._errs(self._q(answer=""))
        self.assertTrue(any("'answer' must be the correct option letter" in e for e in errs), errs)

    def test_stacked_lists_instead_of_a_table_rejected(self):
        errs = self._errs(self._q(text=MatchingQuestionTest.STACKED))
        self.assertTrue(any("two-column Markdown table" in e for e in errs), errs)

    def test_prompt_demands_four_pairs_and_four_options(self):
        from core import section_generator as sg
        prompt = sg.build_section_prompt(self._wo())
        self.assertIn("MATCH THE FOLLOWING", prompt)
        self.assertIn("EXACTLY 4 pairs", prompt)
        self.assertIn("A-3, B-1, C-4, D-2", prompt)          # option format shown
        self.assertIn('"subtype": "matching"', prompt)        # JSON example present


class MaterialIntelTest(TestCase):
    """Intelligent ingestion: chapter-name cleaning, catalog snapping, book page-range splitting,
    and LLM-based unit naming."""

    def test_clean_name_strips_chapter_prefix(self):
        from core.material_intel import _clean_name
        self.assertEqual(_clean_name("Chapter 4 - Chemical Bonding"), "Chemical Bonding")
        self.assertEqual(_clean_name("UNIT 2: Structure of Atom"), "Structure of Atom")
        self.assertEqual(_clean_name("  Thermodynamics  "), "Thermodynamics")

    def test_snap_to_catalog(self):
        from core.material_intel import _snap_to_catalog
        # Physics has a catalog → a close/substring name snaps to the official one.
        self.assertEqual(_snap_to_catalog("current electricity", "Physics"), "Current Electricity")
        # No catalog for the subject → returned unchanged.
        self.assertEqual(_snap_to_catalog("My Custom Topic", "Hindi"), "My Custom Topic")

    def test_ranges_from_starts(self):
        from core.material_intel import _ranges_from_starts
        chapters = _ranges_from_starts([("Intro", 0), ("Atoms", 5), ("Bonds", 12)], 20)
        self.assertEqual(len(chapters), 3)
        self.assertEqual((chapters[0]["start_page"], chapters[0]["end_page"]), (0, 5))
        self.assertEqual((chapters[2]["start_page"], chapters[2]["end_page"]), (12, 20))
        self.assertEqual(chapters[1]["unit"], "Atoms")

    def test_detect_unit_name_snaps_to_catalog(self):
        from unittest import mock
        from core import material_intel as mi
        with mock.patch.object(mi.mantle_client, "converse",
                               return_value=('{"chapter": "current electricity"}', 1, 1)):
            name = mi.detect_unit_name(None, "12", "Physics", sample_text="...physics text...")
        self.assertEqual(name, "Current Electricity")

    def test_detect_unit_name_empty_sample_returns_none(self):
        from core import material_intel as mi
        self.assertIsNone(mi.detect_unit_name(None, "12", "Physics", sample_text="   "))

    def test_legacy_font_gibberish_is_unreadable(self):
        """A Hindi PDF typeset in a legacy non-Unicode font (Walkman-Chanakya905) extracts as
        ASCII keystroke gibberish with no Devanagari — must be flagged unreadable so naming falls
        back to the filename instead of returning confident garbage ('dchj')."""
        from core import material_intel as mi
        chanakya = "i| [kaM ,d jk\"Vªh; vfLerk vkSj jk\"Vªh; pfj=k dk fodkl"  # real Chanakya extraction
        self.assertTrue(mi.text_is_unreadable_for_subject(chanakya, "Hindi"))
        # And detect_unit_name bails to None on it (no LLM call reached).
        self.assertIsNone(mi.detect_unit_name(None, "10", "Hindi", sample_text=chanakya))

    def test_real_devanagari_is_readable(self):
        from core import material_intel as mi
        self.assertFalse(mi.text_is_unreadable_for_subject("दो बैलों की कथा — प्रेमचंद की कहानी", "Hindi"))

    def test_latin_subject_never_flagged_unreadable(self):
        from core import material_intel as mi
        # English/maths content is ASCII by nature — the guard applies only to Indic-script subjects.
        self.assertFalse(mi.text_is_unreadable_for_subject("Chapter 1: Real Numbers", "Mathematics"))
        self.assertFalse(mi.text_is_unreadable_for_subject("The Sound of Music", "English"))

    def test_extract_html_chapters_by_heading(self):
        from core.material_intel import extract_html_chapters
        html = (
            "<html><body><h1>Book Title</h1>"
            "<h2>இயல் ஒன்று</h2>"
            "<p>தமிழ் உரை ஒன்று. இது முதல் பாடத்தின் முழுமையான உரை ஆகும் இங்கே.</p>"
            "<h2>இயல் இரண்டு</h2>"
            "<p>இரண்டாம் பாடத்தின் உரை. போதுமான நீளம் கொண்ட உரை இங்கே உள்ளது.</p>"
            "</body></html>"
        )
        chs = extract_html_chapters(html, "Tamil")
        units = [c["unit"] for c in chs]
        self.assertIn("இயல் ஒன்று", units)
        self.assertIn("இயல் இரண்டு", units)
        # Body text is captured; the <h1> book title before the first chapter is excluded.
        self.assertTrue(all(c["text"] and "Book Title" not in c["text"] for c in chs))

    def test_extract_html_single_chapter_when_no_headings(self):
        from core.material_intel import extract_html_chapters
        html = "<html><body><p>" + ("உரை " * 40) + "</p></body></html>"
        chs = extract_html_chapters(html, "Tamil")
        self.assertEqual(len(chs), 1)
        self.assertIsNone(chs[0]["unit"])   # caller will name it

    def test_preview_url_endpoint_returns_chapters_without_ingesting(self):
        from unittest import mock
        from rest_framework.test import APIClient
        u = User.objects.create_user("previewer", "p@x.com", "pw")
        api = APIClient()
        api.force_authenticate(u)
        html = ("<html><body>"
                "<h2>இயல் ஒன்று</h2><p>" + ("உரை " * 30) + "</p>"
                "<h2>இயல் இரண்டு</h2><p>" + ("உரை " * 30) + "</p></body></html>")
        with mock.patch("core.material_intel.fetch_url", return_value=html):
            r = api.post('/api/materials/preview-url/',
                         {'url': 'https://example.com/book.html', 'subject': 'Tamil'}, format='json')
        self.assertEqual(r.status_code, 200)
        body = r.json()
        self.assertEqual(body['count'], 2)
        self.assertIn("இயல் ஒன்று", [c['unit'] for c in body['chapters']])


class ImageStemGuardTest(TestCase):
    """A generic 'observe the diagram and answer' stem must NOT be turned into an image prompt
    (that produced a second, irrelevant image — e.g. an org chart on a chemistry CBQ)."""

    def test_generic_pointer_stems_are_flagged(self):
        from core.generator import _is_generic_image_stem
        for s in [
            "Observe the diagram carefully and answer the following questions:",
            "Study the figure below and answer:",
            "Refer to the graph shown and identify the trend.",
            "Based on the diagram, explain the periodic trend.",
            "Read the source/case and answer the questions that follow:",
        ]:
            self.assertTrue(_is_generic_image_stem(s), f"should be generic: {s!r}")

    def test_real_descriptions_are_not_flagged(self):
        from core.generator import _is_generic_image_stem
        for s in [
            "A line graph with Atomic Number on the X-axis and Atomic Radius on the Y-axis showing a decreasing trend.",
            "The Lewis structure of carbon dioxide with two double bonds drawn around the central atom.",
            "A titration apparatus with a burette clamped above a conical flask containing an indicator.",
        ]:
            self.assertFalse(_is_generic_image_stem(s), f"should NOT be generic: {s!r}")


class RerenderImageCacheTest(TestCase):
    """cache_only skips uncached images by default, but rerender can opt in to
    generate a missing image on demand."""

    @staticmethod
    def _boom(*a, **k):
        raise AssertionError("external image API was called in cache_only mode")

    def test_cache_only_raises_without_any_network_call(self):
        from unittest import mock
        from core import generator as g
        with mock.patch.object(g, "_openai_image_bytes", self._boom), \
             mock.patch.object(g, "_together_image_bytes", self._boom), \
             mock.patch.object(g, "_pollinations_image_bytes", self._boom):
            with self.assertRaises(g.ImageNotCached):
                g.generate_ai_image("a unique uncached prompt 12345", cache_only=True)

    def test_materialize_images_cache_only_skips_uncached_image(self):
        from unittest import mock
        from core import generator as g
        items = [("q", "1. A question"), ("image_gen", "another unique uncached prompt 67890")]
        with mock.patch.object(g, "_openai_image_bytes", self._boom), \
             mock.patch.object(g, "_together_image_bytes", self._boom), \
             mock.patch.object(g, "_pollinations_image_bytes", self._boom):
            out = g.materialize_images(items, allow=True, cache_only=True)
        self.assertIn(("q", "1. A question"), out)            # text preserved
        self.assertFalse(any(t == "image" for t, _ in out))   # uncached image dropped, not hung on

    def test_materialize_images_can_generate_missing_image_on_rerender(self):
        from unittest import mock
        from core import generator as g

        items = [("q", "1. A question"), ("image_gen", "another unique uncached prompt 67890")]
        calls = []

        def fake_generate(prompt, cache_only=False, **kwargs):
            calls.append((prompt, cache_only))
            if cache_only:
                raise g.ImageNotCached(prompt)
            return "generated_images/on-demand.png"

        with mock.patch.object(g, "generate_ai_image", side_effect=fake_generate):
            out = g.materialize_images(
                items,
                allow=True,
                cache_only=True,
                generate_missing_images=True,
            )

        self.assertIn(("q", "1. A question"), out)
        self.assertIn(("image", "generated_images/on-demand.png"), out)
        self.assertEqual(calls, [
            ("another unique uncached prompt 67890", True),
            ("another unique uncached prompt 67890", False),
        ])


class PaperMarksAuditTest(TestCase):
    """audit_paper_marks: per-question marks (OR pair counted once) must add up to the pattern."""

    class _Pat:
        def __init__(self, sections, total_marks):
            self.sections = sections
            self.total_marks = total_marks

    def _pat(self):
        secs = [
            {"name": "Section A", "marks": 4, "questions_count": 4, "marks_per_question": 1.0},
            {"name": "Section E", "marks": 10, "questions_count": 2, "marks_per_question": 5.0},
        ]
        return self._Pat(secs, 14)

    def test_clean_paper_passes(self):
        from core.paper_audit import audit_paper_marks
        pd = {
            "Section A": {"marks": 4, "questions": [{"marks": 1} for _ in range(4)]},
            "Section E": {"marks": 10, "questions": [{"marks": 5}, {"marks": 5}]},
        }
        res = audit_paper_marks(pd, self._pat())
        self.assertTrue(res["ok"], res["issues"])
        self.assertEqual(res["actual_total"], 14)

    def test_under_produced_section_is_flagged(self):
        from core.paper_audit import audit_paper_marks
        pd = {  # Section A only has 2 of 4 questions
            "Section A": {"marks": 4, "questions": [{"marks": 1}, {"marks": 1}]},
            "Section E": {"marks": 10, "questions": [{"marks": 5}, {"marks": 5}]},
        }
        res = audit_paper_marks(pd, self._pat())
        self.assertFalse(res["ok"])
        self.assertEqual(res["actual_total"], 12)
        self.assertTrue(any("Section A" in i and "questions" in i for i in res["issues"]))

    def test_over_marked_section_is_flagged(self):
        from core.paper_audit import audit_paper_marks
        pd = {  # right count (4) but inflated marks (2 each → 8, expected 4)
            "Section A": {"marks": 4, "questions": [{"marks": 2} for _ in range(4)]},
            "Section E": {"marks": 10, "questions": [{"marks": 5}, {"marks": 5}]},
        }
        res = audit_paper_marks(pd, self._pat())
        self.assertFalse(res["ok"])
        self.assertTrue(any("Section A" in i and "marks" in i for i in res["issues"]))

    def test_or_alternative_field_counts_once(self):
        from core.paper_audit import audit_paper_marks
        pd = {  # Section E questions carry an OR alternative as a FIELD — still 5m each, not 10
            "Section A": {"marks": 4, "questions": [{"marks": 1} for _ in range(4)]},
            "Section E": {"marks": 10, "questions": [
                {"marks": 5, "or_alternative": "Alternative question text"},
                {"marks": 5, "or_alternative": {"text": "Alt"}},
            ]},
        }
        res = audit_paper_marks(pd, self._pat())
        self.assertEqual(res["actual_total"], 14)
        self.assertTrue(res["ok"], res["issues"])

    def test_or_alternative_as_separate_entry_is_excluded(self):
        from core.paper_audit import audit_paper_marks
        pd = {  # an OR alternative leaked in as its own entry + an "OR" marker entry
            "Section A": {"marks": 4, "questions": [{"marks": 1} for _ in range(4)]},
            "Section E": {"marks": 10, "questions": [
                {"marks": 5},
                {"text": "OR"},                              # separator entry → excluded
                {"marks": 5, "is_or_alternative": True},      # alternative entry → excluded
                {"marks": 5},
            ]},
        }
        res = audit_paper_marks(pd, self._pat())
        # Section E should count only the 2 primary 5-mark questions = 10
        self.assertEqual(res["actual_total"], 14)
        self.assertTrue(res["ok"], res["issues"])


class MCQAnswerCorrectionTest(TestCase):
    """V4 auto-corrects a wrong MCQ key only when two independent high-confidence passes agree."""

    def _mcq(self, ans):
        return [{
            "type": "MCQ", "text": "2 + 2 = ?",
            "options": {"a": "3", "b": "4", "c": "5", "d": "6"},
            "answer": ans, "marks": 1,
        }]

    def _reply(self, letter, conf):
        import json as _json
        return (_json.dumps([{"q": 1, "answer": letter, "confidence": conf}]), 0, 0)

    def test_two_high_confidence_passes_correct_the_key(self):
        from unittest import mock
        from core import section_generator as sg
        qs = self._mcq("a")  # stored wrong; verifier insists on 'b' twice, high
        with mock.patch.object(sg.mantle_client, "converse",
                               side_effect=[self._reply("b", "high"), self._reply("b", "high")]):
            res = sg.verify_mcq_answers(qs, "10", "Mathematics")
        self.assertEqual(qs[0]["answer"], "b")          # key fixed in place
        self.assertTrue(res[0]["corrected"])
        self.assertFalse(res[0]["suspect"])             # fixed → not also flagged

    def test_single_disagreement_only_flags_not_corrects(self):
        from unittest import mock
        from core import section_generator as sg
        qs = self._mcq("a")  # first says 'b' high, second pass agrees with stored 'a' → don't flip
        with mock.patch.object(sg.mantle_client, "converse",
                               side_effect=[self._reply("b", "high"), self._reply("a", "high")]):
            res = sg.verify_mcq_answers(qs, "10", "Mathematics")
        self.assertEqual(qs[0]["answer"], "a")          # key unchanged
        self.assertFalse(res[0]["corrected"])
        self.assertTrue(res[0]["suspect"])              # flagged for human review

    def test_medium_confidence_never_corrects(self):
        from unittest import mock
        from core import section_generator as sg
        qs = self._mcq("a")
        with mock.patch.object(sg.mantle_client, "converse",
                               side_effect=[self._reply("b", "medium")]):
            res = sg.verify_mcq_answers(qs, "10", "Mathematics")
        self.assertEqual(qs[0]["answer"], "a")          # medium is never enough to overwrite
        self.assertFalse(res[0]["corrected"])
        self.assertTrue(res[0]["suspect"])


class FirstLoginPasswordTest(TestCase):
    """First-login flow: set a new password OR skip — both clear require_password_change."""

    def setUp(self):
        from rest_framework.test import APIClient
        self.user = User.objects.create_user("newbie", "n@x.com", "temppass123")
        p = self.user.profile
        p.require_password_change = True
        p.save()
        self.api = APIClient()
        self.api.force_authenticate(self.user)

    def test_set_password_clears_flag(self):
        r = self.api.post('/api/auth/first-login-password/', {'new_password': 'brandnew123'}, format='json')
        self.assertEqual(r.status_code, 200)
        self.user.refresh_from_db()
        self.user.profile.refresh_from_db()
        self.assertFalse(self.user.profile.require_password_change)
        self.assertTrue(self.user.check_password('brandnew123'))

    def test_skip_clears_flag_keeps_password(self):
        r = self.api.post('/api/auth/first-login-password/', {'skip': True}, format='json')
        self.assertEqual(r.status_code, 200)
        self.user.profile.refresh_from_db()
        self.assertFalse(self.user.profile.require_password_change)
        self.assertTrue(self.user.check_password('temppass123'))

    def test_set_rejected_when_flag_already_false(self):
        p = self.user.profile
        p.require_password_change = False
        p.save()
        r = self.api.post('/api/auth/first-login-password/', {'new_password': 'brandnew123'}, format='json')
        self.assertEqual(r.status_code, 400)

    def test_short_password_rejected(self):
        r = self.api.post('/api/auth/first-login-password/', {'new_password': 'short'}, format='json')
        self.assertEqual(r.status_code, 400)


class TeamUsagePersistenceTest(TestCase):
    """The team-usage page reads each user's tokens/cost from the UsageEvent log, so the figures
    survive paper deletion (old behaviour: recomputed from live papers → reset to 0 on delete)."""

    def setUp(self):
        from rest_framework.test import APIClient
        self.school = School.objects.create(name="Test School")

        def mk(uname, role):
            u = User.objects.create_user(uname, f"{uname}@x.com", "pw")
            p = u.profile
            p.school = self.school
            p.role = role
            p.save()
            return u

        self.admin = mk("admin1", "school_admin")
        self.teacher = mk("teach1", "teacher")
        self.api = APIClient()
        self.api.force_authenticate(self.admin)

    def _row(self, username):
        r = self.api.get(f'/api/admin/schools/{self.school.id}/user-usage/')
        self.assertEqual(r.status_code, 200)
        return next(x for x in r.json()['users'] if x['username'] == username)

    def test_usage_comes_from_events_not_live_papers(self):
        from core.models import UsageEvent
        UsageEvent.record(user=self.teacher, input_tokens=1000, output_tokens=500, cost=2)
        row = self._row("teach1")
        self.assertEqual(row['total_tokens'], 1500)      # from the event …
        self.assertEqual(float(row['total_cost']), 2.0)
        self.assertEqual(row['total_papers'], 1)
        self.assertEqual(row['current_papers'], 0)       # … even with no surviving paper

    def test_monthly_excludes_prior_months(self):
        from datetime import timedelta
        from django.utils import timezone
        from core.models import UsageEvent
        UsageEvent.record(user=self.teacher, input_tokens=100, cost=1)           # this month
        old = UsageEvent.record(user=self.teacher, input_tokens=900, cost=9)
        UsageEvent.objects.filter(pk=old.pk).update(created_at=timezone.now() - timedelta(days=45))
        row = self._row("teach1")
        self.assertEqual(row['total_tokens'], 1000)      # all-time
        self.assertEqual(row['monthly_tokens'], 100)     # only the current-month event


class HierarchicalVisibilityTest(TestCase):
    """Validates the `_owner_scope` helper used for QuestionPaper visibility: a teacher sees only
    their OWN papers; a school_admin sees the whole school's; superadmin sees all. (Patterns and
    blueprints are intentionally school-wide-shared and do NOT use this helper.) ExamPattern rows
    are used here only as a convenient model with a `created_by` field to exercise the helper."""

    def setUp(self):
        from api.views import _owner_scope
        self._scope = _owner_scope
        self.school = School.objects.create(name="S")
        def mk(uname, role):
            u = User.objects.create_user(uname, f"{uname}@x.com", "pw")
            p = u.profile
            p.school = self.school
            p.role = role
            p.save()
            return u
        self.admin = mk("adm", "school_admin")
        self.t1 = mk("t1", "teacher")
        self.t2 = mk("t2", "teacher")
        secs = [{"name": "A", "marks": 10, "questions_count": 5}]
        self.p_admin = ExamPattern.objects.create(name="pa", sections=secs, created_by=self.admin)
        self.p_t1 = ExamPattern.objects.create(name="p1", sections=secs, created_by=self.t1)
        self.p_t2 = ExamPattern.objects.create(name="p2", sections=secs, created_by=self.t2)

    def test_teacher_sees_only_own(self):
        ids = set(self._scope(ExamPattern.objects.all(), self.t1).values_list("id", flat=True))
        self.assertEqual(ids, {self.p_t1.id})  # not the admin's, not the other teacher's

    def test_admin_sees_whole_school(self):
        ids = set(self._scope(ExamPattern.objects.all(), self.admin).values_list("id", flat=True))
        self.assertEqual(ids, {self.p_admin.id, self.p_t1.id, self.p_t2.id})


class BlueprintIsolationTest(TestCase):
    """A teacher must not see another school's ExamBlueprints (IDOR fix #3)."""

    def setUp(self):
        self.sa = School.objects.create(name="A")
        self.sb = School.objects.create(name="B")
        self.ua = User.objects.create_user("ta", "a@x.com", "pw")
        self.ub = User.objects.create_user("tb", "b@x.com", "pw")
        for u, s in ((self.ua, self.sa), (self.ub, self.sb)):
            p = u.profile          # auto-created by signal
            p.school = s
            p.role = "teacher"
            p.save()
        self.bp_a = ExamBlueprint.objects.create(class_name="10", subject="Science", blueprint={}, created_by=self.ua)
        self.bp_b = ExamBlueprint.objects.create(class_name="10", subject="Science", blueprint={}, created_by=self.ub)

    def test_blueprints_scoped_to_own_school(self):
        ids_a = set(_scoped_blueprints(self.ua).values_list("id", flat=True))
        self.assertIn(self.bp_a.id, ids_a)
        self.assertNotIn(self.bp_b.id, ids_a)   # School A cannot see School B's blueprint


class MCQHardeningTest(TestCase):
    """MCQs must have exactly 4 non-empty a/b/c/d options and a valid answer key."""

    def _opts(self, options, answer="a"):
        return sg._validate_mcq_options({"options": options, "answer": answer}, 1, "MCQ/standard")

    def test_valid_four_pass(self):
        self.assertEqual(self._opts({"a": "1", "b": "2", "c": "3", "d": "4"}), [])

    def test_three_options_flagged(self):
        self.assertTrue(any("EXACTLY 4" in e for e in self._opts({"a": "1", "b": "2", "c": "3"})))

    def test_five_options_flagged(self):
        self.assertTrue(any("EXACTLY 4" in e for e in self._opts({"a": "1", "b": "2", "c": "3", "d": "4", "e": "5"})))

    def test_empty_option_flagged(self):
        self.assertTrue(any("EXACTLY 4" in e for e in self._opts({"a": "1", "b": "  ", "c": "3", "d": "4"})))

    def test_wrong_keys_flagged(self):
        self.assertTrue(any("keys must be" in e for e in self._opts({"a": "1", "b": "2", "c": "3", "e": "4"})))

    def test_missing_answer_flagged(self):
        self.assertTrue(any("answer" in e for e in self._opts({"a": "1", "b": "2", "c": "3", "d": "4"}, answer="")))

    def test_bad_answer_flagged(self):
        self.assertTrue(any("a/b/c/d" in e for e in self._opts({"a": "1", "b": "2", "c": "3", "d": "4"}, answer="z")))

    def test_missing_type_flagged(self):
        pattern = ExamPattern(name="mt", subject="Science", class_name="10", sections=[{
            "name": "A", "marks": 2, "questions_count": 2, "marks_per_question": 1, "question_types": ["MCQ"],
            "subsections": [{"name": "MCQ", "question_types": ["MCQ"], "marks": 2, "questions_count": 2, "marks_per_question": 1}],
        }])
        wo = sg.build_work_orders({s["name"]: s for s in pattern.sections}, pattern, {}, "M", "10", "Science", ["x"])[0]
        errs = sg._validate_by_subtype({"text": "q", "marks": 1}, 1, wo)   # no 'type'
        self.assertTrue(any("missing 'type'" in e for e in errs))


class InlineImageSalvageTest(TestCase):
    """Models sometimes write '(Image description: …)' inline instead of an image_prompt field,
    so no image was generated. The renderer must salvage it into a real image."""

    def test_extracts_neuron_description(self):
        t = ("(Image description: A simple diagram showing a neuron with labelled parts A, B, C. "
             "Part A is a branched structure, B is a long fibre, and C is a gap between two neurons.) "
             "Which label in the diagram shows the synapse?")
        clean, prompt = sg_gen._extract_inline_image(t)
        self.assertEqual(clean, "Which label in the diagram shows the synapse?")
        self.assertIn("neuron", prompt)

    def test_bracketed_and_midsentence(self):
        clean, prompt = sg_gen._extract_inline_image("[Image: bar chart of rainfall] Which city is wettest?")
        self.assertEqual(clean, "Which city is wettest?")
        self.assertIn("bar chart", prompt)

    def test_diagram_placeholder_marker_is_salvaged(self):
        t = ("The diagram below shows a vapour pressure graph.\n\n"
             "[DIAGRAM PLACEHOLDER: A graph with two curves, one labelled Pure solvent and one "
             "labelled Solution, showing boiling point elevation ΔTb.]\n\n"
             "(i) Name the colligative property.")
        clean, prompt = sg_gen._extract_inline_image(t)
        self.assertNotIn("PLACEHOLDER", clean)
        self.assertIn("vapour pressure graph", clean)
        self.assertIn("boiling point elevation", prompt)

    def test_plain_question_untouched(self):
        t = "What is the chemical formula of water?"
        clean, prompt = sg_gen._extract_inline_image(t)
        self.assertEqual(clean, t)
        self.assertIsNone(prompt)

    def test_empty_marker_not_treated_as_image(self):
        clean, prompt = sg_gen._extract_inline_image("Observe the figure (figure) below.")
        self.assertIsNone(prompt)


class DiagramSectionPromptSalvageTest(TestCase):
    """Diagram sections should still render images when the model forgets image_prompt."""

    def test_placeholder_marker_becomes_single_question_image(self):
        out = []
        q = {
            "type": "SA",
            "subtype": "standard",
            "marks": 5,
            "text": (
                "The diagram below shows a vapour pressure graph.\n\n"
                "[DIAGRAM PLACEHOLDER: A graph with two curves, one labelled Pure solvent and "
                "one labelled Solution, showing boiling point elevation ΔTb.]\n\n"
                "(i) Name the colligative property."
            ),
        }
        sg_gen._emit_section_questions(
            out,
            [q],
            {"_section_name": "Diagram Based Questions", "title": ""},
            1,
        )
        self.assertEqual(len([1 for kind, _ in out if kind == "image_gen"]), 1)
        q_lines = [text for kind, text in out if kind == "q"]
        self.assertTrue(q_lines)
        self.assertNotIn("PLACEHOLDER", q_lines[0])

    def test_render_injects_image_prompt_for_descriptive_diagram_question(self):
        out = []
        q = {
            "type": "SA",
            "subtype": "standard",
            "marks": 5,
            "text": (
                "A diagram illustrates the concept of osmosis using a U-tube setup. "
                "The left arm contains a pure solvent, the right arm contains a solution, "
                "and a semi-permeable membrane separates them. "
                "The liquid level in the solution arm is higher.\n\n"
                "(i) Define osmotic pressure."
            ),
        }
        sg_gen._emit_section_questions(
            out,
            [q],
            {"_section_name": "Diagram Based Questions", "title": ""},
            1,
        )
        image_prompts = [text for kind, text in out if kind == "image_gen"]
        self.assertEqual(len(image_prompts), 1)
        self.assertIn("U-tube setup", image_prompts[0])
        self.assertIn("semi-permeable membrane", image_prompts[0])

    def test_generic_diagram_stem_still_does_not_hallucinate_prompt(self):
        out = []
        q = {
            "type": "SA",
            "subtype": "standard",
            "marks": 5,
            "text": "Observe the diagram carefully and answer the following questions.",
        }
        sg_gen._emit_section_questions(
            out,
            [q],
            {"_section_name": "Diagram Based Questions", "title": ""},
            1,
        )
        self.assertFalse(any(kind == "image_gen" for kind, _ in out))


class MissingCountDerivationTest(TestCase):
    """AI-generated patterns sometimes leave questions_count / marks_per_question null with
    only the section marks set. The work-order builder must derive a non-zero count from
    marks + question type, otherwise the section generates nothing and comes out empty."""

    def _pattern(self):
        # Mirrors the real broken pattern 342: only Section A had counts.
        return ExamPattern(name="pt", subject="Mathematics", class_name="6", sections=[
            {"name": "Section A — Objective Type", "marks": 20, "questions_count": 20,
             "marks_per_question": 1, "question_types": ["MCQ", "True/False", "Fill in the Blanks"]},
            {"name": "Section B — Very Short Answer (VSA)", "marks": 12, "questions_count": None,
             "marks_per_question": None, "question_types": ["Very Short Answer"]},
            {"name": "Section C — Short Answer (SA)", "marks": 18, "questions_count": None,
             "marks_per_question": None, "question_types": ["Short Answer"]},
            {"name": "Section D — Long Answer (LA)", "marks": 30, "questions_count": None,
             "marks_per_question": None, "question_types": ["Long Answer"]},
        ])

    def test_typical_marks(self):
        self.assertEqual(sg._typical_marks_for_types(["Very Short Answer"]), 2.0)
        self.assertEqual(sg._typical_marks_for_types(["Short Answer"]), 3.0)
        self.assertEqual(sg._typical_marks_for_types(["Long Answer"]), 5.0)
        self.assertEqual(sg._typical_marks_for_types(["MCQ"]), 1.0)

    def test_work_orders_derive_nonzero_counts(self):
        p = self._pattern()
        bp = {s["name"]: s for s in p.sections}
        wos = sg.build_work_orders(bp, p, {}, "Medium", "6", "Mathematics", ["Patterns"])
        counts = {wo.section_name.split("—")[1].strip()[:3]: wo.questions_count for wo in wos}
        # No section may ask for 0 questions; derived from marks ÷ per-type marks.
        for wo in wos:
            self.assertGreater(wo.questions_count, 0, f"{wo.section_name} got 0 questions")
        self.assertEqual(counts["Obj"], 20)   # A unchanged
        self.assertEqual(counts["Ver"], 6)     # B: 12 / 2
        self.assertEqual(counts["Sho"], 6)     # C: 18 / 3
        self.assertEqual(counts["Lon"], 6)     # D: 30 / 5


class InconsistentCountReconcileTest(TestCase):
    """Pattern 353 (English Core, cls 6, 40m) shipped a 40-mark paper as 89 marks. Its sections
    carried a section-level questions_count that disagreed with their per-type breakdown:
    questions_count=10 for a section whose two subsections describe just 2 questions worth 5m
    each. The generator asked for 10×1m while the per-type validator demanded EXACTLY
    MCQ×1+OTHER×1 at 5m — impossible to satisfy both, so the section shipped partial and blew
    the marks total (a 10-mark section rendered as 38m). The work-order builder must treat the
    per-type breakdown as authoritative for the count and marks."""

    def _pattern(self):
        return ExamPattern(name="p353", subject="English Core", class_name="6", sections=[
            # Reading: plain-string types, genuinely 10×1m — must stay UNCHANGED (no breakdown).
            {"name": "Reading Skills", "marks": 10, "questions_count": 10,
             "marks_per_question": 1, "question_types": ["MCQ"]},
            # Writing: section CLAIMS 10 questions, but subsections describe 2 (5m each, uniform).
            {"name": "Writing Skills and Applied Grammar", "marks": 10, "questions_count": 10,
             "marks_per_question": 1, "question_types": ["Short Answer"], "subsections": [
                 _sub("MCQ", "MCQ", 5, 1, 5), _sub("Letter Writing", "Letter Writing", 5, 1, 5)]},
            # Literature: section CLAIMS 20 questions; subsections describe 3 (5+10+5) → mixed.
            {"name": "Language through Literature", "marks": 20, "questions_count": 20,
             "marks_per_question": 1, "question_types": ["Short Answer"], "subsections": [
                 _sub("Extract Based", "Extract Based", 5, 1, 5),
                 _sub("Short Answer", "Short Answer", 10, 1, 10),
                 _sub("Long Answer", "Long Answer", 5, 1, 5)]},
        ])

    def _wos(self):
        p = self._pattern()
        bp = {s["name"]: s for s in p.sections}
        return {wo.section_name: wo for wo in
                sg.build_work_orders(bp, p, {}, "Medium", "6", "English Core", ["Fables"])}

    def test_count_follows_per_type_breakdown_not_stale_field(self):
        wos = self._wos()
        # Reading: no per-type breakdown → declared count/marks stand.
        self.assertEqual(wos["Reading Skills"].questions_count, 10)
        self.assertEqual(wos["Reading Skills"].marks_per_question, 1)
        # Writing: 2 questions (not 10), 5m each, section total still 10, uniform marks.
        w = wos["Writing Skills and Applied Grammar"]
        self.assertEqual(w.questions_count, 2)
        self.assertEqual(w.marks, 10)
        self.assertEqual(w.marks_per_question, 5.0)
        self.assertFalse(w.mixed_marks)
        # Literature: 3 questions (not 20), section total 20, mixed 5/10/5 marks.
        lit = wos["Language through Literature"]
        self.assertEqual(lit.questions_count, 3)
        self.assertEqual(lit.marks, 20)
        self.assertTrue(lit.mixed_marks)

    def test_reconciled_work_order_has_no_self_conflicting_validation(self):
        # A submission that matches the reconciled typed counts must satisfy BOTH the count
        # check and the marks-total check — they used to contradict each other and fail forever.
        w = self._wos()["Writing Skills and Applied Grammar"]
        data = {"questions": [
            {"qnum": 1, "type": "MCQ", "subtype": "standard", "text": "Pick one.",
             "options": {"a": "1", "b": "2", "c": "3", "d": "4"}, "answer": "b",
             "answer_explanation": "b is right.", "marks": 5, "competency_type": "recall"},
            {"qnum": 2, "type": "Letter Writing", "subtype": "standard",
             "text": "Write a letter to your friend describing your school.",
             "answer_explanation": "Format + content points.", "marks": 5,
             "competency_type": "constructed"},
        ]}
        errs = sg.validate_section_output(data, w)
        joined = " ".join(errs)
        self.assertNotIn("Expected", joined, f"count conflict remains: {errs}")
        self.assertNotIn("Section marks total", joined, f"marks-total conflict remains: {errs}")
        self.assertNotIn("distribution wrong", joined, f"type conflict remains: {errs}")


class TrimOverfullSectionTest(TestCase):
    """A section that generates MORE questions than its blueprint must be trimmed back, or the
    paper's marks overshoot (the Biology 33/30, Physics 28/25 case). Mirror of [Refill]."""

    def _wo(self):
        # Biology from _compound_sections: caps mcq7, ar2, vsa3, sa2, cbq1, la1 = 16q / 30m.
        p = ExamPattern(name="ov", subject="Science", class_name="10",
                        sections=[_compound_sections()[0]])
        bp = {s["name"]: s for s in p.sections}
        return sg.build_work_orders(bp, p, {}, "Medium", "10", "Science", ["Life Processes"])[0]

    def test_trims_excess_of_a_capped_type(self):
        wo = self._wo()
        qs = ([{"type": "MCQ", "subtype": "standard", "marks": 1} for _ in range(7)]
              + [{"type": "MCQ", "subtype": "assertion_reason", "marks": 1} for _ in range(2)]
              + [{"type": "VSA", "marks": 2} for _ in range(3)]
              + [{"type": "SA", "marks": 3} for _ in range(3)]        # one too many (cap 2)
              + [{"type": "CBQ", "subtype": "source_based", "marks": 4}]
              + [{"type": "LA", "marks": 5}])
        for i, q in enumerate(qs, 1):
            q["qnum"] = i
        out = sg.trim_overfull_sections({"Biology": {"questions": qs}}, [wo])
        kept = out["Biology"]["questions"]
        self.assertEqual(len(kept), 16)                                  # 17 → 16
        self.assertEqual(len([q for q in kept if q["type"] == "SA"]), 2)  # capped at 2
        self.assertEqual(sum(q["marks"] for q in kept), 30)              # back to section budget

    def test_leaves_on_count_section_untouched(self):
        wo = self._wo()
        qs = ([{"type": "MCQ", "subtype": "standard", "marks": 1} for _ in range(7)]
              + [{"type": "MCQ", "subtype": "assertion_reason", "marks": 1} for _ in range(2)]
              + [{"type": "VSA", "marks": 2} for _ in range(3)]
              + [{"type": "SA", "marks": 3} for _ in range(2)]
              + [{"type": "CBQ", "subtype": "source_based", "marks": 4}]
              + [{"type": "LA", "marks": 5}])
        for i, q in enumerate(qs, 1):
            q["qnum"] = i
        out = sg.trim_overfull_sections({"Biology": {"questions": qs}}, [wo])
        self.assertEqual(len(out["Biology"]["questions"]), 16)
        self.assertNotIn("_trimmed_overfull", out["Biology"])


class ParseChapterListTest(TestCase):
    """The upload view accepts the 'chapters' field as a JSON array (what the form sends), a
    single string, or a repeated form field — all normalise to a de-duplicated name list."""

    def setUp(self):
        from api.views import _parse_chapter_list
        self.parse = _parse_chapter_list

    def test_json_array_string(self):
        self.assertEqual(self.parse('["Light", "Electricity"]'), ["Light", "Electricity"])

    def test_single_plain_string(self):
        self.assertEqual(self.parse("Light"), ["Light"])

    def test_python_list(self):
        self.assertEqual(self.parse(["Light", "Electricity"]), ["Light", "Electricity"])

    def test_dedupes_case_insensitively_preserving_order(self):
        self.assertEqual(self.parse('["Light", "light", "Sound"]'), ["Light", "Sound"])

    def test_empty_and_none(self):
        self.assertEqual(self.parse(None), [])
        self.assertEqual(self.parse(""), [])
        self.assertEqual(self.parse("[]"), [])

    def test_malformed_json_falls_back_to_literal(self):
        self.assertEqual(self.parse("[oops"), ["[oops"])


class _FakePage:
    def __init__(self, text):
        self._t = text
    def extract_text(self):
        return self._t


class _FakeReader:
    def __init__(self, text):
        self.pages = [_FakePage(text)]


class ChapterIngestTest(TestCase):
    """ingest_pdf stores each chunk ONCE and links it to every chapter via ChunkChapter — a note
    spanning multiple chapters is no longer duplicated per chapter (pgvector many-to-many)."""

    def _run_ingest(self, **kwargs):
        from unittest import mock
        from core import embeddings as emb
        unit = kwargs.pop("unit", "ignored")
        with mock.patch.object(emb, "PdfReader", lambda *_a, **_k: _FakeReader("x" * 1600)), \
             mock.patch.object(emb, "get_embeddings_batch",
                               side_effect=lambda chunks, provider: [[0.0] * 768 for _ in chunks]):
            return emb.ingest_pdf("10", "Physics", unit, "C:/fake.pdf", **kwargs)

    def test_multi_chapter_note_stored_once(self):
        from core.models import MaterialChunk, ChunkChapter
        n = self._run_ingest(units=["Light", "Electricity"], material_type="notes")
        # 1600 chars → 2 chunks, stored ONCE (not 4) even though linked to 2 chapters.
        self.assertEqual(n, 2)
        self.assertEqual(MaterialChunk.objects.count(), 2)
        self.assertEqual(set(ChunkChapter.objects.values_list("unit", flat=True)), {"light", "electricity"})
        self.assertEqual(ChunkChapter.objects.count(), 4)  # 2 chunks × 2 chapters (cheap links, not chunk copies)

    def test_single_chapter_links_once(self):
        from core.models import MaterialChunk, ChunkChapter
        self._run_ingest(unit="Light", material_type="textbook")
        self.assertEqual(MaterialChunk.objects.count(), 2)
        self.assertEqual(ChunkChapter.objects.count(), 2)
        self.assertTrue(all(c.embedding_local is not None for c in MaterialChunk.objects.all()))

    def test_note_and_textbook_coexist_under_same_chapter(self):
        from core.models import MaterialChunk, ChunkChapter
        self._run_ingest(unit="Light", material_type="textbook")
        self._run_ingest(units=["Light"], material_type="notes")
        self.assertEqual(MaterialChunk.objects.count(), 4)            # separate rows — no id collision
        self.assertEqual(ChunkChapter.objects.filter(unit="light").count(), 4)


class DeleteMaterialEmbeddingsTest(TestCase):
    """delete_material_embeddings removes only one material's chunks (cascading their chapter
    links), never a textbook chapter's chunks that share the same unit label."""

    def _chunk(self, material, unit):
        from core.models import MaterialChunk, ChunkChapter
        c = MaterialChunk.objects.create(material=material, class_name="10", subject="physics",
                                         content="x", chunk_index=0, provider="local")
        ChunkChapter.objects.create(chunk=c, unit=unit)
        return c

    def test_deletes_only_that_material(self):
        from core.models import Material, MaterialChunk, ChunkChapter
        from core import embeddings as emb
        note = Material.objects.create(class_name="10", subject="Physics", title="notes", type="notes", file="")
        book = Material.objects.create(class_name="10", subject="Physics", title="book", type="textbook", file="")
        self._chunk(note, "light"); self._chunk(note, "light"); self._chunk(book, "light")
        emb.delete_material_embeddings("10", "Physics", note.id)
        self.assertEqual(MaterialChunk.objects.filter(material=note).count(), 0)
        self.assertEqual(MaterialChunk.objects.filter(material=book).count(), 1)
        self.assertEqual(ChunkChapter.objects.count(), 1)            # deleted chunks' links cascaded away

    def test_noop_when_material_id_none(self):
        from core import embeddings as emb
        emb.delete_material_embeddings("10", "Physics", None)        # must not raise / delete nothing


class EnrichmentTest(TestCase):
    """LLM chunk enrichment (core/enrichment.py): closed-enum label validation, per-chunk
    chapter re-linking, summary chunks, sibling-copy mirroring, idempotent skip, fail-open."""

    def setUp(self):
        from core.models import Material, MaterialChunk, ChunkChapter
        self.mat = Material.objects.create(
            class_name="10", subject="Physics", title="Light+Electricity notes", type="notes",
            file="", metadata={"chapters": ["Light", "Electricity"]})
        self.chunks = []
        for i, text in enumerate(["about refraction of light", "circuits exercise questions",
                                  "garbled symbol soup"]):
            c = MaterialChunk.objects.create(material=self.mat, class_name="10", subject="physics",
                                             content=text, chunk_index=i, provider="local")
            # Mimic ingestion's over-linking: every chunk linked to ALL declared chapters.
            ChunkChapter.objects.create(chunk=c, unit="light")
            ChunkChapter.objects.create(chunk=c, unit="electricity")
            self.chunks.append(c)

    def _reply(self):
        import json
        return json.dumps({
            "chunks": {
                # over-long "clean" (way beyond 2x the original) is a hallucination — dropped
                "c0": {"unit": "Light", "garbled": False, "clean": "X" * 3000},
                # legitimate selective cleanup — stored on content_clean, content untouched
                "c1": {"unit": "Electricity", "garbled": False, "clean": "circuits concept text only"},
                # unknown unit dropped (keeps original links), garbled kept
                "c2": {"unit": "Nonexistent", "garbled": True},
                "c99": {"unit": "Light"},      # hallucinated id — must be ignored
            },
            "summaries": {"Light": "L" * 300, "Electricity": "E" * 300, "Hallucinated": "H" * 300},
        })

    def _enrich(self, reply=None, **kwargs):
        from unittest import mock
        from core import enrichment
        with mock.patch.object(enrichment.mantle_client, "converse",
                               return_value=(reply or self._reply(), 100, 50)) as conv, \
             mock.patch("core.enrichment.get_embeddings_batch",
                        side_effect=lambda texts, provider: [[0.0] * 768 for _ in texts]):
            stats = enrichment.enrich_material(self.mat.id, **kwargs)
        return stats, conv

    def test_labels_relinks_and_summaries(self):
        from core.models import MaterialChunk, ChunkChapter
        stats, conv = self._enrich()
        # one batch, one call — Physics is not a language subject, so no chapter-kind call
        self.assertEqual(conv.call_count, 1)
        self.assertEqual(stats["chunks_labeled"], 3)
        self.assertEqual(stats["garbled"], 1)
        self.assertFalse(stats["skipped"])

        c0, c1, c2 = [MaterialChunk.objects.get(pk=c.pk) for c in self.chunks]
        self.assertTrue(c2.garbled and not c0.garbled)
        self.assertTrue(all(c.enriched_at for c in (c0, c1, c2)))
        # Selective clean copy: stored for c1, hallucinated over-long clean on c0 dropped,
        # absent on c2; originals never mutated.
        self.assertEqual(c1.content_clean, "circuits concept text only")
        self.assertEqual(c0.content_clean, "")
        self.assertEqual(c2.content_clean, "")
        self.assertEqual(c1.content, "circuits exercise questions")

        # Per-chunk chapter attribution replaces the over-linking — except c2, whose LLM
        # unit wasn't in the declared list, so it keeps its original links (fail open).
        self.assertEqual(set(ChunkChapter.objects.filter(chunk=c0).values_list("unit", flat=True)), {"light"})
        self.assertEqual(set(ChunkChapter.objects.filter(chunk=c1).values_list("unit", flat=True)), {"electricity"})
        self.assertEqual(ChunkChapter.objects.filter(chunk=c2).count(), 2)

        # Two summary chunks (hallucinated chapter dropped), negative index, linked to units.
        summaries = MaterialChunk.objects.filter(kind="summary")
        self.assertEqual(summaries.count(), 2)
        self.assertTrue(all(s.chunk_index < 0 for s in summaries))
        self.assertEqual(set(ChunkChapter.objects.filter(chunk__in=summaries)
                             .values_list("unit", flat=True)), {"light", "electricity"})

    def test_second_run_skips_without_llm_call(self):
        self._enrich()
        stats, conv = self._enrich()
        self.assertEqual(conv.call_count, 0)
        self.assertTrue(stats["skipped"])

    def test_force_reprocesses(self):
        self._enrich()
        stats, conv = self._enrich(force=True)
        self.assertEqual(conv.call_count, 1)   # non-language subject: labels only
        self.assertEqual(stats["chunks_labeled"], 3)

    def test_mirrors_labels_to_school_copy_without_second_llm_call(self):
        from core.models import School, MaterialChunk
        school = School.objects.create(name="S1")
        for c in self.chunks:                                     # textbook double-ingest twin copy
            MaterialChunk.objects.create(material=self.mat, school=school, class_name="10",
                                         subject="physics", content=c.content,
                                         chunk_index=c.chunk_index, provider="local")
        stats, conv = self._enrich()
        self.assertEqual(conv.call_count, 1)   # mirror is free — no second LLM pass
        self.assertEqual(stats["chunks_labeled"], 6)
        school_rows = MaterialChunk.objects.filter(school=school, kind="body")
        self.assertTrue(all(r.enriched_at for r in school_rows))
        # The selective clean copy mirrors too.
        self.assertEqual(school_rows.get(chunk_index=1).content_clean, "circuits concept text only")
        # The school copy gets its own summary chunks (scoped queries filter by school).
        self.assertEqual(MaterialChunk.objects.filter(kind="summary", school=school).count(), 2)
        self.assertEqual(MaterialChunk.objects.filter(kind="summary", school__isnull=True).count(), 2)

    def test_llm_failure_fails_open(self):
        from core.models import MaterialChunk
        stats, conv = self._enrich(reply="not json at all")
        self.assertEqual(conv.call_count, 2)                      # one corrective retry
        self.assertEqual(stats["chunks_labeled"], 0)
        self.assertTrue(stats["errors"])
        self.assertTrue(all(c.enriched_at is None
                            for c in MaterialChunk.objects.filter(kind="body")))

    def test_terminal_run_tasks_are_pure_noops(self):
        # A task whose run is already CLOSED (stopped/failed/done — e.g. redelivered after
        # a Redis restart) must do no LLM work AND leave the closed run's counters alone.
        from unittest import mock
        from core.models import EnrichmentRun
        from core import tasks as t, enrichment
        for terminal in ('stopped', 'failed', 'done'):
            run = EnrichmentRun.objects.create(status=terminal, total_groups=1, drained_groups=0)
            with mock.patch.object(enrichment, "enrich_material",
                                   side_effect=AssertionError("must not be called")):
                res = t.enrich_material_task.apply(args=(self.mat.id,), kwargs={"run_id": run.id}).result
            self.assertTrue(res.get("stopped"))
            run.refresh_from_db()
            self.assertEqual((run.done_groups, run.drained_groups), (0, 0))
            self.assertEqual(run.status, terminal)

    def test_stopping_run_drains_queued_tasks_and_flips_to_stopped(self):
        from unittest import mock
        from core.models import EnrichmentRun
        from core import tasks as t, enrichment
        run = EnrichmentRun.objects.create(status='stopping', total_groups=2)
        with mock.patch.object(enrichment, "enrich_material",
                               side_effect=AssertionError("must not be called")):
            t.enrich_material_task.apply(args=(self.mat.id,), kwargs={"run_id": run.id})
            run.refresh_from_db()
            self.assertEqual((run.status, run.drained_groups), ('stopping', 1))
            t.enrich_material_task.apply(args=(self.mat.id,), kwargs={"run_id": run.id})
        run.refresh_from_db()
        self.assertEqual(run.drained_groups, 2)
        self.assertEqual(run.status, 'stopped')       # fully drained → terminal

    def test_mid_material_stop_is_atomic_and_resumable(self):
        # Stop lands between LLM batches: nothing partial is persisted (the copy stays
        # fully pending), the spent tokens are still reported, and a later resume run
        # enriches the material completely.
        from unittest import mock
        from core.models import EnrichmentRun, MaterialChunk
        from core import enrichment

        run = EnrichmentRun.objects.create(status='running', total_groups=1)

        def reply_then_stop(**kwargs):
            EnrichmentRun.objects.filter(id=run.id).update(status='stopping')
            return (self._reply(), 100, 50)

        with mock.patch.object(enrichment.mantle_client, "converse",
                               side_effect=reply_then_stop) as conv, \
             mock.patch("core.enrichment._make_batches",
                        side_effect=lambda rows: [rows[:1], rows[1:]]), \
             mock.patch("core.enrichment.get_embeddings_batch",
                        side_effect=lambda texts, provider: [[0.0] * 768 for _ in texts]):
            stats = enrichment.enrich_material(self.mat.id, run_id=run.id)

        self.assertTrue(stats["stopped"])
        self.assertEqual(conv.call_count, 1)          # second batch never sent
        self.assertEqual(stats["input_tokens"], 100)  # first batch still billed
        self.assertEqual(stats["chunks_labeled"], 0)  # per-copy atomic: nothing persisted
        self.assertTrue(all(c.enriched_at is None
                            for c in MaterialChunk.objects.filter(kind="body")))
        self.assertEqual(MaterialChunk.objects.filter(kind="summary").count(), 0)

        # Resume (fresh run) processes the whole material cleanly.
        stats2, conv2 = self._enrich()
        self.assertFalse(stats2["stopped"])
        self.assertEqual(stats2["chunks_labeled"], 3)

    def test_heartbeat_bumps_updated_at_during_work(self):
        # Long single-material runs must not trip the 15-min staleness auto-fail: every
        # batch check doubles as a heartbeat on the run row.
        from datetime import timedelta
        from django.utils import timezone as tz
        from core.models import EnrichmentRun
        run = EnrichmentRun.objects.create(status='running', total_groups=1)
        stale = tz.now() - timedelta(seconds=1200)
        EnrichmentRun.objects.filter(id=run.id).update(updated_at=stale)
        self._enrich(run_id=run.id)
        run.refresh_from_db()
        self.assertGreater(run.updated_at, stale + timedelta(seconds=1100))

    def test_task_counts_mid_material_stop_as_drained(self):
        # Stop pressed while a material is in flight: the task reports the pause as a
        # drain (not done), and — being the only task — closes the run as 'stopped'.
        from unittest import mock
        from core.models import EnrichmentRun
        from core import tasks as t, enrichment
        run = EnrichmentRun.objects.create(status='running', total_groups=1)

        def enrich_paused(material_id, force=False, run_id=None):
            EnrichmentRun.objects.filter(id=run_id).update(status='stopping')
            return {"material_id": material_id, "chunks_labeled": 0, "summaries_created": 0,
                    "garbled": 0, "input_tokens": 80, "output_tokens": 20, "skipped": False,
                    "stopped": True, "errors": []}

        with mock.patch.object(enrichment, "enrich_material", side_effect=enrich_paused):
            res = t.enrich_material_task.apply(args=(self.mat.id,), kwargs={"run_id": run.id}).result
        self.assertTrue(res["stopped"])
        run.refresh_from_db()
        self.assertEqual((run.done_groups, run.drained_groups), (0, 1))
        self.assertEqual(run.status, 'stopped')
        self.assertEqual(run.input_tokens, 80)        # aborted work is still billed

    def test_summary_chunks_excluded_from_span_and_query_scope(self):
        from core.models import MaterialChunk
        from core import embeddings as emb
        self._enrich()
        summary = MaterialChunk.objects.filter(kind="summary").first()
        # A summary seed returns only itself — never spliced with body neighbours.
        self.assertEqual(emb.fetch_contiguous_span(summary.id), summary.content.strip())
        # And a body seed's span never absorbs summary rows (negative index + kind filter).
        body_span = emb.fetch_contiguous_span(self.chunks[0].id, before=2000, after=2000)
        self.assertNotIn(summary.content[:50], body_span)


class ChapterKindTest(TestCase):
    """Chapter-LEVEL kind classification (ChapterInfo): ONE kind per chapter judged from
    name + summary + content sample — the redesign of the rejected per-chunk taxonomy
    (a prose lesson with back-exercises must be 'prose', not a bag of chunk kinds)."""

    def setUp(self):
        from core.models import Material, MaterialChunk, ChunkChapter
        self.mat = Material.objects.create(
            class_name="10", subject="English", title="First Flight combo", type="textbook",
            file="", metadata={"chapters": ["Dust of Snow", "A Letter to God"]})
        self.chunks = []
        for i, (text, unit) in enumerate([
                ("The way a crow\nShook down on me\nThe dust of snow", "dust_of_snow"),
                ("Lencho wrote a letter to God asking for a hundred pesos.", "a_letter_to_god")]):
            c = MaterialChunk.objects.create(material=self.mat, class_name="10",
                                             subject="english", content=text,
                                             chunk_index=i, provider="local")
            ChunkChapter.objects.create(chunk=c, unit=unit)
            self.chunks.append(c)

    def _label_reply(self):
        import json
        return json.dumps({
            "chunks": {"c0": {"unit": "Dust of Snow", "garbled": False},
                       "c1": {"unit": "A Letter to God", "garbled": False}},
            "summaries": {"Dust of Snow": "P" * 300, "A Letter to God": "L" * 300},
        })

    def _kinds_reply(self, mapping):
        import json
        return json.dumps({"kinds": mapping})

    def test_classifier_validates_enum_and_prompt_carries_evidence(self):
        from unittest import mock
        from core import enrichment
        entries = [
            {"class_name": "10", "subject": "english", "unit": "dust_of_snow",
             "display": "Dust of Snow", "material_title": "First Flight",
             "summary": "A short poem about a crow and snow.", "sample": "The way a crow\nShook down"},
            {"class_name": "10", "subject": "english", "unit": "a_letter_to_god",
             "display": "A Letter to God", "material_title": "First Flight",
             "summary": "Lencho's faith.", "sample": "Lencho wrote"},
        ]
        with mock.patch.object(enrichment.mantle_client, "converse",
                               return_value=(self._kinds_reply({"c0": "poem", "c1": "banana"}), 10, 5)) as conv:
            kinds, tin, tout = enrichment.classify_chapter_kinds(entries)
        self.assertEqual(kinds, {("10", "english", "dust_of_snow"): "poem"})   # bad enum dropped
        prompt = conv.call_args.kwargs["prompt"]
        self.assertIn("Dust of Snow", prompt)                     # chapter name
        self.assertIn("First Flight", prompt)                     # book title (supplementary signal)
        self.assertIn("A short poem about a crow", prompt)        # summary
        self.assertIn("Shook down", prompt)                       # content sample (verse shape)

    def test_enrich_material_classifies_its_chapters(self):
        from unittest import mock
        from core.models import ChapterInfo
        from core import enrichment
        # todo is sorted(declared_norm): c0 = "a_letter_to_god", c1 = "dust_of_snow"
        replies = [(self._label_reply(), 100, 50),
                   (self._kinds_reply({"c0": "prose", "c1": "poem"}), 20, 10)]
        with mock.patch.object(enrichment.mantle_client, "converse", side_effect=replies) as conv, \
             mock.patch("core.enrichment.get_embeddings_batch",
                        side_effect=lambda texts, provider: [[0.0] * 768 for _ in texts]):
            stats = enrichment.enrich_material(self.mat.id)
        self.assertEqual(conv.call_count, 2)
        self.assertEqual(stats["chapters_classified"], 2)
        self.assertEqual(ChapterInfo.objects.get(unit="dust_of_snow").kind, "poem")
        self.assertEqual(ChapterInfo.objects.get(unit="a_letter_to_god").kind, "prose")
        self.assertEqual(stats["input_tokens"], 120)              # classify tokens billed too

    def test_already_classified_chapters_skip_the_call(self):
        from unittest import mock
        from django.utils import timezone as tz
        from core.models import ChapterInfo
        from core import enrichment
        for unit, kind in [("dust_of_snow", "poem"), ("a_letter_to_god", "prose")]:
            ChapterInfo.objects.create(class_name="10", subject="english", unit=unit,
                                       kind=kind, classified_at=tz.now())
        with mock.patch.object(enrichment.mantle_client, "converse",
                               return_value=(self._label_reply(), 100, 50)) as conv, \
             mock.patch("core.enrichment.get_embeddings_batch",
                        side_effect=lambda texts, provider: [[0.0] * 768 for _ in texts]):
            enrichment.enrich_material(self.mat.id)
        self.assertEqual(conv.call_count, 1)                      # labels only — no classify call

    def test_backfill_task_classifies_from_store(self):
        from unittest import mock
        from core.models import ChapterInfo, MaterialChunk, ChunkChapter
        from core import tasks as t, enrichment
        # A stored summary row gives the backfill its evidence without re-reading chunks.
        s = MaterialChunk.objects.create(material=self.mat, class_name="10", subject="english",
                                         kind="summary", content="A poem about a crow.",
                                         chunk_index=-1000, provider="local")
        ChunkChapter.objects.create(chunk=s, unit="dust_of_snow")
        ChapterInfo.objects.create(class_name="10", subject="english", unit="a_letter_to_god",
                                   kind="prose")                  # already done → skipped
        def fake_classify(entries):
            self.assertEqual(len(entries), 1)                     # only the pending chapter
            e = entries[0]
            self.assertEqual(e["unit"], "dust_of_snow")
            self.assertEqual(e["summary"], "A poem about a crow.")
            self.assertIn("crow", e["sample"])                    # sample chunk found
            return {("10", "english", "dust_of_snow"): "poem"}, 15, 5
        with mock.patch.object(enrichment, "classify_chapter_kinds", side_effect=fake_classify):
            res = t.classify_all_chapters_task.apply(kwargs={"force": False}).result
        self.assertEqual(res["classified"], 1)
        self.assertEqual(ChapterInfo.objects.get(unit="dust_of_snow").kind, "poem")

    def test_extract_kind_detection_and_routing(self):
        from core.models import ChapterInfo
        from core import section_generator as secgen
        ChapterInfo.objects.create(class_name="10", subject="english",
                                   unit="dust_of_snow", kind="poem")
        ChapterInfo.objects.create(class_name="10", subject="english",
                                   unit="a_letter_to_god", kind="prose")
        slots = [{"type": "extract", "format": "Poetry extract", "marks": 5},
                 {"type": "sa", "format": "prose question", "marks": 3}]   # non-extract ignored
        wanted = secgen._extract_kinds_wanted(slots, {"extract_instruction": ""})
        self.assertEqual(wanted, {"poem"})
        self.assertEqual(secgen._extract_kind_needs(slots), ["poem"])
        ordered = secgen._chapters_for_extract_needs(
            "10", "English", ["A Letter to God", "Dust of Snow"], ["poem"])
        self.assertEqual(ordered[0], "Dust of Snow")              # poem chapter leads
        self.assertEqual(len(ordered), 2)                         # nothing dropped
        # Fail-open: no kind data for this subject → original order untouched.
        same = secgen._chapters_for_extract_needs(
            "10", "Tamil", ["Ch A", "Ch B"], ["poem"])
        self.assertEqual(same, ["Ch A", "Ch B"])
        # No explicit kind in the pattern → no reordering at all.
        self.assertEqual(secgen._extract_kinds_wanted(
            [{"type": "extract", "format": "", "marks": 5}], {}), set())

    def test_mixed_extract_needs_cover_every_alternative(self):
        # The SQP shape: Q6 = internal-choice extract with NO kind signal, Q7 = poetry
        # extract with internal choice. Four printed passages are needed; the poem slots
        # must get poem chapters and the unsignaled slots must NOT be starved onto poems.
        from core.models import ChapterInfo
        from core import section_generator as secgen
        for unit, kind in [("poem_one", "poem"), ("poem_two", "poem"),
                           ("prose_one", "prose"), ("prose_two", "prose")]:
            ChapterInfo.objects.create(class_name="10", subject="english_core",
                                       unit=unit, kind=kind)
        slots = [
            {"type": "extract", "marks": 5, "choice": "internal",
             "format": "Best explanation selection"},              # unsignaled (Q6)
            {"type": "extract", "marks": 5, "choice": "internal",
             "format": "Poetry extract"},                          # poem (Q7)
        ]
        needs = secgen._extract_kind_needs(slots)
        self.assertEqual(needs, ["", "", "poem", "poem"])
        picked = secgen._chapters_for_extract_needs(
            "10", "English Core",
            ["prose_one", "poem_one", "prose_two", "poem_two"], needs)
        # First two picks (Q6's alternatives) avoid the poem chapters reserved for Q7;
        # picks three and four are the poems.
        self.assertEqual(picked[:4], ["prose_one", "prose_two", "poem_one", "poem_two"])

    def test_kind_keywords_do_not_misfire_on_ordinary_pattern_prose(self):
        from core import section_generator as secgen
        def wanted(text):
            return secgen._extract_kinds_wanted(
                [{"type": "extract", "condition": text, "marks": 5}], {})
        # Substring traps (confirmed review findings): common words must not route.
        self.assertEqual(wanted("Questions should cover diverse themes"), set())      # not 'verse'
        self.assertEqual(wanted("explain how the writer plays with words"), set())    # not 'play'
        self.assertEqual(wanted("describe the key moments in the passage"), set())    # not the Moments reader
        self.assertEqual(wanted("display the universe of ideas"), set())
        # ...while genuine signals still fire, including plurals.
        self.assertEqual(wanted("Extract from the poems prescribed"), {"poem"})
        self.assertEqual(wanted("A verse from the poetry section"), {"poem"})
        self.assertEqual(wanted("prose extract of 100 words"), {"prose"})
        self.assertEqual(wanted("an extract from the play"), {"drama"})
        self.assertEqual(wanted("from the supplementary reader Footprints"), {"supplementary"})

    def test_classify_endpoint_pending_ignores_stale_rows(self):
        # ChapterInfo rows for chapters whose materials were deleted must not mask the
        # live corpus's unclassified chapters (confirmed review finding).
        from unittest import mock
        from django.utils import timezone as tz
        from rest_framework.test import APIClient
        from core.models import ChapterInfo
        for i in range(5):   # stale: no corresponding chunks exist
            ChapterInfo.objects.create(class_name="9", subject="tamil", unit=f"gone_{i}",
                                       kind="prose", classified_at=tz.now())
        admin = User.objects.create_user("kind_admin2", "ka2@x.com", "pw")
        p = admin.profile
        p.role = "superadmin"
        p.save()
        api = APIClient()
        api.force_authenticate(admin)
        with mock.patch("core.tasks.classify_all_chapters_task") as mtask:
            r = api.post("/api/admin/enrichment/classify/", {}, format="json")
        self.assertEqual(r.status_code, 202)
        self.assertEqual(r.json()["queued"], 2)      # the two LIVE chapters, stale rows ignored
        mtask.delay.assert_called_once()

    def test_kinds_are_language_subject_only(self):
        # The user's original objection: a Mathematics chapter must NEVER be "prose".
        from unittest import mock
        from core.models import Material, MaterialChunk, ChunkChapter, ChapterInfo
        from core import enrichment, tasks as t
        self.assertTrue(enrichment.is_language_subject("English Core"))
        self.assertTrue(enrichment.is_language_subject("hindi_course_b"))
        self.assertTrue(enrichment.is_language_subject("Tamil"))
        self.assertFalse(enrichment.is_language_subject("Mathematics"))
        self.assertFalse(enrichment.is_language_subject("Science"))
        self.assertFalse(enrichment.is_language_subject("Social Science"))
        # Backfill skips non-language chapters entirely — no LLM call, no rows.
        mat = Material.objects.create(class_name="6", subject="Mathematics", title="Ganita",
                                      type="textbook", file="", metadata={"chapters": ["Fractions"]})
        c = MaterialChunk.objects.create(material=mat, class_name="6", subject="mathematics",
                                         content="a half plus a quarter", chunk_index=0,
                                         provider="local")
        ChunkChapter.objects.create(chunk=c, unit="fractions")
        with mock.patch.object(enrichment, "classify_chapter_kinds",
                               side_effect=lambda entries: (
                                   {(e['class_name'], e['subject'], e['unit']): "poem"
                                    for e in entries}, 1, 1)):
            t.classify_all_chapters_task.apply(kwargs={"force": True})
        self.assertFalse(ChapterInfo.objects.filter(subject="mathematics").exists())

    DUST_OF_SNOW = ("The way a crow\nShook down on me\nThe dust of snow\n"
                    "From a hemlock tree\nHas given my heart\nA change of mood\n"
                    "And saved some part\nOf a day I had rued.")
    PROSE_CHUNK = ("Lencho was an ox of a man, working like an animal in the fields, "
                   "but still he knew how to write. The following Sunday, at daybreak, "
                   "he began to write a letter which he himself would carry to town and "
                   "place in the mail.\nAll through the years the rain had been kind.")

    def test_verse_shape_detection(self):
        from core import section_generator as secgen
        self.assertTrue(secgen._looks_like_verse(self.DUST_OF_SNOW))
        self.assertFalse(secgen._looks_like_verse(self.PROSE_CHUNK))
        self.assertFalse(secgen._looks_like_verse("One short line"))

    def test_verse_hunter_finds_poems_inside_prose_chapters(self):
        # Class 10 First Flight has NO poem chapters — the poems live inside the prose
        # chapters' chunks. The hunter must pull verse-shaped chunks out by similarity.
        from unittest import mock
        from core import section_generator as secgen
        def fake_query(**kw):
            docs = [self.PROSE_CHUNK, self.DUST_OF_SNOW] if kw.get("unit") == "A Letter to God" else []
            return {"ids": [[str(i) for i in range(len(docs))]], "documents": [docs],
                    "metadatas": [[{} for _ in docs]], "distances": [[0.0] * len(docs)]}
        with mock.patch.object(secgen.embeddings, "query", side_effect=fake_query):
            out = secgen._verse_passages("10", "English Core",
                                         ["A Letter to God", "Bholi"], None, 1, [])
        self.assertEqual(len(out), 1)
        self.assertEqual(out[0]["kind"], "poem")
        self.assertIn("dust of snow", out[0]["text"].lower())     # the verse chunk, not the prose

    def test_classify_endpoint_queues_task(self):
        from unittest import mock
        from rest_framework.test import APIClient
        admin = User.objects.create_user("kind_admin", "ka@x.com", "pw")
        p = admin.profile
        p.role = "superadmin"
        p.save()
        api = APIClient()
        api.force_authenticate(admin)
        with mock.patch("core.tasks.classify_all_chapters_task") as mtask:
            r = api.post("/api/admin/enrichment/classify/", {}, format="json")
        self.assertEqual(r.status_code, 202)
        self.assertGreater(r.json()["queued"], 0)
        mtask.delay.assert_called_once_with(force=False, user_id=admin.id)


class EnrichmentParallelGroupTest(TransactionTestCase):
    """enrich_materials_group_task runs its materials CONCURRENTLY (one thread per
    material, ENRICH_PARALLEL_PER_KEY per API key) while keeping all bookkeeping per
    material. TransactionTestCase: the pool threads open their own DB connections, so
    the test data must be really committed, not held in a test transaction."""

    def _mk_material(self, title):
        from core.models import Material, MaterialChunk
        mat = Material.objects.create(class_name="10", subject="Physics", title=title,
                                      type="notes", file="", metadata={"chapters": ["Light"]})
        for i, text in enumerate(["light refraction text", "light reflection text"]):
            MaterialChunk.objects.create(material=mat, class_name="10", subject="physics",
                                         content=text, chunk_index=i, provider="local")
        return mat

    def test_group_enriches_concurrently_with_per_material_bookkeeping(self):
        import json
        import threading
        from unittest import mock
        from core.models import EnrichmentRun, MaterialChunk
        from core import tasks as t, enrichment

        m1, m2 = self._mk_material("A"), self._mk_material("B")
        run = EnrichmentRun.objects.create(status="running", total_groups=2)

        reply = json.dumps({
            "chunks": {"c0": {"unit": "Light", "garbled": False},
                       "c1": {"unit": "Light", "garbled": False}},
            "summaries": {"Light": "L" * 300},
        })
        # Both materials' LLM calls must be IN FLIGHT at the same time to pass this
        # barrier — a serial implementation times out here and fails loudly.
        rendezvous = threading.Barrier(2, timeout=15)

        def concurrent_converse(**kwargs):
            rendezvous.wait()
            return (reply, 10, 5)

        with mock.patch.object(enrichment.mantle_client, "converse",
                               side_effect=concurrent_converse), \
             mock.patch("core.enrichment.get_embeddings_batch",
                        side_effect=lambda texts, provider: [[0.0] * 768 for _ in texts]), \
             mock.patch.object(enrichment.mantle_client, "num_keys", return_value=2):
            res = t.enrich_materials_group_task.apply(
                args=([m1.id, m2.id],), kwargs={"run_id": run.id}).result

        self.assertEqual(res["ok"], 2)
        run.refresh_from_db()
        self.assertEqual(run.done_groups, 2)              # bookkeeping stayed per material
        self.assertEqual(run.chunks_labeled, 4)
        self.assertEqual(run.status, "done")
        self.assertEqual(MaterialChunk.objects.filter(
            kind="body", enriched_at__isnull=True).count(), 0)


class EnrichmentConcurrencyConfigTest(TestCase):
    """Pool sizing (3 per Mantle key) and material grouping at the enqueue points."""

    def test_pool_size_scales_with_keys(self):
        from unittest import mock
        from core import enrichment
        with mock.patch.object(enrichment, "ENRICH_PARALLEL_PER_KEY", 3):
            with mock.patch.object(enrichment.mantle_client, "num_keys", return_value=2):
                self.assertEqual(enrichment.enrich_concurrency(), 6)
            with mock.patch.object(enrichment.mantle_client, "num_keys", return_value=0):
                self.assertEqual(enrichment.enrich_concurrency(), 3)   # keyless env: assume 1

    def test_ingest_enqueue_groups_materials_to_pool_size(self):
        from unittest import mock
        from core import tasks as t, enrichment
        with mock.patch.object(enrichment, "ENRICH_PARALLEL_PER_KEY", 3), \
             mock.patch.object(enrichment.mantle_client, "num_keys", return_value=2), \
             mock.patch.object(t, "enrich_materials_group_task") as mtask:
            t._enqueue_enrichment(list(range(1, 15)))              # 14 materials, pool 6
        groups = [c.args[0] for c in mtask.delay.call_args_list]
        self.assertEqual([len(g) for g in groups], [6, 6, 2])
        self.assertEqual(sorted(sum(groups, [])), list(range(1, 15)))


class EnrichmentControlEndpointTest(TestCase):
    """Run/stop lifecycle endpoints: two-phase stop (running → stopping → stopped),
    idempotent stop, no concurrent runs while draining, broker-failure safety, and
    staleness auto-close for both live statuses."""

    def setUp(self):
        from rest_framework.test import APIClient
        from core.models import Material, MaterialChunk
        self.user = User.objects.create_user("enr_admin", "ea@x.com", "pw")
        p = self.user.profile
        p.role = "superadmin"
        p.save()
        self.api = APIClient()
        self.api.force_authenticate(self.user)
        mat = Material.objects.create(class_name="10", subject="Physics", title="notes",
                                      type="notes", file="")
        MaterialChunk.objects.create(material=mat, class_name="10", subject="physics",
                                     content="pending chunk", chunk_index=0, provider="local")

    def _mk_run(self, **kw):
        from core.models import EnrichmentRun
        return EnrichmentRun.objects.create(**{"status": "running", "total_groups": 3, **kw})

    def _backdate(self, run, seconds):
        from datetime import timedelta
        from django.utils import timezone as tz
        from core.models import EnrichmentRun
        EnrichmentRun.objects.filter(id=run.id).update(
            updated_at=tz.now() - timedelta(seconds=seconds))

    def test_stop_is_two_phase_and_idempotent(self):
        run = self._mk_run()
        r = self.api.post("/api/admin/enrichment/stop/")
        self.assertEqual(r.status_code, 200)
        self.assertEqual(r.json()["run"]["status"], "stopping")   # drains async, not instant
        r = self.api.post("/api/admin/enrichment/stop/")          # second click: no error
        self.assertEqual(r.status_code, 200)
        self.assertEqual(r.json()["run"]["status"], "stopping")

    def test_stop_without_live_run_is_400(self):
        r = self.api.post("/api/admin/enrichment/stop/")
        self.assertEqual(r.status_code, 400)
        self._mk_run(status="done")
        r = self.api.post("/api/admin/enrichment/stop/")
        self.assertEqual(r.status_code, 400)

    def test_stop_on_stale_run_closes_immediately(self):
        # Worker dead → nothing will ever drain; the stop must not hang in 'stopping'.
        run = self._mk_run()
        self._backdate(run, 1000)
        r = self.api.post("/api/admin/enrichment/stop/")
        self.assertEqual(r.status_code, 200)
        self.assertEqual(r.json()["run"]["status"], "stopped")

    def test_no_new_run_while_draining(self):
        from unittest import mock
        self._mk_run(status="stopping")
        with mock.patch("core.tasks.enrich_materials_group_task") as mtask:
            r = self.api.post("/api/admin/enrichment/run/", {}, format="json")
        self.assertEqual(r.status_code, 409)
        self.assertIn("stopping", r.json()["error"])
        mtask.delay.assert_not_called()

    def test_stale_run_does_not_block_new_run(self):
        from unittest import mock
        old = self._mk_run()
        self._backdate(old, 1000)
        with mock.patch("core.tasks.enrich_materials_group_task") as mtask:
            mtask.delay.return_value = None
            r = self.api.post("/api/admin/enrichment/run/", {}, format="json")
        self.assertEqual(r.status_code, 202)
        old.refresh_from_db()
        self.assertEqual(old.status, "failed")

    def test_dead_broker_leaves_no_zombie_run(self):
        from unittest import mock
        from core.models import EnrichmentRun
        with mock.patch("core.tasks.enrich_materials_group_task") as mtask:
            mtask.delay.side_effect = OSError("connection refused")
            r = self.api.post("/api/admin/enrichment/run/", {}, format="json")
        self.assertEqual(r.status_code, 503)
        self.assertIn("Redis", r.json()["error"])
        self.assertEqual(EnrichmentRun.objects.count(), 0)   # no stuck 'running' row

    def test_stats_auto_closes_stale_runs(self):
        stopping = self._mk_run(status="stopping")
        self._backdate(stopping, 1000)
        r = self.api.get("/api/admin/enrichment/stats/")
        self.assertEqual(r.status_code, 200)
        self.assertEqual(r.json()["latest_run"]["status"], "stopped")
        stopping.refresh_from_db()
        self.assertEqual(stopping.status, "stopped")


class SerialQueueSlotRecoveryTest(TestCase):
    """Per-user serial generation queue self-healing: a ghost 'active' paper (Celery task
    lost in a worker/broker restart) must be auto-failed so it can't hold the user's slot
    forever; deleting an active paper must free the slot like cancel does; and a ghost
    task redelivery must never resurrect a closed paper. (Production defect 2026-07-16:
    a 6-day-old queued+dispatched paper wedged every later paper as eternally waiting.)"""

    def setUp(self):
        from rest_framework.test import APIClient
        self.user = User.objects.create_user("slot_t", "st@x.com", "pw")
        self.api = APIClient()
        self.api.force_authenticate(self.user)
        self.pattern = ExamPattern.objects.create(
            name="PT-1", sections=[{"name": "A", "marks": 2, "questions_count": 1}])

    def _paper(self, status="queued", task_id=None, age_seconds=0):
        from datetime import timedelta
        from django.utils import timezone as tz
        p = QuestionPaper.objects.create(
            class_name="10", subject="English", pattern=self.pattern, chapters=["1"],
            status=status, task_id=task_id, created_by=self.user, paper_data={})
        if age_seconds:
            QuestionPaper.objects.filter(id=p.id).update(
                updated_at=tz.now() - timedelta(seconds=age_seconds))
            p.refresh_from_db()
        return p

    def test_ghost_dispatched_paper_is_reaped_and_next_promoted(self):
        from unittest import mock
        from core import tasks as t
        ghost = self._paper(task_id="dead-task", age_seconds=6 * 24 * 3600)   # 6 days silent
        waiting = self._paper()
        with mock.patch("core.tasks.generate_paper_task") as mtask:
            mtask.delay.return_value = mock.Mock(id="task-new-1")
            t.dispatch_next_queued_paper(self.user.id)
        ghost.refresh_from_db(); waiting.refresh_from_db()
        self.assertEqual(ghost.status, "failed")
        self.assertIn("auto-failed", ghost.status_detail)
        self.assertEqual(waiting.task_id, "task-new-1")      # slot freed → promoted

    def test_fresh_dispatched_paper_keeps_the_slot(self):
        from unittest import mock
        from core import tasks as t
        active = self._paper(task_id="live-task")            # fresh — legitimately active
        waiting = self._paper()
        with mock.patch("core.tasks.generate_paper_task") as mtask:
            self.assertIsNone(t.dispatch_next_queued_paper(self.user.id))
        mtask.delay.assert_not_called()
        active.refresh_from_db(); waiting.refresh_from_db()
        self.assertEqual(active.status, "queued")            # not reaped
        self.assertIsNone(waiting.task_id)                   # still waiting its turn

    def test_reap_windows_and_waiting_papers(self):
        from core.tasks import reap_stale_papers
        stale_gen = self._paper(status="generating", task_id="t1", age_seconds=31 * 60)
        fresh_gen = self._paper(status="generating", task_id="t2", age_seconds=10 * 60)
        old_waiting = self._paper(age_seconds=10 * 24 * 3600)   # no task_id → just in line
        reaped = reap_stale_papers(self.user.id)
        self.assertEqual(reaped, [stale_gen.id])
        fresh_gen.refresh_from_db(); old_waiting.refresh_from_db()
        self.assertEqual(fresh_gen.status, "generating")
        self.assertEqual(old_waiting.status, "queued")       # waiting papers are never reaped

    def test_delete_active_paper_frees_slot(self):
        from unittest import mock
        active = self._paper(task_id="live-task")
        waiting = self._paper()
        with mock.patch("core.tasks.generate_paper_task") as mtask:
            mtask.delay.return_value = mock.Mock(id="task-new-2")
            r = self.api.delete(f"/api/papers/{active.id}/")
        self.assertEqual(r.status_code, 204)
        self.assertFalse(QuestionPaper.objects.filter(id=active.id).exists())
        waiting.refresh_from_db()
        self.assertEqual(waiting.task_id, "task-new-2")      # delete promotes like cancel

    def test_delete_finished_paper_does_not_dispatch(self):
        from unittest import mock
        done = self._paper(status="done", task_id="old")
        waiting = self._paper()
        with mock.patch("core.tasks.generate_paper_task") as mtask:
            r = self.api.delete(f"/api/papers/{done.id}/")
        self.assertEqual(r.status_code, 204)
        mtask.delay.assert_not_called()                      # no slot was held — nothing to free

    def test_ghost_task_never_resurrects_closed_paper(self):
        from core import tasks as t
        for closed_status in ("failed", "cancelled", "done"):
            closed = self._paper(status=closed_status, task_id="old")
            res = t.generate_paper_task.apply(args=(closed.id,)).result
            self.assertTrue(res.get("skipped"))
            closed.refresh_from_db()
            self.assertEqual(closed.status, closed_status)   # untouched
        res = t.generate_paper_task.apply(args=(999999,)).result
        self.assertTrue(res.get("skipped"))                  # deleted paper → quiet no-op


class ChunkingTest(TestCase):
    """Structure-aware chunker: natural boundaries, size bound, overlap, tiny-tail merge."""

    def test_short_text_is_single_chunk(self):
        from core.embeddings import _chunk_text
        out = _chunk_text("A short paragraph about photosynthesis.")
        self.assertEqual(len(out), 1)
        self.assertEqual(out[0][1], "A short paragraph about photosynthesis.")

    def test_empty_text(self):
        from core.embeddings import _chunk_text
        self.assertEqual(_chunk_text("   "), [])

    def test_respects_size_bound(self):
        from core.embeddings import _chunk_text
        text = ". ".join(f"Sentence number {i} explains an idea" for i in range(400))
        out = _chunk_text(text, size=500, overlap=60)
        self.assertGreater(len(out), 1)
        # No chunk grossly exceeds size (allow overlap + one boundary token of slack).
        self.assertTrue(all(len(c) <= 500 + 120 for _, c in out))

    def test_prefers_sentence_boundaries(self):
        from core.embeddings import _chunk_text
        text = " ".join(f"This is sentence {i}." for i in range(200))
        out = _chunk_text(text, size=300, overlap=40)
        # Most chunks should end at a sentence terminator, not mid-word.
        ends_clean = sum(1 for _, c in out if c.endswith("."))
        self.assertGreaterEqual(ends_clean, len(out) - 1)

    def test_overlap_carries_context(self):
        from core.embeddings import _chunk_text
        text = " ".join(f"word{i}" for i in range(400))
        out = [c for _, c in _chunk_text(text, size=300, overlap=60)]
        self.assertGreater(len(out), 1)
        # The tail of one chunk should reappear at the head of the next (context overlap).
        first_tail = out[0].split()[-1]
        self.assertIn(first_tail, out[1].split()[:8])

    def test_indices_are_sequential(self):
        from core.embeddings import _chunk_text
        text = ". ".join(f"Idea {i} is described here" for i in range(120))
        out = _chunk_text(text, size=400, overlap=50)
        self.assertEqual([i for i, _ in out], list(range(len(out))))


class PgvectorQueryTest(TestCase):
    """query() returns the Chroma-shaped dict, orders by cosine distance, and scopes to chapter."""

    def setUp(self):
        from core.models import Material
        # Chunks must belong to a material for the visibility join; a shared material is visible
        # to the default (school_id=None) query context.
        self.mat = Material.objects.create(class_name="10", subject="Physics", title="t",
                                           type="notes", file="", school=None, visibility="shared")

    def _chunk(self, content, vec, unit="light", **kw):
        from core.models import MaterialChunk, ChunkChapter
        c = MaterialChunk.objects.create(material=self.mat, class_name="10", subject="physics",
                                         content=content, chunk_index=0, provider="local",
                                         embedding_local=vec, **kw)
        ChunkChapter.objects.create(chunk=c, unit=unit)
        return c

    def test_orders_by_similarity_and_shape(self):
        from unittest import mock
        from core import embeddings as emb
        near = [1.0] + [0.0] * 767
        far = [0.0] * 767 + [1.0]
        self._chunk("near-doc", near)
        self._chunk("far-doc", far)
        with mock.patch.object(emb, "get_embedding", return_value=near):
            res = emb.query("10", "Physics", "Light", "q", n_results=5)
        self.assertEqual(len(res["documents"]), 1)                   # list-of-one-list (Chroma shape)
        self.assertEqual(res["documents"][0], ["near-doc", "far-doc"])
        self.assertIn("embeddings", res)

    def test_chapter_scoping(self):
        from unittest import mock
        from core import embeddings as emb
        self._chunk("light-doc", [1.0] + [0.0] * 767, unit="light")
        self._chunk("sound-doc", [1.0] + [0.0] * 767, unit="sound")
        with mock.patch.object(emb, "get_embedding", return_value=[1.0] + [0.0] * 767):
            res = emb.query("10", "Physics", "Sound", "q", n_results=5)
        self.assertEqual(res["documents"][0], ["sound-doc"])

    def test_garbled_chunks_excluded_from_retrieval(self):
        from unittest import mock
        from core import embeddings as emb
        near = [1.0] + [0.0] * 767
        self._chunk("clean-doc", near)
        self._chunk("garbled-doc", near, garbled=True)   # enrichment-flagged extraction noise
        with mock.patch.object(emb, "get_embedding", return_value=near):
            res = emb.query("10", "Physics", "Light", "q", n_results=5)
        self.assertEqual(res["documents"][0], ["clean-doc"])

    def test_cleaned_copy_preferred_in_documents(self):
        from unittest import mock
        from core import embeddings as emb
        near = [1.0] + [0.0] * 767
        self._chunk("actual content 42 RUNNING-HEADER page noise", near,
                    content_clean="actual content")
        with mock.patch.object(emb, "get_embedding", return_value=near):
            res = emb.query("10", "Physics", "Light", "q", n_results=5)
        self.assertEqual(res["documents"][0], ["actual content"])

    def test_chapter_summary_lookup(self):
        from core.models import MaterialChunk, ChunkChapter
        from core import embeddings as emb
        row = MaterialChunk.objects.create(
            material=self.mat, class_name="10", subject="physics", kind="summary",
            content="Light covers reflection and refraction.", chunk_index=-1000,
            provider="local")
        ChunkChapter.objects.create(chunk=row, unit="light")
        self.assertEqual(emb.get_chapter_summary("10", "Physics", "Light"),
                         "Light covers reflection and refraction.")
        self.assertEqual(emb.get_chapter_summary("10", "Physics", "Sound"), "")
        self.assertEqual(emb.get_chapter_summary("10", "Physics", None), "")


class SectionContextInterleaveTest(TestCase):
    """get_section_context must give EVERY chapter a proportional slice of the window,
    headed by its enrichment chapter summary — instead of appending chapter-by-chapter
    and truncating the tail, which made the context ~100% the FIRST chapter (the
    same-unit clustering defect)."""

    def _fake_query(self, docs_by_unit):
        def q(**kwargs):
            docs = list(docs_by_unit.get(kwargs.get("unit"), []))[:kwargs.get("n_results", 5)]
            return {"ids": [[str(i) for i in range(len(docs))]], "documents": [docs],
                    "metadatas": [[{} for _ in docs]], "distances": [[0.0] * len(docs)]}
        return q

    def test_every_chapter_represented_with_summary_header(self):
        from unittest import mock
        docs = {
            "Alpha": [f"alpha passage {i} " + "a" * 580 for i in range(20)],   # ~12k chars alone
            "Beta":  [f"beta passage {i} " + "b" * 580 for i in range(5)],
        }
        with mock.patch.object(sg.embeddings, "query", side_effect=self._fake_query(docs)), \
             mock.patch.object(sg.embeddings, "get_chapter_summary",
                               side_effect=lambda c, s, u, school_id=None: f"All about {u}."):
            ctx = sg.get_section_context("10", "Physics", ["Alpha", "Beta"],
                                         ["waves"], max_chars=4000)
        self.assertIn("=== CHAPTER: Alpha ===", ctx)
        self.assertIn("=== CHAPTER: Beta ===", ctx)                # tail chapter SURVIVES
        self.assertIn("[About this chapter: All about Alpha.]", ctx)
        self.assertIn("alpha passage 0", ctx)
        self.assertIn("beta passage 0", ctx)                       # real content, not just header
        self.assertLessEqual(len(ctx), 5000)                       # 1.25 x max_chars ceiling

    def test_empty_retrieval_returns_empty_string(self):
        from unittest import mock
        with mock.patch.object(sg.embeddings, "query", side_effect=self._fake_query({})):
            self.assertEqual(sg.get_section_context("10", "Physics", ["Alpha"], ["q"]), "")

    def test_no_chapter_mode_has_no_headers(self):
        from unittest import mock
        docs = {None: ["general doc one", "general doc two"]}
        with mock.patch.object(sg.embeddings, "query", side_effect=self._fake_query(docs)):
            ctx = sg.get_section_context("10", "Physics", [], ["q"], max_chars=4000)
        self.assertIn("general doc one", ctx)
        self.assertNotIn("=== CHAPTER", ctx)


class MaterialVisibilityScopeTest(TestCase):
    """visibility_q (the single access rule): a school sees its own materials ∪ shared (only if
    granted) ∪ institutional (any school), and NEVER another school's private material."""

    def setUp(self):
        from core.models import School, Material
        self.A = School.objects.create(name="A", access_shared_vector_store=True)
        self.B = School.objects.create(name="B", access_shared_vector_store=False)
        mk = lambda **kw: Material.objects.create(class_name="10", subject="Physics",
                                                  title="t", type="notes", file="", **kw)
        self.shared = mk(school=None, visibility="shared")
        self.a_priv = mk(school=self.A, visibility="private")
        self.b_priv = mk(school=self.B, visibility="private")
        self.b_inst = mk(school=self.B, visibility="institutional")

    def _visible(self, school):
        from core.models import Material
        from core.access import visibility_q
        return set(Material.objects.filter(visibility_q(school)).values_list("id", flat=True))

    def test_school_with_shared_access(self):
        v = self._visible(self.A)
        self.assertEqual(v, {self.shared.id, self.a_priv.id, self.b_inst.id})
        self.assertNotIn(self.b_priv.id, v)                  # never another school's private

    def test_school_without_shared_access(self):
        v = self._visible(self.B)
        self.assertEqual(v, {self.b_priv.id, self.b_inst.id})
        self.assertNotIn(self.shared.id, v)                  # not granted shared store
        self.assertNotIn(self.a_priv.id, v)

    def test_superadmin_or_no_school(self):
        self.assertEqual(self._visible(None), {self.shared.id, self.b_inst.id})

    def test_flip_to_institutional_exposes_cross_school(self):
        self.assertNotIn(self.b_priv.id, self._visible(self.A))
        self.b_priv.visibility = "institutional"
        self.b_priv.save()
        self.assertIn(self.b_priv.id, self._visible(self.A))   # instant, no re-ingest


class ChunkVisibilityScopeTest(TestCase):
    """The retrieval layer (_scoped_chunks) enforces the same visibility rule via the chunk→material
    join, so generation can only pull context a school is allowed to see."""

    def setUp(self):
        from core.models import School, Material, MaterialChunk, ChunkChapter
        self.A = School.objects.create(name="A", access_shared_vector_store=True)
        self.B = School.objects.create(name="B", access_shared_vector_store=False)

        def mk(school, vis):
            m = Material.objects.create(class_name="10", subject="Physics", title="t",
                                        type="notes", file="", school=school, visibility=vis)
            c = MaterialChunk.objects.create(material=m, school=school, class_name="10",
                                             subject="physics", content="c", chunk_index=0,
                                             provider="local", embedding_local=[0.0] * 768)
            ChunkChapter.objects.create(chunk=c, unit="light")
            return c

        self.shared = mk(None, "shared")
        self.a_priv = mk(self.A, "private")
        self.b_priv = mk(self.B, "private")
        self.b_inst = mk(self.B, "institutional")

    def _scoped(self, school_id):
        from core.embeddings import _scoped_chunks
        return set(_scoped_chunks("10", "physics", "embedding_local", school_id)
                   .values_list("id", flat=True))

    def test_with_access(self):
        self.assertEqual(self._scoped(self.A.id), {self.shared.id, self.a_priv.id, self.b_inst.id})

    def test_without_access(self):
        self.assertEqual(self._scoped(self.B.id), {self.b_priv.id, self.b_inst.id})

    def test_superadmin(self):
        self.assertEqual(self._scoped(None), {self.shared.id, self.b_inst.id})


class CrossSchoolLinkTest(TestCase):
    """SchoolVectorLink: a viewer school reads a source school's private materials ONLY when a
    link exists, and only in the granted direction."""

    def setUp(self):
        from core.models import School, Material
        self.A = School.objects.create(name="A")
        self.B = School.objects.create(name="B")
        mk = lambda school: Material.objects.create(class_name="10", subject="Physics", title="t",
                                                    type="notes", file="", school=school, visibility="private")
        self.a_priv = mk(self.A)
        self.b_priv = mk(self.B)

    def _visible(self, school):
        from core.models import Material
        from core.access import visibility_q
        return set(Material.objects.filter(visibility_q(school)).values_list("id", flat=True))

    def test_no_link_is_isolated(self):
        self.assertEqual(self._visible(self.A), {self.a_priv.id})
        self.assertEqual(self._visible(self.B), {self.b_priv.id})

    def test_link_is_directional(self):
        from core.models import SchoolVectorLink
        SchoolVectorLink.objects.create(viewer=self.A, source=self.B)   # A → B only
        self.assertEqual(self._visible(self.A), {self.a_priv.id, self.b_priv.id})  # A sees B's
        self.assertEqual(self._visible(self.B), {self.b_priv.id})                  # B still can't see A's

    def test_mutual_link(self):
        from core.models import SchoolVectorLink
        SchoolVectorLink.objects.create(viewer=self.A, source=self.B)
        SchoolVectorLink.objects.create(viewer=self.B, source=self.A)
        self.assertIn(self.b_priv.id, self._visible(self.A))
        self.assertIn(self.a_priv.id, self._visible(self.B))

    def test_chunk_scope_respects_link(self):
        from core.models import MaterialChunk, ChunkChapter, SchoolVectorLink
        from core.embeddings import _scoped_chunks
        c = MaterialChunk.objects.create(material=self.b_priv, school=self.B, class_name="10",
                                         subject="physics", content="x", chunk_index=0,
                                         provider="local", embedding_local=[0.0] * 768)
        ChunkChapter.objects.create(chunk=c, unit="light")
        scoped = lambda: set(_scoped_chunks("10", "physics", "embedding_local", self.A.id).values_list("id", flat=True))
        self.assertNotIn(c.id, scoped())                                # not linked → invisible
        SchoolVectorLink.objects.create(viewer=self.A, source=self.B)
        self.assertIn(c.id, scoped())                                   # linked → visible

    def test_no_self_link_allowed(self):
        from core.models import SchoolVectorLink
        from django.db import IntegrityError, transaction
        with self.assertRaises(IntegrityError):
            with transaction.atomic():
                SchoolVectorLink.objects.create(viewer=self.A, source=self.A)


# ─────────────────────────────────────────────
# Per-question structure (question_slots) — docs/PER_QUESTION_STRUCTURE.md
# ─────────────────────────────────────────────

from core import pattern_structure as psx
from core.tasks import _fill_section_counts


def _slot_sections():
    """Compact PT-1-style pattern: a uniform Grammar section with per-slot topics
    (incl. the teacher's real 9×1-vs-7 marks conflict) and a mixed Literature
    section with an extract (sub-parts), an open-choice group and internal choice."""
    return [
        {"id": "SEC_C", "name": "Grammar", "marks": 7, "question_slots": [
            {"qnum": 1, "type": "mcq", "topic": "Homophones", "format": "Homophones MCQ", "marks": 1},
            {"qnum": 2, "type": "fill_blank", "topic": "Conjunctions", "marks": 1},
            {"qnum": 3, "type": "fill_blank", "topic": "Conjunctions", "marks": 1},
            {"qnum": 4, "type": "punctuation", "topic": "Contracted words", "marks": 1},
            {"qnum": 5, "type": "mcq", "topic": "Present progressive", "marks": 1},
            {"qnum": 6, "type": "rewrite", "topic": "Punctuation", "marks": 1},
            {"qnum": 7, "type": "error_correction", "topic": "Past tense", "marks": 1},
            {"qnum": 8, "type": "mcq", "topic": "Past progressive", "marks": 1},
            {"qnum": 9, "type": "error_correction", "topic": "Past perfect", "marks": 1},
        ]},
        {"id": "SEC_D", "name": "Literature", "marks": 19, "question_slots": [
            {"qnum": 10, "type": "extract", "marks": 5, "source": "textbook",
             "parts": [{"label": "A", "type": "mcq", "marks": 1},
                       {"label": "B", "type": "mcq", "marks": 1},
                       {"label": "C", "type": "sa", "marks": 1},
                       {"label": "D", "type": "sa", "marks": 1},
                       {"label": "E", "type": "sa", "marks": 1}]},
            {"qnum": 11, "type": "sa", "marks": 10, "choice": "open", "attempt": 5,
             "parts": [{"label": l, "type": "sa", "marks": 2} for l in "ABCDEF"]},
            {"qnum": 12, "type": "la", "marks": 4, "choice": "internal",
             "alternatives": ["critical from prose/poem", "direct from supplementary reader"]},
        ]},
    ]


class SlotStructureValidatorTest(TestCase):
    def test_normalize_coerces_and_canonicalises(self):
        secs = [{"name": "A", "question_slots": [
            {"qnum": "3", "type": "Fill in the blanks", "marks": "1", "choice": None},
        ]}]
        psx.normalize_slots(secs)
        s = secs[0]["question_slots"][0]
        self.assertEqual(s["qnum"], 3)
        self.assertEqual(s["type"], "fill_blank")
        self.assertEqual(s["marks"], 1)
        self.assertEqual(s["choice"], "none")

    def test_flags_teacher_marks_conflict(self):
        secs = psx.normalize_slots(_slot_sections())
        errors = psx.validate_pattern_structure(secs, declared_total=26)
        msgs = " | ".join(e["msg"] for e in errors)
        self.assertIn("slot marks sum to 9 but the section declares 7", msgs)
        # paper sum uses slot sums (9 + 19 = 28) vs the declared 26
        self.assertIn("paper marks sum to 28", msgs)

    def test_clean_pattern_validates(self):
        secs = psx.normalize_slots(_slot_sections())
        secs[0]["marks"] = 9
        self.assertEqual(psx.validate_pattern_structure(secs, declared_total=28), [])

    def test_qnum_duplicates_and_gaps(self):
        secs = [{"name": "A", "marks": 2, "question_slots": [
            {"qnum": 1, "type": "mcq", "marks": 1},
            {"qnum": 3, "type": "mcq", "marks": 1},
        ]}]
        msgs = " | ".join(e["msg"] for e in psx.validate_pattern_structure(secs))
        self.assertIn("no gaps", msgs)
        secs[0]["question_slots"][1]["qnum"] = 1
        msgs = " | ".join(e["msg"] for e in psx.validate_pattern_structure(secs))
        self.assertIn("ascending", msgs)
        self.assertIn("duplicate", msgs)

    def test_open_choice_rules(self):
        base = {"qnum": 1, "type": "sa", "marks": 4, "choice": "open",
                "parts": [{"label": "A", "marks": 2}, {"label": "B", "marks": 2}]}
        # attempt missing
        errs = psx.validate_pattern_structure([{"name": "A", "question_slots": [dict(base)]}])
        self.assertTrue(any("requires 'attempt'" in e["msg"] for e in errs))
        # attempt == len(parts) → not a choice
        errs = psx.validate_pattern_structure([{"name": "A", "question_slots": [dict(base, attempt=2)]}])
        self.assertTrue(any("less than the number of parts" in e["msg"] for e in errs))
        # valid: attempt 2 of 3 → 2×2 == marks 4
        ok = dict(base, attempt=2, parts=[{"label": c, "marks": 2} for c in "ABC"])
        self.assertEqual(psx.validate_pattern_structure([{"name": "A", "question_slots": [ok]}]), [])

    def test_parts_sum_rule(self):
        secs = [{"name": "A", "question_slots": [
            {"qnum": 1, "type": "extract", "marks": 5,
             "parts": [{"label": "A", "marks": 1}, {"label": "B", "marks": 1}]},
        ]}]
        msgs = " | ".join(e["msg"] for e in psx.validate_pattern_structure(secs))
        self.assertIn("parts marks sum to 2 but slot marks is 5", msgs)


class SlotAggregateDerivationTest(TestCase):
    def test_derive_uniform_and_mixed(self):
        secs = psx.normalize_slots(_slot_sections())
        secs[0]["marks_per_question"] = [1] * 9        # stale mpq-list must be replaced
        psx.derive_aggregates_from_slots(secs)
        g, l = secs
        self.assertEqual((g["questions_count"], g["marks"], g["marks_per_question"]), (9, 9, 1))
        self.assertEqual((l["questions_count"], l["marks"]), (3, 19))
        self.assertNotIn("marks_per_question", l)       # mixed marks → no scalar mpq
        # typed dicts with real qnum ranges, contiguous same-(type,marks) runs merged
        self.assertEqual(g["question_types"][1],
                         {"type": "Fill in the blank", "count": 2, "marks_each": 1, "range": "Q2-Q3"})
        self.assertTrue(all(isinstance(qt, dict) for qt in l["question_types"]))

    def test_model_totals_prefer_slots(self):
        p = ExamPattern(name="slots", sections=psx.normalize_slots(_slot_sections()))
        self.assertEqual(p.get_total_marks(), 28)       # 9 + 19 from slots (not section 7)
        self.assertEqual(p.get_total_questions(), 12)

    def test_fill_section_counts_skips_slot_sections(self):
        secs = psx.normalize_slots(_slot_sections())
        psx.derive_aggregates_from_slots(secs)
        out = _fill_section_counts(secs)
        self.assertNotIn("marks_per_question", out[1])  # no bogus avg re-invented for mixed


class SlotTypeCategoryTest(TestCase):
    def test_slot_types_never_other(self):
        for t, cat in psx.SLOT_TYPE_CATEGORY.items():
            self.assertEqual(sg._type_category(t), cat, t)

    def test_free_text_spellings(self):
        self.assertEqual(sg._type_category("Fill in the blanks"), "vsa")
        self.assertEqual(sg._type_category("Error correction"), "vsa")
        self.assertEqual(sg._type_category("Rewrite sentences"), "vsa")
        self.assertEqual(sg._type_category("Descriptive Paragraph"), "la")
        self.assertEqual(sg._type_category("Story Writing"), "la")

    def test_legacy_classifications_unchanged(self):
        self.assertEqual(sg._type_category("MCQ"), "mcq")
        self.assertEqual(sg._type_category("Short Answer"), "sa")
        self.assertEqual(sg._type_category("Long Answer"), "la")
        self.assertEqual(sg._type_category("Source-Based/CBQ"), "cbq")
        self.assertEqual(sg._type_category("Map work"), "map")


class SlotWorkOrderTest(TestCase):
    def _wos(self):
        secs = psx.normalize_slots(_slot_sections())
        secs[0]["marks"] = 9
        psx.derive_aggregates_from_slots(secs)
        pattern = ExamPattern(name="p", sections=secs)
        blueprint = sg_gen.pattern_sections_to_blueprint_dict(pattern)
        return sg.build_work_orders(blueprint, pattern, {}, "Medium", "6", "English", ["Ch1"])

    def test_counts_marks_and_slots(self):
        g, l = self._wos()
        self.assertEqual((g.questions_count, g.marks, g.marks_per_question, g.mixed_marks),
                         (9, 9, 1.0, False))
        self.assertEqual((l.questions_count, l.marks, l.mixed_marks), (3, 19, True))
        self.assertEqual(len(g.slots), 9)
        self.assertEqual(len(l.slots), 3)
        self.assertEqual(l.subsections, [])

    def test_parts_slots_count_as_cbq(self):
        _, l = self._wos()
        cats = [sg._fine_category(qt["type"]) for qt in l.question_types]
        self.assertEqual(cats, ["cbq", "cbq", "la"])    # extract + open-choice group + LA

    def test_slot_prompt_block(self):
        g, l = self._wos()
        pg, pl = sg.build_section_prompt(g), sg.build_section_prompt(l)
        self.assertIn("PER-QUESTION SPECIFICATION", pg)
        self.assertIn("Homophones", pg)                          # topics reach the LLM
        self.assertNotIn("QUESTION-POSITION BLUEPRINT", pl)      # superseded for slot sections
        self.assertIn("INTERNAL CHOICE", pl)
        self.assertIn("attempt any 5", pl.lower())
        self.assertIn('"source_text"', pl)
        # No blanket LA OR-rule: only slot-flagged questions get or_alternative
        self.assertIn('Do NOT add "or_alternative" to any other question', pl)

    def test_or_gating_in_validation(self):
        _, l = self._wos()
        qs = [
            {"type": "CBQ", "subtype": "source_based", "marks": 5, "text": "Read the extract " * 10,
             "source_text": ("The mongoose ran across the sunlit garden and the boy chased it. " * 3).strip(),
             "competency_type": "application",
             "sub_questions": ([{"text": "pick one a) w b) x c) y d) z", "marks": 1}] * 2
                               + [{"text": "a", "marks": 1}] * 3)},
            {"type": "CBQ", "subtype": "standard", "marks": 10, "text": "Attempt any 5 " * 5,
             "competency_type": "constructed",
             "sub_questions": [{"text": "b", "marks": 2}] * 6},   # 12 provided vs 10 attempted
            {"type": "LA", "marks": 4, "text": "Discuss the theme of the poem in detail.",
             "competency_type": "constructed", "answer_explanation": "points",
             "or_alternative": "Describe the character of the mongoose."},
        ]
        errors = sg.validate_section_output({"questions": qs}, l)
        self.assertEqual(errors, [])
        # internal-choice slot without or_alternative must fail
        del qs[2]["or_alternative"]
        errors = sg.validate_section_output({"questions": qs}, l)
        self.assertTrue(any("or_alternative" in e for e in errors))
        # open-choice slot: a sum that matches NEITHER attempted nor provided fails
        qs[2]["or_alternative"] = "alt"
        qs[1]["sub_questions"] = [{"text": "b", "marks": 2}] * 4
        errors = sg.validate_section_output({"questions": qs}, l)
        self.assertTrue(any("sub_question marks sum" in e for e in errors))


class SlotRenderTest(TestCase):
    def test_regroup_skips_slot_sections(self):
        sec_info = {"question_slots": [{"qnum": 1}], "subsections": [
            {"name": "SA", "question_types": ["SA"], "marks": 10, "questions_count": 5,
             "marks_per_question": 2},
        ]}
        qs = [{"type": "CBQ", "marks": 5}, {"type": "SA", "marks": 10}, {"type": "LA", "marks": 4}]
        groups = sg_gen._regroup_section(list(qs), sec_info)
        self.assertEqual(groups, [(None, qs)])           # authored order, marks untouched

    def test_blueprint_dict_passes_slots(self):
        secs = psx.normalize_slots(_slot_sections())
        bp = sg_gen.pattern_sections_to_blueprint_dict(ExamPattern(name="p", sections=secs))
        self.assertEqual(len(bp["Grammar"]["question_slots"]), 9)


class SlotRepairGuardTest(TestCase):
    """The repair round must never be accepted when it deletes question slots or
    lowers slot marks to force the teacher's declared totals to match (observed in
    production: Q19/Q20 dropped and a 4m LA cut to 3m). repair_preserves_slots is
    the deterministic guard in generate_pattern_task."""

    def _orig(self):
        return psx.normalize_slots(_slot_sections())

    def test_rejects_deleted_slots(self):
        import copy
        orig = self._orig()
        mutilated = copy.deepcopy(orig)
        mutilated[0]["question_slots"] = mutilated[0]["question_slots"][:7]  # drop Q8, Q9
        self.assertFalse(psx.repair_preserves_slots(orig, mutilated))

    def test_rejects_lowered_slot_marks(self):
        import copy
        orig = self._orig()
        devalued = copy.deepcopy(orig)
        devalued[1]["question_slots"][2]["marks"] = 3     # LA 4m -> 3m
        self.assertFalse(psx.repair_preserves_slots(orig, devalued))

    def test_accepts_section_marks_correction(self):
        import copy
        orig = self._orig()
        good = copy.deepcopy(orig)
        good[0]["marks"] = 9                              # totals fixed, slots untouched
        self.assertTrue(psx.repair_preserves_slots(orig, good))

    def test_accepts_added_slots_and_filled_marks(self):
        import copy
        orig = self._orig()
        orig[0]["question_slots"][0]["marks"] = 0         # pass 1 left it blank
        richer = copy.deepcopy(orig)
        richer[0]["question_slots"][0]["marks"] = 1       # repair filled it in
        richer[1]["question_slots"].append({"qnum": 13, "type": "sa", "marks": 2})
        self.assertTrue(psx.repair_preserves_slots(orig, richer))


class SlotEditRederiveTest(TestCase):
    """PUT /api/patterns/{id}/ with edited slots must re-derive the aggregates and
    refresh structure warnings server-side (ExamPatternViewSet.perform_update) —
    the frontend slot editor sends slots only, never recomputed totals."""

    def test_put_rederives_aggregates_and_warnings(self):
        import copy
        from rest_framework.test import APIClient
        user = User.objects.create_user("slotedit", "s@x.com", "pw")
        secs = psx.normalize_slots(_slot_sections())
        secs[0]["marks"] = 9
        psx.derive_aggregates_from_slots(secs)
        p = ExamPattern.objects.create(name="se", sections=secs, created_by=user)

        api = APIClient()
        api.force_authenticate(user)
        edited = copy.deepcopy(p.sections)
        edited[0]["question_slots"][0]["marks"] = 2      # Grammar Q1: 1m -> 2m
        r = api.put(f"/api/patterns/{p.id}/", {
            "name": "se", "class_name": "6", "subject": "English",
            "description": "", "ai_prompt": "", "sections": edited,
        }, format="json")
        self.assertEqual(r.status_code, 200, r.content)

        p.refresh_from_db()
        g = p.sections[0]
        self.assertEqual(g["marks"], 10)                  # re-derived from edited slots
        self.assertEqual(g["questions_count"], 9)
        self.assertNotIn("marks_per_question", g)         # now mixed (2m + 1m×8)
        self.assertEqual(p.total_marks, 29)               # 10 + 19, recomputed
        # slot sum (10) no longer matches the stale section 'marks' we sent (9)?
        # No — derive overwrote it; but the section-marks warning must NOT appear:
        self.assertFalse(any("_structure_warnings" in s and
                             any("declares" in w for w in s["_structure_warnings"])
                             for s in p.sections))


class ARRenderMergeTest(TestCase):
    """Assertion-Reason questions are MCQ variants: they render INSIDE the 'Multiple
    Choice Questions' display group (after the plain MCQs), never under a separate
    'II. Assertion–Reason Questions' subheader. Reported on a Social Science paper
    whose Section A printed its 3 AR questions as a standalone trailing group."""

    SEC_INFO = {
        "name": "Section A — Objective Type", "marks": 14,
        "subsections": [
            {"name": "MCQ", "question_types": ["MCQ"], "marks": 7, "questions_count": 7, "marks_per_question": 1},
            {"name": "Assertion-Reason", "question_types": ["Assertion-Reason"], "marks": 3, "questions_count": 3, "marks_per_question": 1},
            {"name": "VSA", "question_types": ["VSA"], "marks": 4, "questions_count": 2, "marks_per_question": 2},
        ],
    }

    def _questions(self):
        # ARs deliberately interleaved among the MCQs — the merge must still put them last.
        mk = lambda st: {"type": "MCQ", "subtype": st, "marks": 1}
        return ([mk("assertion_reason"), mk("standard"), mk("standard"), mk("assertion_reason"),
                 mk("standard"), mk("standard"), mk("standard"), mk("assertion_reason"),
                 mk("standard"), mk("standard")]
                + [{"type": "VSA", "subtype": "standard", "marks": 2} for _ in range(2)])

    def test_no_separate_ar_group(self):
        groups = sg_gen._regroup_section(self._questions(), self.SEC_INFO)
        labels = [lbl for lbl, _ in groups if lbl]
        self.assertFalse(any("Assertion" in l for l in labels), labels)
        # MCQ group holds all 10: the 7 standard first, the 3 AR after them
        mcq_qs = groups[0][1]
        self.assertEqual([q["subtype"] for q in mcq_qs],
                         ["standard"] * 7 + ["assertion_reason"] * 3)
        # Roman numbering stays contiguous: I. MCQ, II. VSA
        self.assertTrue(labels[0].startswith("I.") and "Multiple Choice" in labels[0])
        self.assertTrue(labels[1].startswith("II.") and "Very Short" in labels[1])

    def test_mcq_plus_ar_only_section_has_no_subheaders(self):
        qs = ([{"type": "MCQ", "subtype": "standard", "marks": 1} for _ in range(7)]
              + [{"type": "MCQ", "subtype": "assertion_reason", "marks": 1} for _ in range(3)])
        sec = {"name": "A", "subsections": self.SEC_INFO["subsections"][:2]}
        groups = sg_gen._regroup_section(qs, sec)
        self.assertEqual([lbl for lbl, _ in groups if lbl], [])   # one category → no labels
        self.assertEqual(len(groups[0][1]), 10)

    def test_ar_keeps_distinct_marks_when_pattern_differs(self):
        sec = {"name": "A", "subsections": [
            {"name": "MCQ", "question_types": ["MCQ"], "marks": 4, "questions_count": 4, "marks_per_question": 1},
            {"name": "Assertion-Reason", "question_types": ["Assertion-Reason"], "marks": 4, "questions_count": 2, "marks_per_question": 2},
        ]}
        qs = [{"type": "MCQ", "subtype": "standard", "marks": 9.9},
              {"type": "MCQ", "subtype": "assertion_reason", "marks": 9.9}]
        merged = sg_gen._regroup_section(qs, sec)[0][1]
        self.assertEqual([q["marks"] for q in merged], [1, 2])    # stamped per fine category

    def test_ar_bucket_surfaces_under_mcq_heading(self):
        # AR + LA mix, no plain MCQs: the AR questions take the MCQ group's slot/label.
        qs = [{"type": "MCQ", "subtype": "assertion_reason", "marks": 1},
              {"type": "LA", "subtype": "standard", "marks": 5}]
        groups = sg_gen._regroup_section(qs, {"name": "A", "subsections": []})
        labels = [lbl for lbl, _ in groups if lbl]
        self.assertTrue(labels[0].startswith("I.") and "Multiple Choice" in labels[0])
        self.assertFalse(any("Assertion" in l for l in labels))


class OrAlternativeRenderTest(TestCase):
    """A dict or_alternative is a COMPLETE second question (its own passage and its
    own sub-questions) — the renderer must print all of it after the OR separator,
    not just its 'text' line. Production bug: Q21's second extract printed as a bare
    '21. Read the source above…' header with no passage and no sub-questions, and
    Q22's parts vanished when the stem was stored under 'question'."""

    def _q21(self):
        return {
            "qnum": 21, "text": "Read the source above and answer the following:",
            "type": "CBQ", "subtype": "source_based", "marks": 5,
            "source_text": "First extract line one. " * 10,
            "sub_questions": [{"text": f"first sub {i}", "marks": 1} for i in range(3)],
            "or_alternative": {
                "text": "Read the source above and answer the following:",
                "source_text": "Second extract line one. " * 10,
                "sub_questions": [{"text": f"second sub {i}", "marks": 1} for i in range(3)],
            },
        }

    def test_dict_alternative_renders_passage_and_subquestions(self):
        out = []
        sg_gen.process_question(out, self._q21(), 21)
        self.assertEqual(len([1 for k, _ in out if k == "passage"]), 2)   # one per option
        or_i = out.index(("or", "OR"))
        after = out[or_i:]
        self.assertTrue(any(k == "passage" and "Second extract" in str(t) for k, t in after))
        self.assertEqual(len([1 for k, _ in after if k == "subq"]), 3)
        self.assertTrue(any(k == "q" and str(t).startswith("21.") for k, t in after))

    def test_string_alternative_unchanged(self):
        q = {"qnum": 23, "text": "Discuss the theme.", "type": "LA", "marks": 4,
             "or_alternative": "Describe the character."}
        out = []
        sg_gen.process_question(out, q, 23)
        or_i = out.index(("or", "OR"))
        self.assertIn("Describe the character.", out[or_i + 1][1])
        self.assertIn("[4 marks]", out[or_i + 1][1])

    def test_question_key_branch_renders_sub_questions(self):
        q = {"qnum": 22, "question": "Attempt any 5 of the following 6:", "marks": 10,
             "sub_questions": [{"text": f"part {i}", "marks": 2} for i in range(6)]}
        out = []
        sg_gen.process_question(out, q, 22)
        self.assertEqual(len([1 for k, _ in out if k == "subq"]), 6)

    def test_sub_question_alt_keys_not_dropped(self):
        q = {"qnum": 5, "text": "Attempt all:", "marks": 4,
             "sub_questions": [{"q": "keyed q", "marks": 2},
                               {"question": "keyed question", "marks": 2}]}
        out = []
        sg_gen.process_question(out, q, 5)
        subs = [t for k, t in out if k == "subq"]
        self.assertEqual(len(subs), 2)
        self.assertIn("keyed q", subs[0])
        self.assertIn("keyed question", subs[1])


class GeneralSourceSlotTest(TestCase):
    """source='general' slots ("give in general, NOT from the text book") must reach
    the generation prompt as an explicit original-question instruction, and an
    all-general section must receive NO RAG context at all."""

    def _general_sections(self):
        return [{"id": "SEC_B", "name": "Writing", "marks": 5, "question_slots": [
            {"qnum": 1, "type": "writing", "marks": 5, "source": "general", "choice": "internal",
             "alternatives": ["descriptive paragraph", "informal letter"]},
        ]}]

    def test_normalize_canonicalises_source(self):
        secs = [{"name": "A", "question_slots": [
            {"qnum": 1, "type": "sa", "marks": 2, "source": "General Knowledge"},
            {"qnum": 2, "type": "sa", "marks": 2, "source": ""},
        ]}]
        psx.normalize_slots(secs)
        self.assertEqual(secs[0]["question_slots"][0]["source"], "general")
        self.assertNotIn("source", secs[0]["question_slots"][1])

    def test_slots_all_general(self):
        self.assertTrue(sg._slots_all_general({"question_slots": [
            {"qnum": 1, "source": "general"}, {"qnum": 2, "source": "general"}]}))
        self.assertFalse(sg._slots_all_general({"question_slots": [
            {"qnum": 1, "source": "general"}, {"qnum": 2, "source": "textbook"}]}))
        self.assertFalse(sg._slots_all_general({"question_slots": []}))
        self.assertFalse(sg._slots_all_general({"question_slots": [{"qnum": 1}]}))

    def test_all_general_wo_gets_no_context(self):
        secs = psx.normalize_slots(self._general_sections())
        psx.derive_aggregates_from_slots(secs)
        pattern = ExamPattern(name="p", sections=secs)
        bp = sg_gen.pattern_sections_to_blueprint_dict(pattern)
        ctx_map = {"Writing": "textbook chunk " * 50,
                   "__context_by_type__": {"Writing": {"la": "chunk"}}}
        wo = sg.build_work_orders(bp, pattern, ctx_map, "Medium", "6", "English", ["Ch1"])[0]
        self.assertEqual(wo.context_text, "")
        self.assertEqual(wo.context_by_type, {})
        prompt = sg.build_section_prompt(wo)
        self.assertIn("GENERAL KNOWLEDGE", prompt)
        self.assertIn("Compose original questions", prompt)
        self.assertNotIn("Draw question content from the reference material", prompt)

    def test_mixed_source_prompt_scopes_reference_material(self):
        secs = psx.normalize_slots([{"id": "SEC_D", "name": "Mixed", "marks": 7, "question_slots": [
            {"qnum": 1, "type": "extract", "marks": 5, "source": "textbook",
             "parts": [{"label": c, "type": "sa", "marks": 1} for c in "ABCDE"]},
            {"qnum": 2, "type": "mcq", "marks": 2, "source": "general"},
        ]}])
        psx.derive_aggregates_from_slots(secs)
        pattern = ExamPattern(name="p", sections=secs)
        bp = sg_gen.pattern_sections_to_blueprint_dict(pattern)
        wo = sg.build_work_orders(bp, pattern, {"Mixed": "textbook chunk " * 50},
                                  "Medium", "6", "English", ["Ch1"])[0]
        self.assertEqual(wo.context_text, "textbook chunk " * 50)   # mixed keeps context
        prompt = sg.build_section_prompt(wo)
        self.assertIn("applies ONLY to questions marked textbook/unseen", prompt)
        self.assertIn("GENERAL KNOWLEDGE — do NOT", prompt)
        self.assertIn("VERBATIM", prompt)                # textbook extract must be quoted
        # mixed sections keep the chapter plan but exempt their general questions
        self.assertIn("EXCEPTION: questions marked GENERAL KNOWLEDGE", prompt)

    def test_all_general_prompt_has_no_chapter_assignment(self):
        # Production leak: Grammar was all-general, context was blanked, yet the chapter
        # block still ordered "draw the 9 questions from these chapters" — and Q19 came
        # back as a 'Wit and Humour' comprehension question.
        secs = psx.normalize_slots(self._general_sections())
        psx.derive_aggregates_from_slots(secs)
        pattern = ExamPattern(name="p", sections=secs)
        bp = sg_gen.pattern_sections_to_blueprint_dict(pattern)
        wo = sg.build_work_orders(bp, pattern, {}, "Medium", "6", "English",
                                  ["Learning Together", "Wit and Humour"])[0]
        prompt = sg.build_section_prompt(wo)
        self.assertIn("CHAPTER ASSIGNMENT: NONE", prompt)
        self.assertNotIn("CHAPTER ASSIGNMENT — MANDATORY", prompt)
        self.assertNotIn("no chapter monopoly", prompt)
        self.assertIn("Do NOT reference the textbook", prompt)

    def test_general_question_mentioning_chapter_fails_validation(self):
        secs = psx.normalize_slots([{"id": "SEC_C", "name": "Grammar", "marks": 2,
            "question_slots": [
                {"qnum": 1, "type": "rewrite", "marks": 1, "topic": "Present perfect",
                 "source": "general"},
                {"qnum": 2, "type": "rewrite", "marks": 1, "topic": "Present perfect",
                 "source": "general"},
            ]}])
        psx.derive_aggregates_from_slots(secs)
        pattern = ExamPattern(name="p", sections=secs)
        bp = sg_gen.pattern_sections_to_blueprint_dict(pattern)
        wo = sg.build_work_orders(bp, pattern, {}, "Medium", "6", "English",
                                  ["Wit and Humour"])[0]
        leaky = {"type": "VSA", "marks": 1, "answer_explanation": "a",
                 "text": "In the story 'The Open Window' from the chapter 'Wit and Humour', "
                         "what does Framton Nuttel suffer from?"}
        clean = {"type": "VSA", "marks": 1, "answer_explanation": "a",
                 "text": "Rewrite in the present perfect tense: 'She writes a letter.'"}
        errors = sg.validate_section_output({"questions": [leaky, clean]}, wo)
        self.assertTrue(any("GENERAL KNOWLEDGE" in e and "Wit and Humour" in e
                            for e in errors), errors)
        errors = sg.validate_section_output({"questions": [clean, clean]}, wo)
        self.assertFalse(any("GENERAL KNOWLEDGE" in e for e in errors), errors)

    def test_context_map_skips_retrieval_for_all_general(self):
        bp = {"Writing": {"question_slots": [{"qnum": 1, "source": "general"}],
                          "questions_count": 1, "marks": 5}}
        cmap = sg.get_section_context_map("6", "English", ["Ch1"], bp, ["SA"])
        self.assertEqual(cmap["Writing"], "")
        self.assertEqual(cmap["__context_by_type__"]["Writing"], {})


class EnglishGrammarOwnKnowledgeTest(TestCase):
    """English GRAMMAR questions and CREATIVE WRITING tasks must be composed from the model's OWN
    knowledge — NOTHING may come from the reference material handed to the LLM.

    Two live leaks, one mechanism. NCERT English readers carry no grammar LESSONS, so
    identify_grammar_chapters found nothing to route a grammar section to and it retrieved prose
    instead: "gap filling" and "editing" questions came back built out of story sentences and
    tagged to literature chapters. And a Creative Writing section opened BOTH options of its
    internal choice with "After reading 'The Laburnum Top', you are inspired by the theme of
    nature's vitality. Write an article …" — a composition brief hung off a retrieved poem."""

    def _pattern(self, sections):
        secs = psx.normalize_slots(sections)
        psx.derive_aggregates_from_slots(secs)
        return ExamPattern(name="p", subject="English", class_name="10", sections=secs)

    def _grammar_slots(self):
        return [{"qnum": 3, "type": "fill_blank", "marks": 1, "topic": "Tenses"},
                {"qnum": 4, "type": "error_correction", "marks": 1},
                {"qnum": 5, "type": "mcq", "marks": 1, "topic": "Prepositions"}]

    def _writing_slots(self):
        return [{"qnum": 2, "type": "writing", "marks": 4, "topic": "Article",
                 "choice": "internal", "alternatives": ["magazine article", "newspaper article"]},
                {"qnum": 3, "type": "writing", "marks": 4, "topic": "Classified advertisement"},
                {"qnum": 4, "type": "la", "marks": 4, "topic": "Letter to the editor"}]

    def _wo(self, sections, ctx_map, chapters=("A Letter to God", "Dust of Snow")):
        p = self._pattern(sections)
        bp = sg_gen.pattern_sections_to_blueprint_dict(p)
        return {wo.section_name: wo for wo in sg.build_work_orders(
            bp, p, ctx_map, "Medium", "10", "English", list(chapters))}

    # ── scope detection ──────────────────────────────────────────────────────────
    def test_scope_grammar_section_is_own_knowledge_only(self):
        sd = {"instructions": ["Attempt all questions."], "question_slots": self._grammar_slots()}
        kinds, only, slot_kinds = sg.english_own_scope(
            "English", "", "B — Grammar", sd, sd["question_slots"])
        self.assertEqual(kinds, ("grammar",))
        self.assertTrue(only)
        self.assertEqual(slot_kinds, {0: "grammar", 1: "grammar", 2: "grammar"})

    def test_scope_creative_writing_section_is_own_knowledge_only(self):
        # The reported leak: "After reading 'The Laburnum Top', … write an article" — the whole
        # Creative Writing section must be fenced off from the retrieved poem.
        sd = {"instructions": ["Attempt all questions."], "question_slots": self._writing_slots()}
        kinds, only, slot_kinds = sg.english_own_scope(
            "English", "", "SECTION – Creative Writing Skills", sd, sd["question_slots"])
        self.assertEqual(kinds, ("writing",))
        self.assertTrue(only)
        self.assertEqual(set(slot_kinds.values()), {"writing"})

    def test_scope_writing_section_names(self):
        for name in ("SECTION – Creative Writing Skills", "B — Writing", "Writing Skills",
                     "C — Composition", "Section B — Writing and Grammar"):
            kinds, only, _ = sg.english_own_scope("English", "", name, {"instructions": []}, [])
            self.assertTrue(kinds, name)
            self.assertTrue(only, name)

    def test_scope_literature_section_untouched(self):
        sd = {"instructions": [], "question_slots": [
            {"qnum": 8, "type": "extract", "marks": 5, "source": "textbook"},
            {"qnum": 9, "type": "sa", "marks": 3, "topic": "A Letter to God"}]}
        self.assertEqual(
            sg.english_own_scope("English", "", "D — Literature", sd, sd["question_slots"]),
            ((), False, {}))

    def test_scope_skips_non_english_subjects(self):
        # Tamil/Hindi grammar keeps its existing grammar-LESSON routing untouched.
        slots = [{"type": "fill_blank", "marks": 1, "topic": "எழுத்து"}]
        self.assertEqual(sg.english_own_scope("Tamil", "", "இலக்கணம்", {}, slots), ((), False, {}))
        self.assertEqual(sg.english_own_scope("Hindi", "", "खंड-ख — व्याकरण (Grammar)", {}, []),
                         ((), False, {}))

    def test_scope_mixed_grammar_and_writing_section(self):
        # The reported section held both: an article and an advertisement (writing) alongside a
        # word-rearrangement and conditional-clause question (grammar). All four are fenced.
        slots = [{"type": "writing", "marks": 4, "topic": "Article"},
                 {"type": "writing", "marks": 4, "topic": "Classified advertisement"},
                 {"type": "rewrite", "marks": 2, "topic": "Rearrange the words"},
                 {"type": "mcq", "marks": 2, "topic": "Conditional clauses"}]
        kinds, only, slot_kinds = sg.english_own_scope(
            "English", "", "SECTION – Creative Writing Skills", {}, slots)
        self.assertEqual(sorted(kinds), ["grammar", "writing"])
        self.assertTrue(only)
        self.assertEqual(slot_kinds,
                         {0: "writing", 1: "writing", 2: "grammar", 3: "grammar"})

    def test_scope_exempts_literature_wording(self):
        # A hybrid section must not cost its literature questions the material they need.
        slots = [{"type": "sa", "marks": 3, "topic": "The poem Dust of Snow"},
                 {"type": "fill_blank", "marks": 1, "topic": "Articles"}]
        kinds, only, slot_kinds = sg.english_own_scope(
            "English", "", "E — Literature and Grammar", {}, slots)
        self.assertEqual((kinds, only, slot_kinds), (("grammar",), False, {1: "grammar"}))
        # slot-less (legacy subsection) hybrid: rule applies, context stays
        self.assertEqual(
            sg.english_own_scope("English", "", "E — Literature and Grammar", {}, []),
            (("grammar",), False, {}))
        # slot-less pure grammar section: no context at all
        self.assertEqual(
            sg.english_own_scope("English Core", "", "B — Grammar", {"instructions": []}, []),
            (("grammar",), True, {}))

    def test_writing_form_beats_a_bare_literature_word(self):
        # "story writing" is a composition task; "story" alone is a literature topic. An
        # explicit form is the stronger signal, so it must win the exemption.
        _, _, kinds = sg.english_own_scope("English", "", "B — Writing", {}, [
            {"type": "sa", "marks": 5, "topic": "Story writing"},
            {"type": "sa", "marks": 5, "topic": "The story of Lencho and his cornfield"}])
        self.assertEqual(kinds.get(0), "writing")
        self.assertNotIn(1, kinds)      # literature question keeps its context

    # ── retrieval + work order ───────────────────────────────────────────────────
    def test_context_map_skips_retrieval_for_grammar_section(self):
        bp = {"B — Grammar": {"question_slots": self._grammar_slots(),
                              "questions_count": 3, "marks": 3}}
        cmap = sg.get_section_context_map("10", "English", ["Ch1"], bp, ["MCQ"])
        self.assertEqual(cmap["B — Grammar"], "")
        self.assertEqual(cmap["__context_by_type__"]["B — Grammar"], {})

    def test_grammar_wo_gets_no_context_and_general_slots(self):
        wos = self._wo(
            [{"id": "SEC_B", "name": "B — Grammar", "marks": 3,
              "question_slots": self._grammar_slots()}],
            {"B — Grammar": "Lencho story chunk " * 40,
             "__context_by_type__": {"B — Grammar": {"mcq": "chunk"}}})
        wo = wos["B — Grammar"]
        self.assertTrue(wo.is_english_grammar)
        self.assertTrue(wo.english_own_only)
        self.assertEqual(wo.context_text, "")          # withheld even when the map has one
        self.assertEqual(wo.context_by_type, {})
        self.assertEqual([s.get("source") for s in wo.slots], ["general"] * 3)

    def test_context_map_skips_retrieval_for_creative_writing_section(self):
        bp = {"SECTION – Creative Writing Skills": {
            "question_slots": self._writing_slots(), "questions_count": 3, "marks": 12}}
        cmap = sg.get_section_context_map("11", "English Core", ["The Laburnum Top"], bp, ["LA"])
        self.assertEqual(cmap["SECTION – Creative Writing Skills"], "")

    def test_creative_writing_wo_gets_no_context(self):
        wos = self._wo(
            [{"id": "SEC_B", "name": "SECTION – Creative Writing Skills", "marks": 12,
              "question_slots": self._writing_slots()}],
            {"SECTION – Creative Writing Skills": "The Laburnum Top awake " * 40},
            chapters=("The Laburnum Top", "The Address"))
        wo = wos["SECTION – Creative Writing Skills"]
        self.assertTrue(wo.is_english_writing)
        self.assertTrue(wo.english_own_only)
        self.assertEqual(wo.context_text, "")          # withheld even when the map has one
        self.assertEqual([s.get("source") for s in wo.slots], ["general"] * 3)

    def test_mixed_section_keeps_context_but_fences_own_knowledge_slot(self):
        # A LITERATURE question alongside a grammar one — the literature side keeps the material
        # it needs, the grammar side is fenced off per-question.
        wos = self._wo(
            [{"id": "SEC_C", "name": "E — Literature and Grammar", "marks": 6,
              "question_slots": [
                  {"qnum": 1, "type": "sa", "marks": 5, "topic": "The poem Dust of Snow"},
                  {"qnum": 2, "type": "rewrite", "marks": 1, "topic": "Passive voice"}]}],
            {"E — Literature and Grammar": "textbook chunk " * 40})
        wo = wos["E — Literature and Grammar"]
        self.assertTrue(wo.is_english_grammar)
        self.assertFalse(wo.is_english_writing)
        self.assertFalse(wo.english_own_only)
        self.assertEqual(wo.context_text, "textbook chunk " * 40)
        self.assertEqual([s.get("source") for s in wo.slots], [None, "general"])

    # ── prompt ───────────────────────────────────────────────────────────────────
    def test_grammar_prompt_bans_reference_material_and_chapters(self):
        wo = self._wo([{"id": "SEC_B", "name": "B — Grammar", "marks": 3,
                        "question_slots": self._grammar_slots()}], {})["B — Grammar"]
        prompt = sg.build_section_prompt(wo)
        self.assertIn("ENGLISH GRAMMAR — ABSOLUTE RULE", prompt)
        self.assertIn("Take NOTHING from the REFERENCE MATERIAL", prompt)
        self.assertIn("CHAPTER ASSIGNMENT: NONE", prompt)
        self.assertNotIn("CHAPTER ASSIGNMENT — MANDATORY", prompt)
        self.assertIn('its "chapter_tag" to "Grammar"', prompt)
        # the chapter list must not appear anywhere — naming it re-invites what the rule forbids
        self.assertNotIn("A Letter to God", prompt)
        self.assertNotIn("Dust of Snow", prompt)

    def test_creative_writing_prompt_bans_after_reading_openers(self):
        wo = self._wo(
            [{"id": "SEC_B", "name": "SECTION – Creative Writing Skills", "marks": 12,
              "question_slots": self._writing_slots()}], {},
            chapters=("The Laburnum Top", "The Address"))["SECTION – Creative Writing Skills"]
        prompt = sg.build_section_prompt(wo)
        self.assertIn("CREATIVE WRITING — ABSOLUTE RULE", prompt)
        self.assertIn("SELF-CONTAINED", prompt)
        self.assertIn('never open with "After reading', prompt)
        self.assertIn("BOTH options must be independent briefs", prompt)
        self.assertIn('its "chapter_tag" to "Writing"', prompt)
        self.assertIn("CHAPTER ASSIGNMENT: NONE", prompt)
        # the poem the leaked questions were built on must not be named anywhere
        self.assertNotIn("The Laburnum Top", prompt)
        self.assertNotIn("The Address", prompt)

    def test_mixed_grammar_and_writing_prompt_carries_both_rules(self):
        wo = self._wo(
            [{"id": "SEC_B", "name": "SECTION – Creative Writing Skills", "marks": 12,
              "question_slots": [
                  {"qnum": 2, "type": "writing", "marks": 4, "topic": "Article"},
                  {"qnum": 4, "type": "rewrite", "marks": 2, "topic": "Rearrange the words"}]}],
            {}, chapters=("The Laburnum Top",))["SECTION – Creative Writing Skills"]
        prompt = sg.build_section_prompt(wo)
        self.assertTrue(wo.is_english_writing)
        self.assertTrue(wo.is_english_grammar)
        self.assertIn("CREATIVE WRITING — ABSOLUTE RULE", prompt)
        self.assertIn("ENGLISH GRAMMAR — ABSOLUTE RULE", prompt)
        self.assertIn("ENGLISH GRAMMAR and CREATIVE WRITING, set from your own knowledge", prompt)

    def test_mixed_prompt_states_the_rule_and_keeps_chapters(self):
        wo = self._wo(
            [{"id": "SEC_C", "name": "E — Literature and Grammar", "marks": 6,
              "question_slots": [
                  {"qnum": 1, "type": "sa", "marks": 5, "topic": "The poem Dust of Snow"},
                  {"qnum": 2, "type": "rewrite", "marks": 1, "topic": "Passive voice"}]}],
            {"E — Literature and Grammar": "textbook chunk " * 40})["E — Literature and Grammar"]
        prompt = sg.build_section_prompt(wo)
        self.assertIn("ENGLISH GRAMMAR — ABSOLUTE RULE", prompt)
        self.assertIn("GENERAL KNOWLEDGE — do NOT take this question", prompt)
        self.assertIn("A Letter to God", prompt)     # literature slot still needs its chapters

    def test_non_english_grammar_prompt_unchanged(self):
        wo = sg.SectionWorkOrder(
            section_name="இலக்கணம்", section_id="C", title="", marks=5, questions_count=5,
            marks_per_question=1, question_types=["vsa"], instructions=[], constraints={},
            context_text="tamil grammar chunk", difficulty="Medium", subject="Tamil",
            class_name="6", chapters=["மொழிமுதல் எழுத்துகள்"], is_grammar=True)
        prompt = sg.build_section_prompt(wo)
        self.assertIn("GRAMMAR SECTION — MANDATORY", prompt)
        self.assertNotIn("ENGLISH GRAMMAR", prompt)
        self.assertIn("tamil grammar chunk", prompt)

    # ── validation ───────────────────────────────────────────────────────────────
    def test_lifted_span_ignores_shared_instruction_stems(self):
        ctx = ("Lencho had said it: the raindrops looked like new silver coins falling from "
               "the sky, and the boys ran out to collect them.")
        self.assertTrue(sg._lifted_span(
            "Rewrite in the passive voice: the raindrops looked like new silver coins "
            "falling from the sky.", ctx))
        self.assertIsNone(sg._lifted_span(
            "Rewrite in the passive voice: The gardener waters the plants every morning.", ctx))
        # a question stem that also appears on a textbook exercise page is NOT copied material
        self.assertIsNone(sg._lifted_span(
            "Fill in the blanks with the correct form of the verb given in brackets.",
            "Fill in the blanks with the correct form of the verb given in brackets below."))
        self.assertIsNone(sg._lifted_span("Correct the error.", ctx))
        self.assertIsNone(sg._lifted_span("anything at all here", ""))

    def test_grammar_question_copying_context_fails_validation(self):
        ctx = ("Lencho had said it: the raindrops looked like new silver coins falling from "
               "the sky, and the boys ran out to collect them.")
        wo = self._wo(
            [{"id": "SEC_C", "name": "E — Literature and Grammar", "marks": 6,
              "question_slots": [
                  {"qnum": 1, "type": "sa", "marks": 5, "topic": "The poem Dust of Snow"},
                  {"qnum": 2, "type": "rewrite", "marks": 1, "topic": "Passive voice"}]}],
            {"E — Literature and Grammar": ctx})["E — Literature and Grammar"]
        lifted = {"type": "VSA", "marks": 1, "answer_explanation": "a",
                  "text": "Rewrite in the passive voice: the raindrops looked like new "
                          "silver coins falling from the sky."}
        clean = {"type": "VSA", "marks": 1, "answer_explanation": "a",
                 "text": "Rewrite in the passive voice: The gardener waters the plants."}
        errs = sg._validate_by_subtype(lifted, 2, wo)
        self.assertTrue(any("copies the reference material" in e for e in errs), errs)
        self.assertEqual(sg._validate_by_subtype(clean, 2, wo), [])

    def test_writing_task_naming_a_chapter_fails_validation(self):
        # Exactly the reported leak, as the validator sees it.
        wo = self._wo(
            [{"id": "SEC_B", "name": "SECTION – Creative Writing Skills", "marks": 8,
              "question_slots": [
                  {"qnum": 2, "type": "writing", "marks": 4, "topic": "Article"},
                  {"qnum": 3, "type": "writing", "marks": 4, "topic": "Advertisement"}]}],
            {}, chapters=("The Laburnum Top", "The Address"))[
                "SECTION – Creative Writing Skills"]
        leaky = {"type": "LA", "marks": 4, "answer_explanation": "a",
                 "text": "After reading 'The Laburnum Top', you are inspired by the theme of "
                         "nature's vitality. Write an article in about 150 words on 'The "
                         "Healing Power of Nature' for your school magazine."}
        clean = {"type": "LA", "marks": 4, "answer_explanation": "a",
                 "text": "Your locality's park has been neglected for months. Write an article "
                         "in about 150 words on 'Reclaiming Our Green Spaces' for your school "
                         "magazine."}
        errs = sg._validate_by_subtype(leaky, 1, wo)
        self.assertTrue(any("The Laburnum Top" in e for e in errs), errs)
        self.assertEqual(sg._validate_by_subtype(clean, 1, wo), [])

    def test_chapter_name_check_is_word_bounded(self):
        wo = self._wo([{"id": "SEC_B", "name": "B — Grammar", "marks": 1, "question_slots": [
            {"qnum": 1, "type": "fill_blank", "marks": 1, "topic": "Tenses"}]}],
            {}, chapters=["Water"])["B — Grammar"]
        hit = {"type": "VSA", "marks": 1, "answer_explanation": "waters",
               "text": "Fill in the blank: In the chapter Water, she ______ the plants."}
        miss = {"type": "VSA", "marks": 1, "answer_explanation": "waters",
                "text": "Fill in the blank: The gardener ______ (watering) the plants."}
        self.assertTrue(any("chapter 'Water'" in e for e in sg._validate_by_subtype(hit, 1, wo)))
        self.assertEqual(sg._validate_by_subtype(miss, 1, wo), [])


class AccountancySumsCompositionTest(SimpleTestCase):
    """An Accountancy paper must be 80% SUMS (numerical/practical problems) and 20% QUIZ
    (definitions, concepts, formats) BY MARKS. The blueprint fixes each section's marks and
    counts, so the ratio is met by choosing WHICH questions are sums — plan_sums_allocation
    spends the 20% quiz budget on the cheapest objective questions and declares the rest sums,
    which is what turns the leftover objective marks into NUMERICAL MCQs."""

    def _wo(self, name, n, mpq, types_list, subject="Accountancy"):
        return sg.SectionWorkOrder(
            section_name=name, section_id=name[-1], title="", marks=int(n * mpq),
            questions_count=n, marks_per_question=mpq, question_types=types_list,
            instructions=[], constraints={}, context_text="ledger context",
            difficulty="Medium", subject=subject, class_name="12", chapters=["Partnership"])

    def _official(self):
        """The real CBSE Accountancy 055 pattern: 80 marks = 20m objective + 60m written."""
        return [
            self._wo("Part A — Q1-16", 16, 1, ["MCQ"]),
            self._wo("Part A — Q17-20", 4, 3, ["Short Answer"]),
            self._wo("Part A — Q21-22", 2, 4, ["Short Answer"]),
            self._wo("Part A — Q23-26", 4, 6, ["Long Answer"]),
            self._wo("Part B — Q27-30", 4, 1, ["MCQ"]),
            self._wo("Part B — Q31-32", 2, 3, ["Short Answer"]),
            self._wo("Part B — Q33", 1, 4, ["Short Answer"]),
            self._wo("Part B — Q34", 1, 6, ["Long Answer"]),
        ]

    @staticmethod
    def _paper(spec):
        """spec: {section: [(marks, 'sums'|'quiz'), ...]} → paper_data with matching wording."""
        out = {}
        for sec, items in spec.items():
            qs = []
            for i, (mk, nature) in enumerate(items):
                text = (f"Prepare the Revaluation Account for case {i}. Stock was overvalued by "
                        f"₹12,000 and furniture undervalued by ₹8,000."
                        if nature == "sums" else
                        f"Define the term goodwill as used in case {i}.")
                qs.append({"qnum": i + 1, "type": "SA", "marks": mk, "text": text,
                           "answer_explanation": "..."})
            out[sec] = {"questions": qs}
        return out

    # ── target + allocation ──────────────────────────────────────────────────────
    def test_share_lookup_is_accountancy_scoped(self):
        for subject in ("Accountancy", "accountancy", "Accounts",
                        "Book Keeping", "Book-Keeping and Accountancy"):
            self.assertEqual(sg._sums_share_for_subject(subject), 0.80, subject)
        for subject in ("Mathematics", "English Core", "Economics", "Business Studies", ""):
            self.assertEqual(sg._sums_share_for_subject(subject), 0.0, subject)

    def test_allocation_hits_80_percent_of_marks(self):
        wos = sg.plan_sums_allocation(self._official())
        sums_marks = sum(w.sums_count * w.marks_per_question for w in wos)
        total = sum(w.questions_count * w.marks_per_question for w in wos)
        self.assertEqual(total, 80)
        self.assertAlmostEqual(sums_marks / total, 0.80, delta=0.02)
        # every WRITTEN question is a sum; the shortfall is made up by numerical MCQs
        for w in wos:
            if w.marks_per_question > 1:
                self.assertEqual(w.sums_count, w.questions_count, w.section_name)
        objective = [w for w in wos if w.marks_per_question == 1]
        self.assertTrue(any(w.sums_count for w in objective),
                        "no numerical MCQs were planned")
        self.assertTrue(all(w.sums_share == 0.80 for w in wos))

    def test_allocation_spreads_numerical_mcqs_across_objective_sections(self):
        # Spending the quiz budget section-by-section would leave one objective section pure
        # theory and another pure numerical; the budget is spent round-robin instead.
        wos = sg.plan_sums_allocation(self._official())
        big = next(w for w in wos if w.section_name == "Part A — Q1-16")
        self.assertGreater(big.sums_count, 0)
        self.assertLess(big.sums_count, big.questions_count)

    def test_allocation_is_deterministic(self):
        a = [w.sums_count for w in sg.plan_sums_allocation(self._official())]
        b = [w.sums_count for w in sg.plan_sums_allocation(self._official())]
        self.assertEqual(a, b)

    def test_allocation_reads_slot_marks(self):
        slots = ([{"qnum": i + 1, "type": "mcq", "marks": 1} for i in range(10)]
                 + [{"qnum": 11, "type": "la", "marks": 6}, {"qnum": 12, "type": "la", "marks": 6}])
        wo = self._wo("Part A", 12, 1, [])
        wo.slots = slots
        plan = sg._question_marks_plan(wo)
        self.assertEqual([m for m, _ in plan], [1.0] * 10 + [6.0, 6.0])
        self.assertEqual([o for _, o in plan], [True] * 10 + [False, False])
        sg.plan_sums_allocation([wo])
        self.assertGreaterEqual(wo.sums_count, 2)      # both long answers at minimum

    def test_other_subjects_get_no_plan_and_no_prompt_block(self):
        wos = sg.plan_sums_allocation([self._wo("Section A", 20, 1, ["MCQ"], subject="Mathematics")])
        self.assertEqual(wos[0].sums_count, 0)
        self.assertEqual(wos[0].sums_share, 0.0)
        self.assertNotIn("COMPOSITION — MANDATORY", sg.build_section_prompt(wos[0]))

    # ── classification ───────────────────────────────────────────────────────────
    def test_question_nature_needs_both_a_verb_and_figures(self):
        sums = [
            {"text": "Pass the journal entries: A brought in capital of ₹1,50,000 on 1 April."},
            {"text": "Prepare the Revaluation Account. Stock was overvalued by ₹12,000."},
            {"text": "Compute the Current Ratio if current assets are ₹1,20,000 and current "
                     "liabilities are ₹60,000."},
            # figures living in the OPTIONS still make it a sum
            {"text": "Calculate the interest on capital.",
             "options": {"a": "₹4,000", "b": "₹4,500", "c": "₹5,000", "d": "₹6,000"}},
            # ...and so do figures in a case passage read by sub-questions
            {"text": "Read the case and answer.",
             "source_text": "X and Y share profits 3:2. Z is admitted and brings ₹80,000.",
             "sub_questions": [{"text": "Calculate the new ratio", "marks": 2}]},
        ]
        quiz = [
            {"text": "Define goodwill."},
            {"text": "Explain any three features of a partnership firm."},
            {"text": "State the formula for the Debt-Equity Ratio."},
            # a preparation VERB with no figures is a format question, not a sum
            {"text": "Prepare the format of a Balance Sheet as per Schedule III."},
            {"text": "Which of the following is not a fixed asset?",
             "options": {"a": "Machinery", "b": "Land", "c": "Debtors", "d": "Buildings"}},
        ]
        for q in sums:
            self.assertEqual(sg._question_nature(q), "sums", q["text"][:60])
        for q in quiz:
            self.assertEqual(sg._question_nature(q), "quiz", q["text"][:60])

    # ── prompt ───────────────────────────────────────────────────────────────────
    def test_prompt_states_this_sections_own_quota(self):
        by = {w.section_name: w for w in sg.plan_sums_allocation(self._official())}
        mixed = sg.build_section_prompt(by["Part A — Q1-16"])
        allsums = sg.build_section_prompt(by["Part A — Q23-26"])
        self.assertIn(f"EXACTLY {by['Part A — Q1-16'].sums_count} of the 16", mixed)
        self.assertIn("ALL 4 questions in this section MUST be SUMS", allsums)
        self.assertIn("80% SUMS and 20% QUIZ by marks", allsums)
        self.assertIn("actual amounts in ₹", allsums)
        # an MCQ section that must carry sums is told how a numerical MCQ looks
        self.assertIn("plausible computed amounts", mixed)

    # ── report ───────────────────────────────────────────────────────────────────
    def test_report_flags_a_theory_heavy_paper(self):
        good = self._paper({"A": [(1, "quiz")] * 16 + [(1, "sums")] * 4, "B": [(6, "sums")] * 10})
        report = sg.validate_sums_distribution(good, "Accountancy")
        self.assertTrue(report["compliant"])
        self.assertEqual(report["sums_pct"], 80.0)

        bad = self._paper({"A": [(1, "quiz")] * 20,
                           "B": [(6, "quiz")] * 5 + [(6, "sums")] * 5})
        report = sg.validate_sums_distribution(bad, "Accountancy")
        self.assertFalse(report["compliant"])
        self.assertIn("theory-heavy", report["violations"][0])
        # free for every other subject
        self.assertEqual(sg.validate_sums_distribution(good, "Mathematics"), {})

    # ── enforcement ──────────────────────────────────────────────────────────────
    def _wos_for_enforce(self):
        return sg.plan_sums_allocation([self._wo("A", 20, 1, ["MCQ"]),
                                        self._wo("B", 10, 6, ["Long Answer"])])

    def test_enforcement_converts_theory_questions_to_sums(self):
        replacement = {"type": "SA", "marks": 6,
                       "text": "Prepare the Partners' Capital Accounts. A's opening capital was "
                               "₹2,00,000 and drawings were ₹24,000 during the year.",
                       "answer_explanation": "Working shown."}
        paper = self._paper({"A": [(1, "quiz")] * 20,
                             "B": [(6, "quiz")] * 5 + [(6, "sums")] * 5})
        before = sg._sums_marks_split(paper)
        with mock.patch.object(sg.mantle_client, "converse",
                               return_value=(json.dumps(replacement), 5, 6)):
            out, in_tok, out_tok = sg.enforce_sums_distribution(
                paper, self._wos_for_enforce(), "Accountancy")
        after = sg._sums_marks_split(out)
        self.assertGreater(after[0], before[0])              # more sums marks than before
        self.assertEqual(sum(after), sum(before))            # paper total unchanged
        self.assertEqual([len(s["questions"]) for s in out.values()], [20, 10])
        self.assertGreater(in_tok + out_tok, 0)

    def test_a_replacement_that_is_still_theory_is_discarded(self):
        paper = self._paper({"A": [(1, "quiz")] * 20,
                             "B": [(6, "quiz")] * 5 + [(6, "sums")] * 5})
        snapshot = copy.deepcopy(paper)
        theory_again = {"type": "SA", "marks": 6, "text": "Explain the meaning of goodwill.",
                        "answer_explanation": "..."}
        with mock.patch.object(sg.mantle_client, "converse",
                               return_value=(json.dumps(theory_again), 5, 6)) as conv:
            out, _, _ = sg.enforce_sums_distribution(
                paper, self._wos_for_enforce(), "Accountancy")
        self.assertEqual(out, snapshot)   # never make the split worse than it already is
        # the cap counts ATTEMPTS, not successes — a model that keeps returning theory must not
        # be retried on every candidate in the paper
        self.assertLessEqual(conv.call_count, sg.SUMS_MAX_REGENS)

    def test_compliant_paper_and_other_subjects_cost_nothing(self):
        good = self._paper({"A": [(1, "quiz")] * 16 + [(1, "sums")] * 4, "B": [(6, "sums")] * 10})
        snapshot = copy.deepcopy(good)
        with mock.patch.object(sg.mantle_client, "converse") as conv:
            out, in_tok, out_tok = sg.enforce_sums_distribution(
                good, self._wos_for_enforce(), "Accountancy")
        self.assertEqual(conv.call_count, 0)
        self.assertEqual((in_tok, out_tok), (0, 0))
        self.assertEqual(out, snapshot)

        maths = self._paper({"A": [(1, "quiz")] * 20})
        snapshot = copy.deepcopy(maths)
        maths_wos = sg.plan_sums_allocation(
            [self._wo("A", 20, 1, ["MCQ"], subject="Mathematics")])
        with mock.patch.object(sg.mantle_client, "converse") as conv:
            out, _, _ = sg.enforce_sums_distribution(maths, maths_wos, "Mathematics")
        self.assertEqual(conv.call_count, 0)
        self.assertEqual(out, snapshot)


class McqAnswerKeyBalanceTest(SimpleTestCase):
    """Teachers reported answer keys with a visible shape — "aaabbbccc", or every answer (a).
    STRICT RULE 6 asked the model for a spread and validate_section_output rejects a letter used
    in >65% of a section, but neither catches a RUN: "aaabbbccc" puts each letter at 33%, so it
    passed, and "never more than 2 consecutive" was never enforced anywhere. balance_mcq_answer_keys
    fixes the key deterministically after generation instead."""

    OPTS = ["Chlorophyll traps the light energy", "Mitochondria release the stored energy",
            "Ribosomes assemble the protein chains", "Vacuoles store the cell sap"]

    def _q(self, i, answer="a", **kw):
        vals = [f"{v} in case {i}" for v in self.OPTS]
        q = {"qnum": i + 1, "type": "MCQ", "text": f"Question number {i} about plant cells",
             "options": dict(zip("abcd", vals)), "answer": answer, "marks": 1,
             "answer_explanation": f"Option ({answer}) is correct for case {i}."}
        q.update(kw)
        return q

    def _paper(self, key):
        return {"Section A": {"questions": [self._q(i, a) for i, a in enumerate(key)]}}

    @staticmethod
    def _max_run(s):
        best = cur = 1
        for i in range(1, len(s)):
            cur = cur + 1 if s[i] == s[i - 1] else 1
            best = max(best, cur)
        return best if s else 0

    @staticmethod
    def _repeats_block(s):
        return any(len(s) % b == 0 and s == s[:b] * (len(s) // b)
                   for b in range(1, len(s) // 2 + 1))

    def test_balanced_letters_have_no_run_or_block_pattern(self):
        for seed in (1, 7, 128, 231, 360, 9999):
            for n in (6, 8, 9, 12, 16, 20):
                seq = sg._balanced_answer_letters(n, seed)
                s = "".join(seq)
                self.assertEqual(len(seq), n)
                counts = collections.Counter(seq)
                self.assertLessEqual(max(counts.values()) - min(counts.values()), 1, s)
                self.assertLessEqual(self._max_run(s), 2, s)
                # an honest shuffle lands on a repeated block ~1% of the time at n=8 —
                # those seeds (128/231/360 hit it) must be rejected and reseeded
                self.assertFalse(self._repeats_block(s), f"seed={seed} n={n}: {s}")

    def test_balanced_letters_are_deterministic_but_paper_specific(self):
        self.assertEqual(sg._balanced_answer_letters(16, 42), sg._balanced_answer_letters(16, 42))
        self.assertNotEqual(sg._balanced_answer_letters(16, 42), sg._balanced_answer_letters(16, 43))

    def test_all_a_and_run_patterns_are_broken_up(self):
        for key in ("aaaaaaaaa", "aaabbbccc", "abcdabcdabcd"):
            paper = self._paper(key)
            before = [q["options"][q["answer"]] for q in paper["Section A"]["questions"]]
            sg.balance_mcq_answer_keys(paper)
            qs = paper["Section A"]["questions"]
            after_key = "".join(q["answer"] for q in qs)
            # the correct answer TEXT must survive untouched — only its letter moves
            self.assertEqual([q["options"][q["answer"]] for q in qs], before, key)
            self.assertLessEqual(self._max_run(after_key), 2, f"{key} → {after_key}")
            self.assertFalse(self._repeats_block(after_key), f"{key} → {after_key}")
            for q in qs:
                self.assertEqual(set(q["options"]), set("abcd"))
                self.assertEqual(len(set(q["options"].values())), 4)
                self.assertIn(f"({q['answer']})", q["answer_explanation"].lower())

    def test_explanation_naming_all_four_letters_is_remapped_bijectively(self):
        paper = self._paper("aaaaaaaa")
        originals = []
        for q in paper["Section A"]["questions"]:
            q["answer_explanation"] = ("Option (a) is correct. Option (b), option (c) and "
                                       "option (d) are all wrong.")
            originals.append(dict(q["options"]))
        sg.balance_mcq_answer_keys(paper)
        for q, orig in zip(paper["Section A"]["questions"], originals):
            new_of = {v: k for k, v in q["options"].items()}
            mapping = {old: new_of[val] for old, val in orig.items()}
            self.assertEqual(q["answer"], mapping["a"])
            self.assertEqual(
                q["answer_explanation"],
                f"Option ({mapping['a']}) is correct. Option ({mapping['b']}), "
                f"option ({mapping['c']}) and option ({mapping['d']}) are all wrong.")

    def test_order_bound_and_ar_questions_are_left_alone(self):
        # Assertion-Reason prints four canonical options in a fixed order; "All of the above"
        # must stay last; an articles MCQ ("a"/"an"/"the") makes an explanation "(a)" ambiguous;
        # a sorted numeric set keeps the ascending convention.
        plain = self._q(0)
        self.assertTrue(sg._mcq_is_permutable(plain))
        for label, q in (
            ("AR subtype", self._q(1, subtype="assertion_reason")),
            ("AR type", self._q(2, type="Assertion-Reason")),
            ("matching", self._q(3, subtype="matching")),
            ("all of the above", self._q(4, options=dict(zip("abcd", [
                "Chlorophyll traps light", "Stomata allow exchange",
                "Cuticle limits water loss", "All of the above"])))),
            ("both (a) and (b)", self._q(5, options=dict(zip("abcd", [
                "Chlorophyll traps light", "Stomata allow exchange",
                "Both (a) and (b) are true", "Neither is true"])))),
            ("articles", self._q(6, options=dict(zip("abcd", ["a", "an", "the", "no article"])))),
            ("sorted numeric", self._q(7, options=dict(zip("abcd", ["2 cm", "4 cm", "6 cm", "8 cm"])))),
            ("duplicate values", self._q(8, options=dict(zip("abcd", [
                "Same option text", "Same option text", "Third option text", "Fourth option text"])))),
        ):
            self.assertFalse(sg._mcq_is_permutable(q), label)
        # unsorted numbers are fine to reorder
        self.assertTrue(sg._mcq_is_permutable(
            self._q(9, options=dict(zip("abcd", ["6 cm", "2 cm", "8 cm", "4 cm"])))))

    def test_ineligible_questions_pass_through_untouched(self):
        paper = self._paper("aaaaaa")
        paper["Section A"]["questions"][2]["subtype"] = "assertion_reason"
        paper["Section A"]["questions"][4]["options"]["d"] = "All of the above"
        snapshot = copy.deepcopy([paper["Section A"]["questions"][i] for i in (2, 4)])
        sg.balance_mcq_answer_keys(paper)
        self.assertEqual(paper["Section A"]["questions"][2], snapshot[0])
        self.assertEqual(paper["Section A"]["questions"][4], snapshot[1])

    def test_too_few_mcqs_is_a_noop(self):
        paper = self._paper("aa")
        snapshot = copy.deepcopy(paper)
        sg.balance_mcq_answer_keys(paper)
        self.assertEqual(paper, snapshot)

    def test_letter_remap_does_not_double_apply_on_a_swap(self):
        self.assertEqual(
            sg._remap_answer_letters("Option (a) is right, option (b) is wrong.",
                                     {"a": "b", "b": "a"}),
            "Option (b) is right, option (a) is wrong.")
        self.assertEqual(sg._remap_answer_letters("Option B is correct.", {"b": "d"}),
                         "Option D is correct.")
        self.assertEqual(sg._remap_answer_letters("Correct answer: c", {"c": "d"}),
                         "Correct answer: d")
        self.assertEqual(sg._remap_answer_letters("Because chlorophyll absorbs light.", {"a": "c"}),
                         "Because chlorophyll absorbs light.")


class CrossSectionDuplicateFixTest(TestCase):
    """Cross-section duplicates were detected and then only LOGGED — cross_section_validate
    stored them for the final LLM audit, so a paper asking the same thing in Section A and
    Section D shipped with both. Sections are generated by independent parallel prompts, so
    neither can see the other's questions; this is the only place the overlap can be caught."""

    DUP = "Explain the process of photosynthesis in green plants clearly"

    def _wo(self, name):
        return sg.SectionWorkOrder(
            section_name=name, section_id=name[-1], title="", marks=2, questions_count=2,
            marks_per_question=1, question_types=["vsa"], instructions=[], constraints={},
            context_text="chapter context", difficulty="Medium", subject="Science",
            class_name="10", chapters=["Life Processes"])

    def _paper(self):
        def q(n, text):
            return {"qnum": n, "type": "VSA", "marks": 1, "text": text,
                    "answer_explanation": "..."}
        return {
            "Section A": {"questions": [
                q(1, self.DUP), q(2, "Describe the structure of the human digestive system")]},
            "Section D": {"questions": [
                q(3, self.DUP), q(4, "State three uses of concave mirrors in appliances")]},
        }

    def test_pairs_finds_cross_section_only(self):
        pairs = sg._cross_section_dup_pairs(self._paper())
        self.assertEqual([(a, b, c, d) for a, b, c, d, _ in pairs],
                         [("Section A", 0, "Section D", 0)])
        # two identical questions inside ONE section are validate_uniqueness's job, not this one
        self.assertEqual(sg._cross_section_dup_pairs(
            {"S": {"questions": [{"qnum": 1, "text": self.DUP}, {"qnum": 2, "text": self.DUP}]}}), [])

    def test_confirmed_duplicate_is_replaced_in_the_later_section(self):
        replacement = {"type": "VSA", "marks": 1,
                       "text": "Name the pigment that traps light energy in leaves",
                       "answer_explanation": "Chlorophyll."}

        def fake(**kw):
            if "same_concept" in kw.get("prompt", ""):
                return (json.dumps({"same_concept": True, "reason": "r"}), 3, 4)
            return (json.dumps(replacement), 5, 6)

        paper = self._paper()
        with mock.patch.object(sg.mantle_client, "converse", side_effect=fake):
            out, in_tok, out_tok = sg.fix_cross_section_duplicates(
                paper, [self._wo("Section A"), self._wo("Section D")])
        self.assertTrue(out["Section D"]["questions"][0]["text"].startswith("Name the pigment"))
        self.assertEqual(out["Section A"]["questions"][0]["text"], self.DUP)  # earlier one kept
        self.assertEqual(out["Section D"]["questions"][0]["qnum"], 3)          # numbering kept
        self.assertEqual(out["Section D"]["questions"][0]["marks"], 1)         # marks kept
        self.assertEqual([len(s["questions"]) for s in out.values()], [2, 2])  # count unchanged
        self.assertGreater(in_tok + out_tok, 0)
        self.assertNotIn("_cross_section_duplicates", out["Section A"])

    def test_llm_cleared_false_positive_keeps_both(self):
        paper = self._paper()
        with mock.patch.object(sg.mantle_client, "converse",
                               return_value=(json.dumps({"same_concept": False}), 1, 1)):
            out, _, _ = sg.fix_cross_section_duplicates(
                paper, [self._wo("Section A"), self._wo("Section D")])
        self.assertEqual(out["Section A"]["questions"][0]["text"], self.DUP)
        self.assertEqual(out["Section D"]["questions"][0]["text"], self.DUP)
        self.assertNotIn("_cross_section_duplicates", out["Section A"])

    def test_cbq_duplicate_is_kept_and_flagged_not_regenerated(self):
        # A CBQ carries a passage and sub-questions; a single-shot regen can't reproduce them
        # safely, so both survive and the overlap is reported to the audit instead.
        def cbq(n):
            return {"qnum": n, "type": "CBQ", "subtype": "source_based", "marks": 4,
                    "text": "Read the source and answer",
                    "source_text": "Photosynthesis in green plants converts light energy",
                    "sub_questions": [{"text": "Name the pigment", "marks": 4}]}
        paper = {"Section A": {"questions": [cbq(1)]}, "Section D": {"questions": [cbq(2)]}}
        with mock.patch.object(sg.mantle_client, "converse",
                               return_value=(json.dumps({"same_concept": True}), 1, 1)):
            out, _, _ = sg.fix_cross_section_duplicates(
                paper, [self._wo("Section A"), self._wo("Section D")])
        self.assertEqual(len(out["Section A"]["questions"]), 1)
        self.assertEqual(len(out["Section D"]["questions"]), 1)
        self.assertTrue(out["Section A"]["_cross_section_duplicates"])

    def test_missing_work_order_drops_nothing(self):
        paper = self._paper()
        with mock.patch.object(sg.mantle_client, "converse",
                               return_value=(json.dumps({"same_concept": True}), 1, 1)):
            out, _, _ = sg.fix_cross_section_duplicates(paper, [])
        self.assertEqual([len(s["questions"]) for s in out.values()], [2, 2])
        self.assertTrue(out["Section A"]["_cross_section_duplicates"])


class InternalChoiceCbqValidationTest(TestCase):
    """An internal-choice extract/CBQ slot must ship a COMPLETE or_alternative object
    (own passage + own sub-questions, matching count and marks); sub-question count
    must match the slot's declared parts; a textbook extract must be a verbatim
    quotation of the reference material. Production shipped a grammar meta-summary
    with a bare 'OR' and a 1+1+3 part split against a declared 5×1."""

    SENT = "The wind blows strongly and the tall trees bow to it every evening. "
    CONTEXT = SENT * 25

    def _wo(self, context=""):
        secs = psx.normalize_slots([{"id": "SEC_D", "name": "Literature", "marks": 5,
            "question_slots": [
                {"qnum": 1, "type": "extract", "marks": 5, "source": "textbook",
                 "choice": "internal", "alternatives": ["prose or poem", "supplementary reader"],
                 "parts": [{"label": "A", "type": "mcq", "marks": 1},
                           {"label": "B", "type": "mcq", "marks": 1},
                           {"label": "C", "type": "sa", "marks": 1},
                           {"label": "D", "type": "sa", "marks": 1},
                           {"label": "E", "type": "sa", "marks": 1}]},
            ]}])
        psx.derive_aggregates_from_slots(secs)
        pattern = ExamPattern(name="p", sections=secs)
        bp = sg_gen.pattern_sections_to_blueprint_dict(pattern)
        ctx_map = {"Literature": context} if context else {}
        return sg.build_work_orders(bp, pattern, ctx_map, "Medium", "6", "English", ["Ch1"])[0]

    @staticmethod
    def _subs(prefix):
        # Parts A/B are declared MCQ, so their sub-questions carry inline options.
        out = []
        for i, pt in enumerate(("mcq", "mcq", "sa", "sa", "sa")):
            t = f"{prefix}{i}"
            if pt == "mcq":
                t += " — a) one b) two c) three d) four"
            out.append({"text": t, "marks": 1})
        return out

    def _q(self, **over):
        q = {"type": "CBQ", "subtype": "source_based", "marks": 5,
             "text": "Read the extract and answer the following:",
             "competency_type": "application",
             "source_text": (self.SENT * 6).strip(),
             "sub_questions": self._subs("s"),
             "or_alternative": {
                 "text": "Read the extract and answer the following:",
                 "source_text": (self.SENT * 6).strip(),
                 "sub_questions": self._subs("a"),
             }}
        q.update(over)
        return q

    def test_complete_alternative_passes(self):
        errors = sg.validate_section_output({"questions": [self._q()]}, self._wo(self.CONTEXT))
        self.assertEqual(errors, [])

    def test_missing_or_alternative_fails(self):
        errors = sg.validate_section_output({"questions": [self._q(or_alternative=None)]},
                                            self._wo())
        self.assertTrue(any("or_alternative" in e and "OBJECT" in e for e in errors), errors)

    def test_string_alternative_fails_for_cbq(self):
        errors = sg.validate_section_output({"questions": [self._q(or_alternative="just text")]},
                                            self._wo())
        self.assertTrue(any("bare string" in e for e in errors), errors)

    def test_alternative_without_own_passage_or_parts_fails(self):
        alt = {"text": "Read and answer:"}
        errors = sg.validate_section_output({"questions": [self._q(or_alternative=alt)]},
                                            self._wo())
        self.assertTrue(any("OWN 'source_text'" in e for e in errors), errors)
        self.assertTrue(any("OWN 'sub_questions'" in e for e in errors), errors)

    def test_parts_count_enforced(self):
        q = self._q(sub_questions=[{"text": f"s{i}", "marks": 1} for i in range(3)])
        errors = sg.validate_section_output({"questions": [q]}, self._wo())
        self.assertTrue(any("declares 5 sub-parts" in e and "got 3" in e for e in errors), errors)

    def test_positional_part_marks_enforced(self):
        subs = [{"text": f"s{i}", "marks": m} for i, m in enumerate([2, 1, 1, 0.5, 0.5])]
        errors = sg.validate_section_output({"questions": [self._q(sub_questions=subs)]},
                                            self._wo())
        self.assertTrue(any("sub-question 1 marks=2" in e for e in errors), errors)

    def test_hallucinated_extract_flagged(self):
        q = self._q(source_text="This chapter explores the dynamics of a classroom where "
                                "education is a shared journey and humour makes the lessons "
                                "memorable for every student in the room during the day.")
        errors = sg.validate_section_output({"questions": [q]}, self._wo(self.CONTEXT))
        self.assertTrue(any("verbatim" in e for e in errors), errors)

    def test_mcq_parts_require_inline_options(self):
        q = self._q()
        q["sub_questions"][0]["text"] = "What did the man take out?"   # part A is MCQ
        errors = sg.validate_section_output({"questions": [q]}, self._wo())
        self.assertTrue(any("declared MCQ" in e for e in errors), errors)
        # alternatives are held to the same parts spec
        q = self._q()
        q["or_alternative"]["sub_questions"][1]["text"] = "No options here."
        errors = sg.validate_section_output({"questions": [q]}, self._wo())
        self.assertTrue(any("declared MCQ" in e for e in errors), errors)

    def test_stitched_extract_flagged(self):
        # Overall shingle overlap passes (all fragments verbatim) but the splice
        # sentence does not exist in the material as written.
        src = ((self.SENT * 3).strip()
               + " The wind blows strongly and the tall trees bow to it every morning.")
        errors = sg.validate_section_output({"questions": [self._q(source_text=src)]},
                                            self._wo(self.CONTEXT))
        self.assertTrue(any("stitched" in e for e in errors), errors)

    def test_cross_option_quote_flagged(self):
        # Production bug: option 2 asked "What is the 'currant bun' being used as in
        # this extract?" — but the currant bun is in option 1's passage.
        q = self._q()
        q["source_text"] = ("He took out a currant bun and held it to my nose while "
                            "everyone laughed at the funny man in the classroom that day.")
        alt = q["or_alternative"]
        alt["source_text"] = ("The old man walked slowly to the river bank and watched "
                              "the boats sail away into the golden evening light.")
        alt["sub_questions"][2]["text"] = \
            "What is the 'currant bun' being used as in this extract?"
        errors = sg.validate_section_output({"questions": [q]}, self._wo())
        self.assertTrue(any("OWN source_text" in e and "currant bun" in e for e in errors),
                        errors)

    def test_overlap_helper(self):
        self.assertTrue(sg._text_overlaps_context(self.CONTEXT[:300], self.CONTEXT))
        self.assertFalse(sg._text_overlaps_context(
            "An entirely different composed summary about classroom dynamics and learning "
            "together in a collaborative environment fostered by the teacher every day.",
            self.CONTEXT))

    def test_prompt_demands_object_alternative(self):
        prompt = sg.build_section_prompt(self._wo())
        self.assertIn("MUST be a JSON OBJECT", prompt)
        self.assertIn('its OWN "source_text"', prompt)

    def test_internal_choice_with_parts_is_valid_pattern(self):
        # Q21's canonical shape — an extract printed twice (OR), A-E under each option.
        # The old rule flagged parts+internal as a conflict; it is now the documented way
        # to express a two-passage extract.
        secs = psx.normalize_slots([{"name": "Literature", "marks": 5, "question_slots": [
            {"qnum": 1, "type": "extract", "marks": 5, "choice": "internal",
             "source": "textbook", "alternatives": ["prose or poem", "SR"],
             "parts": [{"label": c, "type": "sa", "marks": 1} for c in "ABCDE"]},
        ]}])
        self.assertEqual(psx.validate_pattern_structure(secs), [])


class MultiAlternativeChoiceTest(TestCase):
    """Q11-style 3-way internal choice (paragraph OR letter OR notice): 'or_alternative'
    becomes a JSON ARRAY with one entry per extra option, and the renderer prints every
    alternative after its own OR separator. Production shipped only 2 of the teacher's
    3 options because the schema knew a single alternative."""

    def _wo(self):
        secs = psx.normalize_slots([{"id": "SEC_B", "name": "Writing", "marks": 5,
            "question_slots": [
                {"qnum": 1, "type": "writing", "marks": 5, "choice": "internal",
                 "source": "general",
                 "alternatives": ["Descriptive Paragraph", "Letter - informal",
                                  "Notice Writing"]},
            ]}])
        psx.derive_aggregates_from_slots(secs)
        pattern = ExamPattern(name="p", sections=secs)
        bp = sg_gen.pattern_sections_to_blueprint_dict(pattern)
        return sg.build_work_orders(bp, pattern, {}, "Medium", "6", "English", ["Ch1"])[0]

    def test_prompt_demands_array(self):
        prompt = sg.build_section_prompt(self._wo())
        self.assertIn("JSON ARRAY", prompt)
        self.assertIn("3 options", prompt)

    def test_single_string_alternative_fails(self):
        q = {"type": "LA", "marks": 5, "text": "Write a descriptive paragraph on rain.",
             "competency_type": "constructed", "answer_explanation": "pts",
             "or_alternative": "Write an informal letter."}
        errors = sg.validate_section_output({"questions": [q]}, self._wo())
        self.assertTrue(any("ARRAY" in e for e in errors), errors)

    def test_array_of_two_passes(self):
        q = {"type": "LA", "marks": 5, "text": "Write a descriptive paragraph on rain.",
             "competency_type": "constructed", "answer_explanation": "pts",
             "or_alternative": ["Write an informal letter to a friend.",
                                "Write a notice about the lost library book."]}
        errors = sg.validate_section_output({"questions": [q]}, self._wo())
        self.assertEqual(errors, [])

    def test_short_array_fails(self):
        q = {"type": "LA", "marks": 5, "text": "Write a descriptive paragraph on rain.",
             "competency_type": "constructed", "answer_explanation": "pts",
             "or_alternative": ["Write an informal letter to a friend."]}
        errors = sg.validate_section_output({"questions": [q]}, self._wo())
        self.assertTrue(any("ARRAY of 2" in e for e in errors), errors)

    def test_renderer_prints_every_alternative(self):
        q = {"qnum": 11, "text": "Write a descriptive paragraph on rain.", "type": "LA",
             "marks": 5,
             "or_alternative": ["Write an informal letter to a friend.",
                                "Write a notice about the lost library book."]}
        out = []
        sg_gen.process_question(out, q, 11)
        self.assertEqual(len([1 for k, _ in out if k == "or"]), 2)
        qlines = [t for k, t in out if k == "q"]
        self.assertEqual(len(qlines), 3)                 # primary + 2 alternatives
        self.assertTrue(all(t.startswith("11.") for t in qlines))
        self.assertIn("notice", qlines[2])

    def test_umbrella_stem_and_letter_labels_fail(self):
        # Production layout bug: "11. Attempt any ONE… OR 11. A. Write… OR 11. B. …" —
        # the stem must BE the first option and options must not carry A/B/C labels.
        q = {"type": "LA", "marks": 5, "text": "Attempt any ONE of the following:",
             "competency_type": "constructed", "answer_explanation": "pts",
             "or_alternative": ["A. Write a descriptive paragraph on rain.",
                                "B. Write an informal letter to a friend."]}
        errors = sg.validate_section_output({"questions": [q]}, self._wo())
        self.assertTrue(any("umbrella" in e for e in errors), errors)
        self.assertTrue(any("prefix options" in e for e in errors), errors)


class StandalonePartsNoPassageTest(TestCase):
    """An SA parts group ('A to F, attempt any 5') is not an extract — production
    attached a skill-box source_text to it and the paper printed an unwanted
    'Read the source/case…' passage above plain short-answer questions."""

    def _wo(self):
        secs = psx.normalize_slots([{"id": "SEC_D", "name": "Literature", "marks": 10,
            "question_slots": [
                {"qnum": 1, "type": "sa", "marks": 10, "choice": "open", "attempt": 5,
                 "source": "textbook",
                 "parts": [{"label": l, "type": "sa", "marks": 2} for l in "ABCDEF"]},
            ]}])
        psx.derive_aggregates_from_slots(secs)
        pattern = ExamPattern(name="p", sections=secs)
        bp = sg_gen.pattern_sections_to_blueprint_dict(pattern)
        return sg.build_work_orders(bp, pattern, {}, "Medium", "6", "English", ["Ch1"])[0]

    def test_source_text_on_sa_parts_group_fails(self):
        q = {"type": "CBQ", "subtype": "source_based", "marks": 10,
             "text": "Attempt any 5 of the following 6:", "competency_type": "constructed",
             "source_text": "The way we use stress and intonation can change meaning. " * 3,
             "sub_questions": [{"text": f"q{i}", "marks": 2} for i in range(6)]}
        errors = sg.validate_section_output({"questions": [q]}, self._wo())
        self.assertTrue(any("standalone sub-questions" in e for e in errors), errors)

    def test_clean_sa_parts_group_passes(self):
        q = {"type": "CBQ", "subtype": "standard", "marks": 10,
             "text": "Attempt any 5 of the following 6:", "competency_type": "constructed",
             "sub_questions": [{"text": f"Explain point {i} from the chapter.", "marks": 2}
                               for i in range(6)]}
        errors = sg.validate_section_output({"questions": [q]}, self._wo())
        self.assertEqual(errors, [])

    def test_prompt_forbids_passages(self):
        prompt = sg.build_section_prompt(self._wo())
        self.assertIn("NO PASSAGE", prompt)
        self.assertIn('Do NOT output a section-level "passage" key', prompt)


class SlotSectionPassageRenderTest(TestCase):
    """Renderer: a section-level 'passage' on a slot section is generator planning
    junk ('This section tests your understanding…') unless the section is an
    unseen-reading one — those genuinely share one passage across their slots."""

    def _render(self, slots, source):
        blueprint = {"Sec": {"marks": 5, "question_slots": slots}}
        data = {"Sec": {"passage": "Some section level text here.",
                        "questions": [{"qnum": 1, "type": "SA", "subtype": "standard",
                                       "marks": 5, "text": "Answer this."}]}}
        out = []
        sg_gen.render_section_questions(out, data, blueprint)
        return [t for k, t in out if k == "passage"]

    def test_textbook_slot_section_skips_section_passage(self):
        slots = [{"qnum": 1, "type": "sa", "marks": 5, "source": "textbook"}]
        self.assertEqual(self._render(slots, "textbook"), [])

    def test_unseen_slot_section_keeps_section_passage(self):
        slots = [{"qnum": 1, "type": "sa", "marks": 5, "source": "unseen"}]
        self.assertEqual(self._render(slots, "unseen"), ["Some section level text here."])


class ExtractCoherenceTest(TestCase):
    """Literature extracts must be coherent literary text quoted whole — production
    shipped a fill-in-the-blanks worksheet and a bullet-pointed pronunciation skill
    box clipped mid-thought as 'extracts'."""

    def test_worksheet_markup_flagged(self):
        self.assertIsNotNone(sg._extract_text_issue(
            "To communicate (i) ___________ (effective), choose your words and topics "
            "(ii) ___________ (wise)."))
        self.assertIsNotNone(sg._extract_text_issue(
            "→ Content words like book, run are stressed. • Use falling intonation "
            "towards the end of the sentence."))

    def test_clipped_fragments_flagged(self):
        self.assertIsNotNone(sg._extract_text_issue(
            "and then the mongoose ran into the garden. The boy followed it home."))
        self.assertIsNotNone(sg._extract_text_issue(
            "The boy followed the mongoose into the"))

    def test_clean_passage_passes(self):
        self.assertIsNone(sg._extract_text_issue(
            '"Where are you going?" asked the old man. The boy pointed to the river '
            "and smiled at him."))
        self.assertIsNone(sg._extract_text_issue(""))

    def test_exercise_question_page_flagged(self):
        # The exact production failure: the poem's own exercise page (page header +
        # numbered question + activity item) quoted as the OR alternative's "passage".
        src = ("Poorvi62\n"
               "   2. Why has the poet used phrases like 'funny sounding sight' and "
               "'funny feeling sound' with reference to the funny man?\n"
               " VI Can you think of any real-world situations where people do similar "
               "things for fun, entertainment, or performance? Share with your classmates "
               "and the teacher.")
        self.assertIsNotNone(sg._extract_text_issue(src))
        # each artefact alone is enough
        self.assertIsNotNone(sg._extract_text_issue("Poorvi62\nHe held it to my nose."))
        self.assertIsNotNone(sg._extract_text_issue(
            "2. Why has the poet used these phrases in the poem about the funny man?"))
        self.assertIsNotNone(sg._extract_text_issue(
            "Think of similar situations and share with your classmates and the teacher."))

    def test_poem_quote_with_slashes_passes(self):
        self.assertIsNone(sg._extract_text_issue(
            'He said, "Allow me to present / Your Highness with a rose." / And taking '
            "out a currant bun / He held it to my nose."))



class RegenerateAllPatternsTest(TestCase):
    """POST /api/patterns/regenerate-all/ re-queues every AI-generated pattern (or just the
    given ids) from its own saved prompt, and skips manual patterns, promptless ones and
    ones already queued/generating — without ever touching non-AI patterns."""

    def setUp(self):
        from rest_framework.test import APIClient
        self.school = School.objects.create(name="Regen School")
        self.user = User.objects.create_user("regen_t", "rt@x.com", "pw")
        p = self.user.profile
        p.school = self.school
        p.role = "teacher"
        p.save()
        self.api = APIClient()
        self.api.force_authenticate(self.user)

        def mk(name, source, prompt, status="done"):
            return ExamPattern.objects.create(
                name=name, subject="English", class_name="6",
                sections=[{"name": "A"}], total_marks=40, total_questions=10,
                pattern_source=source, ai_prompt=prompt, status=status,
                created_by=self.user)

        self.p_ai       = mk("AI ok",        "ai_generated", "PT-1 paper, 40 marks")
        self.p_running  = mk("AI running",   "ai_generated", "same prompt", status="generating")
        self.p_manual   = mk("Manual",       "manual",       "")
        self.p_noprompt = mk("AI no prompt", "ai_generated", "")

    def _post(self, body=None):
        from unittest import mock
        with mock.patch("core.tasks.generate_pattern_task") as mtask:
            mtask.delay.return_value = mock.Mock(id="task-regen-1")
            r = self.api.post("/api/patterns/regenerate-all/", body or {}, format="json")
        return r, mtask

    def test_queues_only_eligible_ai_patterns(self):
        r, mtask = self._post()
        self.assertEqual(r.status_code, 202)
        data = r.json()
        self.assertEqual(data["queued_ids"], [self.p_ai.id])
        self.assertIn(self.p_manual.id, data["skipped_not_ai"])
        self.assertIn(self.p_noprompt.id, data["skipped_no_prompt"])
        self.assertIn(self.p_running.id, data["skipped_active"])
        mtask.delay.assert_called_once_with(self.p_ai.id)
        self.p_ai.refresh_from_db()
        self.assertEqual(self.p_ai.status, "queued")
        self.assertEqual(self.p_ai.sections, [])           # rebuilt from the prompt
        self.assertEqual(self.p_ai.task_id, "task-regen-1")
        self.p_manual.refresh_from_db()
        self.assertEqual(self.p_manual.status, "done")     # untouched

    def test_ids_restricts_scope(self):
        r, mtask = self._post({"ids": [self.p_manual.id]})
        self.assertEqual(r.status_code, 202)
        data = r.json()
        self.assertEqual(data["queued"], 0)
        self.assertEqual(data["skipped_not_ai"], [self.p_manual.id])
        mtask.delay.assert_not_called()
        self.p_ai.refresh_from_db()
        self.assertEqual(self.p_ai.status, "done")         # not in ids → untouched

    def test_bad_ids_rejected(self):
        r, _ = self._post({"ids": "everything"})
        self.assertEqual(r.status_code, 400)
        r, _ = self._post({"ids": []})
        self.assertEqual(r.status_code, 400)
        r, _ = self._post({"ids": ["not-a-number"]})
        self.assertEqual(r.status_code, 400)

    def test_other_schools_patterns_excluded(self):
        other = User.objects.create_user("other_t", "ot@x.com", "pw")
        op = other.profile
        op.school = School.objects.create(name="Other School")
        op.role = "teacher"
        op.save()
        foreign = ExamPattern.objects.create(
            name="Foreign AI", subject="English", class_name="6", sections=[{"name": "A"}],
            total_marks=40, total_questions=10, pattern_source="ai_generated",
            ai_prompt="their prompt", status="done", created_by=other)
        r, mtask = self._post({"ids": [foreign.id, self.p_ai.id]})
        self.assertEqual(r.status_code, 202)
        self.assertEqual(r.json()["queued_ids"], [self.p_ai.id])
        foreign.refresh_from_db()
        self.assertEqual(foreign.status, "done")           # out of scope → untouched


class PatternPdfImportTest(TestCase):
    """Import-a-sample-paper-PDF flow: the view stages the upload, extracts its text,
    rejects scanned/non-PDF uploads with a clear 400, and queues generate_pattern_task
    on an 'imported' pattern; the task routes imported patterns to the SQP schema
    extractor (never the teacher-text parser); the extraction prompt forbids copying
    the paper's content into the pattern."""

    # Long enough to pass the ≥200-char scanned-PDF guard.
    SQP_TEXT = ("SAMPLE QUESTION PAPER — CLASS X ENGLISH\n"
                "Section A Reading 20 marks. Section B Grammar and Writing 20 marks.\n"
                "Section C Literature 40 marks.\n") * 5

    def setUp(self):
        from rest_framework.test import APIClient
        self.user = User.objects.create_user("pdf_t", "pdf@x.com", "pw")
        self.api = APIClient()
        self.api.force_authenticate(self.user)

    def _upload(self, filename="EnglishL-SQP.pdf", extracted=None):
        from unittest import mock
        from django.core.files.uploadedfile import SimpleUploadedFile
        f = SimpleUploadedFile(filename, b"%PDF-1.4 fake body", content_type="application/pdf")
        with mock.patch("core.material_intel.extract_pages_text",
                        return_value=self.SQP_TEXT if extracted is None else extracted) as mex, \
             mock.patch("core.tasks.generate_pattern_task") as mtask:
            mtask.delay.return_value = mock.Mock(id="task-pdf-1")
            r = self.api.post(
                "/api/patterns/import-from-pdf/",
                {"file": f, "class_name": "10", "subject": "English", "name": "SQP 2025-26"},
                format="multipart")
        return r, mex, mtask

    def test_import_creates_queued_imported_pattern(self):
        r, mex, mtask = self._upload()
        self.assertEqual(r.status_code, 202)
        pattern = ExamPattern.objects.get(id=r.json()["id"])
        self.assertEqual(pattern.pattern_source, "imported")
        self.assertEqual(pattern.status, "queued")
        self.assertEqual(pattern.ai_prompt, self.SQP_TEXT)   # task reads the SQP text from here
        self.assertEqual(pattern.name, "SQP 2025-26")
        self.assertEqual(pattern.task_id, "task-pdf-1")
        mex.assert_called_once()
        mtask.delay.assert_called_once_with(pattern.id)

    def test_scanned_pdf_rejected(self):
        r, _, mtask = self._upload(extracted="")   # image-based PDF → no text layer
        self.assertEqual(r.status_code, 400)
        self.assertIn("scanned", r.json()["error"])
        self.assertEqual(ExamPattern.objects.count(), 0)
        mtask.delay.assert_not_called()

    def test_non_pdf_rejected(self):
        r, _, mtask = self._upload(filename="paper.docx")
        self.assertEqual(r.status_code, 400)
        self.assertEqual(ExamPattern.objects.count(), 0)
        mtask.delay.assert_not_called()

    def test_missing_file_rejected(self):
        r = self.api.post("/api/patterns/import-from-pdf/",
                          {"class_name": "10", "subject": "English"}, format="multipart")
        self.assertEqual(r.status_code, 400)

    def test_task_routes_imported_pattern_to_sqp_extractor(self):
        from unittest import mock
        pattern = ExamPattern.objects.create(
            name="SQP 2025-26", subject="English", class_name="10", sections=[],
            total_marks=0, total_questions=0, pattern_source="imported",
            ai_prompt=self.SQP_TEXT, status="queued", created_by=self.user)
        extracted = {
            "sections": [{"id": "SEC_A", "name": "Reading", "marks": 2,
                          "question_slots": [{"qnum": 1, "type": "mcq", "marks": 1},
                                             {"qnum": 2, "type": "mcq", "marks": 1}]}],
            "total_marks": 2, "total_questions": 2,
        }
        with mock.patch("api.ai_service.extract_pattern_from_sqp_via_api",
                        return_value=extracted) as mextract, \
             mock.patch("api.ai_service.generate_pattern_via_api") as mgen:
            from core.tasks import generate_pattern_task
            generate_pattern_task.apply(args=[pattern.id])
        mextract.assert_called_once()
        self.assertEqual(mextract.call_args.kwargs["sqp_text"], self.SQP_TEXT)
        mgen.assert_not_called()
        pattern.refresh_from_db()
        self.assertEqual(pattern.status, "done")
        self.assertEqual(pattern.total_marks, 2)
        self.assertEqual(len(pattern.sections), 1)

    def test_extraction_prompt_forbids_copying_content(self):
        from unittest import mock
        from api import ai_service
        reply = '{"sections": [], "total_marks": 80, "total_questions": 11}'
        with mock.patch.object(ai_service.mantle_client, "converse",
                               return_value=(reply, 10, 10)) as mconv:
            data = ai_service.extract_pattern_from_sqp_via_api(
                self.SQP_TEXT, "10", "English", "SQP 2025-26")
        prompt = mconv.call_args.kwargs["prompt"]
        self.assertIn("NEVER copy", prompt)                       # abstraction rule present
        self.assertIn("SAMPLE QUESTION PAPER — CLASS X", prompt)  # paper text embedded
        self.assertIn("QUESTION SLOT RULES", prompt)              # shared slot schema included
        self.assertEqual(data["total_marks"], 80)


class PoemPassageRenderTest(TestCase):
    """A poem quoted WITH its line breaks must PRINT as verse: the DOCX passage box
    must turn every "\\n" in source_text into a real <w:br/> (python-docx does this in
    its run-text setter — this test pins that guarantee so a renderer change can never
    silently flatten poems again). The breaks being PRESENT in source_text is the
    generation prompt's job (POEM FORMATTING rule in section_generator)."""

    POEM = ("He should be lurking in shadow,\nSliding through long grass\n"
            "Near the water hole\nWhere plump deer pass.")

    def test_poem_line_breaks_survive_docx_rendering(self):
        from docx import Document
        doc = Document()
        sg_gen._add_passage_box(doc, self.POEM)
        xml = doc.tables[-1].rows[0].cells[0]._tc.xml
        self.assertEqual(xml.count("<w:br/>"), 3)               # one per verse line break
        for line in self.POEM.split("\n"):
            self.assertIn(line, xml)                            # every line's text intact


class MarkdownTableRenderTest(TestCase):
    """LLM question/passage text may embed data tables as markdown "| a | b |" blocks
    (e.g. a conduction-tester observation table). The DOCX renderer must turn those into
    real Word tables — not print the raw pipe rows as text (production defect)."""

    QTEXT = (
        "17. Study the diagram of the conduction tester shown above. A student uses this "
        "tester to check various objects and creates the following observation table:\n"
        "| Object Tested | Bulb Glows? |\n"
        "|----------------|-------------|\n"
        "| Iron Nail | Yes |\n"
        "| Plastic Spoon | No |\n"
        "| Common Salt Solution in Water | Yes (Dimly) |\n"
        "Which objects are conductors of electricity? [3 marks]"
    )

    def test_segments_parse_table_with_header(self):
        segs = sg_gen._md_table_segments(self.QTEXT)
        self.assertEqual([k for k, _ in segs], ["text", "table", "text"])
        table = segs[1][1]
        self.assertTrue(table["header"])
        self.assertEqual(len(table["rows"]), 4)            # separator row dropped
        self.assertEqual(table["rows"][0], ["Object Tested", "Bulb Glows?"])
        self.assertEqual(table["rows"][3], ["Common Salt Solution in Water", "Yes (Dimly)"])

    def test_plain_text_returns_none(self):
        self.assertIsNone(sg_gen._md_table_segments("What is 2 + 2? Explain your answer."))
        # determinant-style pipes are not table rows (lines don't start AND end with |)
        self.assertIsNone(sg_gen._md_table_segments("Find |A| = 5\nand |B| = 3 for the matrices."))

    def test_single_pipe_line_stays_text(self):
        self.assertIsNone(sg_gen._md_table_segments(
            "Read this:\n| just one decorated line |\nand answer."))

    def test_question_renders_real_docx_table(self):
        import re as _re
        from docx import Document as Doc
        mp = _re.compile(r"\s*\[(\d+)\s*marks?\]", _re.IGNORECASE)
        doc = Doc()
        segs = sg_gen._md_table_segments(mp.sub("", self.QTEXT).rstrip())
        self.assertIsNotNone(segs)
        sg_gen._render_question_segments(doc, segs, self.QTEXT, mp)
        self.assertEqual(len(doc.tables), 1)
        tbl = doc.tables[0]
        self.assertEqual(tbl.cell(0, 0).text, "Object Tested")
        self.assertEqual(tbl.cell(2, 1).text, "No")
        self.assertTrue(tbl.cell(0, 0).paragraphs[0].runs[0].bold)   # header row bold
        body = [p.text for p in doc.paragraphs]
        self.assertTrue(any(t.endswith("[3]") for t in body))        # marks kept on the stem
        self.assertTrue(any("Which objects are conductors" in t for t in body))
        self.assertFalse(any("| Iron Nail" in t for t in body))      # no raw pipe rows

    def test_marks_tag_glued_to_last_row_still_parses(self):
        # process_question appends " [N marks]" to the END of the text — which lands on the
        # last table row. render_docx strips it before segmentation; verify that path.
        import re as _re
        mp = _re.compile(r"\s*\[(\d+)\s*marks?\]", _re.IGNORECASE)
        segs = sg_gen._md_table_segments(mp.sub("", self.QTEXT.rstrip()
                                                .replace("Which objects are conductors of electricity? [3 marks]", ""))
                                         .rstrip())
        self.assertIsNotNone(segs)
        self.assertEqual(segs[-1][0], "table")             # table can end the question

    def test_passage_box_nests_table(self):
        from docx import Document as Doc
        doc = Doc()
        sg_gen._add_passage_box(doc, "A student records:\n| Metal | Conducts |\n| Copper | Yes |\nStudy it.")
        outer = doc.tables[0]
        cell = outer.cell(0, 0)
        self.assertIn("A student records:", cell.paragraphs[0].text)
        self.assertEqual(len(cell.tables), 1)
        self.assertEqual(cell.tables[0].cell(0, 0).text, "Metal")
        self.assertEqual(cell.tables[0].cell(1, 1).text, "Yes")
        self.assertFalse(any("| Copper" in p.text for p in cell.paragraphs))


# ── Answer key (teacher copy) ──────────────────────────────────────────────────

from core import answer_key_generator as akg
from core.answer_key_docx import render_answer_key_docx
from core.models import QuestionPaper, AnswerKey


class AnswerTargetsTest(TestCase):
    """answer_targets must expose every independently answerable part of a question:
    the main text, each sub-question, and each internal-choice alternative."""

    def test_simple_question_single_target(self):
        t = akg.answer_targets({"text": "Define photosynthesis.", "marks": 3})
        self.assertEqual([x["id"] for x in t], ["main"])
        self.assertEqual(t[0]["marks"], 3)

    def test_sub_questions_become_parts(self):
        q = {"text": "Read the extract.", "marks": 4,
             "sub_questions": [{"text": "Who wrote it?", "marks": 1},
                               {"text": "Explain the theme.", "marks": 3}]}
        t = akg.answer_targets(q)
        self.assertEqual([x["id"] for x in t], ["part_1", "part_2"])
        self.assertEqual([x["marks"] for x in t], [1, 3])
        self.assertIn("(a)", t[0]["label"])

    def test_or_alternative_dict_and_string(self):
        q = {"text": "Main question", "marks": 5,
             "or_alternative": [{"text": "Alt question", "marks": 5}, "Plain-text alternative"]}
        ids = [x["id"] for x in akg.answer_targets(q)]
        self.assertEqual(ids, ["main", "alternative_1", "alternative_2"])

    def test_mcq_options_passed_through(self):
        q = {"text": "Pick one.", "marks": 1, "options": {"a": "x", "b": "y"}}
        t = akg.answer_targets(q)
        self.assertEqual(t[0]["options"], {"a": "x", "b": "y"})

    def test_textless_targets_dropped(self):
        self.assertEqual(akg.answer_targets({"text": "", "marks": 1}), [])


class AnswerKeyNormaliseTest(TestCase):
    """_normalise_response must keep only supplied evidence ids, flag marking schemes
    that don't total the question's marks, and reject missing/empty answers."""

    EVIDENCE = [{"chunk_id": "c1", "material_id": 1, "title": "t", "unit": "u",
                 "excerpt": "e", "distance": 0.1}]

    def _payload(self, **over):
        answer = {"target": "main", "answer": "Green plants make food using sunlight.",
                  "correct_option": "", "marking_scheme": [{"point": "definition", "marks": 2},
                                                           {"point": "equation", "marks": 1}],
                  "concept": {"name": "Photosynthesis", "chapter": "Nutrition", "explanation": "x"},
                  "insight": {"explanation": "e", "common_misconception": "m", "revision_tip": "r"},
                  "evidence_chunk_ids": ["c1", "made-up"], "confidence": "high"}
        answer.update(over)
        return {"answers": [answer]}

    def test_valid_answer_keeps_only_supplied_evidence(self):
        targets = akg.answer_targets({"text": "Define photosynthesis.", "marks": 3})
        answers, issues = akg._normalise_response(self._payload(), targets, self.EVIDENCE)
        self.assertEqual(issues, [])
        self.assertEqual([e["chunk_id"] for e in answers[0]["evidence"]], ["c1"])

    def test_marks_mismatch_recorded_as_issue(self):
        targets = akg.answer_targets({"text": "Define photosynthesis.", "marks": 5})
        _, issues = akg._normalise_response(self._payload(), targets, self.EVIDENCE)
        self.assertTrue(any("instead of 5" in i for i in issues), issues)

    def test_missing_target_raises(self):
        targets = akg.answer_targets({"text": "Define photosynthesis.", "marks": 3})
        with self.assertRaises(ValueError):
            akg._normalise_response({"answers": []}, targets, self.EVIDENCE)

    def test_invalid_mcq_option_flagged_and_cleared(self):
        targets = akg.answer_targets({"text": "Pick one.", "marks": 1,
                                      "options": {"a": "x", "b": "y", "c": "z", "d": "w"}})
        payload = self._payload(correct_option="e",
                                marking_scheme=[{"point": "correct choice", "marks": 1}])
        answers, issues = akg._normalise_response(payload, targets, self.EVIDENCE)
        self.assertEqual(answers[0]["correct_option"], "")
        self.assertTrue(any("invalid MCQ option" in i for i in issues), issues)


class PaperRevisionHashTest(TestCase):
    def test_key_order_independent_but_content_sensitive(self):
        h1 = akg.paper_revision_hash({"a": 1, "b": [1, 2]})
        h2 = akg.paper_revision_hash({"b": [1, 2], "a": 1})
        self.assertEqual(h1, h2)
        self.assertNotEqual(h1, akg.paper_revision_hash({"a": 2, "b": [1, 2]}))


def _mk_paper(paper_data=None):
    pattern = ExamPattern.objects.create(
        name="PT-1", sections=[{"name": "A", "marks": 2, "questions_count": 1}])
    return QuestionPaper.objects.create(
        class_name="6-A", subject="Science", pattern=pattern, chapters=["1"],
        status="done", paper_data=paper_data or {})


class AnswerKeyDocxRenderTest(TestCase):
    """The on-demand DOCX renderer must produce a readable key: title, question,
    correct option (with the option's text), marking scheme, low-confidence flag,
    and the failed-questions appendix."""

    def test_render_contains_all_parts(self):
        paper = _mk_paper()
        key_data = {
            "paper": {"id": paper.id},
            "sections": [{"name": "Section A", "questions": [{
                "qnum": 1, "text": "Which gas do plants release?", "type": "MCQ", "subtype": "",
                "marks": 1, "chapter_tag": "Ch1",
                "options": {"a": "CO2", "b": "Oxygen", "c": "Nitrogen", "d": "Hydrogen"},
                "answers": [{"target": "main", "label": "Answer",
                             "question": "Which gas do plants release?", "marks": 1,
                             "answer": "Plants release oxygen during photosynthesis.",
                             "correct_option": "b",
                             "marking_scheme": [{"point": "names oxygen", "marks": 1}],
                             "concept": {"name": "Photosynthesis", "chapter": "Nutrition",
                                         "explanation": "x"},
                             "insight": {"explanation": "e", "common_misconception": "confuses CO2",
                                         "revision_tip": "recall the equation"},
                             "confidence": "low", "evidence": [], "warnings": []}],
                "warnings": [],
            }]}],
            "errors": [{"section": "Section A", "qnum": 2, "error": "model failed"}],
            "generated_questions": 1,
        }
        from docx import Document as Doc
        doc = Doc(render_answer_key_docx(paper, key_data, school_name="Test School"))
        text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("ANSWER KEY", text)
        self.assertIn("Q1. Which gas do plants release?", text)
        self.assertIn("Correct option: (b) Oxygen", text)
        self.assertIn("names oxygen", text)
        self.assertIn("Low confidence", text)
        self.assertIn("Revision tip: recall the equation", text)
        self.assertIn("model failed", text)


class AnswerKeyStalenessTest(TestCase):
    """A finished key must flip to 'stale' the moment the paper's stored JSON no longer
    matches the hash the key was generated from — covering ai_edit, restore, rerender
    and whole-paper regenerate without per-path hooks."""

    def test_done_key_flips_stale_when_paper_changes(self):
        from api.views import _sync_answer_key_staleness
        pd = {"Section A": {"questions": [{"qnum": 1, "text": "Q?", "marks": 2}]}}
        paper = _mk_paper(pd)
        key = AnswerKey.objects.create(paper=paper, status="done", data={"sections": []},
                                       source_revision_hash=akg.paper_revision_hash(pd))
        self.assertEqual(_sync_answer_key_staleness(paper).status, "done")
        paper.paper_data = {"Section A": {"questions": [{"qnum": 1, "text": "Edited?", "marks": 2}]}}
        paper.save(update_fields=["paper_data"])
        self.assertEqual(_sync_answer_key_staleness(paper).status, "stale")
        key.refresh_from_db()
        self.assertEqual(key.status, "stale")

    def test_no_key_returns_none_and_payload_says_none(self):
        from api.views import _sync_answer_key_staleness, _answer_key_payload
        paper = _mk_paper()
        self.assertIsNone(_sync_answer_key_staleness(paper))
        self.assertEqual(_answer_key_payload(None), {"status": "none"})

    def test_payload_reports_counts(self):
        from api.views import _answer_key_payload
        paper = _mk_paper()
        key = AnswerKey.objects.create(
            paper=paper, status="done",
            data={"generated_questions": 12, "errors": [{"qnum": 3, "error": "x"}]})
        payload = _answer_key_payload(key)
        self.assertEqual(payload["status"], "done")
        self.assertEqual(payload["generated_questions"], 12)
        self.assertEqual(payload["failed_questions"], 1)


class AnswerKeyApiTest(TestCase):
    """End-to-end viewset flow: status 'none' → POST queues the Celery task (mocked)
    → repeat POST is idempotent → finished key streams a valid DOCX."""

    def setUp(self):
        from rest_framework.test import APIClient
        self.user = User.objects.create_user(username="key-teacher", password="x")
        self.paper = _mk_paper({"Section A": {"questions": [{"qnum": 1, "text": "Q?", "marks": 2}]}})
        self.paper.created_by = self.user
        self.paper.save(update_fields=["created_by"])
        self.api = APIClient()
        self.api.force_authenticate(self.user)

    def test_status_queue_idempotence_and_download(self):
        from unittest.mock import MagicMock, patch

        r = self.api.get(f"/api/papers/{self.paper.id}/answer_key/")
        self.assertEqual(r.status_code, 200)
        self.assertEqual(r.json()["status"], "none")

        with patch("core.tasks.generate_answer_key_task") as task:
            task.delay.return_value = MagicMock(id="tid-123")
            r = self.api.post(f"/api/papers/{self.paper.id}/answer_key/")
        self.assertEqual(r.status_code, 202)
        self.assertEqual(r.json()["status"], "queued")
        key = AnswerKey.objects.get(paper=self.paper)
        self.assertEqual(key.task_id, "tid-123")

        with patch("core.tasks.generate_answer_key_task") as task2:
            r = self.api.post(f"/api/papers/{self.paper.id}/answer_key/")
            self.assertEqual(r.status_code, 200)   # already in flight — reported, not re-queued
            task2.delay.assert_not_called()

        key.status = "done"
        key.source_revision_hash = akg.paper_revision_hash(self.paper.paper_data)
        key.data = {"sections": [{"name": "Section A", "questions": [{
            "qnum": 1, "text": "Q?", "marks": 2, "options": {},
            "answers": [{"target": "main", "label": "Answer", "question": "Q?", "marks": 2,
                         "answer": "The answer.", "correct_option": "",
                         "marking_scheme": [{"point": "p", "marks": 2}],
                         "concept": {}, "insight": {}, "confidence": "high",
                         "evidence": [], "warnings": []}],
            "warnings": []}]}], "errors": [], "generated_questions": 1}
        key.save()
        r = self.api.get(f"/api/papers/{self.paper.id}/answer_key_docx/")
        self.assertEqual(r.status_code, 200)
        self.assertEqual(b"".join(r.streaming_content)[:2], b"PK")   # a real ZIP/DOCX

    def test_post_without_paper_data_rejected(self):
        paper = _mk_paper({})
        paper.created_by = self.user
        paper.save(update_fields=["created_by"])
        r = self.api.post(f"/api/papers/{paper.id}/answer_key/")
        self.assertEqual(r.status_code, 400)
        self.assertEqual(r.json()["error"], "no_paper_data")

    def test_download_without_key_is_404(self):
        r = self.api.get(f"/api/papers/{self.paper.id}/answer_key_docx/")
        self.assertEqual(r.status_code, 404)

    def test_list_serializer_reports_answer_key_status(self):
        r = self.api.get("/api/papers/")
        row = next(p for p in r.json()["results"] if p["id"] == self.paper.id)
        self.assertEqual(row["answer_key_status"], "none")
        AnswerKey.objects.create(paper=self.paper, status="done")
        r = self.api.get("/api/papers/")
        row = next(p for p in r.json()["results"] if p["id"] == self.paper.id)
        self.assertEqual(row["answer_key_status"], "done")


# ── Open-choice slot marks + unseen comprehension (Tamil PT-1 incident) ────────

def _tamil_sec_d(q18_marks=2, q19_marks=4, part_marks=(2, 4)):
    """The exact SEC_D shape from the Tamil PT-1 incident: open-choice slots whose
    'marks' carry the per-part value (2 and 4) instead of the earnable totals
    (attempt 4 x 2m = 8 and attempt 2 x 4m = 8), shrinking the 40-mark paper to 30."""
    m18, m19 = part_marks
    # qnums 1/2 rather than the real paper's 18/19 — a lone section must still
    # satisfy the run-1..N qnum check, and the marks logic is qnum-independent.
    return [{"id": "SEC_D", "name": "பகுதி - ஈ", "marks": 16,
             "instructions": ["எவையேனும் நான்கு வினாக்களுக்கு குறு விடையளி.",
                              "எவையேனும் இரண்டு வினாக்களுக்கு மட்டும் விடையளி."],
             "question_slots": [
                 {"qnum": 1, "type": "vsa", "marks": q18_marks, "choice": "open", "attempt": 4,
                  "parts": [{"label": c, "type": "vsa", "marks": m18} for c in "ABCDEF"]},
                 {"qnum": 2, "type": "sa", "marks": q19_marks, "choice": "open", "attempt": 2,
                  "parts": [{"label": c, "type": "sa", "marks": m19} for c in "ABCD"]},
             ]}]


class OpenChoiceMarksReconcileTest(TestCase):
    """normalize_slots must promote an open-choice slot's marks to the earnable
    total (attempt x per-part marks) when the LLM put the per-part value or the
    all-parts sum there — deterministically, before validation, so the repair
    round is never needed for this error class."""

    def _slots(self, secs):
        return secs[0]["question_slots"]

    def test_per_part_marks_promoted_to_earnable_total(self):
        secs = psx.normalize_slots(_tamil_sec_d())
        self.assertEqual([s["marks"] for s in self._slots(secs)], [8, 8])
        errors = psx.validate_pattern_structure(secs)
        self.assertEqual(errors, [], errors)
        psx.derive_aggregates_from_slots(secs)
        self.assertEqual(secs[0]["marks"], 16)   # was shrunk to 6 in production

    def test_all_parts_sum_also_reconciled(self):
        # The other plausible confusion: marks = every part summed (6 x 2 = 12).
        secs = psx.normalize_slots(_tamil_sec_d(q18_marks=12, q19_marks=16))
        self.assertEqual([s["marks"] for s in self._slots(secs)], [8, 8])

    def test_missing_marks_derived(self):
        secs = psx.normalize_slots(_tamil_sec_d(q18_marks=0, q19_marks=0))
        self.assertEqual([s["marks"] for s in self._slots(secs)], [8, 8])

    def test_correct_marks_untouched(self):
        secs = psx.normalize_slots(_tamil_sec_d(q18_marks=8, q19_marks=8))
        self.assertEqual([s["marks"] for s in self._slots(secs)], [8, 8])
        self.assertEqual(psx.validate_pattern_structure(secs), [])

    def test_ambiguous_conflict_left_for_validator(self):
        # marks 6 matches neither the per-part value (2) nor the parts sum (12) —
        # rewriting it could destroy a correct teacher total, so it must survive
        # normalize and be flagged by the validator instead.
        secs = psx.normalize_slots(_tamil_sec_d(q18_marks=6, q19_marks=8))
        self.assertEqual(self._slots(secs)[0]["marks"], 6)
        errors = psx.validate_pattern_structure(secs)
        self.assertTrue(any("attempt (4) x part marks" in e["msg"] for e in errors), errors)

    def test_declared_paper_total_now_consistent(self):
        # With the reconcile, the whole-paper sum check passes against the
        # teacher's declared 16 for this section (was "30 but declared 40").
        secs = psx.normalize_slots(_tamil_sec_d())
        errors = psx.validate_pattern_structure(secs, declared_total=16)
        self.assertEqual([e for e in errors if "declared total" in e["msg"]], [])

    def test_repair_may_fix_conflicted_open_choice_marks(self):
        # The ambiguous case reaches the repair round; the guard must accept a
        # repair that changes the conflicted slot's marks to the earnable total
        # (it used to reject ANY slot-marks change, making this class unfixable)…
        original = psx.normalize_slots(_tamil_sec_d(q18_marks=6, q19_marks=8))
        repaired = psx.normalize_slots(_tamil_sec_d(q18_marks=8, q19_marks=8))
        self.assertTrue(psx.repair_preserves_slots(original, repaired))
        # …while still rejecting marks changes on slots that were NOT conflicted.
        clean = psx.normalize_slots(_tamil_sec_d(q18_marks=8, q19_marks=8))
        devalued = psx.normalize_slots(_tamil_sec_d(q18_marks=8, q19_marks=8))
        devalued[0]["question_slots"][1]["marks"] = 6
        devalued[0]["question_slots"][1]["parts"] = [
            {"label": c, "type": "sa", "marks": 3} for c in "ABCD"]
        self.assertFalse(psx.repair_preserves_slots(clean, devalued))


class UnseenComprehensionPromptTest(TestCase):
    """A cbq slot with source 'unseen' and parts (the encoding for 'read the
    passage/poem and answer' exercises) must instruct the generator to COMPOSE
    a new passage in source_text — not to quote the textbook, and never to emit
    bare standalone questions with no passage."""

    def _prompt(self):
        secs = psx.normalize_slots([{"id": "SEC_A", "name": "பகுதி - அ", "marks": 8,
            "instructions": ["உரைப்பத்தியை படித்து பொருளுணர்ந்து வினாக்களுக்கு விடையளி."],
            "question_slots": [
                {"qnum": 1, "type": "cbq", "marks": 4, "source": "unseen", "topic": "உரைப்பத்தி",
                 "parts": [{"label": c, "type": "mcq", "marks": 1} for c in "ABCD"]},
                {"qnum": 2, "type": "cbq", "marks": 4, "source": "unseen", "topic": "பாடல்",
                 "parts": [{"label": c, "type": "mcq", "marks": 1} for c in "ABCD"]},
            ]}])
        psx.derive_aggregates_from_slots(secs)
        pattern = ExamPattern(name="p", sections=secs)
        bp = sg_gen.pattern_sections_to_blueprint_dict(pattern)
        wo = sg.build_work_orders(bp, pattern, {}, "Medium", "6", "Tamil", ["Ch1"])[0]
        return sg.build_section_prompt(wo)

    def test_unseen_cbq_prompt_demands_new_passage(self):
        prompt = self._prompt()
        self.assertIn("COMPOSE a NEW, original passage", prompt)
        self.assertIn('"source_text"', prompt)
        self.assertIn("answerable ONLY from that passage", prompt)

    def test_unseen_cbq_prompt_does_not_demand_verbatim_quote(self):
        # VERBATIM quoting is the textbook-extract instruction — an unseen
        # comprehension passage must not inherit it.
        self.assertNotIn("VERBATIM", self._prompt())

    def test_pattern_prompt_rules_state_both_conventions(self):
        self.assertIn("READING COMPREHENSION", psx.SLOT_SCHEMA_PROMPT_RULES)
        self.assertIn("attempt x per-part marks", psx.SLOT_SCHEMA_PROMPT_RULES)


class NormalizeLabelUnicodeTest(TestCase):
    """normalize_label must preserve non-ASCII (Indic-script) chapter names — the old
    ASCII-only regex reduced them to "", so Tamil/Hindi chunks got no ChunkChapter links
    and the unit filter in embeddings.query silently vanished (whole-book retrieval)."""

    def test_tamil_chapter_name_survives(self):
        from core.embeddings import normalize_label
        self.assertEqual(normalize_label("மொழிமுதல் எழுத்துகள்"), "மொழிமுதல்_எழுத்துகள்")

    def test_hindi_chapter_name_survives(self):
        from core.embeddings import normalize_label
        self.assertEqual(normalize_label("क्षितिज - बालगोबिन भगत"), "क्षितिज_बालगोबिन_भगत")

    def test_ascii_behavior_unchanged(self):
        from core.embeddings import normalize_label
        # Existing ChunkChapter rows store labels produced by the OLD function — these
        # exact outputs must never change or every stored link goes stale.
        self.assertEqual(normalize_label("Light - Reflection & Refraction"),
                         "light_reflection_refraction")
        self.assertEqual(normalize_label("English Language & Literature"),
                         "english_language_literature")
        self.assertEqual(normalize_label("Class 10"), "class_10")
        self.assertEqual(normalize_label("chapter_1"), "chapter_1")

    def test_ascii_punctuation_still_stripped(self):
        from core.embeddings import normalize_label
        self.assertEqual(normalize_label("நூலகம் நோக்கி ..."), "நூலகம்_நோக்கி")
        self.assertIsNone(normalize_label(""))
        self.assertEqual(normalize_label("..."), "")

    def test_idempotent(self):
        from core.embeddings import normalize_label
        once = normalize_label("மொழிமுதல் எழுத்துகள்")
        self.assertEqual(normalize_label(once), once)


class UnicodeChapterLinkTest(TestCase):
    """_store_chunks must create ChunkChapter links for non-ASCII unit labels (they were
    dropped entirely before), and the query-side unit filter must match them."""

    def _run_ingest(self, **kwargs):
        from unittest import mock
        from core import embeddings as emb
        unit = kwargs.pop("unit", "ignored")
        with mock.patch.object(emb, "PdfReader", lambda *_a, **_k: _FakeReader("x" * 1600)), \
             mock.patch.object(emb, "get_embeddings_batch",
                               side_effect=lambda chunks, provider: [[0.0] * 768 for _ in chunks]):
            return emb.ingest_pdf("6", "Tamil", unit, "C:/fake.pdf", **kwargs)

    def test_tamil_unit_creates_links(self):
        from core.models import MaterialChunk, ChunkChapter
        n = self._run_ingest(unit="மொழிமுதல் எழுத்துகள்", material_type="textbook")
        self.assertEqual(n, 2)
        self.assertEqual(set(ChunkChapter.objects.values_list("unit", flat=True)),
                         {"மொழிமுதல்_எழுத்துகள்"})
        # and the same normalization on the query side finds them
        from core.embeddings import normalize_label
        u = normalize_label("மொழிமுதல் எழுத்துகள்")
        self.assertEqual(MaterialChunk.objects.filter(chapter_links__unit=u).count(), 2)


class GrammarSectionDetectTest(TestCase):
    """_is_grammar_section: grammar sections are detected from name/instructions in any
    supported script; ordinary sections never match."""

    def test_tamil_grammar_name(self):
        self.assertTrue(sg._is_grammar_section("பகுதி - இ (இலக்கணம்)"))

    def test_english_grammar_in_instructions(self):
        self.assertTrue(sg._is_grammar_section("Section B", ["Writing & Grammar"]))

    def test_hindi_grammar_name(self):
        self.assertTrue(sg._is_grammar_section("खंड ब (व्याकरण)"))

    def test_plain_sections_do_not_match(self):
        self.assertFalse(sg._is_grammar_section("பகுதி - அ", ["பகுதி - அ"]))
        self.assertFalse(sg._is_grammar_section("Section C — Literature"))
        self.assertFalse(sg._is_grammar_section("Section A", ["Answer all questions"]))

    def test_blueprint_gate(self):
        self.assertTrue(sg._blueprint_has_grammar_section(
            {"பகுதி - அ": {}, "பகுதி - இ (இலக்கணம்)": {"instructions": []}}))
        self.assertFalse(sg._blueprint_has_grammar_section(
            {"Section A": {"instructions": ["Answer all"]}, "Section B": {}}))


class GrammarChapterIdentifyTest(TestCase):
    """identify_grammar_chapters: keyword titles skip the LLM; unknown titles get ONE
    cached LLM classification; hallucinated titles are dropped; LLM failure fails open."""

    def setUp(self):
        sg._grammar_chapters_cache.clear()

    def tearDown(self):
        sg._grammar_chapters_cache.clear()

    def test_keyword_titles_skip_llm(self):
        from unittest import mock
        with mock.patch.object(sg.mantle_client, "converse") as conv:
            out = sg.identify_grammar_chapters("6", "tamil", ["மொழி இலக்கணம்", "Grammar Basics"])
        conv.assert_not_called()
        self.assertEqual(set(out), {"மொழி இலக்கணம்", "Grammar Basics"})

    def test_llm_classifies_unknown_titles_and_caches(self):
        from unittest import mock
        chapters = ["இன்பத்தமிழ்", "மொழிமுதல் எழுத்துகள்", "திருக்குறள்"]
        reply = '["மொழிமுதல் எழுத்துகள்", "Hallucinated Chapter"]'
        with mock.patch.object(sg.mantle_client, "converse",
                               return_value=(reply, 0, 0)) as conv:
            out1 = sg.identify_grammar_chapters("6", "tamil", chapters)
            out2 = sg.identify_grammar_chapters("6", "tamil", chapters)
        self.assertEqual(out1, ["மொழிமுதல் எழுத்துகள்"])   # hallucination filtered out
        self.assertEqual(out2, out1)
        conv.assert_called_once()                            # second call served from cache

    def test_llm_failure_fails_open(self):
        from unittest import mock
        with mock.patch.object(sg.mantle_client, "converse", side_effect=RuntimeError("down")):
            out = sg.identify_grammar_chapters("6", "tamil", ["இன்பத்தமிழ்", "திருக்குறள்"])
        self.assertEqual(out, [])

    def test_empty_chapters(self):
        self.assertEqual(sg.identify_grammar_chapters("6", "tamil", []), [])


class GrammarChapterRoutingTest(TestCase):
    """_route_grammar_chapters + build_work_orders wiring: grammar sections keep only the
    grammar lessons, other sections drop them, and nothing changes when no grammar section
    exists or no grammar chapters were identified."""

    CHAPTERS = ["இன்பத்தமிழ்", "மொழிமுதல் எழுத்துகள்", "திருக்குறள்"]
    GRAM = ["மொழிமுதல் எழுத்துகள்"]

    def setUp(self):
        sg._grammar_chapters_cache.clear()

    def tearDown(self):
        sg._grammar_chapters_cache.clear()

    def test_route_helper(self):
        self.assertEqual(sg._route_grammar_chapters(True, self.CHAPTERS, self.GRAM), self.GRAM)
        self.assertEqual(sg._route_grammar_chapters(False, self.CHAPTERS, self.GRAM),
                         ["இன்பத்தமிழ்", "திருக்குறள்"])
        # fail-open fallbacks: no grammar chapters identified / routing would starve the section
        self.assertEqual(sg._route_grammar_chapters(True, ["a", "b"], []), ["a", "b"])
        self.assertEqual(sg._route_grammar_chapters(True, ["a", "b"], ["c"]), ["a", "b"])
        self.assertEqual(sg._route_grammar_chapters(False, ["c"], ["c"]), ["c"])

    def _blueprint(self):
        return {
            "பகுதி - அ": {"marks": 8, "questions_count": 2, "marks_per_question": 4,
                          "question_types": ["Short Answer"], "instructions": ["பகுதி - அ"]},
            "பகுதி - இ (இலக்கணம்)": {"marks": 8, "questions_count": 2, "marks_per_question": 4,
                                     "question_types": ["MCQ"],
                                     "instructions": ["பகுதி - இ (இலக்கணம்)"]},
        }

    def _work_orders(self):
        from unittest import mock
        with mock.patch.object(sg.mantle_client, "converse",
                               return_value=('["மொழிமுதல் எழுத்துகள்"]', 0, 0)):
            return sg.build_work_orders(self._blueprint(), None, {}, "Medium",
                                        "6", "Tamil", list(self.CHAPTERS))

    def test_work_orders_scope_chapters(self):
        wos = {w.section_name: w for w in self._work_orders()}
        gram_wo = wos["பகுதி - இ (இலக்கணம்)"]
        lit_wo = wos["பகுதி - அ"]
        self.assertTrue(gram_wo.is_grammar)
        self.assertFalse(lit_wo.is_grammar)
        self.assertEqual(gram_wo.chapters, self.GRAM)
        self.assertEqual(lit_wo.chapters, ["இன்பத்தமிழ்", "திருக்குறள்"])
        # chapter_plan follows the scoped lists
        self.assertTrue(set(gram_wo.chapter_plan) <= set(self.GRAM))
        self.assertTrue(set(lit_wo.chapter_plan) <= {"இன்பத்தமிழ்", "திருக்குறள்"})

    def test_grammar_prompt_rule(self):
        wos = {w.section_name: w for w in self._work_orders()}
        gram_prompt = sg.build_section_prompt(wos["பகுதி - இ (இலக்கணம்)"])
        lit_prompt = sg.build_section_prompt(wos["பகுதி - அ"])
        self.assertIn("GRAMMAR SECTION", gram_prompt)
        self.assertIn("Do NOT ask reading-comprehension", gram_prompt)
        self.assertNotIn("GRAMMAR SECTION", lit_prompt)

    def test_no_grammar_section_means_no_llm_and_no_scoping(self):
        from unittest import mock
        bp = {"Section A": {"marks": 10, "questions_count": 2, "marks_per_question": 5,
                            "question_types": ["Short Answer"], "instructions": ["Answer all"]}}
        with mock.patch.object(sg.mantle_client, "converse") as conv:
            wos = sg.build_work_orders(bp, None, {}, "Medium", "10", "Science",
                                       ["Light", "Electricity"])
        conv.assert_not_called()
        self.assertEqual(wos[0].chapters, ["Light", "Electricity"])
        self.assertFalse(wos[0].is_grammar)


class UnitVariantMatchTest(SimpleTestCase):
    """embeddings._unit_variants: fuzzy fallback so a chapter whose stored label drifted from
    the planned name still retrieves (the 'no question anywhere for <chapter>' coverage gap).
    Both the query unit and the stored labels are already normalize_label()-ed."""

    def test_number_prefix_variant_matches(self):
        from core.embeddings import _unit_variants
        stored = ["13_the_sermon_at_benares", "the_proposal", "amanda"]
        self.assertEqual(_unit_variants("the_sermon_at_benares", stored),
                         ["13_the_sermon_at_benares"])

    def test_shortened_stored_label_matches(self):
        from core.embeddings import _unit_variants
        # Plan carries the article, the ingested label dropped it.
        self.assertEqual(_unit_variants("the_sermon_at_benares", ["sermon_at_benares"]),
                         ["sermon_at_benares"])

    def test_unrelated_chapters_do_not_match(self):
        from core.embeddings import _unit_variants
        self.assertEqual(_unit_variants("amanda", ["the_proposal", "dust_of_snow"]), [])

    def test_empty_unit_matches_nothing(self):
        from core.embeddings import _unit_variants
        self.assertEqual(_unit_variants("", ["amanda", "the_proposal"]), [])
        self.assertEqual(_unit_variants(None, ["amanda"]), [])

    def test_result_is_sorted_and_deduped(self):
        from core.embeddings import _unit_variants
        stored = ["sermon_at_benares", "13_the_sermon_at_benares", "sermon_at_benares"]
        self.assertEqual(_unit_variants("the_sermon_at_benares", stored),
                         ["13_the_sermon_at_benares", "sermon_at_benares"])


class AnswerLeakAuditTest(SimpleTestCase):
    """Teachers reported Social Science papers where an Assertion-Reason stem stated the answer
    to a 2-mark question, and where a case-based passage contained the answer to another 2-mark
    question. Neither is a DUPLICATE, so _cross_section_dup_pairs cannot see them: leakage is
    asymmetric (the answer to Q7 sits in the STEM of Q3, sharing almost no tokens) and it
    dilutes (a 2-mark answer inside a 150-word passage scores near zero on overlap normalised
    across the whole passage). V8 is given only type/chapter summaries and V10 only the
    accumulated warning strings, so neither paper-level audit ever reads question text either."""

    AR = ("Assertion (A): The Dandi March was a satyagraha launched against the British salt "
          "tax, which affected the poorest Indians most.\nReason (R): Salt was consumed by every "
          "household.\n(a) Both A and R are true and R explains A\n(b) Both true, R does not "
          "explain A\n(c) A true, R false\n(d) A false, R true")
    AR_OPTS = {"a": "Both A and R are true and R explains A",
               "b": "Both true, R does not explain A",
               "c": "A true, R false", "d": "A false, R true"}
    SALT_Q = "Why did Mahatma Gandhi choose salt as the symbol of his satyagraha in 1930?"
    SALT_A = ("Salt was taxed by the British and consumed by every household, so the salt tax "
              "affected the poorest Indians most, making it a powerful unifying symbol.")
    PASSAGE = (
        "In the years after 1858 the British reorganised their administration in India. Print "
        "culture spread rapidly and vernacular newspapers multiplied in every province. The "
        "Vernacular Press Act of 1878 was passed to control the vernacular press and gave the "
        "government the power to censor reports and editorials in Indian-language newspapers. "
        "Despite this, nationalist papers continued to grow in circulation across the country.")
    PRESS_Q = "State the purpose of the Vernacular Press Act of 1878."
    PRESS_A = ("It was passed to control the vernacular press and gave the government power to "
               "censor reports and editorials in Indian-language newspapers.")
    AR_SPAN = "The Dandi March was a satyagraha launched against the British salt tax"
    PASSAGE_SPAN = "to control the vernacular press and gave the government the power to censor"

    def _wo(self, name, mpq=2):
        return sg.SectionWorkOrder(
            section_name=name, section_id=name[:1], title="", marks=4, questions_count=2,
            marks_per_question=mpq, question_types=[], instructions=[], constraints={},
            context_text="", difficulty="Medium", subject="Social Science", class_name="10",
            chapters=["Nationalism"])

    def _wos(self):
        return [self._wo(n) for n in ("Section A", "Section B", "Section C")]

    def _paper(self):
        return {
            "Section A": {"marks": 2, "questions": [
                {"qnum": 1, "type": "MCQ", "subtype": "assertion_reason", "marks": 1,
                 "text": self.AR, "options": dict(self.AR_OPTS), "answer": "a",
                 "answer_explanation": "Both statements are true and R explains A."},
                {"qnum": 2, "type": "MCQ", "marks": 1,
                 "text": "In which year was the Rowlatt Act passed?",
                 "options": {"a": "1919", "b": "1920", "c": "1921", "d": "1922"},
                 "answer": "a", "answer_explanation": "The Rowlatt Act was passed in 1919."}]},
            "Section B": {"marks": 4, "questions": [
                {"qnum": 3, "type": "SA", "marks": 2, "text": self.SALT_Q,
                 "answer_explanation": self.SALT_A},
                {"qnum": 4, "type": "SA", "marks": 2, "text": self.PRESS_Q,
                 "answer_explanation": self.PRESS_A}]},
            "Section C": {"marks": 4, "passage": self.PASSAGE, "questions": [
                {"qnum": 5, "type": "CBQ", "marks": 4,
                 "text": "Read the passage above and answer the following:",
                 "sub_questions": [
                     {"text": "Name the Act mentioned in the passage.", "marks": 1,
                      "answer_explanation": "The Vernacular Press Act of 1878."},
                     {"text": "How did nationalist papers respond?", "marks": 3,
                      "answer_explanation": "They continued to grow in circulation."}]}]},
        }

    def _audit(self, *leaks):
        return json.dumps({"leaks": list(leaks)})

    def _leak(self, victim, leaker, span, why="w"):
        return {"victim": victim, "leaker": leaker, "leaked_span": span, "why": why}

    # ── surfaces ──────────────────────────────────────────────────────────────────────
    def test_inventory_ids_are_positional_and_passage_is_its_own_surface(self):
        inv = sg._leak_inventory(self._paper())
        self.assertEqual([it["id"] for it in inv],
                         ["S1Q1", "S1Q2", "S2Q1", "S2Q2", "S3P", "S3Q1"])
        by = {it["id"]: it for it in inv}
        self.assertEqual(by["S3P"]["kind"], "p")
        self.assertEqual(by["S1Q1"]["cat"], "ar")
        self.assertEqual(by["S3Q1"]["cat"], "cbq")
        # ids are POSITIONAL, so cross_section_validate's renumbering cannot invalidate them
        self.assertEqual(by["S2Q2"]["qnum"], 4)

    def test_short_passage_is_not_a_surface(self):
        paper = self._paper()
        paper["Section C"]["passage"] = "Too short."
        self.assertNotIn("S3P", [it["id"] for it in sg._leak_inventory(paper)])

    def test_answer_key_is_protected_but_never_a_disclosure_surface(self):
        # A marking key is never printed for students, so a key restating another question's
        # answer cannot help anyone in the exam hall — including it would flood the audit.
        q = {"text": "Name the Act.", "answer_explanation": "The Vernacular Press Act of 1878."}
        self.assertNotIn("1878", sg._disclosure_text(q))
        self.assertIn("1878", sg._victim_answer_text(q))

    def test_mcq_correct_option_text_is_the_protected_answer(self):
        inv = {it["id"]: it for it in sg._leak_inventory(self._paper())}
        self.assertIn("1919", inv["S1Q2"]["answer"])       # the letter 'a' alone reveals nothing
        self.assertIn("salt tax", inv["S1Q1"]["disclosure"])
        self.assertIn("circulation", inv["S3Q1"]["answer"])  # CBQ sub-answers count too

    def test_disclosure_covers_options_or_alternative_and_sub_questions(self):
        text = sg._disclosure_text({
            "text": "Stem here", "or_alternative": "Alternative brief",
            "options": {"a": "First option", "b": "Second option"},
            "source_text": "Source paragraph", "sub_questions": [{"text": "Sub one"}, "Sub two"]})
        for part in ("Stem here", "Alternative brief", "First option", "Second option",
                     "Source paragraph", "Sub one", "Sub two"):
            self.assertIn(part, text)

    # ── by-design exclusions ──────────────────────────────────────────────────────────
    def test_passage_never_leaks_to_its_own_section(self):
        by = {it["id"]: it for it in sg._leak_inventory(self._paper())}
        # V6 REQUIRES every sub-question to be answerable from the passage alone
        self.assertFalse(sg._leak_pair_allowed(by["S3Q1"], by["S3P"]))
        self.assertFalse(sg._leak_pair_allowed(by["S2Q1"], by["S2Q1"]))
        self.assertFalse(sg._leak_pair_allowed(by["S3P"], by["S1Q1"]))   # passage has no answer
        self.assertTrue(sg._leak_pair_allowed(by["S2Q2"], by["S3P"]))    # cross-section IS a leak
        self.assertTrue(sg._leak_pair_allowed(by["S2Q1"], by["S1Q1"]))

    # ── prefilter ─────────────────────────────────────────────────────────────────────
    def test_prefilter_finds_verbatim_leak_but_not_the_paraphrased_one(self):
        hits = {v: lk for v, lk, _ in sg._leak_prefilter(sg._leak_inventory(self._paper()))}
        self.assertEqual(hits.get("S2Q2"), "S3P")
        # The reported A-R case PARAPHRASES, which is exactly why the scan is a hint and never
        # a gate — gating on it would miss the defect that prompted this whole check.
        self.assertNotIn("S2Q1", hits)

    # ── span verification: the load-bearing guard ─────────────────────────────────────
    def test_audit_confirms_both_reported_leaks(self):
        with mock.patch.object(sg.mantle_client, "converse", return_value=(self._audit(
                self._leak("S2Q1", "S1Q1", self.AR_SPAN),
                self._leak("S2Q2", "S3P", self.PASSAGE_SPAN)), 9, 9)):
            found, in_tok, out_tok = sg.run_answer_leak_audit(
                self._paper(), "10", "Social Science")
        self.assertEqual([(f["victim"], f["leaker"]) for f in found],
                         [("S2Q1", "S1Q1"), ("S2Q2", "S3P")])
        self.assertGreater(in_tok + out_tok, 0)

    def test_hallucinated_span_is_dropped(self):
        with mock.patch.object(sg.mantle_client, "converse", return_value=(self._audit(
                self._leak("S2Q1", "S1Q1",
                           "Gandhi chose salt because it united rich and poor alike")), 1, 1)):
            self.assertEqual(
                sg.run_answer_leak_audit(self._paper(), "10", "Social Science")[0], [])

    def test_span_quoted_from_an_answer_key_is_dropped(self):
        # The key is not printed, so it cannot be the vehicle of a leak.
        with mock.patch.object(sg.mantle_client, "converse", return_value=(self._audit(
                self._leak("S2Q1", "S1Q1",
                           "Both statements are true and R explains A")), 1, 1)):
            self.assertEqual(
                sg.run_answer_leak_audit(self._paper(), "10", "Social Science")[0], [])

    def test_short_unknown_and_by_design_findings_are_dropped(self):
        for leak in (self._leak("S2Q1", "S1Q1", "salt tax"),
                     self._leak("S9Q9", "S1Q1", self.AR_SPAN),
                     self._leak("S2Q1", "S9P", self.AR_SPAN),
                     self._leak("S3Q1", "S3P",
                                "nationalist papers continued to grow in circulation")):
            with mock.patch.object(sg.mantle_client, "converse",
                                   return_value=(self._audit(leak), 1, 1)):
                self.assertEqual(
                    sg.run_answer_leak_audit(self._paper(), "10", "Social Science")[0], [],
                    f"should have been dropped: {leak}")

    def test_span_verification_ignores_case_whitespace_and_punctuation(self):
        # Collapsing whitespace BEFORE stripping punctuation turns 'press, and' into a double
        # space and loses the match, throwing away a real leak over one comma.
        with mock.patch.object(sg.mantle_client, "converse", return_value=(self._audit(
                self._leak("S2Q2", "S3P", "TO CONTROL   the Vernacular Press, and gave the "
                                          "government the power to censor")), 1, 1)):
            self.assertEqual(
                len(sg.run_answer_leak_audit(self._paper(), "10", "Social Science")[0]), 1)

    def test_duplicate_findings_are_collapsed(self):
        with mock.patch.object(sg.mantle_client, "converse", return_value=(self._audit(
                self._leak("S2Q1", "S1Q1", self.AR_SPAN),
                self._leak("S2Q1", "S1Q1", self.AR_SPAN)), 1, 1)):
            self.assertEqual(
                len(sg.run_answer_leak_audit(self._paper(), "10", "Social Science")[0]), 1)

    def test_paper_with_one_question_makes_no_llm_call(self):
        with mock.patch.object(sg.mantle_client, "converse") as m:
            self.assertEqual(sg.run_answer_leak_audit(
                {"S": {"questions": [{"qnum": 1, "type": "SA", "marks": 2, "text": "x",
                                      "answer_explanation": "y"}]}}, "10", "Science"),
                ([], 0, 0))
            m.assert_not_called()

    def test_unparseable_response_changes_nothing(self):
        with mock.patch.object(sg.mantle_client, "converse",
                               return_value=("I could not find any leaks, sorry.", 1, 1)):
            self.assertEqual(
                sg.run_answer_leak_audit(self._paper(), "10", "Social Science")[0], [])

    # ── model selection ───────────────────────────────────────────────────────────────
    def test_audit_prefers_audit_model_and_falls_back_to_gen_model(self):
        self.assertNotEqual(sg.mantle_client.AUDIT_MODEL, sg.GEN_MODEL)  # self-audit is weaker
        seen = []

        def fake(**kw):
            seen.append(kw["model_id"])
            if len(seen) == 1:
                raise RuntimeError("HTTP 404 model not found")
            return ('{"leaks": []}', 1, 1)

        with mock.patch.object(sg.mantle_client, "converse", side_effect=fake):
            sg.run_answer_leak_audit(self._paper(), "10", "Social Science")
        self.assertEqual(seen, [sg.mantle_client.AUDIT_MODEL, sg.GEN_MODEL])

    def test_total_model_outage_leaves_the_paper_untouched(self):
        paper = self._paper()
        with mock.patch.object(sg.mantle_client, "converse",
                               side_effect=RuntimeError("endpoint down")):
            out, _, _ = sg.fix_answer_leaks(paper, self._wos(), "10", "Social Science")
        self.assertEqual(out["Section A"]["questions"][0]["text"], self.AR)
        self.assertEqual(out["Section B"]["questions"][0]["text"], self.SALT_Q)
        self.assertNotIn("_answer_leaks", out["Section B"])

    # ── the fixer's target policy ─────────────────────────────────────────────────────
    def _fixer(self, audit_responses, replacement):
        """Serve the queued audit responses in order, and `replacement` for every regen call."""
        queue = list(audit_responses)

        def fake(**kw):
            if "answer-leak auditor" in kw.get("prompt", ""):
                return (queue.pop(0) if queue else '{"leaks": []}', 5, 5)
            return (json.dumps(replacement), 3, 3)
        return fake

    def test_cheaper_leaker_is_rewritten_and_the_passage_is_never_touched(self):
        new_sa = {"type": "SA", "marks": 2,
                  "text": "Describe two features of the growth of vernacular newspapers.",
                  "answer_explanation": "Rapid multiplication; rising circulation."}
        paper = self._paper()
        with mock.patch.object(sg.mantle_client, "converse", side_effect=self._fixer(
                [self._audit(self._leak("S2Q2", "S3P", self.PASSAGE_SPAN))], new_sa)):
            out, in_tok, out_tok = sg.fix_answer_leaks(
                paper, self._wos(), "10", "Social Science")
        # A passage is section-level scaffolding every question beside it depends on, so the
        # VICTIM is rewritten instead — and the passage survives byte-identical.
        self.assertIn("vernacular newspapers", out["Section B"]["questions"][1]["text"])
        self.assertEqual(out["Section C"]["passage"], self.PASSAGE)
        self.assertEqual(out["Section B"]["questions"][1]["qnum"], 4)     # numbering kept
        self.assertEqual(out["Section B"]["questions"][1]["marks"], 2)    # marks kept
        self.assertEqual([len(s["questions"]) for s in out.values()], [2, 2, 1])
        self.assertNotIn("_answer_leaks", out["Section B"])
        self.assertGreater(in_tok + out_tok, 0)

    def test_one_mark_leaker_is_preferred_over_the_two_mark_victim(self):
        new_ar = {"type": "MCQ", "subtype": "assertion_reason", "marks": 1,
                  "text": ("Assertion (A): The Non-Cooperation Movement was withdrawn in 1922.\n"
                           "Reason (R): Violence broke out at Chauri Chaura.\n"
                           "(a) Both A and R are true and R explains A\n(b) Both true, R does "
                           "not explain A\n(c) A true, R false\n(d) A false, R true"),
                  "options": dict(self.AR_OPTS), "answer": "a",
                  "answer_explanation": "Chauri Chaura caused the withdrawal."}
        paper = self._paper()
        with mock.patch.object(sg.mantle_client, "converse", side_effect=self._fixer(
                [self._audit(self._leak("S2Q1", "S1Q1", self.AR_SPAN))], new_ar)):
            out, _, _ = sg.fix_answer_leaks(paper, self._wos(), "10", "Social Science")
        self.assertIn("Chauri Chaura", out["Section A"]["questions"][0]["text"])
        self.assertEqual(out["Section B"]["questions"][0]["text"], self.SALT_Q)  # victim kept
        self.assertEqual(out["Section A"]["questions"][0]["marks"], 1)

    def test_regen_prompt_carries_the_right_rule_for_each_side(self):
        prompts = []

        def fake(**kw):
            p = kw.get("prompt", "")
            if "answer-leak auditor" in p:
                return (self._audit(self._leak("S2Q2", "S3P", self.PASSAGE_SPAN)), 5, 5)
            prompts.append(p)
            return (json.dumps({"type": "SA", "marks": 2, "text": "A fresh 2-mark question here",
                                "answer_explanation": "Key points."}), 3, 3)

        with mock.patch.object(sg.mantle_client, "converse", side_effect=fake):
            sg.fix_answer_leaks(self._paper(), self._wos(), "10", "Social Science")
        self.assertTrue(prompts)
        self.assertIn("must test DIFFERENT knowledge", prompts[0])   # victim rule
        self.assertIn(self.PASSAGE_SPAN[:40], prompts[0])

    def test_unfixable_leak_is_reported_rather_than_dropped(self):
        def fake(**kw):
            if "answer-leak auditor" in kw.get("prompt", ""):
                return (self._audit(self._leak("S2Q1", "S1Q1", self.AR_SPAN)), 5, 5)
            raise RuntimeError("regen endpoint down")

        paper = self._paper()
        with mock.patch.object(sg.mantle_client, "converse", side_effect=fake):
            out, _, _ = sg.fix_answer_leaks(paper, self._wos(), "10", "Social Science")
        leaks = out["Section B"]["_answer_leaks"]
        self.assertTrue(any("S2Q1 answered by S1Q1" in x for x in leaks))
        self.assertEqual(out["Section A"]["questions"][0]["text"], self.AR)  # both survive

    def test_missing_work_order_flags_instead_of_crashing(self):
        with mock.patch.object(sg.mantle_client, "converse", side_effect=self._fixer(
                [self._audit(self._leak("S2Q1", "S1Q1", self.AR_SPAN))], {})):
            out, _, _ = sg.fix_answer_leaks(self._paper(), [], "10", "Social Science")
        self.assertTrue(out["Section B"]["_answer_leaks"])
        self.assertEqual([len(s["questions"]) for s in out.values()], [2, 2, 1])

    def test_surviving_leak_is_caught_by_the_single_re_audit(self):
        new_sa = {"type": "SA", "marks": 2, "text": "A fresh 2-mark question about the press.",
                  "answer_explanation": "Key points."}
        with mock.patch.object(sg.mantle_client, "converse", side_effect=self._fixer(
                [self._audit(self._leak("S2Q2", "S3P", self.PASSAGE_SPAN)),
                 self._audit(self._leak("S2Q1", "S1Q1", self.AR_SPAN))], new_sa)):
            out, _, _ = sg.fix_answer_leaks(self._paper(), self._wos(), "10", "Social Science")
        self.assertTrue(any("still answered by" in x
                            for x in out["Section A"]["_answer_leaks"]))

    def test_regen_cap_is_honoured_and_the_excess_is_flagged(self):
        new_sa = {"type": "SA", "marks": 2, "text": "A fresh 2-mark question about the press.",
                  "answer_explanation": "Key points."}
        audit = self._audit(self._leak("S2Q1", "S3P", self.PASSAGE_SPAN),
                            self._leak("S2Q2", "S3P", self.PASSAGE_SPAN))
        with mock.patch.object(sg.mantle_client, "converse",
                               side_effect=self._fixer([audit], new_sa)):
            with mock.patch.object(sg, "ANSWER_LEAK_MAX_REGENS", 1):
                out, _, _ = sg.fix_answer_leaks(
                    self._paper(), self._wos(), "10", "Social Science")
        self.assertEqual(len(out["Section B"]["_answer_leaks"]), 1)

    def test_clean_paper_costs_exactly_one_call_and_changes_nothing(self):
        calls = []

        def fake(**kw):
            calls.append(kw["model_id"])
            return ('{"leaks": []}', 5, 5)

        paper = self._paper()
        with mock.patch.object(sg.mantle_client, "converse", side_effect=fake):
            out, _, _ = sg.fix_answer_leaks(paper, self._wos(), "10", "Social Science")
        self.assertEqual(len(calls), 1)              # no re-audit when nothing was replaced
        self.assertEqual(out["Section A"]["questions"][0]["text"], self.AR)
        self.assertNotIn("_answer_leaks", out["Section A"])

    def test_leak_warnings_reach_the_final_audit(self):
        with mock.patch.object(sg.mantle_client, "converse", return_value=(json.dumps({
                "quality_score": 6, "ready_to_issue": "needs-minor-fix", "top_issues": [],
                "verdict": "v"}), 1, 1)):
            report = sg.run_final_paper_audit(
                {"S1": {"questions": [{"qnum": 1, "marks": 1}],
                        "_answer_leaks": ["S2Q1 answered by S1Q1: assertion states the answer"]}},
                "10", "Social Science", "Medium")
        self.assertEqual(report["total_warnings"], 1)


class MantleObservabilityTest(SimpleTestCase):
    """Celery logs must say which model is doing what, on which key, and how far a run got.

    The load-bearing test here is the first one: an API key must NEVER reach the log. Keys are
    identified by position plus a 4-hex SHA-256 fingerprint, which is enough to tell which key is
    throttled or dead and to correlate that across lines, but useless to anyone who reads the log
    or ships it to support.
    """

    K1 = "sk-mantle-NEVERLOGGED-key-one-abcdef123456"
    K2 = "sk-mantle-NEVERLOGGED-key-two-987654fedcba"

    def setUp(self):
        self.env = mock.patch.dict(
            __import__("os").environ, {"MANTLE_API_KEYS": f"{self.K1},{self.K2}"})
        self.env.start()
        mc.reset_run_stats()

    def tearDown(self):
        self.env.stop()
        mc.reset_run_stats()

    # ── plumbing ──────────────────────────────────────────────────────────────────────
    class _Resp:
        def __init__(self, status, payload):
            self.status_code, self._p = status, payload

        def raise_for_status(self):
            if self.status_code >= 400:
                raise mc.requests.exceptions.HTTPError(response=self)

        def json(self):
            return self._p

    @classmethod
    def _ok(cls, out_tokens=120, in_tokens=3120):
        return cls._Resp(200, {"choices": [{"message": {"content": '{"ok": true}'}}],
                               "usage": {"prompt_tokens": in_tokens,
                                         "completion_tokens": out_tokens}})

    def _run(self, responses, **kw):
        """Drive converse() against a scripted list of responses/exceptions, capturing stdout."""
        import contextlib, io
        queue = list(responses)

        def post(url, headers=None, json=None, timeout=None):
            # The key must be on the wire...
            self.assertTrue(headers["Authorization"].startswith("Bearer sk-mantle-"))
            r = queue.pop(0) if queue else self._ok()
            if isinstance(r, Exception):
                raise r
            return r

        buf = io.StringIO()
        with mock.patch.object(mc.requests, "post", side_effect=post):
            with contextlib.redirect_stdout(buf):
                try:
                    mc.converse(model_id=kw.pop("model_id", "test-model"),
                                prompt=kw.pop("prompt", "p" * 500), **kw)
                except Exception:
                    pass
        return buf.getvalue()

    # ── the invariant that must never regress ─────────────────────────────────────────
    def test_no_api_key_is_ever_printed(self):
        import re
        out = "".join((
            self._run([self._ok()], stage="ok-path"),
            self._run([self._Resp(401, {}), self._ok()], stage="dead-key"),
            self._run([self._Resp(429, {}), self._ok()], stage="throttled"),
            self._run([ConnectionError("TLS handshake timed out")], retries=1, stage="dead"),
        ))
        self.assertNotIn(self.K1, out)
        self.assertNotIn(self.K2, out)
        self.assertNotIn("NEVERLOGGED", out)
        # ...and no key-shaped substring survives anywhere in the log
        self.assertIsNone(re.search(r"sk-mantle-\w{4,}", out))
        # but the key is still identifiable
        self.assertRegex(out, r"key=[12]/2:[0-9a-f]{4}")

    def test_both_keys_get_distinct_fingerprints(self):
        import re
        out = self._run([self._Resp(401, {}), self._ok()], stage="s")
        self.assertEqual(len(set(re.findall(r"key=([12]/2:[0-9a-f]{4})", out))), 2)

    def test_keys_summary_never_shows_a_key(self):
        s = mc.keys_summary()
        self.assertIn("2 key(s) rotating", s)
        self.assertNotIn("NEVERLOGGED", s)
        self.assertRegex(s, r"1/2:[0-9a-f]{4}")

    def test_keys_summary_says_so_when_none_configured(self):
        with mock.patch.dict(__import__("os").environ, {"MANTLE_API_KEYS": ""}, clear=False):
            with mock.patch.object(mc, "_get_keys", return_value=[]):
                self.assertIn("NO KEYS CONFIGURED", mc.keys_summary())

    # ── "what is going on currently" ──────────────────────────────────────────────────
    def test_start_is_logged_before_the_request_so_a_hang_is_visible(self):
        # A stalled call used to produce no log line at all until it finally failed.
        out = self._run([ConnectionError("stalled")], retries=1, stage="hang")
        self.assertIn("[Mantle] START", out)
        self.assertIn("stage=hang", out)
        self.assertIn("[Mantle] FAIL", out)
        self.assertLess(out.index("START"), out.index("FAIL"))

    def test_retry_failover_and_failure_are_each_labelled(self):
        self.assertIn("RETRY", self._run([self._Resp(429, {}), self._ok()], stage="s"))
        self.assertIn("KEYDEAD", self._run([self._Resp(401, {}), self._ok()], stage="s"))
        self.assertIn("FAIL", self._run([self._Resp(500, {})], retries=1, stage="s"))

    def test_truncation_is_flagged_when_output_hits_the_cap(self):
        self.assertIn("TRUNCATED?", self._run(
            [self._ok(out_tokens=900)], max_tokens=900, stage="s"))
        self.assertNotIn("TRUNCATED?", self._run(
            [self._ok(out_tokens=100)], max_tokens=900, stage="s"))

    # ── stage labels ──────────────────────────────────────────────────────────────────
    def test_stage_labels_compose_and_carry_no_whitespace(self):
        # Section names and pipeline titles contain spaces; an unquoted space inside a key=value
        # field makes the whole line unsplittable, so labels are slugged.
        with mc.stage("Section D — Literature"):
            with mc.stage("cbq"):
                out = self._run([self._ok()], stage="v6-cbq-passage")
        label = out.split("stage=")[1].split(" ")[0]
        self.assertEqual(label, "Section_D_—_Literature/cbq/v6-cbq-passage")
        self.assertNotIn(" ", label)

    def test_stage_argument_alone_is_not_placeholder_prefixed(self):
        self.assertEqual(mc.current_stage(), "-")
        out = self._run([self._ok()], stage="v10-final")
        self.assertIn("stage=v10-final ", out)
        self.assertNotIn("stage=-/", out)

    def test_stage_stack_unwinds_even_when_the_body_raises(self):
        with self.assertRaises(ValueError):
            with mc.stage("a"):
                raise ValueError("boom")
        self.assertEqual(mc.current_stage(), "-")

    def test_stage_is_thread_local_so_parallel_sections_never_mix(self):
        # Sections are generated on a ThreadPoolExecutor; a shared global would scramble labels.
        from concurrent.futures import ThreadPoolExecutor
        import re
        seen = {}

        def worker(name):
            with mc.stage(name):
                seen[name] = re.findall(r"stage=(\S+)", self._run([self._ok()], stage="gen"))

        with ThreadPoolExecutor(max_workers=3) as ex:
            list(ex.map(worker, ["Section A", "Section B", "Section C"]))
        for name, labels in seen.items():
            self.assertEqual(set(labels), {name.replace(" ", "_") + "/gen"})

    # ── per-paper accounting ──────────────────────────────────────────────────────────
    def test_run_stats_tally_by_model_key_and_stage(self):
        self._run([self._ok(out_tokens=500)], model_id="model-x", stage="a")
        self._run([self._ok(out_tokens=700)], model_id="model-y", stage="b")
        s = mc.run_stats()
        self.assertEqual(s["calls"], 2)
        self.assertEqual(s["out"], 1200)
        self.assertEqual(set(s["by_model"]), {"model-x", "model-y"})
        self.assertEqual(set(s["by_stage"]), {"a", "b"})
        self.assertEqual(len(s["by_key"]), 2)          # rotation actually spread the load
        self.assertFalse(any(" " in k for k in s["by_stage"]))

    def test_errors_are_attributed_to_the_key_that_caused_them(self):
        self._run([self._Resp(401, {}), self._ok()], stage="s")
        by_key = mc.run_stats()["by_key"]
        self.assertEqual(sum(v["errors"] for v in by_key.values()), 1)

    def test_reset_run_stats_makes_the_tally_per_paper(self):
        self._run([self._ok()], stage="s")
        self.assertEqual(mc.run_stats()["calls"], 1)
        mc.reset_run_stats()
        s = mc.run_stats()
        self.assertEqual((s["calls"], s["in"], s["out"]), (0, 0, 0))
        self.assertEqual(s["by_model"], {})

    def test_run_stats_lines_are_printable_and_leak_nothing(self):
        self._run([self._Resp(429, {}), self._ok()], model_id="model-x", stage="a")
        lines = mc.run_stats_lines()
        self.assertTrue(any("calls=1" in x for x in lines))
        self.assertTrue(any("by model:" in x for x in lines))
        self.assertTrue(any("by key:" in x for x in lines))
        self.assertNotIn("NEVERLOGGED", " ".join(lines))

    def test_models_summary_names_all_three_roles(self):
        s = mc.models_summary()
        for part in ("gen=", "val=", "audit="):
            self.assertIn(part, s)


class PipelineStepLogTest(SimpleTestCase):
    """_StepLog turns the paper-assembly passes into a numbered progress trace with per-pass LLM
    cost, and makes each pass the mantle stage so every '[Mantle]' line inside is attributable."""

    def _capture(self, fn):
        import contextlib, io
        buf = io.StringIO()
        with contextlib.redirect_stdout(buf):
            fn()
        return buf.getvalue()

    def test_step_prints_start_and_done_with_numbering(self):
        step = sg._StepLog(total=3)

        def run():
            with step("First pass"):
                pass
            with step("Second pass", "detail here"):
                pass

        out = self._capture(run)
        self.assertIn("[Pipeline] 1/3 First pass — START", out)
        self.assertIn("[Pipeline] 2/3 Second pass — START detail here", out)
        self.assertIn("2/3 Second pass — done", out)
        self.assertIn("no llm", out)          # nothing called out, so say so explicitly

    def test_step_sets_the_mantle_stage_for_calls_inside_it(self):
        step = sg._StepLog(total=1)
        seen = []

        def run():
            with step("Answer-leak audit (V11)"):
                seen.append(mc.current_stage())

        self._capture(run)
        self.assertEqual(seen, ["Answer-leak_audit_(V11)"])
        self.assertEqual(mc.current_stage(), "-")     # popped afterwards

    def test_step_reports_a_raising_pass_and_re_raises(self):
        step = sg._StepLog(total=1)

        def run():
            with self.assertRaises(RuntimeError):
                with step("Explodes"):
                    raise RuntimeError("no work orders")

        out = self._capture(run)
        self.assertIn("Explodes — RAISED RuntimeError", out)
        self.assertIn("no work orders", out)
        self.assertEqual(mc.current_stage(), "-")     # stage still unwinds

    def test_step_attributes_llm_cost_to_the_pass(self):
        # run_stats is read once on entry and once on exit; the DELTA is what this pass spent.
        step = sg._StepLog(total=1)

        def run():
            with step("Costly pass"):
                pass

        with mock.patch.object(mc, "run_stats", side_effect=[
                {"calls": 5, "in": 1000, "out": 500},
                {"calls": 7, "in": 4000, "out": 1500}]):
            out = self._capture(run)
        self.assertIn("2 llm call(s) in=3.0k out=1.0k", out)


class VisionAndExternalCallLoggingTest(SimpleTestCase):
    """Image traffic used to be invisible. image_finder._vision_call was direct HTTP with its own
    unsynchronised `_key_idx` cursor: it logged only failures, kept its own key rotation (so two
    parallel sections could draw the SAME key), had no 401 failover, and its tokens were absent
    from every total. It now delegates to mantle_client.converse_vision, which shares the text
    path's key rotation, failover, logging and counters."""

    K1 = "sk-mantle-NEVERLOGGED-vision-one-abcdef123456"
    K2 = "sk-mantle-NEVERLOGGED-vision-two-987654fedcba"

    def setUp(self):
        self.env = mock.patch.dict(
            __import__("os").environ, {"MANTLE_API_KEYS": f"{self.K1},{self.K2}"})
        self.env.start()
        mc.reset_run_stats()

    def tearDown(self):
        self.env.stop()
        mc.reset_run_stats()

    class _Resp:
        def __init__(self, status, payload):
            self.status_code, self._p = status, payload

        def raise_for_status(self):
            if self.status_code >= 400:
                raise mc.requests.exceptions.HTTPError(response=self)

        def json(self):
            return self._p

    @classmethod
    def _ok(cls, out_tokens=260):
        return cls._Resp(200, {"choices": [{"message": {"content": "verified"}}],
                               "usage": {"prompt_tokens": 3120,
                                         "completion_tokens": out_tokens}})

    def _drive(self, responses, fn):
        import contextlib, io
        queue = list(responses)
        sent = []

        def post(url, headers=None, json=None, timeout=None):
            sent.append(json)
            r = queue.pop(0) if queue else self._ok()
            if isinstance(r, Exception):
                raise r
            return r

        buf = io.StringIO()
        with mock.patch.object(mc.requests, "post", side_effect=post):
            with contextlib.redirect_stdout(buf):
                try:
                    result = fn()
                except Exception:
                    result = None
        return buf.getvalue(), sent, result

    def test_vision_call_is_logged_on_success(self):
        out, sent, res = self._drive([self._ok()], lambda: mc.converse_vision(
            model_id="moonshotai.kimi-k2.5", prompt="Does the diagram show stomata?",
            image_bytes=b"\x89PNG" + b"z" * 40000, mime="image/png", stage="v9-verify"))
        self.assertIn("[Mantle] START", out)
        self.assertIn("[Mantle] OK", out)
        self.assertIn("kind=vision", out)
        self.assertIn("img=39kb", out)                     # the image size is visible
        self.assertIn("stage=v9-verify", out)
        self.assertRegex(out, r"key=[12]/2:[0-9a-f]{4}")
        self.assertEqual(res, ("verified", 3120, 260))

    def test_vision_never_logs_a_key(self):
        import re
        out = "".join((
            self._drive([self._ok()], lambda: mc.converse_vision(
                model_id="m", prompt="p", image_bytes=b"x" * 10, stage="s"))[0],
            self._drive([self._Resp(401, {}), self._ok()], lambda: mc.converse_vision(
                model_id="m", prompt="p", image_bytes=b"x" * 10, stage="s"))[0],
            self._drive([ConnectionError("down")], lambda: mc.converse_vision(
                model_id="m", prompt="p", image_bytes=b"x" * 10, retries=1, stage="s"))[0],
        ))
        self.assertNotIn("NEVERLOGGED", out)
        self.assertIsNone(re.search(r"sk-mantle-\w{4,}", out))

    def test_vision_counts_toward_the_paper_totals(self):
        self._drive([self._ok(out_tokens=260)], lambda: mc.converse_vision(
            model_id="moonshotai.kimi-k2.5", prompt="p", image_bytes=b"x" * 10, stage="s"))
        s = mc.run_stats()
        self.assertEqual((s["calls"], s["out"]), (1, 260))
        self.assertIn("moonshotai.kimi-k2.5", s["by_model"])
        self.assertEqual(len(s["by_key"]), 1)

    def test_vision_fails_over_on_a_dead_key(self):
        # The old _vision_call had NO failover — a 401 just failed the image check outright.
        import re
        out, _, res = self._drive([self._Resp(401, {}), self._ok()], lambda: mc.converse_vision(
            model_id="m", prompt="p", image_bytes=b"x" * 10, stage="s"))
        self.assertIn("KEYDEAD", out)
        self.assertIn("[Mantle] OK", out)
        self.assertEqual(len(set(re.findall(r"key=([12]/2:[0-9a-f]{4})", out))), 2)
        self.assertIsNotNone(res)

    def test_vision_sends_a_data_uri_image_part(self):
        _, sent, _ = self._drive([self._ok()], lambda: mc.converse_vision(
            model_id="m", prompt="describe", image_bytes=b"\x89PNG123", mime="image/png",
            stage="s"))
        content = sent[0]["messages"][0]["content"]
        self.assertEqual(content[0], {"type": "text", "text": "describe"})
        self.assertTrue(content[1]["image_url"]["url"].startswith("data:image/png;base64,"))

    def test_image_finder_vision_call_delegates_and_swallows_errors(self):
        from core import image_finder
        with mock.patch.object(mc, "converse_vision",
                               return_value=("  scored 9  ", 10, 5)) as m:
            self.assertEqual(
                image_finder._vision_call("prompt", b"img", "image/png", max_tokens=200,
                                          stage="v9-wikimedia-score"),
                "scored 9")
        self.assertEqual(m.call_args.kwargs["model_id"], image_finder.VISION_MODEL)
        self.assertEqual(m.call_args.kwargs["stage"], "v9-wikimedia-score")
        # A failed image check must degrade the question, never fail the paper.
        with mock.patch.object(mc, "converse_vision", side_effect=RuntimeError("endpoint down")):
            self.assertEqual(image_finder._vision_call("p", b"i", "image/png"), "")

    def test_image_finder_no_longer_keeps_its_own_key_rotation(self):
        # The duplicate cursor was unlocked and failover-less; it must be gone, not just unused.
        from core import image_finder
        self.assertFalse(hasattr(image_finder, "_next_mantle_key"))
        self.assertFalse(hasattr(image_finder, "_key_idx"))

    # ── external_call: non-model HTTP ────────────────────────────────────────────────
    def test_external_call_logs_start_and_ok(self):
        import contextlib, io
        buf = io.StringIO()
        with contextlib.redirect_stdout(buf):
            with mc.stage("Section C"):
                with mc.external_call("pollinations:ideogram", "candidate=1/3"):
                    pass
        out = buf.getvalue()
        self.assertIn("[HTTP] START stage=Section_C target=pollinations:ideogram candidate=1/3",
                      out)
        self.assertIn("[HTTP] OK    stage=Section_C target=pollinations:ideogram", out)

    def test_external_call_logs_failure_and_re_raises(self):
        import contextlib, io
        buf = io.StringIO()
        with contextlib.redirect_stdout(buf):
            with self.assertRaises(TimeoutError):
                with mc.external_call("wikimedia-search"):
                    raise TimeoutError("read timed out")
        self.assertIn("[HTTP] FAIL", buf.getvalue())
        self.assertIn("TimeoutError", buf.getvalue())

    def test_external_call_does_not_touch_the_llm_counters(self):
        # Pollinations/Wikimedia/Ollama are not completions — they carry no tokens.
        import contextlib, io
        with contextlib.redirect_stdout(io.StringIO()):
            with mc.external_call("ollama-embed:nomic-embed-text", "32 chunk(s)"):
                pass
        self.assertEqual(mc.run_stats()["calls"], 0)


class AnswerLeakNormalisationTest(SimpleTestCase):
    """Two bugs found by replaying the V11 audit against a real Class 11 Mathematics paper
    (paper 235). Both were invisible in English-language tests.

    1. `_leak_norm` kept only [a-z0-9]. Every Tamil/Hindi/Sanskrit span normalised to the EMPTY
       string, so every finding on those papers was dropped as "span too short" — V11 was a
       silent no-op for a large part of the catalogue. On Maths it stripped Greek letters and
       superscripts, collapsing 'cot θ = cos θ / sin θ' to 'cot cos sin'.
    2. `_leak_prefilter` emitted hints below the floor the audit itself applies, so it pointed the
       auditor at shared notation. On the live paper BOTH hints were false positives.
    """

    def test_norm_preserves_non_latin_scripts(self):
        for s in ("இலக்கணம் பாடம் ஒன்று விடை", "खंड-ख व्याकरण संधि विच्छेद कीजिए"):
            self.assertTrue(sg._leak_norm(s), f"{s!r} normalised to empty")
            # combining marks kept, so words are not fragmented into syllable pieces
            self.assertEqual(len(sg._leak_norm(s).split()), len(s.replace("-", " ").split()))

    def test_norm_preserves_greek_and_superscripts(self):
        self.assertEqual(sg._leak_norm("cot θ = cos θ / sin θ"), "cot θ cos θ sin θ")
        self.assertEqual(sg._leak_norm("sin²x + cos²x = 1"), "sin²x cos²x 1")

    def test_span_verification_works_on_a_tamil_paper(self):
        span = "இந்தப் பாடலின் கருத்தை விளக்குக என்று கேட்கப்பட்டது"
        self.assertIn(sg._leak_norm(span), sg._leak_norm("முன்னுரை " + span + " முடிவு"))
        self.assertNotIn(sg._leak_norm(span), sg._leak_norm("வேறு தொடர்பில்லாத வாக்கியம்"))
        # and it now clears the word floor instead of being dropped as empty
        self.assertGreaterEqual(len(sg._leak_norm(span).split()), sg._LEAK_MIN_SPAN_WORDS)

    def test_norm_still_ignores_punctuation_and_dashes(self):
        self.assertIn(sg._leak_norm("to control the vernacular press, and gave the government"),
                      sg._leak_norm("passed to control the vernacular press and gave the "
                                    "government power"))
        self.assertEqual(sg._leak_norm("Section A — Multiple Choice"), "section a multiple choice")

    def test_prefilter_ignores_shared_mathematical_notation(self):
        # Verbatim from paper 235: the answer to a 'solve 2sin²x + sin x - 1 <= 0' question and an
        # MCQ on 'sin x = -1/2 in [0, 2π)' share only notation. Not a leak.
        inv = sg._leak_inventory({
            "Section B": {"questions": [
                {"qnum": 1, "type": "SA", "marks": 2,
                 "text": "Solve the inequality: 2 sin² x + sin x - 1 ≤ 0, where 0 ≤ x < 2π.",
                 "answer_explanation": "Let t = sin x. Then -1 ≤ sin x ≤ 1/2, so the solution "
                                       "set is sin x = -1/2 in [0, 2π) together with x = 3π/2."}]},
            "Section A": {"questions": [
                {"qnum": 2, "type": "MCQ", "marks": 1,
                 "text": "The principal solution of the equation sin x = -1/2 in [0, 2π) are:",
                 "options": {"a": "π/6, 5π/6", "b": "5π/6, 7π/6", "c": "π/6, 11π/6",
                             "d": "7π/6, 11π/6"}, "answer": "d",
                 "answer_explanation": "7π/6 and 11π/6."}]},
        })
        self.assertEqual(sg._leak_prefilter(inv), [])

    def test_prefilter_never_emits_a_hint_the_audit_would_reject(self):
        # Any hint it does emit must clear the audit's own span floor, or it is pointing the
        # auditor at something the verifier will throw away.
        inv = sg._leak_inventory({
            "S1": {"questions": [
                {"qnum": 1, "type": "SA", "marks": 2,
                 "text": "State the purpose of the Vernacular Press Act of 1878.",
                 "answer_explanation": "It was passed to control the vernacular press and gave "
                                       "the government power to censor Indian-language papers."}]},
            "S2": {"questions": [
                {"qnum": 2, "type": "MCQ", "marks": 1,
                 "text": "The Act was passed to control the vernacular press and gave the "
                         "government power to censor reports. In which year?",
                 "options": {"a": "1878", "b": "1879", "c": "1880", "d": "1881"}, "answer": "a",
                 "answer_explanation": "1878."}]},
        })
        hits = sg._leak_prefilter(inv)
        self.assertTrue(hits, "a genuine verbatim leak must still be hinted")
        for _v, _l, span in hits:
            self.assertGreaterEqual(len(sg._leak_norm(span).split()), sg._LEAK_MIN_SPAN_WORDS)


class ChapterWeightScopingTest(SimpleTestCase):
    """Chapter WEIGHTS decide both how many questions a chapter gets and how much of the
    retrieval budget it gets, so a bad weight starves a chapter twice over.

    Found by replaying paper 235 (Class 11 Mathematics, Trigonometric Functions + Linear
    Inequalities) against temp.log. The plan was Trig x19 / LinearIneq x2 out of 21 questions,
    and Linear Inequalities received 900 of 7618 context chars — the teacher picked two
    chapters and got one. Three separate defects combined to produce it.
    """

    def test_key_must_be_contained_in_the_chapter_not_the_reverse(self):
        # THE paper-235 bug: the two-way substring test let the SHORTER chapter name inherit a
        # longer key's weight, so Class 11's 'Trigonometric Functions' collected the weight of
        # Class 12's 'Inverse Trigonometric Functions' (8) and outranked its co-chapter 8:1.
        self.assertEqual(sg._chapter_weight("Mathematics", "Trigonometric Functions", "11"), 1)
        self.assertEqual(sg._chapter_weight("Mathematics", "Trigonometric Functions"), 1)
        self.assertEqual(
            sg._chapter_weight("Mathematics", "Inverse Trigonometric Functions", "12"), 8)

    def test_catalog_key_inside_a_longer_chapter_name_still_matches(self):
        # The table is deliberately written as short fragments of the real NCERT titles, so
        # key-inside-chapter is the direction that must keep working.
        self.assertEqual(sg._chapter_weight("Mathematics", "Matrices and Determinants", "12"), 10)
        self.assertEqual(
            sg._chapter_weight("Physics", "Ray Optics and Optical Instruments", "12"), 18)
        self.assertEqual(
            sg._chapter_weight("Science", "Chemical Reactions and Equations", "10"), 10)

    def test_longest_matching_key_wins(self):
        # Dict order used to decide: 'Applications of Integrals' hit the key 'Integrals' (8)
        # before reaching its own entry (5).
        self.assertEqual(sg._chapter_weight("Mathematics", "Applications of Integrals", "12"), 5)
        self.assertEqual(sg._chapter_weight("Mathematics", "Integrals", "12"), 8)

    def test_weights_are_scoped_to_the_class_the_table_was_compiled_for(self):
        self.assertEqual(sg._chapter_weight("Physics", "Current Electricity", "12"), 17)
        self.assertEqual(sg._chapter_weight("Physics", "Current Electricity", "11"), 1)
        # The Science table covers both 9 and 10; an absent class is unscoped, not blocked.
        self.assertEqual(sg._chapter_weight("Science", "Light", "9"), 8)
        self.assertEqual(sg._chapter_weight("Science", "Light"), 8)

    def test_unlisted_chapter_inherits_the_mean_not_1(self):
        # A weight of 1 beside CBSE weights of 3-18 is not a light weight, it is a penalty of
        # up to 18x for not being in the catalog — applied to question slots AND context chars.
        w = sg._chapter_weights("Physics", ["Electromagnetic Waves", "Ray Optics", "Custom"], "12")
        self.assertEqual(w["Custom"], 11.0)                       # mean of 4 and 18
        self.assertEqual((w["Electromagnetic Waves"], w["Ray Optics"]), (4.0, 18.0))

    def test_no_weightage_data_for_the_class_means_uniform(self):
        w = sg._chapter_weights("Mathematics",
                                ["Trigonometric Functions", "Linear Inequalities"], "11")
        self.assertEqual(set(w.values()), {1.0})

    def test_class_key_reads_any_label_form(self):
        for label in ("11", "Class 11", " class-11 ", "11th"):
            self.assertEqual(sg._class_key(label), "11")
        self.assertEqual(sg._class_key(""), "")       # no number -> unscoped, not class ''
        self.assertEqual(sg._class_key(None), "")


class Paper235ChapterBalanceTest(SimpleTestCase):
    """End-to-end replay of the allocation that produced the 'questions only come from one part
    of the textbook' complaints, using paper 235's real section counts from temp.log."""

    SECTIONS = [10, 4, 3, 2, 2]                       # 21 questions
    CHAPTERS = ["Trigonometric Functions", "Linear Inequalities"]

    def _plan(self):
        covered, totals = {}, collections.Counter()
        for n in self.SECTIONS:
            totals.update(sg._allocate_chapters_to_slots(
                self.CHAPTERS, n, "Mathematics", covered, "11"))
        return totals

    def test_two_chapter_paper_is_split_evenly(self):
        totals = self._plan()
        self.assertEqual(sum(totals.values()), 21)
        self.assertEqual(len(totals), 2)
        # Was 19:2. Neither chapter may fall below ~80% of the other.
        self.assertGreaterEqual(min(totals.values()) / max(totals.values()), 0.8)

    def test_full_portion_paper_still_follows_cbse_weightage(self):
        # The fix must not flatten a real board-pattern paper into uniform coverage.
        chapters = ["Electric Charges and Fields", "Current Electricity",
                    "Moving Charges and Magnetism", "Electromagnetic Waves",
                    "Ray Optics and Optical Instruments", "Semiconductor Electronics"]
        c = collections.Counter(
            sg._allocate_chapters_to_slots(chapters, 30, "Physics", {}, "12"))
        self.assertEqual(len(c), len(chapters))                    # every chapter covered
        self.assertGreater(c["Ray Optics and Optical Instruments"],  # 18m unit
                           c["Electromagnetic Waves"])              # 4m unit


class ContextSpreadOrderTest(SimpleTestCase):
    """_spread_order: the prompt context must sample the WHOLE chapter, not the similarity
    cluster. Adjacent textbook chunks are near-identical in embedding space, so top-k retrieval
    returns a contiguous run — paper 235 sent every section the identical 7618-char block, so
    all 21 questions came from ~8 adjacent excerpts."""

    # Similarity order: positions 0-7 are one cluster, the rest is the remaining chapter.
    CANDS = [(1, i, f"chunk{i}") for i in
             [3, 4, 5, 2, 6, 1, 0, 7, 31, 30, 44, 52, 12, 19, 25, 38]]

    def test_prefix_spans_the_chapter(self):
        first4 = [c[1] for c in sg._spread_order(self.CANDS)[:4]]
        # The similarity top-4 spanned chunk idx 2..5; the spread prefix must reach far wider.
        self.assertGreater(max(first4) - min(first4), 40)

    def test_most_relevant_chunk_is_still_first(self):
        self.assertEqual(sg._spread_order(self.CANDS)[0], self.CANDS[0])

    def test_nothing_is_dropped_or_duplicated(self):
        out = sg._spread_order(self.CANDS)
        self.assertEqual(sorted(out), sorted(self.CANDS))

    def test_deterministic(self):
        self.assertEqual(sg._spread_order(self.CANDS), sg._spread_order(list(self.CANDS)))

    def test_degenerate_inputs_pass_through(self):
        self.assertEqual(sg._spread_order([]), [])
        self.assertEqual(sg._spread_order([(1, 5, "a")]), [(1, 5, "a")])
        two = [(1, 5, "a"), (1, 9, "b")]
        self.assertEqual(sg._spread_order(two), two)

    def test_positions_from_different_materials_do_not_collide(self):
        # Two books on the same chapter both numbered from 0 — ranking by (material, index)
        # keeps them apart instead of treating m1#0 and m2#0 as the same place.
        cands = [(1, 0, "a"), (1, 1, "b"), (2, 0, "c"), (2, 1, "d")]
        out = sg._spread_order(cands)
        self.assertEqual(out[0], cands[0])
        self.assertEqual(out[1][0], 2)             # jumps to the other material, not the neighbour
        self.assertEqual(sorted(out), sorted(cands))
