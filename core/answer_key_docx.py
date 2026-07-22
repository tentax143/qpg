"""Render a stored answer key (AnswerKey.data JSON) into a teacher-facing DOCX.

Rendered on demand at download time from AnswerKey.data — there is no stored file,
so the document can never drift out of sync with the JSON. Reuses the paper
renderer's code-built header (_build_header) and script-font handling so
Tamil/Hindi/Sanskrit keys print with the right complex-script font.
"""

import io
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .generator import (
    _build_header,
    _header_title_line,
    _pick_script_font,
    apply_tamil_document_styles,
    set_tamil_font,
)


def _fmt_marks(value):
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value or "")
    text = str(int(number)) if number.is_integer() else f"{number:g}"
    return text


def _iter_key_text(key_data):
    """Yield every displayable string in the key — used for script-font detection."""
    for section in key_data.get("sections", []):
        yield section.get("name") or ""
        for question in section.get("questions", []):
            yield question.get("text") or ""
            for answer in question.get("answers", []):
                yield answer.get("question") or ""
                yield answer.get("answer") or ""
                for item in answer.get("marking_scheme", []):
                    yield item.get("point") or ""


class _KeyWriter:
    """Small helper that applies consistent typography + the script font to every run."""

    def __init__(self, doc, script_font):
        self.doc = doc
        self.script_font = script_font

    def run(self, paragraph, text, size=11, bold=False, italic=False, color=None):
        r = paragraph.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = RGBColor(*color)
        if self.script_font:
            set_tamil_font(r, self.script_font)
        else:
            r.font.name = 'Times New Roman'
        return r

    def para(self, indent=None, space_before=0, space_after=2):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        if indent:
            p.paragraph_format.left_indent = Inches(indent)
        return p

    def section_header(self, title):
        p = self.para(space_before=10, space_after=6)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(p, title, size=13, bold=True)
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bdr = OxmlElement('w:bottom')
        bdr.set(qn('w:val'), 'single')
        bdr.set(qn('w:sz'), '6')
        bdr.set(qn('w:space'), '4')
        bdr.set(qn('w:color'), '555555')
        pBdr.append(bdr)
        pPr.append(pBdr)


def _render_answer(writer, answer, question_options, multiple):
    """One answer target: label, correct option, model answer, marking scheme, insight."""
    label = str(answer.get("label") or "").strip()
    if multiple and label and label.lower() != "answer":
        p = writer.para(indent=0.25, space_before=4)
        writer.run(p, f"{label}  ", size=11, bold=True)
        marks = _fmt_marks(answer.get("marks"))
        if marks:
            writer.run(p, f"[{marks} mark{'s' if marks != '1' else ''}]", size=10, italic=True,
                       color=(0x55, 0x55, 0x55))
        part_text = str(answer.get("question") or "").strip()
        if part_text:
            p2 = writer.para(indent=0.25)
            writer.run(p2, part_text, size=10, italic=True, color=(0x44, 0x44, 0x44))

    indent = 0.25
    correct_option = str(answer.get("correct_option") or "").strip()
    if correct_option:
        option_text = ""
        if isinstance(question_options, dict):
            option_text = str(question_options.get(correct_option)
                              or question_options.get(correct_option.upper()) or "").strip()
        p = writer.para(indent=indent)
        writer.run(p, "Correct option: ", size=11, bold=True)
        writer.run(p, f"({correct_option}) {option_text}".strip(), size=11)

    answer_text = str(answer.get("answer") or "").strip()
    if answer_text:
        p = writer.para(indent=indent)
        writer.run(p, "Answer: ", size=11, bold=True)
        writer.run(p, answer_text, size=11)

    scheme = answer.get("marking_scheme") or []
    if scheme:
        p = writer.para(indent=indent, space_before=2)
        writer.run(p, "Marking scheme", size=10, bold=True, color=(0x33, 0x33, 0x33))
        for item in scheme:
            point = str(item.get("point") or "").strip()
            if not point:
                continue
            marks = _fmt_marks(item.get("marks"))
            pp = writer.para(indent=indent + 0.2)
            writer.run(pp, f"•  {point}", size=10)
            writer.run(pp, f"  ({marks} mark{'s' if marks != '1' else ''})", size=10, italic=True,
                       color=(0x55, 0x55, 0x55))

    concept = answer.get("concept") or {}
    insight = answer.get("insight") or {}
    concept_name = str(concept.get("name") or "").strip()
    pieces = []
    if concept_name:
        chapter = str(concept.get("chapter") or "").strip()
        pieces.append(("Concept: ", f"{concept_name}{f' ({chapter})' if chapter else ''}"))
    misconception = str(insight.get("common_misconception") or "").strip()
    if misconception:
        pieces.append(("Common error: ", misconception))
    tip = str(insight.get("revision_tip") or "").strip()
    if tip:
        pieces.append(("Revision tip: ", tip))
    for heading, body in pieces:
        p = writer.para(indent=indent, space_after=1)
        writer.run(p, heading, size=9.5, bold=True, italic=True, color=(0x4B, 0x4B, 0x8A))
        writer.run(p, body, size=9.5, italic=True, color=(0x4B, 0x4B, 0x8A))

    if str(answer.get("confidence") or "").lower() == "low":
        p = writer.para(indent=indent, space_after=1)
        writer.run(p, "⚠ Low confidence — verify this answer against the textbook.",
                   size=9.5, italic=True, color=(0xB4, 0x5F, 0x06))


def render_answer_key_docx(paper, key_data, school_name=""):
    """Build the answer-key DOCX for `paper` from AnswerKey.data. Returns io.BytesIO."""
    doc = Document()

    script_font = _pick_script_font(paper.subject, [(None, t) for t in _iter_key_text(key_data)])
    if script_font:
        apply_tamil_document_styles(doc, script_font)

    section = doc.sections[0]
    section.top_margin = Inches(1.2)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    title_line = _header_title_line(paper.pattern.name if paper.pattern else "")
    marks_val = str(paper.pattern.total_marks) if paper.pattern else ""
    try:
        _build_header(section, paper.subject, paper.class_name, "", marks_val,
                      title_line, school_name or "", script_font=script_font)
    except Exception as e:
        print(f"[AnswerKey-DOCX] header build failed: {e}")

    # Keep the school header on the first page only (same OOXML trick as the paper).
    try:
        sectPr = section._sectPr
        if sectPr.find(qn('w:titlePg')) is None:
            sectPr.insert(0, OxmlElement('w:titlePg'))
        for hdr_ref in sectPr.findall(qn('w:headerReference')):
            if hdr_ref.get(qn('w:type')) == 'default':
                hdr_ref.set(qn('w:type'), 'first')
    except Exception as e:
        print(f"[AnswerKey-DOCX] first-page-only header setup failed: {e}")

    try:
        normal = doc.styles['Normal']
        normal.font.size = Pt(11)
        normal.paragraph_format.line_spacing = 1.15
        if not script_font:
            normal.font.name = 'Times New Roman'
    except Exception as e:
        print(f"[AnswerKey-DOCX] Normal-style setup failed: {e}")

    writer = _KeyWriter(doc, script_font)

    title = writer.para(space_before=4, space_after=2)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    writer.run(title, "ANSWER KEY", size=16, bold=True)
    subtitle = writer.para(space_after=8)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    writer.run(subtitle, "Teacher copy — with marking scheme. Not for circulation to students.",
               size=10, italic=True, color=(0x77, 0x77, 0x77))

    for key_section in key_data.get("sections", []):
        name = str(key_section.get("name") or "").strip()
        if name:
            writer.section_header(name)
        for question in key_section.get("questions", []):
            qnum = question.get("qnum")
            question_text = str(question.get("text") or "").strip()
            p = writer.para(space_before=8, space_after=3)
            prefix = f"Q{qnum}. " if qnum not in (None, "") else "Q. "
            writer.run(p, prefix, size=12, bold=True)
            writer.run(p, question_text, size=11.5, bold=True)
            marks = _fmt_marks(question.get("marks"))
            if marks:
                writer.run(p, f"  [{marks} mark{'s' if marks != '1' else ''}]", size=10,
                           italic=True, color=(0x55, 0x55, 0x55))
            answers = question.get("answers", [])
            multiple = len(answers) > 1
            for answer in answers:
                _render_answer(writer, answer, question.get("options"), multiple)
            for warning in question.get("warnings", []):
                p = writer.para(indent=0.25, space_after=1)
                writer.run(p, f"⚠ {warning}", size=9.5, italic=True, color=(0xB4, 0x5F, 0x06))

    errors = key_data.get("errors") or []
    if errors:
        writer.section_header("Questions without generated answers")
        for error in errors:
            p = writer.para(space_after=2)
            qnum = error.get("qnum")
            where = f"{error.get('section') or ''}{f' Q{qnum}' if qnum not in (None, '') else ''}".strip()
            writer.run(p, f"•  {where}: ", size=10, bold=True)
            writer.run(p, str(error.get("error") or "generation failed"), size=10,
                       color=(0x99, 0x33, 0x33))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
