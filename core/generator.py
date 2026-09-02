import os, re, json, random, time, traceback, threading
import hashlib

# Per-request storage for the last validated paper JSON.
# Eliminates the shared temp_clean.json file, which caused cross-user data
# contamination under concurrent Celery workers or gunicorn threads.
_request_state = threading.local()
import numpy as np
from datetime import datetime
from PyPDF2 import PdfReader, PdfWriter
from django.conf import settings
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import simpleSplit
from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from . import embeddings
from . import pattern_structure
from . import mantle_client
from .models import ExamBlueprint, BlueprintTemplate, GeneratedQuestion
from .blueprint_manager import BlueprintManager


# ------------------------------
# Model IDs
# ------------------------------
GEN_MODEL_ID = mantle_client.GEN_MODEL
VAL_MODEL_ID = mantle_client.VAL_MODEL



def _qt_str(qt) -> str:
    """Safely extract a lowercase type string from a question_type that may be str or dict."""
    if isinstance(qt, dict):
        return str(qt.get("type", "")).lower()
    return str(qt).lower()

# Strip embedded LOWERCASE option label prefix from values (e.g. LLM writes "(a) text").
# Lowercase-only: uppercase (A)-(D) in AR options ("A is true but R is false") must NOT be stripped.
_STRIP_OPT_PREFIX = re.compile(r'^\([a-d]\)\s*')

# Inline option list embedded in question text, e.g. "What is X? (a) foo (b) bar (c) baz".
# LOWERCASE-only and requires whitespace + content after the marker, so it can NEVER match
# the "(A):"/"(R):" markers inside Assertion-Reason statements (which are uppercase + colon).
_INLINE_OPTS_RE = re.compile(r'\s*\([a-d]\)\s+\S.*$', re.DOTALL)


def _is_ar_question(q, options) -> bool:
    """True if this question is Assertion-Reason — by declared subtype or AR option content."""
    if str(q.get("subtype", "")).strip().lower() == "assertion_reason":
        return True
    if isinstance(options, dict):
        blob = " ".join(str(v).lower() for v in options.values())
    elif isinstance(options, (list, tuple)):
        blob = " ".join(str(v).lower() for v in options)
    else:
        blob = ""
    return (
        "both a and r" in blob
        or ("a is true" in blob and "r is false" in blob)
        or ("a is false" in blob and "r is true" in blob)
    )


def _strip_inline_options(text, options, q):
    """
    Remove an option list embedded in the question text when options also exist separately.

    Skips Assertion-Reason questions entirely: their text legitimately contains "(A):" and
    "(R):", which the old IGNORECASE [a-dA-D] regex (with DOTALL .*$) matched and deleted —
    turning a full "Assertion (A): … Reason (R): …" into the bare word "Assertion".
    """
    if not (isinstance(options, list) and options and text):
        return text
    if _is_ar_question(q, options):
        return text
    return _INLINE_OPTS_RE.sub('', text).strip()


# ------------------------------
# Mantle helper (replaces call_bedrock)
# ------------------------------
def call_bedrock(
    prompt,
    model_ref,
    max_tokens=None,
    temperature=0.7,
    retries=5,
    model_source="aws",
):
    """
    Drop-in replacement for the old boto3-based call_bedrock.
    Uses Bedrock Converse API with Mantle bearer-token auth.
    Returns (text, input_tokens, output_tokens).
    """
    token_budget = max_tokens if max_tokens is not None else 4096

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    model_name = re.sub(r'[<>:"/\\|?*]', "_", model_ref.split("/")[-1])

    # Log prompt to file (keep existing debug behaviour)
    prompt_file = f"temp_prompt_{model_name}_{timestamp}.txt"
    with open(prompt_file, "w", encoding="utf-8") as f:
        f.write(f"Model: {model_ref}\nTimestamp: {timestamp}\n")
        f.write(f"Max Tokens: {token_budget}\nTemperature: {temperature}\n")
        f.write("=" * 50 + "\n\n")
        f.write(prompt)
    print(f"[Mantle] Prompt saved to: {prompt_file}")

    text, input_tokens, output_tokens = mantle_client.converse(
        model_id=model_ref,
        prompt=prompt,
        max_tokens=token_budget,
        temperature=temperature,
        retries=retries,
    )

    # Log response to file
    response_file = f"temp_response_{model_name}_{timestamp}.txt"
    with open(response_file, "w", encoding="utf-8") as f:
        f.write(f"Model: {model_ref}\nTimestamp: {timestamp}\n")
        f.write(f"Input Tokens: {input_tokens}\nOutput Tokens: {output_tokens}\n")
        f.write("=" * 50 + "\n\n")
        f.write(text)
    print(f"[Mantle] Response saved to: {response_file}")

    return text, input_tokens, output_tokens



# ------------------------------
# Fetch blueprint from DB
# ------------------------------
# DEPRECATED: Use BlueprintManager.normalize_blueprint instead
def convert_blueprint_to_dict(blueprint):
    """
    Convert blueprint from array format to dictionary format if needed.
    Input: {"sections": [{"name": "A", "title": "...", ...}]}
    Output: {"A": {"title": "...", ...}}
    """
    # Handle different blueprint formats
    if isinstance(blueprint, dict):
        if "sections" in blueprint:
            if isinstance(blueprint["sections"], list):
                # New blueprint format with sections array
                blueprint_dict = {}
                for section in blueprint["sections"]:
                    section_name = section.get("name", "")
                    blueprint_dict[section_name] = {
                        "title": section.get("title", ""),
                        "marks": section.get("marks", 0),
                        "question_types": section.get("question_types", []),
                        "subsections": section.get("subsections", {})  # Preserve subsections if they exist
                    }
                return blueprint_dict
            elif isinstance(blueprint["sections"], dict):
                # Old format where sections is a dict - just use it directly
                return blueprint["sections"]
            else:
                # sections key exists but is neither list nor dict
                print(f"[Blueprint-Convert] WARNING: 'sections' is neither list nor dict: {type(blueprint['sections'])}")
                return {}
        else:
            # No sections key - assume it's already in the right format
            # Check if it looks like a proper blueprint dict (has section keys like A, B, C)
            if any(key in ['A', 'B', 'C', 'D', 'E'] for key in blueprint.keys()):
                return blueprint
            else:
                print(f"[Blueprint-Convert] WARNING: Blueprint dict doesn't have expected structure")
                return blueprint
    elif isinstance(blueprint, list):
        # If blueprint is a list, try to convert it
        print(f"[Blueprint-Convert] WARNING: Blueprint is a list, attempting conversion")
        blueprint_dict = {}
        for idx, section in enumerate(blueprint):
            if isinstance(section, dict) and "name" in section:
                section_name = section.get("name", f"Section_{idx}")
                blueprint_dict[section_name] = {
                    "title": section.get("title", ""),
                    "marks": section.get("marks", 0),
                    "question_types": section.get("question_types", []),
                    "subsections": section.get("subsections", {})
                }
        return blueprint_dict if blueprint_dict else {}
    else:
        # Unknown format
        print(f"[Blueprint-Convert] ERROR: Blueprint is neither dict nor list: {type(blueprint)}")
        return {}

def get_blueprint(class_name, subject, section=None):
    """
    Get blueprint using the centralized BlueprintManager.
    This function is kept for backward compatibility but delegates to BlueprintManager.
    """
    return BlueprintManager.get_blueprint(class_name, subject, section)

def get_legacy_blueprint(class_name, subject, section=None):
    """Legacy blueprint system for backward compatibility"""
    normalized_subject = subject.strip().lower()
    
    # Legacy blueprint definitions
    legacy_blueprints = {
        "11": {
            "english core": {
                "sections": [
                    {"name": "A", "title": "Reading", "marks": 20, "question_types": ["unseen_passage", "case_based"]},
                    {"name": "B", "title": "Writing", "marks": 30, "question_types": ["grammar", "writing_tasks"]},
                    {"name": "C", "title": "Literature", "marks": 30, "question_types": ["extract_based", "short_answer", "long_answer"]},
                    {"name": "D", "title": "Flamingo", "marks": 20, "question_types": ["short_answer", "long_answer"]}
                ]
            },
            "biology": {
                "sections": [
                    {"name": "A", "title": "Multiple Choice", "marks": 20, "question_types": ["mcq"]},
                    {"name": "B", "title": "Assertion-Reason", "marks": 10, "question_types": ["assertion_reason"]},
                    {"name": "C", "title": "Very Short Answer", "marks": 20, "question_types": ["very_short_answer"]},
                    {"name": "D", "title": "Short Answer", "marks": 30, "question_types": ["short_answer"]},
                    {"name": "E", "title": "Long Answer", "marks": 20, "question_types": ["long_answer"]}
                ]
            }
        }
    }
    
    blueprint = legacy_blueprints.get(class_name, {}).get(normalized_subject)
    if blueprint:
        print(f"[Blueprint] ✅ Found legacy blueprint for {class_name} {subject}")
        return blueprint
    
    # Final fallback - raise error
    raise ValueError(f"No blueprint found for {class_name} {subject} (section: {section}). Please create a blueprint template or exam blueprint.")

# ------------------------------
# Universal Context Retrieval
# ------------------------------
def get_universal_context(class_name, subject, chapters, question_types=None, variation_offset=0, school_id=None):
    """
    Universal context retrieval system that works for all subjects and question types.
    Uses variation_offset to retrieve different context chunks each time for more variation.
    
    Args:
        class_name: Class (e.g., "11", "12")
        subject: Subject name (e.g., "English", "Biology")
        chapters: List of chapters to include
        question_types: List of question types to optimize context for
        variation_offset: Offset to vary context retrieval (default: random)
    
    Returns:
        dict: Context information with documents and metadata
    """
    print(f"[Context] Retrieving universal context for {class_name} {subject}")
    print(f"[Context] Chapters: {chapters}")
    print(f"[Context] Question types: {question_types}")
    
    # Use variation offset to get different context chunks
    if variation_offset == 0:
        variation_offset = random.randint(0, 10)  # Random offset for variation
    print(f"[Variation] Using context variation offset: {variation_offset}")
    
    all_contexts = []
    context_metadata = {
        "total_chapters": len(chapters),
        "question_types": question_types or [],
        "retrieval_timestamp": datetime.now().isoformat(),
        "variation_offset": variation_offset
    }
    
    # Build optimized queries based on question types - add variation to queries
    base_queries = [f"{subject} important NCERT concepts"]
    
    # Add variation to queries to get different context
    variation_terms = ["detailed", "examples", "applications", "concepts", "principles", "theories", "practical"]
    variation_term = random.choice(variation_terms)
    base_queries.append(f"{subject} {variation_term} NCERT content")
    
    if question_types:
        # Add subject-specific queries based on question types
        for qtype in question_types:
            ql = _qt_str(qtype)
            if ql in ["unseen_passage", "case_based"]:
                base_queries.append(f"{subject} reading comprehension passages")
                base_queries.append(f"{subject} case studies examples")
            elif ql in ["mcq", "assertion_reason"]:
                base_queries.append(f"{subject} multiple choice questions")
                base_queries.append(f"{subject} facts and concepts")
            elif ql in ["writing_tasks", "grammar"]:
                base_queries.append(f"{subject} writing skills")
                base_queries.append(f"{subject} grammar rules")
            elif ql in ["extract_based", "short_answer", "long_answer"]:
                base_queries.append(f"{subject} detailed explanations")
                base_queries.append(f"{subject} important topics")
            elif ql in ["numerical", "proof"]:
                base_queries.append(f"{subject} solved examples")
                base_queries.append(f"{subject} mathematical problems")
    
    # Remove duplicates while preserving order
    unique_queries = list(dict.fromkeys(base_queries))
    # Shuffle queries for variation
    random.shuffle(unique_queries)
    print(f"[Context] Using queries: {unique_queries}")
    
    # Retrieve context for each chapter and query combination
    for chapter in chapters:
        chapter_contexts = []
        
        for query in unique_queries:
            try:
                # Vary n_results slightly for variation
                n_results = 50 + variation_offset % 10
                results = embeddings.query(
                    class_name=class_name,
                    subject=subject,
                    unit=chapter,
                    query_text=query,
                    n_results=n_results,
                    school_id=school_id,
                )
                
                if results and "documents" in results:
                    for docs in results["documents"]:
                        chapter_contexts.extend(docs)
                        
            except Exception as e:
                print(f"[Context] ⚠️ Error retrieving context for {chapter} with query '{query}': {e}")
                continue
        
        # Deduplicate and limit chapter contexts - use offset to vary selection
        chapter_contexts = list(dict.fromkeys(chapter_contexts))  # Remove duplicates
        
        # Vary the context selection by using offset
        if len(chapter_contexts) > 100:
            # Start from different position based on offset
            start_idx = variation_offset % min(50, len(chapter_contexts) - 100)
            chapter_contexts = chapter_contexts[start_idx:start_idx+100]
        else:
            chapter_contexts = chapter_contexts[:100]
        
        all_contexts.extend(chapter_contexts)
        print(f"[Context] Retrieved {len(chapter_contexts)} contexts for chapter {chapter}")
    
    # Final deduplication and limiting - shuffle for more variation
    all_contexts = list(dict.fromkeys(all_contexts))  # Remove duplicates
    random.shuffle(all_contexts)  # Shuffle for variation
    all_contexts = all_contexts[:500]  # Global limit
    
    context_text = "\n".join(all_contexts)
    
    context_metadata.update({
        "total_contexts": len(all_contexts),
        "context_length": len(context_text),
        "queries_used": unique_queries
    })
    
    print(f"[Context] ✅ Retrieved {len(all_contexts)} total contexts ({len(context_text)} characters)")
    
    return {
        "context_text": context_text,
        "contexts": all_contexts,
        "metadata": context_metadata
    }

# ------------------------------
# Question Variation Helpers
# ------------------------------

def get_variation_temperature():
    """Get a slightly varied temperature between 0.75 and 0.85 for more variation"""
    return round(random.uniform(0.75, 0.85), 2)

def calculate_question_hash(question_text):
    """Calculate SHA256 hash of question text for quick duplicate detection"""
    cleaned = re.sub(r'\s+', ' ', question_text.strip().lower())
    return hashlib.sha256(cleaned.encode('utf-8')).hexdigest()

def check_question_similarity(new_question, class_name, subject, chapters, similarity_threshold=0.85):
    """
    Check if a question is too similar to previously generated questions.
    Uses both hash-based exact matching and embedding-based similarity.
    """
    # First, check for exact duplicates using hash
    question_hash = calculate_question_hash(new_question)
    existing_hash = GeneratedQuestion.objects.filter(
        class_name=class_name,
        subject=subject,
        question_hash=question_hash
    ).first()
    
    if existing_hash:
        print(f"[Variation] Exact duplicate detected: {new_question[:50]}...")
        return True, 1.0
    
    # Check for similar questions using embeddings
    try:
        # Get embeddings for the new question
        new_embedding_result = embeddings.query(
            class_name=class_name,
            subject=subject,
            unit="",
            query_text=new_question,
            n_results=1
        )
        
        if new_embedding_result and "embeddings" in new_embedding_result:
            new_embedding = new_embedding_result["embeddings"][0] if new_embedding_result["embeddings"] else None
        else:
            # Fallback: use simple text similarity
            return check_text_similarity(new_question, class_name, subject, chapters), None
        
        if not new_embedding:
            return False, 0.0
        
        # Compare with existing questions
        existing_questions = GeneratedQuestion.objects.filter(
            class_name=class_name,
            subject=subject,
            embedding__isnull=False
        )[:100]  # Limit to recent 100 for performance
        
        for existing in existing_questions:
            if existing.embedding:
                try:
                    # Calculate cosine similarity
                    similarity = cosine_similarity(new_embedding, existing.embedding)
                    if similarity >= similarity_threshold:
                        print(f"[Variation] Similar question detected (similarity: {similarity:.2f}): {new_question[:50]}...")
                        return True, similarity
                except Exception as e:
                    print(f"[Variation] Error calculating similarity: {e}")
                    continue
    except Exception as e:
        print(f"[Variation] Error in similarity check: {e}")
        # Fallback to text-based similarity
        return check_text_similarity(new_question, class_name, subject, chapters), None
    
    return False, 0.0

def cosine_similarity(vec1, vec2):
    """Calculate cosine similarity between two vectors"""
    try:
        vec1 = np.array(vec1)
        vec2 = np.array(vec2)
        dot_product = np.dot(vec1, vec2)
        norm1 = np.linalg.norm(vec1)
        norm2 = np.linalg.norm(vec2)
        if norm1 == 0 or norm2 == 0:
            return 0.0
        return dot_product / (norm1 * norm2)
    except Exception as e:
        print(f"[Variation] Error in cosine similarity calculation: {e}")
        return 0.0

def check_text_similarity(new_question, class_name, subject, chapters):
    """Fallback text-based similarity check"""
    # Simple word overlap check
    new_words = set(re.findall(r'\w+', new_question.lower()))
    
    existing_questions = GeneratedQuestion.objects.filter(
        class_name=class_name,
        subject=subject
    ).values_list('question_text', flat=True)[:50]
    
    for existing_text in existing_questions:
        existing_words = set(re.findall(r'\w+', existing_text.lower()))
        if len(new_words) > 0:
            overlap = len(new_words & existing_words) / len(new_words)
            if overlap > 0.7:  # 70% word overlap
                return True
    
    return False

def save_generated_question(question_text, class_name, subject, chapter, question_type, marks, paper_id=None):
    """Save a generated question to track it and avoid duplicates"""
    try:
        question_hash = calculate_question_hash(question_text)
        
        # Get embedding if possible
        embedding = None
        try:
            embedding_result = embeddings.query(
                class_name=class_name,
                subject=subject,
                unit=chapter or "",
                query_text=question_text,
                n_results=1
            )
            if embedding_result and "embeddings" in embedding_result:
                embedding = embedding_result["embeddings"][0] if embedding_result["embeddings"] else None
        except Exception as e:
            print(f"[Variation] Could not get embedding for question: {e}")
        
        GeneratedQuestion.objects.create(
            class_name=class_name,
            subject=subject,
            chapter=chapter,
            question_text=question_text,
            question_hash=question_hash,
            embedding=embedding,
            question_type=question_type or "",
            marks=marks,
            paper_id=paper_id
        )
    except Exception as e:
        print(f"[Variation] Error saving generated question: {e}")

def get_variation_instructions():
    """Get instructions to ensure question variation"""
    variation_hints = [
        "Generate questions from different angles and perspectives",
        "Use different examples and scenarios than typical questions",
        "Focus on different aspects of the topic",
        "Vary the complexity and approach",
        "Create unique questions that haven't been seen before",
        "Use creative problem-solving approaches",
        "Avoid repeating common question patterns",
        "Generate fresh, original questions each time"
    ]
    return random.sample(variation_hints, 3)  # Return 3 random hints

def get_question_type_instructions(question_types, subject):
    """
    Get specific instructions for different question types.
    
    Args:
        question_types: List of question types
        subject: Subject name
    
    Returns:
        str: Formatted instructions for the question types
    """
    if not question_types:
        return ""
    
    instructions = []
    
    for qtype in question_types:
        ql = _qt_str(qtype)
        if ql == "unseen_passage":
            instructions.append("• Unseen Passage: Include a reading comprehension passage with 5-6 questions based on the passage")
        elif ql == "case_based":
            instructions.append("• Case-based: Create scenario-based questions that test application of concepts")
        elif ql == "mcq":
            instructions.append("• Multiple Choice: Create 4-option MCQs with one correct answer and plausible distractors")
        elif ql == "assertion_reason":
            instructions.append("• Assertion-Reason: Create statements with assertion and reason, test logical relationship")
        elif ql == "grammar":
            instructions.append("• Grammar: Focus on grammatical concepts, sentence correction, and language usage")
        elif ql == "writing_tasks":
            instructions.append("• Writing Tasks: Include essay writing, letter writing, or other composition tasks")
        elif ql == "extract_based":
            instructions.append("• Extract-based: Provide literary extracts and ask questions based on them")
        elif ql == "short_answer":
            instructions.append("• Short Answer: Create questions requiring brief, focused responses (2-3 sentences)")
        elif ql == "long_answer":
            instructions.append("• Long Answer: Create questions requiring detailed, comprehensive responses")
        elif ql == "very_short_answer":
            instructions.append("• Very Short Answer: Create questions requiring one-word or one-sentence responses")
        elif ql == "numerical":
            instructions.append("• Numerical: Create problems requiring mathematical calculations and solutions")
        elif ql == "proof":
            instructions.append("• Proof: Create questions requiring mathematical proofs or logical reasoning")
        elif ql == "or_choice":
            instructions.append("• Either/Or: Provide alternative questions where students choose one to answer")
    
    return "\n".join(instructions) if instructions else ""

# ------------------------------
# JSON enforcement
# ------------------------------
def enforce_json(validated_json):
    print(f"[JSON-Validator] Starting JSON enforcement process...")
    print(f"[JSON-Validator] Input length: {len(validated_json)} characters")
    
    # Save raw response
    with open("temp_raw.json", "w", encoding="utf-8") as f:
        f.write(validated_json)
    print(f"[JSON-Validator] Raw response saved to temp_raw.json")

    # Clean up markdown code blocks
    if validated_json.strip().startswith("```"):
        print(f"[JSON-Validator] Removing markdown code blocks...")
        validated_json = re.sub(r"^```[a-zA-Z]*\n|\n```$", "", validated_json.strip())
        print(f"[JSON-Validator] After markdown cleanup: {len(validated_json)} characters")

    # Try to parse as-is first
    data = None
    try:
        data = json.loads(validated_json)
        print("[JSON-Validator] ✅ Successfully parsed JSON on first attempt")
        print(f"[JSON-Validator] Parsed JSON has {len(data)} top-level keys: {list(data.keys())}")
    except Exception as e:
        print(f"[JSON-Validator] ❌ First attempt failed: {e}")
        print(f"[JSON-Validator] Attempting to fix JSON issues...")
        
        # Try to fix common JSON issues
        fixed_json = fix_json_issues(validated_json)
        
        try:
            data = json.loads(fixed_json)
            print("[JSON-Validator] Successfully parsed JSON after fixing issues")
        except Exception as e2:
            print(f"[JSON-Validator] Fixed JSON still invalid: {e2}")
            
            # Try extracting JSON from text
            try:
                match = re.search(r"\{.*\}", fixed_json, re.S)
                if match:
                    candidate = match.group(0)
                    data = json.loads(candidate)
                    print("[JSON-Validator] Successfully extracted and parsed JSON")
            except Exception as e3:
                print(f"[JSON-Validator] JSON extraction failed: {e3}")

    # Save the processed JSON
    with open("temp_validated.json", "w", encoding="utf-8") as f:
        f.write(validated_json)

    if data:
        _request_state.paper_data = data
        return data

    raise ValueError("Validator did not return valid JSON. Check temp_raw.json and temp_validated.json")

def fix_json_issues(json_text):
    """Fix common JSON formatting issues"""
    print("[JSON-Fixer] Attempting to fix JSON issues...")
    
    # Remove any leading/trailing whitespace
    json_text = json_text.strip()
    
    # Fix common issues
    fixes_applied = []
    
    # 1. Fix missing closing brackets/braces
    open_braces = json_text.count('{')
    close_braces = json_text.count('}')
    open_brackets = json_text.count('[')
    close_brackets = json_text.count(']')
    
    if open_braces > close_braces:
        missing_braces = open_braces - close_braces
        json_text += '}' * missing_braces
        fixes_applied.append(f"Added {missing_braces} missing closing braces")
    
    if open_brackets > close_brackets:
        missing_brackets = open_brackets - close_brackets
        json_text += ']' * missing_brackets
        fixes_applied.append(f"Added {missing_brackets} missing closing brackets")
    
    # 2. Fix trailing commas
    json_text = re.sub(r',(\s*[}\]])', r'\1', json_text)
    if re.search(r',(\s*[}\]])', json_text):
        fixes_applied.append("Removed trailing commas")
    
    # 3. Fix unescaped quotes in strings
    # This is more complex, so we'll try a simpler approach first
    
    # 4. Fix incomplete strings at the end
    if json_text.endswith('"') and not json_text.endswith('""'):
        # Check if we have an incomplete string
        last_quote_pos = json_text.rfind('"', 0, -1)
        if last_quote_pos > 0:
            # Look for the opening quote
            before_last_quote = json_text[:last_quote_pos]
            if before_last_quote.count('"') % 2 == 1:  # Odd number of quotes means incomplete
                json_text = json_text[:-1]  # Remove the incomplete quote
                fixes_applied.append("Removed incomplete string quote")
    
    # 5. Fix incomplete objects/arrays
    # If the JSON ends abruptly, try to close it properly
    if not json_text.endswith(('}', ']', '"')):
        # Find the last complete structure
        lines = json_text.split('\n')
        for i in range(len(lines) - 1, -1, -1):
            line = lines[i].strip()
            if line and not line.endswith(','):
                # This might be the last complete line
                if line.endswith('"') or line.endswith('}') or line.endswith(']'):
                    # Close any open structures
                    if open_braces > close_braces:
                        json_text += '}' * (open_braces - close_braces)
                    if open_brackets > close_brackets:
                        json_text += ']' * (open_brackets - close_brackets)
                    fixes_applied.append("Closed incomplete structures")
                break
    
    # 6. Handle specific case from temp_raw.json - incomplete extract
    # If the JSON ends with an incomplete string or object, try to complete it
    if json_text.endswith('"') and '"' in json_text:
        # Check if we have an incomplete string at the end
        lines = json_text.split('\n')
        last_line = lines[-1].strip()
        
        # If the last line looks like an incomplete string, try to close it
        if last_line.count('"') == 1 and not last_line.endswith('",'):
            # This might be an incomplete string, try to close it properly
            if last_line.endswith('"'):
                json_text = json_text[:-1]  # Remove the incomplete quote
                # Add proper closing for the structure
                if open_braces > close_braces:
                    json_text += '}' * (open_braces - close_braces)
                if open_brackets > close_brackets:
                    json_text += ']' * (open_brackets - close_brackets)
                fixes_applied.append("Fixed incomplete string and closed structures")
    
    # 7. Fix incomplete extract objects (specific to your case)
    if '"extract":' in json_text and not json_text.endswith('}'):
        # Look for incomplete extract objects
        extract_pos = json_text.rfind('"extract":')
        if extract_pos > 0:
            # Find the end of the extract value
            after_extract = json_text[extract_pos + 10:]
            # Look for the next quote or end
            if '"' in after_extract:
                # Find the end of the extract string
                end_quote = after_extract.find('"', 1)
                if end_quote > 0:
                    # Check if we have proper closing after the extract
                    after_extract_value = after_extract[end_quote + 1:].strip()
                    if not after_extract_value.startswith(','):
                        # Add missing comma and closing
                        json_text = json_text[:extract_pos + 10 + end_quote + 1] + '",'
                        # Close any remaining structures
                        if open_braces > close_braces:
                            json_text += '}' * (open_braces - close_braces)
                        if open_brackets > close_brackets:
                            json_text += ']' * (open_brackets - close_brackets)
                        fixes_applied.append("Fixed incomplete extract object")
    
    if fixes_applied:
        print(f"[JSON-Fixer] Applied fixes: {', '.join(fixes_applied)}")
    else:
        print("[JSON-Fixer] No obvious fixes needed")
    
    return json_text


# ------------------------------
# PDF helper
# ------------------------------
def draw_wrapped(can, text, x, y, max_width, font="Helvetica", size=11, line_height=16):
    can.setFont(font, size)
    lines = simpleSplit(text, font, size, max_width)
    for line in lines:
        if y < 100:
            can.showPage()
            can.setFont(font, size)
            y = 760
        can.drawString(x, y, line)
        y -= line_height
    return y


# ------------------------------
# Flexible renderer for English
# ------------------------------
# ── Type-grouped rendering ────────────────────────────────────────────────────
# CBSE papers list questions grouped by type (all MCQs first, then VSA, SA, …), each
# group under its own "I. / II. / III." sub-heading. The generator (especially the
# single-prompt fallback) can emit questions out of type order and stamp the section
# AVERAGE marks (e.g. 1.9) on every question. The helpers below regroup a section's
# questions into canonical type order and restore the correct per-type marks from the
# blueprint's subsections — so an SA never renders inside the MCQ block and a 1-mark
# MCQ never shows "1.9 marks".
_TYPE_GROUP_ORDER = [
    ("mcq", "Multiple Choice Questions"),
    ("vsa", "Very Short Answer Questions"),
    ("sa",  "Short Answer Questions"),
    ("cbq", "Case-Based / Source-Based Questions"),
    ("la",  "Long Answer Questions"),
]
_ROMAN = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII"]


def _fmt_marks_num(v):
    """2.0 -> '2', 2.5 -> '2.5' — marks read as whole numbers on a printed paper."""
    f = _coerce_float(v, 0.0)
    return str(int(f)) if f == int(f) else f"{f:g}"


def _section_calc(sec_info):
    """"6 x 2 = 12" — the arithmetic a teacher writes beside a section heading.

    The multiplier is what the student ANSWERS, not what is printed: an attempt-N-of-M section
    prints eight 2-mark questions but is worth 6 x 2 = 12. Returns "" for a section whose marks
    are not a single product (mixed marks), where no such line is truthful.
    """
    mpq = _coerce_float(sec_info.get('marks_per_question'), 0.0)
    marks = _coerce_float(sec_info.get('marks'), 0.0)
    try:
        count = int(sec_info.get('attempt_count') or sec_info.get('questions_count') or 0)
    except (TypeError, ValueError):
        return ""
    if not count or mpq <= 0 or marks <= 0:
        return ""
    if abs(count * mpq - marks) > 0.5:
        return ""
    return f"{count} x {_fmt_marks_num(mpq)} = {_fmt_marks_num(marks)}"


def _coerce_float(v, default=0.0):
    try:
        return float(v)
    except (TypeError, ValueError):
        return default


def _question_category(q):
    """Bucket a question dict into a canonical type category (mcq/ar/vsa/sa/cbq/la)."""
    t = str(q.get("type", "")).lower()
    st = str(q.get("subtype", "")).lower()
    if "assertion" in st or "assertion" in t:
        return "ar"
    if st in ("source_based", "case_based", "image_based") or "cbq" in t or "case" in t or "source" in t:
        return "cbq"
    if "vsa" in t or "very short" in t:
        return "vsa"
    if t == "la" or "long" in t:
        return "la"
    if "mcq" in t or "multiple" in t or "objective" in t:
        return "mcq"
    if t == "sa" or "short" in t:
        return "sa"
    return "sa"


def _section_type_marks(sec_info):
    """{category: marks_each} from a blueprint section's subsections (compound papers).

    Only a subsection that states a real per-question figure contributes: an explicit
    marks_per_question, or a marks total alongside an explicit question count (AI-authored
    patterns spell these 'count' and a singular 'question_type'). A count-less marks total
    is NOT divided by an assumed 1 — that stamped the subsection TOTAL on every question
    of the type (5 × 2m short answers rendered as 5 × 10m). And when two subsections land
    on the same category with different values (1m extract questions and 2m short answers
    both bucket to 'sa'), the category is ambiguous — it is dropped so the generated
    per-question marks stand."""
    out = {}
    if not isinstance(sec_info, dict):
        return out
    ambiguous = set()
    for ss in (sec_info.get("subsections") or []):
        if not isinstance(ss, dict):
            continue
        qts = ss.get("question_types") or ss.get("question_type") or [ss.get("name", "")]
        if isinstance(qts, str):
            qts = [qts]
        cat = _question_category({"type": qts[0] if qts else ss.get("name", ""),
                                  "subtype": ss.get("name", "")})
        me = ss.get("marks_per_question")
        if me in (None, "", "varies"):
            cnt = _coerce_float(ss.get("questions_count") or ss.get("questions") or ss.get("count"), 0)
            me = _coerce_float(ss.get("marks"), 0) / cnt if cnt > 0 else 0
        me = _coerce_float(me, 0)
        if me > 0:
            if cat in out and abs(out[cat] - me) > 0.01:
                ambiguous.add(cat)
            out[cat] = me
    for cat in ambiguous:
        out.pop(cat, None)
    return out


def _exact_distribution_spec(blueprint_dict):
    """Human-readable, per-section per-type count+marks spec for the fallback prompt.

    The single-prompt fallback otherwise under-delivers (e.g. 4 MCQs when 7 are required),
    so we spell out the exact distribution it must produce.
    """
    lines = []
    for sec_name, sec in blueprint_dict.items():
        if not isinstance(sec, dict):
            continue
        marks = sec.get("marks", 0)
        subs = sec.get("subsections") or []
        if subs:
            # AI-authored patterns store the subsection count as 'count' (and the type as a
            # singular 'question_type') — read those too or the spec asks for 0 questions.
            total_q = sum(int(_coerce_float(s.get("questions_count") or s.get("questions") or s.get("count"), 0))
                          for s in subs if isinstance(s, dict))
            lines.append(f"SECTION {sec_name} ({marks} marks) — generate EXACTLY {total_q} questions:")
            for s in subs:
                if not isinstance(s, dict):
                    continue
                qts = s.get("question_types") or s.get("question_type") or [s.get("name", "")]
                if isinstance(qts, str):
                    qts = [qts]
                typ = qts[0] if qts else s.get("name", "")
                cnt = int(_coerce_float(s.get("questions_count") or s.get("questions") or s.get("count"), 1) or 1)
                me = s.get("marks_per_question")
                if me in (None, "", "varies"):
                    me = _coerce_float(s.get("marks"), 0) / (cnt or 1)
                me = _coerce_float(me, 0)
                mestr = str(int(me)) if float(me).is_integer() else str(me)
                lines.append(f"   - {cnt} × {typ} @ {mestr} mark{'s' if me != 1 else ''} each")
        else:
            qc = int(_coerce_float(sec.get("questions_count") or sec.get("questions"), 0))
            qts = sec.get("question_types") or []
            tnames = ", ".join(qt.get("type", str(qt)) if isinstance(qt, dict) else str(qt) for qt in qts)
            lines.append(f"SECTION {sec_name} ({marks} marks) — {qc} questions; types: {tnames}")
    return "\n".join(lines)


def _regroup_section(questions_list, sec_info):
    """Reorder questions into canonical type groups and fix per-type marks.

    Returns a list of (group_label_or_None, [questions]). Group labels are only emitted
    when the section genuinely mixes types, so single-type sections (and ordinary papers)
    render exactly as before.
    """
    if not isinstance(questions_list, list):
        return [(None, questions_list or [])]

    # Slot-authored sections (question_slots) are generated and validated in the
    # authored per-question order with per-question marks — regrouping by type
    # would destroy that order, and category-level marks stamping would clobber
    # legitimately varied marks (two same-category slots with different marks).
    # Render them exactly as stored.
    if isinstance(sec_info, dict) and sec_info.get("question_slots"):
        return [(None, questions_list)]

    type_marks = _section_type_marks(sec_info)
    buckets, order_seen = {}, []
    for q in questions_list:
        cat = _question_category(q) if isinstance(q, dict) else "_str"
        if isinstance(q, dict) and type_marks.get(cat):
            m = type_marks[cat]
            q["marks"] = int(m) if float(m).is_integer() else round(m, 2)
        if cat not in buckets:
            buckets[cat] = []
            order_seen.append(cat)
        buckets[cat].append(q)

    # Assertion-Reason questions are MCQ variants — they render inside the MCQ group
    # (after the plain MCQs), never under their own subheader. Merged AFTER stamping
    # so AR questions keep their own marks when the pattern gives them a different
    # marks_each than plain MCQs.
    if "ar" in buckets:
        buckets.setdefault("mcq", []).extend(buckets.pop("ar"))
        if "mcq" in order_seen:
            order_seen.remove("ar")
        else:
            order_seen[order_seen.index("ar")] = "mcq"

    multi = len([c for c in buckets if c != "_str"]) > 1
    groups, roman_i, used = [], 0, set()
    for cat, lbl in _TYPE_GROUP_ORDER:
        if cat in buckets:
            used.add(cat)
            label = None
            if multi:
                note = " (answer all questions)" if cat in ("mcq", "vsa") else ""
                label = f"{_ROMAN[roman_i]}. {lbl}{note}"
                roman_i += 1
            groups.append((label, buckets[cat]))
    for cat in order_seen:               # any uncategorised / string questions, in place
        if cat not in used:
            groups.append((None, buckets[cat]))
    return groups


def _emit_section_questions(all_questions, questions_list, sec_info, q_counter,
                            class_name=None, subject=None, chapters=None, paper_id=None):
    """Render a section's questions grouped by type with correct per-type marks."""
    position = 0
    for label, qs in _regroup_section(questions_list, sec_info):
        if label:
            all_questions.append(("subheader", label))
        for q in qs:
            # Slot sections render in authored order (see _regroup_section), so this counter is
            # the slot index the picture spec is attached to.
            position += 1
            if isinstance(q, dict):
                _inject_missing_diagram_prompt(q, sec_info, position)
                q["qnum"] = q_counter
                q_counter = process_question(all_questions, q, q_counter,
                                             class_name, subject, chapters, paper_id)
            else:
                all_questions.append(("q", f"{q_counter}. {str(q)}"))
                if class_name and subject:
                    try:
                        save_generated_question(str(q), class_name, subject,
                                                chapters[0] if chapters else None, "", 1, paper_id)
                    except Exception as e:
                        print(f"[Variation] Error saving question: {e}")
                q_counter += 1
    return q_counter


def render_section_questions(all_questions, data, blueprint, class_name=None, subject=None, chapters=None, paper_id=None):
    q_counter = 1
    seen_instructions = set()

    # Resolve sections_data from the data dict
    if "sections" in data:
        sections_data = data["sections"]
        if isinstance(sections_data, list):
            sections_data = {s["name"]: s for s in sections_data if isinstance(s, dict) and "name" in s}
    else:
        sections_data = data

    # A teacher-authored pattern states each section heading in its own words together with the
    # arithmetic ("II. Answer any SIX of the following    6 x 2 = 12"). When every section
    # carries such an instruction, print the headings the way the teacher wrote them instead of
    # the generic "SECTION – <name> (N MARKS)" banner. All-or-nothing so one section without an
    # instruction can't leave the paper with two different heading styles.
    _teacher_headings = bool(blueprint) and all(
        isinstance(si, dict) and any(str(x).strip() for x in (si.get('instructions') or []))
        for si in blueprint.values()
    )
    _sec_roman = 0

    for sec, sec_info in blueprint.items():
        if not isinstance(sec_info, dict):
            raise ValueError(f"Section '{sec}' has invalid blueprint format: {type(sec_info)}")

        print(f"[DEBUG] Processing section {sec}: {sec_info.get('title', 'NO TITLE')}")
        title = sec_info.get('title', '') or ''
        marks = sec_info.get('marks', 0)
        # M-03: include sub-subject name for compound papers
        sub_subject = sec_info.get('section_subject', '') or ''
        # The blueprint key 'sec' is normally a bare letter ("A"), but AI-generated /
        # teacher-authored patterns often name the section itself "Section A — Objective
        # Type" — prefixing "SECTION – " onto that doubles the word ("SECTION – Section
        # A — Objective Type"). Only add the prefix when 'sec' doesn't already read as one.
        already_labelled = bool(re.match(r'(?i)^\s*section\b', str(sec)))
        if sub_subject:
            # e.g. "SECTION A — BIOLOGY (26 MARKS)"
            prefix = str(sec) if already_labelled else f"SECTION {sec}"
            header_text = f"{prefix} — {sub_subject.upper()} ({marks} MARKS)"
        elif title and title.lower() not in ('section', sec.lower()):
            prefix = str(sec) if already_labelled else f"SECTION – {sec}"
            header_text = f"{prefix}: {title.upper()} ({marks} MARKS)"
        else:
            prefix = str(sec) if already_labelled else f"SECTION – {sec}"
            header_text = f"{prefix} ({marks} MARKS)"

        # The teacher's own heading, when the pattern carries one. A tab separates the heading
        # from its arithmetic; both renderers set a right tab stop so the sum sits at the
        # margin the way it does on a printed paper.
        lead_instruction = None
        if _teacher_headings:
            _instrs = [" ".join(str(x).split()) for x in (sec_info.get('instructions') or [])]
            _instrs = [x for x in _instrs if x]
            if _instrs:
                lead_instruction = _instrs[0]
                numbered = (f"{_ROMAN[_sec_roman]}. {lead_instruction}"
                            if _sec_roman < len(_ROMAN) else lead_instruction)
                calc = _section_calc(sec_info)
                header_text = f"{numbered}\t{calc}" if calc else numbered
                _sec_roman += 1

        all_questions.append(("header", header_text))

        # The teacher's own section instruction ("Answer any SIX of the following") — without it
        # a student meets eight 2-mark questions under a heading that says 12 marks and has no
        # way to know they answer six. Deduplicated across the paper: patterns routinely repeat
        # one general note (the internal-choice line) inside every section's instructions.
        for _instr in (sec_info.get('instructions') or []):
            _instr = " ".join(str(_instr).split())
            if not _instr:
                continue
            # Already printed as this section's heading — don't say it twice.
            if lead_instruction and _instr == lead_instruction:
                continue
            if _instr.lower() not in seen_instructions:
                seen_instructions.add(_instr.lower())
                all_questions.append(("instruction", _instr))

        # Get section data - handle both old and new structures
        # Try exact match first
        section_key = None
        
        # First, try exact match
        if sec in sections_data:
            section_key = sec
            print(f"[DEBUG] Exact match found for '{sec}'")
        else:
            print(f"[DEBUG] No exact match for '{sec}', trying fallback strategies. Available: {list(sections_data.keys())}")
            
            # Get section title for better matching
            section_title = sec_info.get('title', '').lower()
            section_question_types = sec_info.get('question_types', [])
            
            # Extract blueprint section name parts for matching
            sec_lower = sec.lower()
            
            # Try multiple matching strategies in order of preference
            # First, try keyword/content matching before letter matching
            sec_lower = sec.lower()
            
            # Strategy 1: Match by shared keywords in section names (highest priority)
            # Use specific keywords ordered by specificity
            keywords = ['case based', 'case study', 'multiple choice', 'very short answer', 'short answer', 'long answer', 'reading comprehension', 'reading', 'writing', 'grammar', 'literature']
            for keyword in keywords:
                if keyword in sec_lower:
                    # Blueprint section contains this keyword, find JSON section with SAME keyword
                    for key in sections_data.keys():
                        key_lower = key.lower()
                        # BOTH must contain the same keyword for a match
                        if keyword in key_lower:
                            # Extra validation: make sure it's not a substring match of a different section type
                            # e.g., "short answer" shouldn't match "very short answer"
                            if keyword == 'short answer' and 'very short' in key_lower:
                                continue  # Skip this, it's a different type
                            section_key = key
                            break
                    if section_key != sec:
                        break
            
            # Strategy 2: Match by section title (e.g., "Case Based Questions")
            if section_key == sec and section_title:
                for key in sections_data.keys():
                    key_lower = key.lower()
                    if section_title in key_lower:
                        section_key = key
                        break
            
            # Strategy 3: Match by overlapping question types
            if section_key == sec and section_question_types:
                for key in sections_data.keys():
                    key_data = sections_data[key]
                    if isinstance(key_data, dict):
                        json_qtypes = key_data.get('question_types', [])
                        if json_qtypes:
                            # Check if any question types overlap
                            json_qtypes_lower = [_qt_str(qt) for qt in json_qtypes]
                            sec_qtypes_lower = [_qt_str(qt) for qt in section_question_types]
                            overlap = [qt for qt in sec_qtypes_lower if qt in json_qtypes_lower]
                            if overlap:
                                section_key = key
                                break
            
            # Strategy 4: Match by section letter ONLY if content also matches
            # This prevents wrong matches like "Section D - Case Based" with "Section D - Long Answer"
            if section_key == sec:
                match = re.search(r'section\s+([a-z])', sec.lower())
                if match:
                    sec_letter = match.group(1).upper()
                    for key in sections_data.keys():
                        key_lower = key.lower()
                        # Check if JSON key has the same letter
                        key_match = re.search(r'section\s+([a-z])', key_lower)
                        if key_match:
                            key_letter = key_match.group(1).upper()
                            if sec_letter == key_letter:
                                # Also verify content similarity before matching
                                # Check if they share at least one keyword beyond the letter
                                shared_keywords = ['case', 'multiple', 'very', 'short', 'long', 'reading', 'writing', 'grammar', 'literature']
                                content_matches = any(kw in sec_lower and kw in key_lower for kw in shared_keywords)
                                if content_matches:
                                    section_key = key
                                    break
        
        if section_key in sections_data:
            section_data = sections_data[section_key]
            print(f"[DEBUG] Found section data for '{sec}', keys: {list(section_data.keys()) if isinstance(section_data, dict) else type(section_data)}")
            sec_render_info = dict(sec_info)
            sec_render_info["_section_name"] = sec
            sec_render_info["section_title"] = title

            # Check if it's the new structure with subsections
            if "subsections" in section_data:
                subsections = section_data["subsections"]

                # Handle both list and dict formats for subsections
                if isinstance(subsections, dict):
                    # Dictionary format: {"reading": [...], "grammar": [...]}
                    for sub, q_list in subsections.items():
                        if not isinstance(q_list, list):
                            continue
                        for q in q_list:
                            if isinstance(q, dict):
                                q_counter = process_question(all_questions, q, q_counter, class_name, subject, chapters, paper_id)
                            else:
                                all_questions.append(("q", f"{q_counter}. {str(q)}"))
                                q_counter += 1
                elif isinstance(subsections, list):
                    # List format: list of subsection objects
                    for subsection in subsections:
                        if not isinstance(subsection, dict):
                            continue
                        
                        # Handle subsection with passage/extract
                        subsection_name = subsection.get('name', 'Subsection')
                        
                        # Add subsection instruction if present
                        if 'instructions' in subsection and subsection['instructions']:
                            instruction = subsection['instructions'][0] if isinstance(subsection['instructions'], list) else subsection['instructions']
                            all_questions.append(("instruction", instruction))
                        
                        # Handle passage/extract in subsection
                        if 'passage' in subsection:
                            passage_text = subsection.get('passage', '')
                            if passage_text:
                                all_questions.append(("instruction", "Read the passage:"))
                                all_questions.append(("passage", passage_text))

                        if 'extract' in subsection:
                            extract_text = subsection.get('extract', '')
                            if extract_text:
                                instruction_text = subsection.get('extract_instruction') or subsection.get('instruction')
                                all_questions.append(("instruction", instruction_text or "Read the extract:"))
                                all_questions.append(("passage", extract_text))
                        
                        # Handle picture_description in subsection
                        if 'picture_description' in subsection:
                            picture_desc = subsection.get('picture_description', '')
                            if picture_desc:
                                all_questions.append(("instruction", "Observe the picture:"))
                                # Check if it's a [Picture: ...] marker
                                if picture_desc.startswith('[Picture:') or picture_desc.startswith('[Picture '):
                                    all_questions.append(("q", picture_desc))  # Will be processed by materialize_images
                                else:
                                    all_questions.append(("instruction", picture_desc))
                        
                        # Handle questions in subsection
                        if 'questions' in subsection:
                            questions_list = subsection.get('questions', [])
                            if isinstance(questions_list, list):
                                for q in questions_list:
                                    if isinstance(q, dict):
                                        q_counter = process_question(all_questions, q, q_counter, class_name, subject, chapters, paper_id)
                                    else:
                                        all_questions.append(("q", f"{q_counter}. {str(q)}"))
                                        q_counter += 1
                        else:
                            # If no questions key, treat the subsection itself as a question
                            q_counter = process_question(all_questions, subsection, q_counter, class_name, subject, chapters, paper_id)
                else:
                    pass  # unknown subsections format — skip silently
            else:
                # Old structure - check if questions are in a 'questions' key or directly in section_data
                
                # Handle sections with passages or extracts
                if isinstance(section_data, dict):
                    if 'image_path' in section_data:
                        img_path = section_data.get('image_path', '')
                        if img_path and os.path.isfile(img_path):
                            all_questions.append(("image", img_path))
                            sec_render_info["_shared_image_present"] = True

                    if 'passage' in section_data:
                        passage_text = section_data.get('passage', '')
                        # Slot sections carry extract passages per-question (source_text); a
                        # section-level 'passage' there is generator planning junk ("This
                        # section tests your understanding…") — except unseen-reading
                        # sections, whose slots (source='unseen') genuinely share one passage.
                        _slots_here = sec_info.get("question_slots") if isinstance(sec_info, dict) else None
                        if _slots_here and not any(
                                str(s.get("source") or "").lower() == "unseen"
                                for s in _slots_here if isinstance(s, dict)):
                            passage_text = ''
                        if passage_text:
                            all_questions.append(("instruction", "Read the passage:"))
                            all_questions.append(("passage", passage_text))

                    if 'extract' in section_data:
                        extract_text = section_data.get('extract', '')
                        if extract_text:
                            # Prefer instruction from section data if available
                            instruction_text = None
                            section_instructions = section_data.get('instructions')
                            if isinstance(section_instructions, list) and section_instructions:
                                instruction_text = section_instructions[0]
                            elif isinstance(section_instructions, str):
                                instruction_text = section_instructions
                            all_questions.append(("instruction", instruction_text or "Read the extract:"))
                            all_questions.append(("passage", extract_text))
                
                # Handle both cases: questions in 'questions' key or section_data is directly a list
                if isinstance(section_data, dict) and 'questions' in section_data:
                    questions_list = section_data['questions']
                    print(f"[DEBUG] Found questions list with {len(questions_list) if isinstance(questions_list, list) else 'N/A'} items")
                    if isinstance(questions_list, list):
                        # Regroup by type (MCQ → VSA → SA → CBQ → LA) with sub-headings and
                        # restore per-type marks from the section blueprint.
                        q_counter = _emit_section_questions(
                            all_questions, questions_list, sec_render_info, q_counter,
                            class_name, subject, chapters, paper_id)
                elif isinstance(section_data, list):
                    # section_data is directly a list of questions
                    q_counter = _emit_section_questions(
                        all_questions, section_data, sec_render_info, q_counter,
                        class_name, subject, chapters, paper_id)
                else:
                    pass
        else:
            all_questions.append(("q", f"No questions found for section {sec}"))
    
    return all_questions


def _marks_suffix(marks_raw):
    """Return ' [X marks]' / ' [1 mark]' string for appending to question text, or '' if absent."""
    if marks_raw is None:
        return ""
    try:
        val = float(marks_raw)
    except (TypeError, ValueError):
        return ""
    if val <= 0:
        return ""
    int_val = int(val) if val == int(val) else val
    label = "mark" if int_val == 1 else "marks"
    return f" [{int_val} {label}]"


# Models sometimes describe a diagram INLINE — "(Image description: a neuron with parts A, B, C…)"
# — instead of emitting a structured image_prompt field, so no image gets generated. This catches
# that parenthetical/bracketed description so we can render a real image and drop the stray text.
_INLINE_IMG_RE = re.compile(
    r'[\(\[]\s*(?:image\s*description|image|diagram|figure|picture|illustration)'
    r'(?:\s+placeholder)?\b[:\-\s]*'
    r'(.*?)[\)\]]',
    re.IGNORECASE | re.DOTALL,
)


# A generic "observe / study the diagram and answer…" stem POINTS to an image, it does not
# DESCRIBE one. Deriving an image prompt from such a stem makes the generator invent a random,
# meaningless picture (e.g. a CEO/Manager org chart for a chemistry CBQ). The real image is
# already supplied via the section's image_path or an explicit image_prompt — so these stems
# must never be turned into image prompts.
# The definite article is what makes a stem a POINTER: "study the diagram" points at something
# the paper prints, while "osmosis using a U-tube setup with a semi-permeable membrane" describes
# a scene to draw. Without it every descriptive sentence containing "using … a setup" read as a
# pointer and its image was silently dropped.
_GENERIC_IMG_STEM_RE = re.compile(
    r'(observe|study|look\s+at|refer\s+to|examine|consider|based\s+on|using|from|read)\b'
    r'[^.]{0,50}\b(?:the|this|that|these|those)\s+'
    r'(diagram|figure|image|graph|picture|illustration|chart|map|flow\s*chart|'
    r'table|setup|set-up|circuit|structure|source|case|passage|given|above|below|following)\b',
    re.IGNORECASE,
)


def _is_generic_image_stem(text: str) -> bool:
    """True if `text` merely points at an image/source rather than describing one to draw."""
    t = (text or "").strip().lower()
    if not t:
        return True
    if "answer the following" in t or "answer the question" in t:
        return True
    return bool(_GENERIC_IMG_STEM_RE.search(t))


def _extract_inline_image(text):
    """Return (clean_text, image_prompt_or_None). If the question text contains an inline
    '(Image description: …)' style note, pull it out as an image prompt and strip it from
    the visible text."""
    if not text:
        return text, None
    m = _INLINE_IMG_RE.search(text)
    if not m:
        return text, None
    desc = (m.group(1) or "").strip().strip(".")
    clean = (text[:m.start()] + " " + text[m.end():])
    clean = re.sub(r"\s{2,}", " ", clean).strip()
    # Only treat it as an image if the description is substantive enough to render.
    if len(desc) < 12:
        return clean or text, None
    return clean, desc


_DIAGRAM_SECTION_RE = re.compile(
    r'\b(diagram|figure|picture|visual|image)(?:\s+based)?\b',
    re.IGNORECASE,
)
_DIAGRAM_PROMPT_PREFIX_RE = re.compile(
    r'^\s*(?:'
    r'a\s+student\s+is\s+given\s+a\s+diagram\s+showing|'
    r'the\s+diagram(?:\s+(?:below|above|given|shown))?\s+shows|'
    r'the\s+figure(?:\s+(?:below|above|given|shown))?\s+shows|'
    r'a\s+diagram\s+illustrates(?:\s+the\s+concept\s+of)?|'
    r'the\s+diagram\s+illustrates(?:\s+the\s+concept\s+of)?|'
    r'the\s+diagram\s+below\s+shows\s+a?\s*|'
    r'the\s+schematic\s+representation\s+of'
    r')\s*',
    re.IGNORECASE,
)
_DIAGRAM_PROMPT_STOP_RE = re.compile(
    r'(?:\(\s*i+\s*\)|\bidentify\b|\bcalculate\b|\bdefine\b|\bstate\b|\bexplain\b|'
    r'\bgive\s+reason\b|\bbased\s+on\s+the\s+diagram\b|\bwhich\s+label\b|\bwhich\s+side\b)',
    re.IGNORECASE,
)


def _section_wants_question_diagrams(sec_info) -> bool:
    """True when the section metadata signals per-question diagrams."""
    if not isinstance(sec_info, dict):
        return False
    if sec_info.get("_shared_image_present"):
        return False

    blobs = [
        sec_info.get("_section_name", ""),
        sec_info.get("section_name", ""),
        sec_info.get("title", ""),
        sec_info.get("section_title", ""),
    ]

    instr = sec_info.get("instructions")
    if isinstance(instr, (list, tuple)):
        blobs.extend(str(i) for i in instr)
    elif instr:
        blobs.append(str(instr))

    qtypes = sec_info.get("question_types") or []
    if isinstance(qtypes, str):
        blobs.append(qtypes)
    else:
        for qt in qtypes:
            blobs.append(qt.get("type", "") if isinstance(qt, dict) else str(qt))

    return bool(_DIAGRAM_SECTION_RE.search(" ".join(str(b) for b in blobs if b)))


def _derive_diagram_prompt(text: str) -> str | None:
    """Build a usable diagram prompt from descriptive question text."""
    raw = re.sub(r"\s+", " ", str(text or "")).strip()
    if not raw:
        return None

    desc = _DIAGRAM_PROMPT_PREFIX_RE.sub("", raw, count=1).strip(" :,-")
    desc = re.split(r"\n\s*\(\s*i+\s*\)", desc, maxsplit=1, flags=re.IGNORECASE)[0]
    desc = _DIAGRAM_PROMPT_STOP_RE.split(desc, maxsplit=1)[0].strip(" :;,-")

    if not desc or len(desc) < 18 or _is_generic_image_stem(desc):
        return None

    parts = [p.strip() for p in re.split(r'(?<=[.!?])\s+', desc) if p.strip()]
    desc = " ".join(parts[:3]).strip()
    if not desc or len(desc) < 18:
        return None

    return (
        "Scientific textbook diagram, pure white background, black line art, "
        "clean labels where needed: "
        f"{desc}"
    )


def _inject_missing_diagram_prompt(q, sec_info, position=None):
    """Backfill image_prompt when a diagram section question forgot to provide one."""
    if not isinstance(q, dict):
        return q
    if str(q.get("image_prompt", "")).strip():
        return q

    # A slot-authored section names exactly which questions carry a picture. The section-wide
    # salvage below would put an image above every question that merely reads descriptively —
    # so in a section where the teacher asked for ONE picture, it printed several. It is only
    # safe when there is no per-question spec to obey.
    slots = sec_info.get("question_slots") if isinstance(sec_info, dict) else None
    if slots:
        from .section_generator import _slot_wants_image
        if position is None or not (1 <= position <= len(slots)):
            return q
        if not _slot_wants_image(slots[position - 1]):
            return q
    elif not _section_wants_question_diagrams(sec_info):
        return q

    text = str(q.get("text") or q.get("question") or "").strip()
    if not text:
        return q

    clean_text, inline_prompt = _extract_inline_image(text)
    if inline_prompt:
        if "text" in q:
            q["text"] = clean_text
        elif "question" in q:
            q["question"] = clean_text
        q["image_prompt"] = inline_prompt
        return q

    prompt = _derive_diagram_prompt(text)
    if prompt:
        q["image_prompt"] = prompt
    return q


def _emit_sub_questions(all_questions, sub_qs):
    """Render a question's sub_questions as lettered (a)/(b)/(c) sub-lines.

    Accepts dict entries keyed 'text' (canonical) or 'q'/'question' (legacy model
    output) — entries with none of these are skipped rather than lost silently.
    """
    if not isinstance(sub_qs, list):
        return
    _letters = "abcdefghij"
    for _si, _sq in enumerate(sub_qs):
        _lbl = f"({_letters[_si]})" if _si < len(_letters) else f"({_si + 1})"
        if isinstance(_sq, dict):
            _sq_text = str(_sq.get("text") or _sq.get("q") or _sq.get("question") or "").strip()
            if _sq_text:
                all_questions.append(("subq", f"{_lbl} {_sq_text}{_marks_suffix(_sq.get('marks'))}"))
        elif isinstance(_sq, str) and _sq.strip():
            all_questions.append(("subq", f"{_lbl} {_sq.strip()}"))


def process_question(all_questions, q, q_counter, class_name=None, subject=None, chapters=None, paper_id=None):
    # Handle case where q might be a string instead of dict
    if not isinstance(q, dict):
        question_text = str(q)
        all_questions.append(("q", f"{q_counter}. {question_text}"))
        
        # Save question if metadata available
        if class_name and subject:
            try:
                chapter = chapters[0] if chapters else None
                save_generated_question(question_text, class_name, subject, chapter, "", 1, paper_id)
            except Exception as e:
                print(f"[Variation] Error saving question: {e}")
        
        return q_counter + 1
    
    qnum = q.get("qnum", q_counter)
    print(f"[DEBUG] Processing question {qnum}, type: {q.get('type', 'no type')}, keys: {list(q.keys())}")

    # If it has an instruction field, treat it as an instruction (no number)
    if "instruction" in q and q.get("type") not in ["unseen_passage_or_case_based"]:
        all_questions.append(("instruction", q.get('instruction', '')))  # No number for instructions
        
        # Handle gap_filling type with text field
        if q.get("type") == "gap_filling" and "text" in q:
            text_content = q.get("text", "")
            if text_content:
                all_questions.append(("passage", text_content))
        
        # Handle reordering type with sentences array
        elif q.get("type") == "reordering" and "sentences" in q:
            sentences = q.get("sentences", [])
            if isinstance(sentences, list):
                for item in sentences:
                    all_questions.append(("subq", f"- {str(item)}"))
            else:
                pass
        
        # Handle other types with items, sentences, phrases
        else:
            for k in ("items", "sentences", "phrases"):
                if k in q:
                    items = q.get(k, [])
                    if isinstance(items, list):
                        for item in items:
                            all_questions.append(("subq", f"- {str(item)}"))
                    else:
                        pass
        return q_counter + 1

    elif "extract" in q:
        # Handle literature extracts with instruction
        if "instruction" in q:
            all_questions.append(("instruction", q.get('instruction', '')))  # No number for instructions
        else:
            all_questions.append(("instruction", "Read the extract:"))  # No number for instructions
        extract_text = q.get("extract", "")
        if extract_text:
            all_questions.append(("passage", extract_text))
        for i, subq in enumerate(q.get("questions", []), start=1):
            # Handle both string and dict formats for questions
            if isinstance(subq, dict):
                # Check for both "text" and "q" fields
                question_text = subq.get("text", subq.get("q", str(subq)))
                # Also check for marks field
                marks = subq.get("marks")
                question_text = f"{question_text}{_marks_suffix(marks)}"
            else:
                question_text = str(subq)
            all_questions.append(("subq", f"({i}) {question_text}"))

    elif q.get("type") == "unseen_passage_or_case_based":
        options = q.get('options', [])
        if not isinstance(options, list):
            options = []
        all_questions.append(("instruction", q.get('instruction', 'Read the following:')))  # No number for instructions
        for opt in options:
            if not isinstance(opt, dict):
                continue
            opt_questions = opt.get('questions', [])
            if not isinstance(opt_questions, list):
                opt_questions = []
            if 'passage' in opt:
                passage_text = opt.get('passage', '')
                if passage_text:
                    all_questions.append(("passage", f"[{opt.get('kind', '').replace('_',' ').title()}]\n{passage_text}"))
            for i, subq in enumerate(opt_questions, start=1):
                # Handle both string and dict formats for questions
                if isinstance(subq, dict):
                    # Check for both "text" and "q" fields
                    question_text = subq.get("text", subq.get("q", str(subq)))
                    # Also check for marks field
                    marks = subq.get("marks", "")
                    if marks:
                        question_text = f"{question_text} ({marks} marks)"
                else:
                    question_text = str(subq)
                all_questions.append(("subq", f"({i}) {question_text}"))

    elif "passage" in q and "questions" in q:
        all_questions.append(("instruction", "Read the passage:"))  # No number for instructions
        passage_text = q.get("passage", "")
        if passage_text:
            all_questions.append(("passage", passage_text))
        questions_list = q.get("questions", [])
        if not isinstance(questions_list, list):
            questions_list = []
        for i, subq in enumerate(questions_list, start=1):
            # Handle both string and dict formats for questions
            if isinstance(subq, dict):
                # Check for both "text" and "q" fields
                question_text = subq.get("text", subq.get("q", str(subq)))
                # Also check for marks field
                marks = subq.get("marks")
                question_text = f"{question_text}{_marks_suffix(marks)}"
            else:
                question_text = str(subq)
            all_questions.append(("subq", f"({i}) {question_text}"))

    elif "passage" in q:
        instr = q.get("instruction", "Based on passage:")
        all_questions.append(("instruction", instr))  # No number for instructions
        passage_text = q.get("passage", "")
        if passage_text:
            all_questions.append(("passage", passage_text))

    # This case is now handled above

    elif "text" in q and not ("passage" in q or "extract" in q):
        # Handle questions that have only "text" field (new structure)
        text_content = q.get("text", "")
        q_type = q.get("type", "")
        q_subtype = str(q.get("subtype", "")).strip().lower()
        _raw_opts = q.get("options")
        # Normalize dict options {"a": "...", "b": "..."} → list (parallel pipeline format)
        # Also strip any embedded "(x) " prefix the LLM sometimes writes in values —
        # render_docx adds its own "(a)" prefix so we'd get "(a) (a) text" without this.
        if isinstance(_raw_opts, dict):
            options = [_STRIP_OPT_PREFIX.sub('', str(v)).strip() for _, v in sorted(_raw_opts.items())]
        else:
            options = _raw_opts

        # If options exist separately, strip inline options from question text.
        # AR questions are skipped — their "(A):"/"(R):" markers are part of the statements.
        text_content = _strip_inline_options(text_content, options, q)

        # CBQ source passage lives on the question itself (NOT at section level), so it
        # renders immediately before this question and is scoped to it alone.
        _src_text = str(q.get("source_text", "") or "").strip()
        if _src_text:
            all_questions.append(("instruction", "Read the source/case and answer the questions that follow:"))
            all_questions.append(("passage", _src_text))

        # Image-based questions: render a diagram ABOVE the question. An image is produced when
        # the model gives an image_prompt, OR writes an inline "(Image description: …)", OR simply
        # tags the question subtype/type "image_based" — making the tag itself a reliable trigger.
        _img_p = str(q.get('image_prompt', '') or '').strip().strip('.')
        _is_image_q = (str(q.get('subtype', '')).lower() == 'image_based'
                       or _qt_str(q.get('type', '')) == 'image_based')
        if not _img_p:
            _orig_text = text_content
            text_content, _inline = _extract_inline_image(text_content)
            if _inline:
                _img_p = _inline
            elif _is_image_q:
                # Tagged image_based but no prompt/description — derive a prompt from the
                # question text, but ONLY if it actually describes a visual. A generic
                # "observe the diagram and answer…" stem would generate a meaningless image
                # (the real one comes from the section image_path / an explicit image_prompt).
                _derived = re.sub(r'\s+', ' ', _orig_text or '').strip()
                if not _is_generic_image_stem(_derived):
                    _img_p = _derived
        if _img_p and len(_img_p) > 10:
            all_questions.append(('image_gen', _img_p))

        if text_content:
            marks_raw = q.get("marks")
            marks_suffix = _marks_suffix(marks_raw)
            all_questions.append(("q", f"{qnum}. {text_content}{marks_suffix}"))

            # Render sub_questions for image_based / source_based CBQ
            _emit_sub_questions(all_questions, q.get("sub_questions") or [])

            # Save question for tracking and deduplication
            if class_name and subject:
                try:
                    marks = q.get("marks", 1)
                    chapter = chapters[0] if chapters else None
                    save_generated_question(text_content, class_name, subject, chapter, q_type, marks, paper_id)
                except Exception as e:
                    print(f"[Variation] Error saving question: {e}")

        # Render MCQ options if present as a block (for two-column layout later)
        if isinstance(options, list) and options:
            labeled = []
            for i, opt in enumerate(options, start=1):
                label = chr(96 + i)  # a, b, c, d
                labeled.append(f"({label}) {str(opt).strip()}")
            all_questions.append(("opts_block", labeled))

        # M-04: map work note — trigger on type, q_type field, OR subtype
        map_note = str(q.get('map_note', '') or '').strip()
        if map_note or q_type == "map_work" or q_subtype == "map_based":
            all_questions.append(("instruction", map_note or "[Attach outline map — examiner to supply]"))

        # M-02: OR alternatives (internal choice). A string is a bare alternate question;
        # a dict is a FULL alternative carrying its own source_text/passage, sub_questions
        # and options (internal-choice extract/CBQ); a LIST is an N-way choice ("paragraph
        # OR letter OR notice") — every entry prints after its own OR separator. Rendering
        # only the stem left the second alternative as an empty header.
        _or_raw = q.get('or_alternative')
        _or_list = _or_raw if isinstance(_or_raw, list) else ([_or_raw] if _or_raw else [])
        for _alt in _or_list:
            if isinstance(_alt, dict):
                or_alt = str(_alt.get('text', '') or '').strip()
                _or_src = str(_alt.get('source_text', '') or _alt.get('passage', '') or '').strip()
                _or_subs = _alt.get('sub_questions')
                _or_opts = _alt.get('options')
                _or_marks = _alt.get('marks', q.get('marks'))
            elif isinstance(_alt, str):
                or_alt, _or_src, _or_subs, _or_opts, _or_marks = _alt.strip(), '', None, None, q.get('marks')
            else:
                continue
            if not (or_alt or _or_src or (isinstance(_or_subs, list) and _or_subs)):
                continue
            all_questions.append(("or", "OR"))
            if _or_src:
                all_questions.append(("instruction", "Read the source/case and answer the questions that follow:"))
                all_questions.append(("passage", _or_src))
            if or_alt:
                all_questions.append(("q", f"{qnum}. {or_alt}{_marks_suffix(_or_marks)}"))
            _emit_sub_questions(all_questions, _or_subs if isinstance(_or_subs, list) else [])
            if isinstance(_or_opts, dict):
                _or_opts = [_STRIP_OPT_PREFIX.sub('', str(v)).strip() for _, v in sorted(_or_opts.items())]
            if isinstance(_or_opts, list) and _or_opts:
                all_questions.append(("opts_block",
                                      [f"({chr(96 + i)}) {str(opt).strip()}"
                                       for i, opt in enumerate(_or_opts, start=1)]))

    elif "question" in q and "or" in q:
        question_text = q.get("question", "")
        or_text = q.get("or", "")
        marks_suffix = _marks_suffix(q.get("marks"))
        if question_text:
            all_questions.append(("q", f"{qnum}. {question_text}{marks_suffix}"))
        all_questions.append(("or", "OR"))
        if or_text:
            all_questions.append(("q", or_text))

    elif "question" in q:
        question_text = q.get("question", "")
        _raw_opts = q.get("options")
        if isinstance(_raw_opts, dict):
            options = [_STRIP_OPT_PREFIX.sub('', str(v)).strip() for _, v in sorted(_raw_opts.items())]
        else:
            options = _raw_opts

        # If options exist separately, strip inline options from question text.
        # AR questions are skipped — their "(A):"/"(R):" markers are part of the statements.
        question_text = _strip_inline_options(question_text, options, q)

        # A CBQ stored under 'question' (instead of 'text') still carries its own passage.
        _src_text = str(q.get("source_text", "") or "").strip()
        if _src_text:
            all_questions.append(("instruction", "Read the source/case and answer the questions that follow:"))
            all_questions.append(("passage", _src_text))

        # Image-based questions: render a diagram ABOVE the question (see main branch above).
        _img_p = str(q.get('image_prompt', '') or '').strip().strip('.')
        _is_image_q = (str(q.get('subtype', '')).lower() == 'image_based'
                       or _qt_str(q.get('type', '')) == 'image_based')
        if not _img_p:
            _orig_text = question_text
            question_text, _inline = _extract_inline_image(question_text)
            if _inline:
                _img_p = _inline
            elif _is_image_q:
                _derived = re.sub(r'\s+', ' ', _orig_text or '').strip()
                if not _is_generic_image_stem(_derived):
                    _img_p = _derived
        if _img_p and len(_img_p) > 10:
            all_questions.append(('image_gen', _img_p))

        if question_text:
            marks_suffix = _marks_suffix(q.get("marks"))
            all_questions.append(("q", f"{qnum}. {question_text}{marks_suffix}"))

        # Multi-part questions ("Attempt any 5 of the following 6") keep their parts
        # under 'sub_questions' regardless of which key holds the stem.
        _emit_sub_questions(all_questions, q.get("sub_questions") or [])

        # Handle options for MCQ style
        if isinstance(options, list) and options:
            labeled = []
            for i, option in enumerate(options, start=1):
                label = chr(96 + i)  # a, b, c, d
                labeled.append(f"({label}) {str(option).strip()}")
            all_questions.append(("opts_block", labeled))
        elif options and not isinstance(options, list):
            pass

    return q_counter + 1


# ------------------------------
# Science/Maths generator
# ------------------------------
def generate_science_paper(class_name, subject, chapters, difficulty, pattern, blueprint, summary_file=None, model_source='aws', school_id=None):
    contexts = []
    for ch in chapters:
        results = embeddings.query(class_name=class_name, subject=subject, unit=ch,
                                   query_text=f"{subject} important NCERT concepts", n_results=100, school_id=school_id)
        for docs in results["documents"]:
            contexts.extend(docs)
    context_text = "\n".join(contexts)

    # Use pattern sections if available, otherwise fall back to blueprint
    pattern_sections = None
    if hasattr(pattern, 'sections') and pattern.sections:
        # Ensure pattern.sections is a list, not a string
        if isinstance(pattern.sections, list):
            pattern_sections = pattern.sections
        elif isinstance(pattern.sections, str):
            # If it's a string, treat it as a single section description
            pattern_sections = [pattern.sections]
    
    schema = """
{ "sections": { "A": [{"qnum": int, "text": str, "options": [str]}],
                "B": [{"qnum": int, "text": str}],
                "C": [{"qnum": int, "text": str}],
                "D": [{"qnum": int, "text": str}],
                "E": [{"qnum": int, "question": str, "or": str}] } }
"""

    # Build rules based on pattern or blueprint
    if pattern_sections:
        rules_text = f"Follow this exam pattern: {pattern.name}\n"
        for section in pattern_sections:
            # Handle both dict and string formats
            if isinstance(section, dict):
                name = section.get('name', 'Section')
                q_count = section.get('questions', section.get('questions_count', 0))
                marks_each = section.get('marks', 0)
                rules_text += f"- Section {name}: {q_count} questions ({marks_each} marks each)\n"
            elif isinstance(section, str):
                # If section is a string, use it as is
                rules_text += f"- {section}\n"
    else:
        # Handle both old format {'A': (count, marks)} and new format {'A': {'title': ..., 'marks': ...}}
        def get_section_info(sec_key):
            sec_data = blueprint.get(sec_key)
            if isinstance(sec_data, tuple) and len(sec_data) == 2:
                # Old format: (count, marks)
                return sec_data[0], sec_data[1]
            elif isinstance(sec_data, dict):
                # New format: {'title': ..., 'marks': ..., 'question_types': ...}
                # Try to get questions_count, or calculate from marks
                count = sec_data.get('questions_count', 0)
                marks = sec_data.get('marks', 0)
                if count == 0 and marks > 0:
                    # Estimate count from marks (assume 1 mark per question for MCQs)
                    marks_per_q = sec_data.get('marks_per_question', 1)
                    count = marks // marks_per_q if marks_per_q > 0 else marks
                return count, marks_per_q if marks_per_q > 0 else 1
            return 0, 1
        
        a_count, a_marks = get_section_info('A')
        b_count, b_marks = get_section_info('B')
        c_count, c_marks = get_section_info('C')
        d_count, d_marks = get_section_info('D')
        e_count, e_marks = get_section_info('E')
        
        rules_text = f"""- Section A: {a_count} MCQs ({a_marks} mark each), include Assertion–Reason.
- Section B: {b_count} VSA ({b_marks} marks).
- Section C: {c_count} SA ({c_marks} marks).
- Section D: {d_count} Case-based ({d_marks} marks).
- Section E: {e_count} Long with OR choice ({e_marks} marks)."""

    gen_prompt = f"""
You are an expert CBSE paper setter.

Generate exam questions strictly in JSON format.
Schema:
{schema}

Rules:
{rules_text}
- Use only NCERT context from the selected chapters: {', '.join(chapters)}
{_difficulty_directive(difficulty)}
- Output raw JSON only.

Context:
{context_text}
"""
    raw_json, _, _ = call_bedrock(gen_prompt, GEN_MODEL_ID, temperature=0.7, max_tokens=16384)

    validator_prompt = f"""
You are a strict JSON validator.
Input will be JSON text. If it is already valid JSON, return unchanged.
If it is invalid, fix it but preserve all content.
Output only JSON.

{raw_json}
"""
    validated, _, _ = call_bedrock(validator_prompt, VAL_MODEL_ID, temperature=0.3, max_tokens=16384)
    data = enforce_json(validated)
    
    # Log generation results to summary file
    if summary_file:
        with open(summary_file, "a", encoding="utf-8") as f:
            f.write(f"=== GENERATION RESULTS ===\n")
            f.write(f"Raw JSON Length: {len(raw_json)} chars\n")
            f.write(f"Validated JSON Length: {len(validated)} chars\n")
            f.write(f"Questions Generated: {len(data.get('sections', {}))} sections\n")
            f.write(f"{'='*50}\n\n")

    all_questions, summary = [], {}
    q_counter = 1
    
    # Helper function to extract count and marks from blueprint section
    def get_section_info(sec_key):
        sec_data = blueprint.get(sec_key)
        if isinstance(sec_data, tuple) and len(sec_data) == 2:
            # Old format: (count, marks)
            return sec_data[0], sec_data[1]
        elif isinstance(sec_data, dict):
            # New format: {'title': ..., 'marks': ..., 'question_types': ...}
            count = sec_data.get('questions_count', 0)
            marks = sec_data.get('marks', 0)
            marks_per_q = sec_data.get('marks_per_question', 1)
            if count == 0 and marks > 0:
                # Estimate count from marks
                count = marks // marks_per_q if marks_per_q > 0 else marks
            return count, marks_per_q
        return 0, 1
    
    for sec in blueprint.keys():
        count, marks = get_section_info(sec)
        sec_data = data.get("sections", {}).get(sec, [])
        all_questions.append(("header", f"SECTION {sec} ({count} × {marks} = {count*marks})"))
        for i in range(count):
            if i < len(sec_data):
                q = sec_data[i]
                if not isinstance(q, dict):
                    continue
                if sec == "A":
                    text_content = q.get('text', '')
                    if text_content:
                        all_questions.append(("q", f"{q_counter}) {text_content}"))
                    options = q.get("options", [])
                    if isinstance(options, list):
                        for j, opt in enumerate(options):
                            opt_str = str(opt).strip() if opt else ""
                            if opt_str:
                                all_questions.append(("subq", f"   {chr(97+j)}) {opt_str}"))
                elif sec == "E":
                    question_text = q.get('question', '')
                    or_text = q.get('or', '')
                    if question_text:
                        all_questions.append(("q", f"{q_counter}) {question_text}"))
                    all_questions.append(("or", "OR"))
                    if or_text:
                        all_questions.append(("q", or_text))
                else:
                    text_content = q.get('text', '')
                    if text_content:
                        all_questions.append(("q", f"{q_counter}) {text_content}"))
            else:
                all_questions.append(("q", f"{q_counter}) [Placeholder]"))
            q_counter += 1
        summary[sec] = {"questions": count, "marks_each": marks}

    return render_docx(class_name, subject, chapters, all_questions, summary)


# ------------------------------
# English Core generator
# ------------------------------
def generate_english_paper(class_name, subject, chapters, difficulty, pattern, blueprint, summary_file=None, model_source='aws', school_id=None):
    # Blueprint is already normalized by BlueprintManager.get_blueprint()
    print(f"[Blueprint] Blueprint structure (already normalized): {list(blueprint.keys()) if isinstance(blueprint, dict) else 'Not a dict'}")

    from .data.cbse_patterns import get_english_lessons
    lesson_names = get_english_lessons(class_name)

    # Use selected chapters if available, otherwise fall back to full class lesson list
    selected_lessons = chapters if chapters else lesson_names

    print(f"[Embeddings] Querying embeddings for {len(selected_lessons)} chapters...")
    contexts = []
    for ch in selected_lessons:
        print(f"[Embeddings] Querying chapter: {ch}")
        results = embeddings.query(class_name=class_name, subject=subject, unit=ch,
                                   query_text="important NCERT extract content", n_results=50, school_id=school_id)
        print(f"[Embeddings] Found {len(results['documents'])} document chunks for {ch}")
        for docs in results["documents"]:
            contexts.extend(docs)
    
    context_text = "\n".join(contexts)
    print(f"[Embeddings] Total context chunks retrieved: {len(contexts)}")
    print(f"[Embeddings] Total context length: {len(context_text)} characters")

    # Use pattern sections if available, otherwise use blueprint
    pattern_sections = None
    if hasattr(pattern, 'sections') and pattern.sections:
        # Ensure pattern.sections is a list, not a string
        if isinstance(pattern.sections, list):
            pattern_sections = pattern.sections
        elif isinstance(pattern.sections, str):
            # If it's a string, treat it as a single section description
            pattern_sections = [pattern.sections]
    
    if pattern_sections:
        rules_text = f"Follow this exam pattern: {pattern.name}\n"
        for section in pattern_sections:
            # Handle both dict and string formats
            if isinstance(section, dict):
                name = section.get('name', 'Section')
                q_count = section.get('questions', section.get('questions_count', 0))
                marks_each = section.get('marks', 0)
                rules_text += f"- Section {name}: {q_count} questions ({marks_each} marks each)\n"
            elif isinstance(section, str):
                # If section is a string, use it as is
                rules_text += f"- {section}\n"
        rules_text += f"- Focus on selected chapters: {', '.join(selected_lessons)}\n"
        rules_text += _difficulty_directive(difficulty) + "\n"
    else:
        rules_text = f"""- Section A: One unseen passage OR one case-based passage (Answer ANY ONE). Then Note making + Summary.
- Section B: Grammar + Writing (Gap filling, reordering, ad/poster, speech, debate).
- Section C: Extracts, short answers, long answers ONLY from Hornbill and Snapshots NCERT chapters.
- Focus on selected chapters: {', '.join(selected_lessons)}
{_difficulty_directive(difficulty)}"""

    # Clean and limit context text to avoid corruption
    print(f"[Context] Original context length: {len(context_text)} characters")
    
    # Clean the context more carefully
    clean_context = context_text.replace('\n', ' ').replace('\r', ' ')
    # Remove any non-printable characters but preserve more content
    clean_context = re.sub(r'[^\x20-\x7E]', ' ', clean_context)
    clean_context = ' '.join(clean_context.split())  # Normalize whitespace
    
    # Increase limit to get more NCERT content
    clean_context = clean_context[:3000]  # Limit to 3000 chars to get more context
    
    print(f"[Context] Cleaned context length: {len(clean_context)} characters")
    print(f"[Context] Context preview: {clean_context[:200]}...")
    
    gen_prompt = f"""
You are an expert CBSE paper setter for ENGLISH CORE CLASS XI.

Generate exam questions strictly in JSON format following this exact schema:
{json.dumps(blueprint, indent=2)}

IMPORTANT RULES:
{rules_text}

SELECTED CHAPTERS FOR LITERATURE QUESTIONS:
{", ".join(selected_lessons)}

NCERT CONTEXT (use for literature questions only):
{clean_context}

QUESTION REQUIREMENTS:
- Section A: Create ONE unseen passage (300-400 words) OR case-based passage with 4-5 questions
- Section B: Grammar questions (gap filling, reordering) + Writing tasks (advertisement, poster, speech, debate)
- Section C: Literature questions ONLY from the selected chapters above
- All questions must be appropriate for Class XI level
- Difficulty: {difficulty}

OUTPUT FORMAT:
- Return ONLY valid JSON
- Do not include answers or explanations
- Follow the exact schema structure provided
- Ensure all required fields are present

Generate the question paper now:
"""
    # Step 1: Get complete response from generation model
    print(f"[Generation] Starting question paper generation...")
    print(f"[Generation] Using model: {GEN_MODEL_ID}")
    
    raw_json, _, _ = call_bedrock(gen_prompt, GEN_MODEL_ID, temperature=0.7, max_tokens=16384)

    # Log the complete raw response
    print(f"[Generation] Raw response received: {len(raw_json)} characters")
    print(f"[Generation] Raw response preview: {raw_json[:200]}...")
    
    # Check if the response is complete
    if not raw_json.strip().endswith('}') and not raw_json.strip().endswith('```'):
        print(f"[Generation] ⚠️  WARNING: Response appears incomplete!")
        print(f"[Generation] Last 200 chars: {raw_json[-200:]}")
        
        # Try to complete the JSON if it's clearly cut off
        if raw_json.strip().endswith('"') and raw_json.count('{') > raw_json.count('}'):
            print(f"[Generation] Attempting to complete truncated JSON...")
            # Add missing closing brackets
            missing_braces = raw_json.count('{') - raw_json.count('}')
            raw_json += '"}' * missing_braces
            print(f"[Generation] Added {missing_braces} closing brackets")
    
    # Step 2: Validate the complete response with validator model
    print(f"[Validation] Starting JSON validation...")
    print(f"[Validation] Using model: {VAL_MODEL_ID}")
    
    validator_prompt = f"""
You are a JSON validator and fixer for CBSE question papers.

TASK: Fix the following JSON to ensure it's valid and complete.

REQUIREMENTS:
- Must be valid JSON syntax
- Must follow the exact schema structure
- All required fields must be present
- No missing brackets, quotes, or commas
- All question numbers must be sequential
- All marks must add up correctly

INPUT JSON TO FIX:
{raw_json}

OUTPUT: Return ONLY the corrected JSON, no explanations.
"""
    
    validated, _, _ = call_bedrock(validator_prompt, VAL_MODEL_ID, temperature=0.3, max_tokens=16384)

    print(f"[Validation] Validation response received: {len(validated)} characters")
    print(f"[Validation] Validation response preview: {validated[:200]}...")
    
    # Step 3: Final JSON enforcement and cleanup
    print(f"[Final-Validation] Starting final JSON enforcement...")
    data = enforce_json(validated)
    
    print(f"[Final-Validation] JSON enforcement completed successfully!")
    print(f"[Final-Validation] Final data structure: {list(data.keys()) if isinstance(data, dict) else 'Not a dict'}")
    
    # Log generation results to summary file
    if summary_file:
        with open(summary_file, "a", encoding="utf-8") as f:
            f.write(f"=== GENERATION RESULTS ===\n")
            f.write(f"Generation Model: {GEN_MODEL_ID}\n")
            f.write(f"Validation Model: {VAL_MODEL_ID}\n")
            f.write(f"Raw JSON Length: {len(raw_json)} chars\n")
            f.write(f"Validated JSON Length: {len(validated)} chars\n")
            f.write(f"Final JSON Structure: {list(data.keys()) if isinstance(data, dict) else 'Not a dict'}\n")
            f.write(f"Questions Generated: {len(data.get('sections', {}))} sections\n")
            f.write(f"Generation Process: Complete\n")
            f.write(f"{'='*50}\n\n")

    print("=" * 50)
    print("USING DIRECT JSON APPROACH")
    print("=" * 50)

    # Use in-memory paper data stored by enforce_json() — thread-safe and avoids
    # the shared temp_clean.json file that caused cross-user data contamination.
    direct_data = getattr(_request_state, 'paper_data', None) or data
    print(f"[DIRECT] Using paper data with keys: {list(direct_data.keys())}")
    all_questions = render_section_questions([], direct_data, blueprint, class_name, subject, chapters, None)
    print(f"[DIRECT] Generated {len(all_questions)} questions")

    summary = {sec: {"title": sec_info["title"], "marks": sec_info["marks"]}
               for sec, sec_info in blueprint.items()}
    return render_docx(class_name, subject, chapters, all_questions, summary)


# ------------------------------
# Math notation helpers
# ------------------------------
_UNICODE_SUP = {
    '0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴',
    '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹',
    'n': 'ⁿ', 'a': 'ᵃ', 'b': 'ᵇ', 'm': 'ᵐ', 'x': 'ˣ',
    't': 'ᵗ', 'i': 'ⁱ', 'k': 'ᵏ', 'r': 'ʳ',
    '+': '⁺', '-': '⁻', '=': '⁼', '(': '⁽', ')': '⁾',
}

def preprocess_math_text(text):
    """Convert math shorthand to Unicode for plain-text/PDF rendering.
    ^2 → ², ^(n-1) → ⁿ⁻¹, standalone * → ×
    """
    if not isinstance(text, str):
        return text
    def _sup(m):
        exp = m.group(1) or m.group(2) or m.group(3) or ''
        return ''.join(_UNICODE_SUP.get(c, c) for c in exp)
    text = re.sub(r'\^\(([^)]+)\)|\^\{([^}]+)\}|\^([-+]?[A-Za-z0-9])', _sup, text)
    # Standalone * as multiplication → ×  (e.g. 3*4, A*B but not ** or *= )
    text = re.sub(r'(?<=[0-9A-Za-z|)])\*(?=[0-9A-Za-z|(√])', '×', text)
    return text


_MATH_RUN_RE = re.compile(r'(\^(?:\([^)]*\)|\{[^}]*\}|[-+]?[A-Za-z0-9]))')

def _add_math_runs(p, text, qrun_fn, bold=False):
    """Add text to a python-docx paragraph with ^ exponents as superscript runs."""
    text = re.sub(r'(?<=[0-9A-Za-z|)])\*(?=[0-9A-Za-z|(√])', '×', text)
    for part in _MATH_RUN_RE.split(text):
        if not part:
            continue
        if part.startswith('^'):
            exp = part[1:]
            if len(exp) > 1 and exp[0] in '({' and exp[-1] in ')}':
                exp = exp[1:-1]
            r = p.add_run(exp)
            r.font.superscript = True
        else:
            r = p.add_run(part)
        if bold:
            r.bold = True
        qrun_fn(r)


# ------------------------------
# PDF renderer
# ------------------------------
def render_pdf(class_name, subject, chapters, all_questions, summary, header_meta=None):
    writer = PdfWriter()
    packet = BytesIO()
    can = canvas.Canvas(packet, pagesize=A4)
    y = 650  # Much more space from header to avoid clash

    # Draw header fields onto the first page using provided coordinates
    try:
        page_width, page_height = A4

        def flip_y(y_top, height):
            return page_height - y_top - height

        can.setFont("Helvetica", 11)

        # Coordinates spec as provided by user
        coords = {
            "TESTTYPE": {"x": 253, "y": 86, "w": 190, "h": 14},
            "class": {"x": 128, "y": 109, "w": 285, "h": 19},
            "subject": {"x": 259, "y": 102, "w": 174, "h": 17},
            "time": {"x": 506, "y": 107, "w": 40, "h": 22},
            "marks": {"x": 506, "y": 131, "w": 45, "h": 14},
        }

        # Expand test type like PT-2 -> Periodic Test - 2
        def expand_test_type(pattern_name: str) -> str:
            if not pattern_name:
                return ""
            name = pattern_name.strip()
            import re
            m = re.match(r"(?i)\s*pt\s*[- ]?\s*(\d+)", name)
            if m:
                return f"Periodic Test - {m.group(1)}"
            m = re.match(r"(?i)\s*fa\s*[- ]?\s*(\d+)", name)
            if m:
                return f"Formative Assessment - {m.group(1)}"
            m = re.match(r"(?i)\s*sa\s*[- ]?\s*(\d+)", name)
            if m:
                return f"Summative Assessment - {m.group(1)}"
            return name

        if header_meta is None:
            header_meta = {}

        test_type_val = expand_test_type(header_meta.get("test_type", header_meta.get("pattern_name", "")))
        class_val = header_meta.get("class_name", class_name)
        subject_val = header_meta.get("subject", subject)
        time_val = str(header_meta.get("duration", ""))
        marks_val = str(header_meta.get("marks", ""))

        # Draw each value if present
        def draw_in_rect(key, value):
            if not value:
                return
            r = coords[key]
            x = r["x"] + 2
            ty = flip_y(r["y"], r["h"]) + (r["h"] / 3)
            can.drawString(x, ty, str(value))

        draw_in_rect("TESTTYPE", test_type_val)
        draw_in_rect("class", class_val)
        draw_in_rect("subject", subject_val)
        draw_in_rect("time", time_val)
        draw_in_rect("marks", marks_val)

    except Exception as e:
        print(f"[PDF-Header] WARNING: Failed to render header fields: {e}")

    if not all_questions:
        all_questions = [("header", "NO QUESTIONS GENERATED"),
                         ("q", "Check blueprint or JSON parsing")]

    print("[PDF-Render] DEBUG all_questions sample:")
    for i, (typ, text) in enumerate(all_questions[:15]):
        if typ == "passage":
            print(f"  {i}. ({typ}): {text[:70]}... [{len(text)} chars]")
        else:
            print(f"  {i}. ({typ}): {text[:80] if len(text) > 80 else text}")

    for typ, text in all_questions:
        print(f"[PDF-Render-Loop] Processing: typ={typ}, y={y}, text_len={len(str(text))}")
        if typ == "header":
            can.setFont("Helvetica-Bold", 14)
            # Teacher-authored heading: "<heading>\t<6 x 2 = 12>" prints left with the
            # arithmetic flush right, matching the pattern the teacher wrote.
            _left, _tab, _calc = str(text).partition("\t")
            if _calc:
                can.drawString(60, y, _left)
                can.drawRightString(540, y, _calc)
                y -= 30
            else:
                can.drawCentredString(300, y, text)
                y -= 60  # Even more spacing after header to avoid clash

        elif typ == "subheader":
            can.setFont("Helvetica-Bold", 12)
            y = draw_wrapped(can, preprocess_math_text(text), 60, y, 470)
            y -= 10

        elif typ == "instruction":
            can.setFont("Helvetica", 11)
            y = draw_wrapped(can, preprocess_math_text(text), 60, y, 470)

        elif typ == "q":
            can.setFont("Helvetica", 11)
            y = draw_wrapped(can, preprocess_math_text(text), 60, y, 470)

        elif typ == "subq":
            can.setFont("Helvetica", 11)
            y = draw_wrapped(can, preprocess_math_text(text), 80, y, 440)

        elif typ == "passage":
            print(f"[PDF-Render] Rendering passage: {len(text)} chars, y={y}")
            can.setFont("Helvetica-Oblique", 10)
            y = draw_wrapped(can, text, 60, y, 470, size=10, line_height=14)
            print(f"[PDF-Render] After passage, y={y}")
            y -= 5

        elif typ == "opts_block":
            # Draw MCQ options in a 2×2 grid
            opts = list(text) if isinstance(text, (list, tuple)) else []
            col_x = [70, 310]
            for idx, opt_text in enumerate(opts[:4]):
                row, col = divmod(idx, 2)
                x = col_x[col]
                opt_y = y - row * 16
                can.setFont("Helvetica", 11)
                can.drawString(x, opt_y, str(opt_text)[:60])
            rows_used = (min(len(opts), 4) + 1) // 2
            y -= rows_used * 16 + 4

        elif typ == "or":
            can.setFont("Helvetica-Bold", 11)
            can.drawCentredString(300, y, "OR")
            y -= 20

    # Summary Page
    can.showPage()
    can.setFont("Helvetica-Bold", 14)
    can.drawCentredString(300, 800, "SUMMARY OF PAPER GENERATION")
    y = 760
    can.setFont("Helvetica", 11)
    for sec, info in summary.items():
        if "marks_each" in info:
            q_count = info.get('questions', info.get('questions_count', 0))
            line = f"Section {sec}: {q_count} questions | Marks each: {info.get('marks_each', 0)}"
        else:
            line = f"Section {sec}: {info['title']} | Total Marks: {info['marks']}"
        y = draw_wrapped(can, line, 60, y, 470)

    can.save()
    packet.seek(0)
    overlay_reader = PdfReader(packet)

    base_path = os.path.join(os.path.dirname(__file__), 'data', 'base.pdf')
    if os.path.exists(base_path):
        base_reader = PdfReader(base_path)
        for i, overlay_page in enumerate(overlay_reader.pages):
            if i < len(base_reader.pages):
                base_page = base_reader.pages[i]
                overlay_page.merge_page(base_page)
                writer.add_page(overlay_page)
            else:
                writer.add_page(overlay_page)
    else:
        for page in overlay_reader.pages:
            writer.add_page(page)

    output_dir = os.path.join(settings.MEDIA_ROOT, "question_papers")
    os.makedirs(output_dir, exist_ok=True)

    safe_subject = subject.replace(" ", "_")
    filename = f"{safe_subject}_{datetime.now().strftime('%Y%m%d%H%M%S')}.pdf"
    file_path = os.path.join(output_dir, filename)

    with open(file_path, "wb") as f:
        writer.write(f)

    return f"question_papers/{filename}", summary


# ------------------------------
# DOCX renderer (default)
# ------------------------------
TAMIL_DEFAULT_FONT = 'Latha'
# Nirmala UI ships on Windows 8+ and renders Devanagari (Hindi/Sanskrit) cleanly. Used the same
# way as the Tamil font — set on the run's ascii/hAnsi/cs/eastAsia so Word doesn't fall back to
# a Latin font (which has no Devanagari glyphs → boxes).
DEVANAGARI_DEFAULT_FONT = 'Nirmala UI'


def set_tamil_font(run, font_name: str = TAMIL_DEFAULT_FONT):
    """Set a complex-script font on a text run (Tamil 'Latha', Devanagari 'Nirmala UI', …).
    Name kept for back-compat; works for any font passed in."""
    if run is None:
        return
    
    tamil_font = font_name or TAMIL_DEFAULT_FONT
    
    run.font.name = tamil_font
    
    r = run._element
    rPr = r.get_or_add_rPr()
    
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    
    for attr in ['ascii', 'hAnsi', 'cs', 'eastAsia']:
        rFonts.set(qn(f'w:{attr}'), tamil_font)


def has_tamil_text(text):
    """Check if text contains Tamil characters"""
    if not isinstance(text, str):
        return False
    return any('\u0B80' <= char <= '\u0BFF' for char in text)


def has_devanagari_text(text):
    """Check if text contains Devanagari characters (Hindi/Sanskrit)."""
    if not isinstance(text, str):
        return False
    return any('\u0900' <= char <= '\u097F' for char in text)


def _pick_script_font(subject, all_questions):
    """Pick the complex-script font this paper needs: Tamil \u2192 Latha, Hindi/Sanskrit \u2192 Nirmala UI.
    Decided by the subject name first, then by scanning the generated text for either script.
    Returns the font name, or None for ordinary Latin-script papers."""
    s = (subject or "").lower()
    if "tamil" in s or "\u0BA4\u0BAE\u0BBF\u0BB4" in s:
        return TAMIL_DEFAULT_FONT
    if "hindi" in s or "sanskrit" in s or "\u0939\u093F\u0902\u0926" in s or "\u0939\u093F\u0928\u094D\u0926" in s or "\u0938\u0902\u0938\u094D\u0915\u0943\u0924" in s:
        return DEVANAGARI_DEFAULT_FONT
    for _typ, text in all_questions:
        if has_tamil_text(text):
            return TAMIL_DEFAULT_FONT
        if has_devanagari_text(text):
            return DEVANAGARI_DEFAULT_FONT
    return None


def clear_paragraph(paragraph):
    """Remove all runs from a paragraph."""
    if paragraph is None:
        return
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def apply_tamil_document_styles(doc: Document, font_name: str = TAMIL_DEFAULT_FONT):
    """Ensure common styles use a Tamil-compatible font."""
    if doc is None:
        return
    
    style_names = [
        'Normal', 'Heading 1', 'Heading 2', 'Heading 3',
        'List Bullet', 'List Number', 'Table Grid',
        'Title', 'Subtitle', 'Default Paragraph Font'
    ]
    
    for style_name in style_names:
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        
        style.font.name = font_name
        
        try:
            rPr = style._element.get_or_add_rPr()
        except AttributeError:
            rPr = getattr(style._element, 'rPr', None)
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                style._element.append(rPr)
        
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        
        for attr in ['ascii', 'hAnsi', 'cs', 'eastAsia']:
            rFonts.set(qn(f'w:{attr}'), font_name)


def _expand_test_type(pattern_name: str) -> str:
    if not pattern_name:
        return ""
    name = pattern_name.strip()
    m = re.match(r"(?i)\s*pt\s*[- ]?\s*(\d+)", name)
    if m:
        return f"Periodic Test - {m.group(1)}"
    m = re.match(r"(?i)\s*fa\s*[- ]?\s*(\d+)", name)
    if m:
        return f"Formative Assessment - {m.group(1)}"
    m = re.match(r"(?i)\s*sa\s*[- ]?\s*(\d+)", name)
    if m:
        return f"Summative Assessment - {m.group(1)}"
    return name


def _header_title_line(pattern_name=""):
    """The header's second title line: the exam name with the redundant word
    'Pattern' stripped, followed by the current month and year — e.g. a pattern
    named 'Board Exam Pattern' becomes 'Board Exam July - 2026'. Falls back to
    just the month-year when there is no pattern name."""
    name = _expand_test_type(pattern_name or "")
    name = re.sub(r"\s*\bpattern\b\s*", " ", name, flags=re.IGNORECASE).strip()
    period = datetime.now().strftime("%B - %Y")
    return f"{name} {period}".strip() if name else period


# --- Paper header, generated in code ----------------------------------------
# The header — the school name / subject / test type in a bordered box, above a
# CLASS/TIME and EXAM NO/MARKS grid — used to live in a data/base.docx template
# whose placeholder tokens ("SCHOOLNAME", "CLASS       :", …) were string-replaced
# at render time. It is now built programmatically: no external .docx is needed,
# and the values can't be silently lost to whitespace drift in the template.
# Layout mirrors the old base.docx header1.xml exactly — a 1.5pt box border, a
# centred three-line title block, and a borderless 2x2 detail grid.

_HDR_BORDER = {"val": "single", "sz": "12", "color": "000000"}  # sz in eighths = 1.5pt


def _set_cell_border(cell, **edges):
    """Set individual borders on a table cell. Each edge kwarg (top/left/bottom/
    right/insideH/insideV) is either the string 'none' or a spec dict with
    'val'/'sz'/'color' keys."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge, spec in edges.items():
        el = tcBorders.find(qn(f'w:{edge}'))
        if el is None:
            el = OxmlElement(f'w:{edge}')
            tcBorders.append(el)
        if spec == 'none':
            el.set(qn('w:val'), 'none')
        else:
            el.set(qn('w:val'), spec.get('val', 'single'))
            el.set(qn('w:sz'), str(spec.get('sz', '12')))
            el.set(qn('w:color'), spec.get('color', '000000'))


def _fix_table_width(table, col_twips):
    """Pin a table to a fixed layout with explicit column widths (in twips), so
    Word does not autofit-shrink or expand it. `col_twips` is one width per column."""
    table.autofit = False
    table.allow_autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = tblPr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tblPr.append(layout)
    layout.set(qn('w:type'), 'fixed')
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:type'), 'dxa')
    tblW.set(qn('w:w'), str(sum(col_twips)))
    grid_cols = tbl.tblGrid.findall(qn('w:gridCol'))
    for i, w in enumerate(col_twips):
        if i < len(grid_cols):
            grid_cols[i].set(qn('w:w'), str(w))
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(col_twips):
                cell.width = Twips(col_twips[i])


def _set_cell_margins(table, top=0, bottom=0, left=108, right=108):
    """Set the default cell padding (in twips) for every cell in `table`. Zeroing
    top/bottom removes the vertical padding Word otherwise adds inside each cell —
    the main lever for making the header box compact."""
    tblPr = table._tbl.tblPr
    mar = tblPr.find(qn('w:tblCellMar'))
    if mar is None:
        mar = OxmlElement('w:tblCellMar')
        tblPr.append(mar)
    for edge, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        el = mar.find(qn(f'w:{edge}'))
        if el is None:
            el = OxmlElement(f'w:{edge}')
            mar.append(el)
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')


def _build_header(section, subject_val, class_val, time_val, marks_val,
                  test_type_val, school_name_val="", script_font=None):
    """Build the paper header into `section`'s header: a compact single bordered box.
    The top centres the school name (bold 9pt) and the exam-name/period line
    (bold 8pt); beneath it is a 3-column grid — CLASS / TIME on the first line
    and EXAM NO (left, blank for the student to fill) / SUBJECT (centred, bold)
    / MARKS on the second. Line spacing is 0.8 (80%) for compactness.
    Replaces the former base.docx template + _fill_header_placeholders
    substitution. `script_font`, when set, applies the complex-script font
    (Tamil/Devanagari) to every header run so language papers print correctly,
    matching the body."""
    header = section.header
    header.is_linked_to_previous = False
    # Drop python-docx's auto-created empty paragraph so the box sits flush at the top.
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)

    def _style_run(run, size, bold=False):
        run.font.size = Pt(size)
        run.bold = bold
        if script_font:
            set_tamil_font(run, script_font)
        else:
            run.font.name = 'Times New Roman'
        return run

    def _tight(para, exact_pt):
        """Force an EXACT line height (in points). This is the real height lever:
        without WD_LINE_SPACING.EXACTLY, Word adds font-driven leading and ignores
        small line_spacing multipliers, so the box never actually shrinks. Text
        taller than exact_pt would clip, so keep exact_pt >= the run's font size."""
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(exact_pt)
        return para

    def _fix_row_height(row, twips):
        """Pin a table row to an EXACT height. Without hRule='exact', trHeight is
        only a minimum and Word grows the row to fit content — which is why the
        earlier tiny row heights did nothing."""
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        row.height = Twips(twips)

    # Outer 1-column table is the bordered box (9360 twips = 6.5", matching header1.xml).
    outer = header.add_table(rows=2, cols=1, width=Inches(6.5))
    _fix_table_width(outer, [9360])
    _set_cell_margins(outer, top=0, bottom=0)   # no vertical padding → compact box

    # Title block — school name + exam-name/period line, centred. No internal
    # divider: the top cell's bottom border is off so it merges with the grid below.
    top = outer.rows[0].cells[0]
    title_lines = [(school_name_val or "", 11, True),
                   (test_type_val or "", 10, True)]
    for idx, (txt, size, bold) in enumerate(title_lines):
        para = top.paragraphs[0] if idx == 0 else top.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _tight(para, exact_pt=size)   # exact line = font size (minimum, no leading)
        _style_run(para.add_run(txt), size, bold)
    _set_cell_border(top, top=_HDR_BORDER, left=_HDR_BORDER, right=_HDR_BORDER, bottom='none')

    # Compartment 2 — detail block: a 3-column grid. The subject sits centred on the
    # second line, between EXAM NO (left, blank for the student) and MARKS (right):
    #     CLASS : <class>                                  TIME  : <time>
    #     EXAM NO :               <SUBJECT>                MARKS : <marks>
    bot = outer.rows[1].cells[0]
    # Collapse the cell's mandatory leading empty paragraph to ~nothing (1pt exact)
    # instead of leaving a full blank line above the grid.
    _tight(bot.paragraphs[0], exact_pt=1)
    grid = bot.add_table(rows=2, cols=3)   # _Cell.add_table takes no width; pinned below
    _fix_table_width(grid, [3120, 3120, 3120])
    _set_cell_margins(grid, top=0, bottom=0)

    def _detail(cell, text, align="left", bold=False, size=10):
        para = cell.paragraphs[0]
        _tight(para, exact_pt=size)
        if align == "right":
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align == "center":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _style_run(para.add_run(text), size, bold)
        _set_cell_border(cell, top='none', left='none', bottom='none',
                         right='none', insideH='none', insideV='none')

    # Row 1: CLASS | (blank) | TIME
    _detail(grid.rows[0].cells[0], f"CLASS       :  {class_val}" if class_val else "CLASS       :")
    _detail(grid.rows[0].cells[1], "", align="center")
    _detail(grid.rows[0].cells[2], f"TIME        :  {time_val}" if time_val else "TIME        :", align="right")
    # Row 2: EXAM NO | SUBJECT (centred, bold) | MARKS
    _detail(grid.rows[1].cells[0], "EXAM NO  :")
    _detail(grid.rows[1].cells[1], subject_val or "", align="center", bold=True, size=10)
    _detail(grid.rows[1].cells[2], f"MARKS    :  {marks_val}" if marks_val else "MARKS    :", align="right")

    # A table cell may not end on a nested table — close with the smallest possible
    # (1pt) spacer so the detail grid sits right against the box's bottom border.
    tail = bot.add_paragraph()
    _tight(tail, exact_pt=1)
    _set_cell_border(bot, top='none', left=_HDR_BORDER, right=_HDR_BORDER, bottom=_HDR_BORDER)


def _add_passage_box(doc, text, script_font=None):
    """Render a passage in a bordered, shaded single-cell table. Markdown pipe tables
    inside the passage become real nested tables instead of raw "| … |" lines.
    Line breaks in the text are preserved (python-docx renders "\\n" as <w:br/>), so a
    poem quoted with its line breaks PRINTS as verse — the generation prompt's POEM
    FORMATTING rule is what guarantees the breaks are present in source_text."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    cell.text = ""

    def _fill_para(para, chunk):
        run = para.add_run(chunk)
        run.font.size = Pt(14)
        if script_font:
            set_tamil_font(run, script_font)
        else:
            run.font.name = 'Times New Roman'
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.left_indent = Inches(0.1)
        para.paragraph_format.right_indent = Inches(0.1)

    segs = _md_table_segments(text) or [("text", text)]
    first = True
    for kind, payload in segs:
        if kind == "text":
            chunk = payload.strip()
            if not chunk:
                continue
            _fill_para(cell.paragraphs[0] if first else cell.add_paragraph(), chunk)
            first = False
        else:
            inner = cell.add_table(rows=len(payload["rows"]),
                                   cols=max(len(r) for r in payload["rows"]))
            try:
                inner.style = 'Table Grid'
            except Exception:
                pass
            _fill_md_table(inner, payload, script_font)
            cell.add_paragraph()   # Word requires a paragraph after a nested table
            first = False

    # Light grey shading on cell
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F8F8F8')
    tcPr.append(shd)

    doc.add_paragraph("")


def _add_or_separator(doc):
    """Add an OR separator with thin paragraph borders above and below."""
    p = doc.add_paragraph()
    r = p.add_run("OR")
    r.bold = True
    r.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'bottom'):
        bdr = OxmlElement(f'w:{side}')
        bdr.set(qn('w:val'), 'single')
        bdr.set(qn('w:sz'), '6')
        bdr.set(qn('w:space'), '4')
        bdr.set(qn('w:color'), 'AAAAAA')
        pBdr.append(bdr)
    pPr.append(pBdr)


# ── Markdown pipe-table rendering ────────────────────────────────────────────
# LLMs legitimately emit data tables (observation tables, match-the-columns) as
# markdown "| a | b |" blocks inside question/passage text. Rendered as raw text
# they read as broken pipes — these helpers turn them into real Word tables.

_MD_TABLE_ROW_RE = re.compile(r'^\s*\|.+\|\s*$')


def _md_row_cells(line):
    s = line.strip()
    if s.startswith('|'):
        s = s[1:]
    if s.endswith('|'):
        s = s[:-1]
    return [c.strip() for c in s.split('|')]


def _is_md_separator_row(cells):
    """A markdown header separator row: every cell is dashes (with optional colons)."""
    return bool(cells) and all(re.fullmatch(r':?-{2,}:?', c.replace(' ', '')) for c in cells)


def _md_table_segments(text):
    """Split text into ('text', str) and ('table', {'rows': [...], 'header': bool}) segments.
    A table = 2+ consecutive lines that each start AND end with '|'. Separator rows
    (|---|---|) are dropped; one right after the first row marks it as a header row.
    Returns None when the text contains no table (the common fast path)."""
    if not isinstance(text, str) or text.count('|') < 4 or '\n' not in text:
        return None
    segs, text_buf, row_buf = [], [], []

    def _flush_text():
        if text_buf:
            chunk = '\n'.join(text_buf).strip('\n')
            if chunk.strip():
                segs.append(("text", chunk))
            text_buf.clear()

    def _flush_rows():
        if len(row_buf) >= 2:
            rows = [_md_row_cells(l) for l in row_buf]
            header = len(rows) > 1 and _is_md_separator_row(rows[1])
            rows = [r for r in rows if not _is_md_separator_row(r)]
            if rows:
                segs.append(("table", {"rows": rows, "header": header}))
                row_buf.clear()
                return
        text_buf.extend(row_buf)   # 0–1 pipe lines — ordinary text, not a table
        row_buf.clear()

    for line in text.split('\n'):
        if _MD_TABLE_ROW_RE.match(line):
            if not row_buf:
                _flush_text()
            row_buf.append(line)
        else:
            _flush_rows()
            text_buf.append(line)
    _flush_rows()
    _flush_text()

    if not any(k == "table" for k, _ in segs):
        return None
    return segs


def _fill_md_table(tbl, table, script_font=None):
    rows, header = table["rows"], table["header"]
    ncols = max(len(r) for r in rows)
    for i, row in enumerate(rows):
        cells = tbl.rows[i].cells
        for j in range(ncols):
            cell = cells[j]
            p = cell.paragraphs[0]   # fresh cells hold one empty paragraph, no runs
            run = p.add_run(row[j] if j < len(row) else "")
            run.font.size = Pt(12)
            run.bold = bool(header and i == 0)
            if script_font:
                set_tamil_font(run, script_font)
            else:
                run.font.name = 'Times New Roman'
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)


def _add_md_table(doc, table, script_font=None, indent_inches=0.35):
    """Add a real Word table for a parsed markdown table, indented to the question-body column."""
    rows = table["rows"]
    tbl = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    try:
        tbl.style = 'Table Grid'
    except Exception:
        pass
    tbl.autofit = True
    _fill_md_table(tbl, table, script_font)
    try:
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), str(int(indent_inches * 1440)))
        tblInd.set(qn('w:type'), 'dxa')
        tbl._tbl.tblPr.append(tblInd)
    except Exception:
        pass
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)


def _render_question_segments(doc, segs, full_text, marks_pattern, left_indent=None,
                              script_font=None, right_tab_twips=8280):
    """Render a question whose text embeds markdown pipe tables: the stem keeps the
    number + right-aligned marks treatment, each pipe block becomes a real Word table,
    and text after a table continues as body-indented paragraphs."""
    base_in = left_indent.inches if left_indent else 0.0
    m = marks_pattern.search(full_text)
    marks_tag = f" [{m.group(1)} marks]" if m else ""
    stem_done = False
    for kind, payload in segs:
        if kind == "text":
            chunk = payload.strip()
            if not chunk:
                continue
            if not stem_done:
                _add_question_with_marks(doc, chunk + marks_tag, marks_pattern, left_indent,
                                         script_font, right_tab_twips=right_tab_twips)
                stem_done = True
            else:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(base_in + 0.35)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(chunk)
                run.font.size = Pt(14)
                if script_font:
                    set_tamil_font(run, script_font)
                else:
                    run.font.name = 'Times New Roman'
        else:
            _add_md_table(doc, payload, script_font, indent_inches=base_in + 0.35)
            stem_done = True


def _add_question_with_marks(doc, text, marks_pattern, left_indent=None, script_font=None,
                             right_tab_twips=8280):
    """Add a question paragraph.

    The leading number/label ("8.", "(i)") is set in a HANGING INDENT so it sits alone in a
    narrow left column and every line of the question body — the first line and all wrapped
    lines — aligns at a common position to the RIGHT of the number (no continuation line ever
    wraps back underneath the number). Marks, if present, are right-aligned to
    ``right_tab_twips`` (the real text-area width, so they sit flush at the right margin)."""
    match = marks_pattern.search(text)
    marks_str = f"[{match.group(1)}]" if match else ""
    clean_text = marks_pattern.sub("", text).rstrip()

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(2)

    def _qrun(run):
        run.font.size = Pt(14)
        if script_font:
            set_tamil_font(run, script_font)
        else:
            run.font.name = 'Times New Roman'

    # Split the leading number/label ("1. ", "(i) ", "Q3. ") so we can bold it and hang-indent
    # the body. rstrip the captured whitespace — a tab now separates the number from the text.
    num_match = re.match(r'^(\([a-zA-Z0-9]+\)\s*|[ivxIVX]+\.\s*|[0-9]+[.)]\s*)', clean_text)
    if num_match:
        num_part = num_match.group(1).rstrip()
        rest_part = clean_text[num_match.end():]
    else:
        num_part = ""
        rest_part = clean_text

    base_inches = left_indent.inches if left_indent else 0.0
    HANG_IN = 0.35   # width of the number column / hanging indent

    if num_part:
        # Hanging indent: number at base_inches, body text (+ every wrap) at base_inches + HANG.
        p.paragraph_format.left_indent = Inches(base_inches + HANG_IN)
        p.paragraph_format.first_line_indent = Inches(-HANG_IN)
    elif left_indent:
        p.paragraph_format.left_indent = left_indent

    # Tab stops: a LEFT tab at the hang position aligns the body text after the number; a RIGHT
    # tab at the margin right-aligns the marks. Use python-docx's API (not a manual append) so
    # the <w:tabs> element is inserted at its schema-correct position in <w:pPr> — appending it
    # after <w:ind> produces invalid OOXML that Word silently mishandles (the left tab is
    # dropped, throwing the first line far to the right).
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Twips
    tstops = p.paragraph_format.tab_stops
    if num_part:
        tstops.add_tab_stop(Twips(int(round((base_inches + HANG_IN) * 1440))), WD_TAB_ALIGNMENT.LEFT)
    if marks_str:
        tstops.add_tab_stop(Twips(int(right_tab_twips)), WD_TAB_ALIGNMENT.RIGHT)

    if num_part:
        r_num = p.add_run(f"{num_part}\t")   # number + tab → body text snaps to the hang column
        r_num.bold = True
        _qrun(r_num)
    _add_math_runs(p, rest_part, _qrun)
    if marks_str:
        r_marks = p.add_run(f"\t{marks_str}")
        r_marks.bold = True
        _qrun(r_marks)

    return p

    return p


def _parse_edited_text(text):
    """
    Parse AI-corrected plain text back into (type, text) tuples consumable by render_docx.

    Detects:
      header     — SECTION A / SECTION B … lines
      q          — numbered questions  "1. ..."
      subq       — roman sub-questions "(i) ..." / "(ii) ..."
      opts       — MCQ options line    "(a) ... (b) ..."  or single "(a) ..."
      or         — standalone OR
      instruction — everything else (passage lead-ins, general instructions)
    """
    SECTION_RE  = re.compile(r'^SECTION\s+[A-Z]', re.IGNORECASE)
    QUESTION_RE = re.compile(r'^\d+[\.\)]\s+\S')
    ROMAN_RE    = re.compile(r'^\([ivxlIVXL]+\)\s+\S')
    OPTION_RE   = re.compile(r'^\([a-dA-D]\)\s+\S')
    OR_RE       = re.compile(r'^OR$', re.IGNORECASE)
    IMGFILE_RE  = re.compile(r'\[IMG_FILE:\s*([^\]]+?)\]', re.IGNORECASE)

    result = []
    for raw_line in text.split('\n'):
        line = raw_line.strip()
        if not line:
            continue
        # Image-file markers (from the editor round-trip) → ('image', path), kept in place
        # so a diagram still renders above its question. Strip them, then classify the rest.
        if IMGFILE_RE.search(line):
            for _path in IMGFILE_RE.findall(line):
                result.append(('image', _path.strip()))
            line = IMGFILE_RE.sub('', line).strip()
            if not line:
                continue
        if OR_RE.match(line):
            result.append(('or', line))
        elif SECTION_RE.match(line):
            result.append(('header', line))
        elif QUESTION_RE.match(line):
            result.append(('q', line))
        elif ROMAN_RE.match(line):
            result.append(('subq', line))
        elif OPTION_RE.match(line):
            result.append(('opts', line))
        else:
            result.append(('instruction', line))
    return result


def render_docx(class_name, subject, chapters, all_questions, summary, header_meta=None):
    # The header is generated in code (see _build_header) — no base.docx template.
    doc = Document()

    # Pick a complex-script font if this is a Tamil/Hindi/Sanskrit paper (by subject or by
    # scanning the text). None → ordinary Latin font. `is_tamil` kept as the in-body gate, now
    # meaning "needs a complex-script font"; `script_font` says which one.
    script_font = _pick_script_font(subject, all_questions)
    is_tamil = bool(script_font)
    if script_font:
        print(f"[DOCX-Script] complex script detected — applying font '{script_font}'")
        apply_tamil_document_styles(doc, script_font)
    
    # Set document margins
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # Build the header (school / subject / test type / class / time / marks)
    if header_meta is None:
        header_meta = {}

    test_type_val = _header_title_line(header_meta.get("test_type", header_meta.get("pattern_name", "")))
    class_val = header_meta.get("class_name", class_name) or class_name
    subject_val = header_meta.get("subject", subject) or subject
    time_val = str(header_meta.get("duration", "")).strip()
    marks_val = str(header_meta.get("marks", "")).strip()
    school_name_val = str(header_meta.get("school_name", "")).strip()

    try:
        _build_header(section, subject_val, class_val, time_val, marks_val, test_type_val, school_name_val, script_font=script_font)
        print(f"[DOCX-Header] Built header in code — school={school_name_val!r} class={class_val!r} subject={subject_val!r} time={time_val!r} marks={marks_val!r} test_type={test_type_val!r}")
    except Exception as e:
        print(f"[DOCX-Header] WARNING: header build failed: {e}")

    # Set page margins (leave room for the code-built header)
    section = doc.sections[0]
    section.top_margin = Inches(1.2)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # Right-tab position for per-question marks = the actual text-area width, so marks sit flush
    # at the right margin (the old fixed 5.75" left them floating ~1.25" short of the edge).
    # OOXML tab positions are twips; 635 EMU = 1 twip. Small epsilon keeps them off the very edge.
    try:
        _usable_emu = int(section.page_width) - int(section.left_margin) - int(section.right_margin)
        marks_tab_twips = max(1440, int(_usable_emu / 635) - 12)
    except Exception:
        marks_tab_twips = 8280

    # Body typography: Times New Roman, 14pt, 1.15 line spacing. Applied to the Normal style so
    # EVERY paragraph inherits the 1.15 spacing (the code doesn't set line_spacing per paragraph);
    # the explicit run sizes below set 14pt directly. For Tamil/Hindi keep the script font.
    try:
        _normal = doc.styles['Normal']
        _normal.font.size = Pt(14)
        _normal.paragraph_format.line_spacing = 1.15
        if not is_tamil:
            _normal.font.name = 'Times New Roman'
    except Exception as _e:
        print(f"[DOCX] Normal-style typography setup failed: {_e}")

    # Restrict the school header to the first page only.
    # In OOXML: add <w:titlePg/> to sectPr (enables different first-page header),
    # then change any type="default" headerReference to type="first" so pages 2+
    # have no header reference → blank header.
    try:
        sectPr = section._sectPr
        if sectPr.find(qn('w:titlePg')) is None:
            titlePg = OxmlElement('w:titlePg')
            sectPr.insert(0, titlePg)
        for hdr_ref in sectPr.findall(qn('w:headerReference')):
            if hdr_ref.get(qn('w:type')) == 'default':
                hdr_ref.set(qn('w:type'), 'first')
    except Exception as _e:
        print(f"[DOCX-Header] first-page-only setup failed: {_e}")

    marks_pattern = re.compile(r"\s*\[(\d+)\s*marks?\]", re.IGNORECASE)

    for typ, text in all_questions:
        text_str = text if isinstance(text, str) else str(text)
        if typ == "header":
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(0)
            p = doc.add_paragraph()
            # A teacher-authored heading carries its arithmetic after a tab
            # ("II. Answer any SIX of the following\t6 x 2 = 12"): print it as the teacher
            # wrote it — left aligned, sum flush at the right margin, no centering or rule.
            head_left, _tab, head_calc = text_str.partition("\t")
            if head_calc:
                from docx.enum.text import WD_TAB_ALIGNMENT
                from docx.shared import Twips
                p.paragraph_format.tab_stops.add_tab_stop(
                    Twips(int(marks_tab_twips)), WD_TAB_ALIGNMENT.RIGHT)
                r = p.add_run(f"{head_left}\t{head_calc}")
            else:
                r = p.add_run(text_str)
            r.bold = True
            r.font.size = Pt(16)   # section header — larger than the 14pt body
            if not is_tamil:
                r.font.name = 'Times New Roman'
            if is_tamil:
                set_tamil_font(r, script_font)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(6)
            if not head_calc:
                # CBSE-style banner: centered with a bottom rule.
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pPr = p._element.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bdr = OxmlElement('w:bottom')
                bdr.set(qn('w:val'), 'single')
                bdr.set(qn('w:sz'), '6')
                bdr.set(qn('w:space'), '4')
                bdr.set(qn('w:color'), '555555')
                pBdr.append(bdr)
                pPr.append(pBdr)
        elif typ == "subheader":
            p = doc.add_paragraph()
            r = p.add_run(text_str)
            r.bold = True
            r.font.size = Pt(14)
            if not is_tamil:
                r.font.name = 'Times New Roman'
            if is_tamil:
                set_tamil_font(r, script_font)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
        elif typ == "instruction":
            p = doc.add_paragraph()
            r = p.add_run(text_str)
            r.font.size = Pt(14)
            r.italic = True
            if not is_tamil:
                r.font.name = 'Times New Roman'
            if is_tamil:
                set_tamil_font(r, script_font)
            p.paragraph_format.space_after = Pt(4)
        elif typ in ("q", "subq"):
            indent = Inches(0.25) if typ == "subq" else None
            # Strip the marks tag BEFORE table detection — "[3 marks]" appended to the
            # last "| … |" row would otherwise stop that line matching as a table row.
            segs = _md_table_segments(marks_pattern.sub("", text_str).rstrip())
            if segs:
                _render_question_segments(doc, segs, text_str, marks_pattern, indent,
                                          script_font, right_tab_twips=marks_tab_twips)
            else:
                _add_question_with_marks(doc, text_str, marks_pattern, indent, script_font,
                                         right_tab_twips=marks_tab_twips)
        elif typ == "opts":
            p = doc.add_paragraph()
            r = p.add_run(text_str)
            r.font.size = Pt(14)
            if not is_tamil:
                r.font.name = 'Times New Roman'
            if is_tamil:
                set_tamil_font(r, script_font)
            p.paragraph_format.left_indent = Inches(0.6)   # options nest INSIDE the 0.35in question-body column
            p.paragraph_format.space_after = Pt(4)
        elif typ == "opts_block":
            try:
                opts = list(text) if isinstance(text, (list, tuple)) else []
                # Use 2-per-row only when all options are short enough to fit side-by-side.
                # If any option exceeds 45 chars, fall back to one option per line.
                two_col = all(len(o) <= 45 for o in opts)
                step = 2 if two_col else 1
                for row_start in range(0, len(opts), step):
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.6)   # options nest INSIDE the 0.35in question-body column
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(0)
                    if two_col:
                        pPr = p._element.get_or_add_pPr()
                        tabs = OxmlElement('w:tabs')
                        tab = OxmlElement('w:tab')
                        tab.set(qn('w:val'), 'left')
                        tab.set(qn('w:pos'), '4320')  # ~3 inches
                        tabs.append(tab)
                        pPr.append(tabs)
                    opt1 = opts[row_start]
                    r1 = p.add_run(opt1)
                    r1.font.size = Pt(14)
                    if not is_tamil:
                        r1.font.name = 'Times New Roman'
                    if is_tamil:
                        set_tamil_font(r1, script_font)
                    if two_col and row_start + 1 < len(opts):
                        opt2 = opts[row_start + 1]
                        r2 = p.add_run(f"\t{opt2}")
                        r2.font.size = Pt(14)
                        if not is_tamil:
                            r2.font.name = 'Times New Roman'
                        if is_tamil:
                            set_tamil_font(r2, script_font)
                doc.add_paragraph("")
            except Exception as e:
                print(f"[DOCX] opts_block failed: {e}")
        elif typ == "passage":
            _add_passage_box(doc, text_str, script_font)
        elif typ == "or":
            _add_or_separator(doc)
        elif typ == "image":
            try:
                img_path = os.path.join(settings.MEDIA_ROOT, text_str) if not os.path.isabs(text_str) else text_str
                # Smaller display width (source PNG is full-res, so quality is unchanged —
                # a smaller box just raises the effective DPI). Centered above the question.
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(img_path, width=Inches(3.0))
            except Exception as e:
                print(f"[DOCX] Image insert failed: {e}")

    output_dir = os.path.join(settings.MEDIA_ROOT, "question_papers")
    os.makedirs(output_dir, exist_ok=True)

    safe_subject = subject.replace(" ", "_")
    filename = f"{safe_subject}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    file_path = os.path.join(output_dir, filename)
    doc.save(file_path)

    return f"question_papers/{filename}", summary


# ------------------------------
# Image generation helpers
# ------------------------------
def _openai_image_size(width: int, height: int) -> str:
    """Map requested dimensions to a gpt-image-1 supported size string."""
    if width > height:
        return "1536x1024"
    if height > width:
        return "1024x1536"
    return "1024x1024"


# Number of OpenAI attempts before falling back to Together AI → Pollinations.
IMAGE_OPENAI_ATTEMPTS = int(os.environ.get("IMAGE_OPENAI_ATTEMPTS", "2"))


class ImageNotCached(Exception):
    """Raised when an image is requested in cache-only mode but isn't already on disk.

    Callers normally use this to skip missing images in fast synchronous paths.
    Some rerender flows may catch it and opt into on-demand generation explicitly."""


def _decode_image_item(item, _requests):
    """Pull image bytes from a Together/OpenAI 'data' item (b64 or url)."""
    b64 = item.get("b64_json") or item.get("b64")
    if b64:
        import base64
        return base64.b64decode(b64)
    if item.get("url"):
        return _requests.get(item["url"], timeout=60).content
    return None


def _openai_image_bytes(prompt: str, size: str, _requests) -> bytes:
    api_key = os.environ.get('OPENAI_API_KEY') or getattr(settings, 'OPENAI_API_KEY', '')
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY not set in environment")
    model_id = getattr(settings, 'OPENAI_IMAGE_MODEL', 'gpt-image-1')
    with mantle_client.external_call(f"openai-image:{model_id}", f"size={size}"):
        resp = _requests.post(
            "https://api.openai.com/v1/images/generations",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json={"model": model_id, "prompt": prompt, "n": 1, "size": size},
            timeout=180,
        )
    if not resp.ok:
        raise RuntimeError(f"OpenAI image HTTP {resp.status_code}: {resp.text[:300]}")
    data = resp.json().get("data") or []
    if not data:
        raise RuntimeError("No data in OpenAI image response")
    img = _decode_image_item(data[0], _requests)
    if img is None:
        raise RuntimeError("No image data in OpenAI response item")
    return img


def _together_image_bytes(prompt: str, width: int, height: int, _requests) -> bytes:
    key = os.environ.get('TOGETHER_API_KEY') or getattr(settings, 'TOGETHER_API_KEY', '')
    if not key:
        raise RuntimeError("TOGETHER_API_KEY not set in environment")
    model_id = getattr(settings, 'TOGETHER_IMAGE_MODEL', 'google/flash-image-2.5')
    with mantle_client.external_call(f"together-image:{model_id}", f"{width}x{height}"):
        resp = _requests.post(
            "https://api.together.ai/v1/images/generations",
            headers={"Authorization": f"Bearer {key}"},
            json={"model": model_id, "prompt": prompt, "n": 1, "width": width, "height": height},
            timeout=120,
        )
    if not resp.ok:
        raise RuntimeError(f"Together AI HTTP {resp.status_code}: {resp.text[:300]}")
    data = resp.json().get("data") or []
    if not data:
        raise RuntimeError("No data in Together AI response")
    img = _decode_image_item(data[0], _requests)
    if img is None:
        raise RuntimeError("No image data in Together AI response item")
    return img


def _pollinations_image_bytes(prompt: str, _requests) -> bytes:
    import urllib.parse
    url = f"https://image.pollinations.ai/prompt/{urllib.parse.quote(prompt)}"
    params = {"width": 1024, "height": 1024, "nologo": "true", "private": "true"}
    pk = os.environ.get("POLLINATIONS_API_KEY", "").strip()
    if pk:
        params["token"] = pk
    with mantle_client.external_call("pollinations", "1024x1024"):
        resp = _requests.get(url, params=params, timeout=120)
        resp.raise_for_status()
    return resp.content


def generate_ai_image(prompt: str, width: int = 1024, height: int = 1024, cfg_scale: float = 8.0, cache_only: bool = False) -> str:
    """Generate an image and save under MEDIA_ROOT/generated_images. Returns relative media path.

    Provider chain: OpenAI gpt-image-1 (primary, IMAGE_OPENAI_ATTEMPTS tries) →
    Together AI → Pollinations (fallbacks, only after OpenAI exhausts its retries).
    (cfg_scale kept for signature compatibility; unused.)"""
    import hashlib, requests as _requests
    output_dir = getattr(settings, 'IMAGE_OUTPUT_DIR', os.path.join(settings.MEDIA_ROOT, 'generated_images'))
    os.makedirs(output_dir, exist_ok=True)
    if not prompt or not str(prompt).strip():
        raise ValueError("Empty image prompt")
    key = hashlib.sha256(f"{prompt}|{width}|{height}".encode('utf-8')).hexdigest()[:24]
    rel_path = f"generated_images/{key}.png"
    abs_path = os.path.join(settings.MEDIA_ROOT, rel_path)
    if os.path.exists(abs_path):
        return rel_path
    if cache_only:
        # Web-request re-render path: the image isn't cached and we must not block the
        # request on a slow external API call. Skip it (raise → caller drops the image).
        raise ImageNotCached(str(prompt)[:80])

    size = _openai_image_size(width, height)
    img_bytes = None
    last_err = None

    # ── Primary: OpenAI gpt-image-1, with retries ───────────────────────────────
    for attempt in range(1, IMAGE_OPENAI_ATTEMPTS + 1):
        try:
            print(f"[ImageGen] OpenAI gpt-image-1 attempt {attempt}/{IMAGE_OPENAI_ATTEMPTS} | "
                  f"prompt={prompt[:70]}... | size={size}")
            img_bytes = _openai_image_bytes(prompt, size, _requests)
            break
        except Exception as e:
            last_err = e
            print(f"[ImageGen] OpenAI attempt {attempt} failed: {e}")

    # ── Fallback 1: Together AI ──────────────────────────────────────────────────
    if img_bytes is None:
        try:
            print("[ImageGen] OpenAI exhausted — falling back to Together AI")
            img_bytes = _together_image_bytes(prompt, width, height, _requests)
        except Exception as e:
            print(f"[ImageGen] Together AI fallback failed: {e}")

    # ── Fallback 2: Pollinations ─────────────────────────────────────────────────
    if img_bytes is None:
        try:
            print("[ImageGen] Together unavailable — falling back to Pollinations")
            img_bytes = _pollinations_image_bytes(prompt, _requests)
        except Exception as e:
            print(f"[ImageGen] Pollinations fallback failed: {e}")

    if img_bytes is None:
        raise RuntimeError(f"All image providers failed (last OpenAI error: {last_err})")

    with open(abs_path, 'wb') as f:
        f.write(img_bytes)
    return rel_path


def materialize_images(all_questions, allow=True, cache_only=False, generate_missing_images=False):
    """Convert ('image_gen', prompt) and [IMAGE:/[Picture:/[Diagram:] markers (case-insensitive) to ('image', rel_path).
    Supports markers embedded in any text type (q, subq, instruction, passage).

    cache_only=True (synchronous re-render path): reuse images already on disk but never
    call the external image APIs — unless generate_missing_images=True, in which case
    cache misses are generated on demand."""
    if not allow:
        return all_questions
    import re
    # Match [Picture: ...], [Image: ...], or [Diagram: ...] (colon optional), any case
    pattern = re.compile(r"\[(image|picture|diagram)\s*:?\s*(.*?)\]", re.IGNORECASE)
    out = []
    for typ, text in all_questions:
        try:
            # Direct generation tuple
            if typ == 'image_gen':
                print(f"[ImageGen] image_prompt detected: {str(text)[:120]}...")
                try:
                    rel = generate_ai_image(str(text), cache_only=cache_only)
                    out.append(('image', rel))
                except ImageNotCached:
                    if generate_missing_images:
                        print(f"[ImageGen] cache miss during re-render — generating image on demand ({str(text)[:60]})")
                        try:
                            rel = generate_ai_image(str(text), cache_only=False)
                            out.append(('image', rel))
                        except Exception as ge:
                            print(f"[ImageGen] On-demand generation failed for prompt '{str(text)[:60]}': {ge}")
                    else:
                        print(f"[ImageGen] cache-only re-render: skipping uncached image ({str(text)[:60]})")
                continue

            if isinstance(text, str):
                prompts = pattern.findall(text)
                if prompts:
                    cleaned = pattern.sub('', text).strip()
                    appended_text = False
                    # Track whether any image was generated for this text
                    generated_any = False

                    for _kind, prompt in prompts:
                        prompt_clean = prompt.strip()

                        # Skip common labels like "Based" that aren't real prompts
                        if not prompt_clean or prompt_clean.lower() in {"based", "picture based", "diagram based"}:
                            if not appended_text:
                                out.append((typ, text.strip()))
                                appended_text = True
                            continue

                        if not appended_text and cleaned:
                            out.append((typ, cleaned))
                            appended_text = True

                        print(f"[ImageGen] marker detected: {_kind} | {prompt_clean[:120]}...")
                        try:
                            rel = generate_ai_image(prompt_clean, cache_only=cache_only)
                            out.append(('image', rel))
                            generated_any = True
                        except ImageNotCached:
                            if generate_missing_images:
                                print(f"[ImageGen] cache miss during re-render — generating marker image on demand ({prompt_clean[:60]})")
                                try:
                                    rel = generate_ai_image(prompt_clean, cache_only=False)
                                    out.append(('image', rel))
                                    generated_any = True
                                except Exception as ge:
                                    print(f"[ImageGen] On-demand marker generation failed for prompt '{prompt_clean[:60]}': {ge}")
                            else:
                                print(f"[ImageGen] cache-only re-render: skipping uncached image ({prompt_clean[:60]})")
                        except Exception as ge:
                            print(f"[ImageGen] Generation failed for prompt '{prompt_clean[:60]}': {ge}")

                    if not appended_text and not generated_any:
                        out.append((typ, text))
                    continue

            # Default passthrough
            out.append((typ, text))
        except Exception as e:
            print(f"[ImageGen] WARNING: {e}")
            out.append((typ, text))
    return out


# ------------------------------
# Render helper for parallel pipeline
# ------------------------------
_COST_PER_INPUT_1K  = 0.49   # INR per 1k input tokens
_COST_PER_OUTPUT_1K = 1.47   # INR per 1k output tokens

def _render_paper_from_data(paper_data, blueprint, class_name, subject, chapters, additional_context, pattern, total_input_tokens=0, total_output_tokens=0, allow_images=True, cache_only=False, generate_missing_images=False):
    """
    Render a pre-generated paper_data dict to DOCX.
    Used by the parallel pipeline after generate_paper_parallel() succeeds.
    Returns (file_path, summary, total_cost, total_input_tokens, total_output_tokens).
    """
    # Store paper data in thread-local state so tasks.py can persist it
    # without reading from a shared temp file (avoids cross-user data contamination).
    _request_state.paper_data = paper_data

    # Strip internal pipeline metadata fields (prefixed with _) before rendering
    # These hold validation reports and warnings and must not appear in the DOCX
    _INTERNAL_KEYS = {
        "_competency_report", "_coherence_report", "_final_audit",
        "_uniqueness_warnings", "_mcq_answer_warnings", "_mcq_answer_corrections", "_quality_flags",
        "_grounding_issues", "_cbq_passage_issues", "_cross_section_duplicates",
        "_partial", "_errors", "_chapter_plan", "_dropped_wrong_type", "_topped_up",
    }
    render_data = {}
    for sec_name, sec_data in paper_data.items():
        if sec_name.startswith("__"):  # sentinel keys e.g. __context_by_type__
            continue
        clean_sec = {k: v for k, v in sec_data.items() if k not in _INTERNAL_KEYS}
        render_data[sec_name] = clean_sec

    all_questions = render_section_questions([], render_data, blueprint, class_name, subject, chapters, None)
    all_questions = materialize_images(
        all_questions,
        allow=allow_images,
        cache_only=cache_only,
        generate_missing_images=generate_missing_images,
    )

    summary = {sec: {"title": sec, "marks": render_data.get(sec, {}).get("marks", 0)} for sec in render_data.keys()}

    header_meta = {
        "class_name": class_name,
        "subject": subject,
        "pattern_name": getattr(pattern, "name", ""),
        "marks": _paper_total_marks(pattern, blueprint),
    }
    if additional_context:
        try:
            ctx_obj = json.loads(additional_context)
            if isinstance(ctx_obj, dict):
                for _k in ("class_name", "duration", "marks", "test_type", "school_name"):
                    if ctx_obj.get(_k):
                        header_meta[_k] = ctx_obj[_k]
        except Exception:
            pass

    total_cost = (total_input_tokens / 1000) * _COST_PER_INPUT_1K + (total_output_tokens / 1000) * _COST_PER_OUTPUT_1K
    file_path, summary = render_docx(class_name, subject, chapters, all_questions, summary, header_meta=header_meta)
    return file_path, summary, total_cost, total_input_tokens, total_output_tokens


# ------------------------------
# Entrypoint
# ------------------------------
def _paper_total_marks(pattern, blueprint):
    """The marks the printed paper actually adds up to.

    pattern.total_marks is a stored integer, recomputed only when the pattern is saved — one
    saved before attempt-N-of-M was understood still carries the sum of every printed question
    (90 for an 80-mark paper). The blueprint is rebuilt from the slots on every run, so prefer it
    whenever every section states its marks; fall back to the stored value otherwise.
    """
    stored = getattr(pattern, "total_marks", 0) or 0
    if not isinstance(blueprint, dict) or not blueprint:
        return stored
    marks = [float(s.get("marks") or 0) for s in blueprint.values() if isinstance(s, dict)]
    if not marks or any(m <= 0 for m in marks):
        return stored
    total = sum(marks)
    total = int(total) if total == int(total) else round(total, 2)
    if stored and total != stored:
        print(f"[Header] Paper total {total} from the blueprint (pattern.total_marks={stored})")
    return total


def pattern_sections_to_blueprint_dict(pattern):
    """
    Convert pattern.sections (list) to blueprint dict format expected by generators.
    Pattern sections: list of section objects with 'name', 'marks', 'questions_count', 'question_types'
    Blueprint format: dict with section name as key, section data as value
    """
    if not pattern or not pattern.sections:
        return {}

    blueprint_dict = {}
    for section in pattern.sections:
        section_name = section.get('name', 'Section')
        title = section.get('title', '')

        # CBSE-seeded sections store count as 'questions'; user blueprints use 'questions_count'
        q_count = (
            section.get('questions_count')
            or section.get('questions')
            or 0
        )

        # Resolve the section's question type(s), tolerating every shape a pattern can carry:
        #   1. 'question_types' — the canonical plural list
        #   2. 'question_type'  — singular string/list (the shape the generate-page form saves);
        #      without this fallback the type is LOST, leaving the section unconstrained so the
        #      model freely mixes types (MCQs landing in a Short-Answer section, etc.)
        #   3. 'title'          — CBSE-seeded sections store the type there ("MCQ + Assertion-Reason")
        qt = section.get('question_types') or []
        if not qt:
            single = section.get('question_type')
            if single:
                qt = [single] if isinstance(single, str) else list(single)
        if not qt and title:
            qt = [t.strip() for t in re.split(r'[,+&/]|\band\b', title) if t.strip()]

        # What the student can earn: an "answer any SIX of eight" section prints 8 questions but
        # is worth 6 x 2 = 12. Recomputed here rather than trusted from the stored section, so a
        # pattern saved before attempt-N-of-M was understood prints the right marks without a
        # re-save. Falls back to the stored value for sections with no slots.
        marks = pattern_structure.attemptable_marks(section) or section.get('marks', 0)
        marks = int(marks) if float(marks) == int(marks) else marks
        attempt = pattern_structure.section_attempt(section)
        # `marks` is what the section is WORTH (attempt-N-of-M counts only N), so the per-question
        # value divides by the number answered, not the number printed — eight 2-mark questions
        # marked "answer any six" are worth 12, and 12/8 would call each one 1.5m.
        marks_per_q = section.get('marks_per_question') or section.get('marks_each') or (
            round(marks / (attempt or q_count), 2) if (attempt or q_count) else 1
        )

        blueprint_dict[section_name] = {
            'id': section.get('id', ''),
            'title': title,
            'marks': marks,
            'questions_count': q_count,
            'question_types': qt,
            'marks_per_question': marks_per_q,
            'instructions': section.get('instructions', []),
            'constraints': section.get('constraints', {}),
            'subsections': section.get('subsections', []),
            # C-01: preserve per-section sub-subject for compound papers
            'section_subject': section.get('subject', ''),
            # MO-01: preserve attempt-N-of-M counts
            # Seeded format stores 'attempt' directly; raw pattern dict uses 'attempt' too
            'attempt_count': attempt or section.get('attempt'),
            'provided_count': q_count,   # questions_count = the full provided set
            # Store raw question_types detail for mixed-marks detection
            'question_type_details': section.get('question_types', []),
            # Per-question structure (docs/PER_QUESTION_STRUCTURE.md): slots drive the
            # work order and disable render-time regrouping/marks re-stamping.
            'question_slots': section.get('question_slots', []),
        }

    return blueprint_dict if blueprint_dict else None


def generate_universal_paper(class_name, subject, chapters, difficulty, pattern, section=None, model_source='aws', additional_context="", school_id=None,
                            unit_map=None, creative_ratio=0):
    """
    Universal question paper generator that works for all subjects and question types.
    Uses database-driven blueprints and adaptive prompt generation.
    """
    # Create generation summary log
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    summary_file = f"temp_generation_summary_{timestamp}.txt"
    
    with open(summary_file, "w", encoding="utf-8") as f:
        f.write(f"=== UNIVERSAL QUESTION PAPER GENERATION ===\n")
        f.write(f"Timestamp: {datetime.now()}\n")
        f.write(f"Class: {class_name}\n")
        f.write(f"Subject: {subject}\n")
        f.write(f"Chapters: {', '.join(chapters) if chapters else 'All'}\n")
        f.write(f"Difficulty: {difficulty}\n")
        f.write(f"Pattern: {pattern.name}\n")
        f.write(f"Section: {section or 'All'}\n")
        f.write(f"Model Source: {model_source}\n")
        f.write(f"Source Mix: {creative_ratio}% own composition / {100 - int(creative_ratio or 0)}% from the book\n")
        f.write(f"Additional Context Length: {len(additional_context)} chars\n")
        f.write(f"{'='*50}\n\n")
    
    print(f"[Universal-Generator] Starting generation for {class_name} {subject}")
    print(f"[Universal-Generator] Summary saved to: {summary_file}")
    
    try:
        # Use pattern's sections converted to blueprint format
        blueprint = pattern_sections_to_blueprint_dict(pattern)
        if not blueprint:
            # Fallback to old system if pattern conversion fails
            blueprint = get_blueprint(class_name, subject, section)
            print(f"[Universal-Generator] ✅ Blueprint resolved: {len(blueprint)} sections")
        
        # Extract all question types from blueprint
        all_question_types = []
        
        # Blueprint is already normalized by BlueprintManager.get_blueprint()
        blueprint_dict = blueprint

        # Debug: Verify normalized structure
        print(f"[Universal-Generator] Blueprint dict type: {type(blueprint_dict)}")
        print(f"[Universal-Generator] Blueprint dict keys: {list(blueprint_dict.keys())}")

        # Check if this is a complex blueprint (has non-empty subsections)
        is_complex_blueprint = False
        for section_key, section_data in blueprint_dict.items():
            if isinstance(section_data, dict) and section_data.get('subsections'):
                is_complex_blueprint = True
                break

        if is_complex_blueprint:
            # Extract question types from complex blueprint structure
            for section_key, section_data in blueprint_dict.items():
                if isinstance(section_data, dict) and section_data.get('subsections'):
                    subsections = section_data['subsections']
                    # Handle both list and dict formats for subsections
                    if isinstance(subsections, list):
                        for subsection in subsections:
                            if isinstance(subsection, dict):
                                subsec_types = subsection.get('question_types', [])
                                all_question_types.extend(subsec_types)
                    elif isinstance(subsections, dict):
                        for subsection_key, subsection_questions in subsections.items():
                            if isinstance(subsection_questions, list):
                                for question in subsection_questions:
                                    if isinstance(question, dict) and 'type' in question:
                                        all_question_types.append(question['type'])
        else:
            # Extract question types from simple blueprint structure
            for section_name, section_data in blueprint_dict.items():
                section_question_types = section_data.get('question_types', [])
                all_question_types.extend(section_question_types)
        
        # Remove duplicates while preserving order (dict items are not hashable)
        seen_keys, unique_types = set(), []
        for qt in all_question_types:
            key = str(qt)
            if key not in seen_keys:
                seen_keys.add(key)
                unique_types.append(qt)
        all_question_types = unique_types
        print(f"[Universal-Generator] Blueprint type: {'Complex' if is_complex_blueprint else 'Simple'}")
        print(f"[Universal-Generator] Question types: {all_question_types}")
        
        # ── Attempt 1: parallel per-section pipeline ────────────────────────
        try:
            from .section_generator import generate_paper_parallel, get_section_context_map
            print("[Universal-Generator] Trying parallel per-section pipeline...")
            # Superadmin per-school kill switch for AI image generation. Resolved once here
            # (worker threads can't hit a request-scoped state), threaded onto each work order.
            _disable_images = False
            if school_id:
                try:
                    from .models import School
                    _disable_images = School.objects.filter(pk=school_id).values_list(
                        'disable_image_generation', flat=True).first() or False
                except Exception as _de:
                    print(f"[Universal-Generator] Could not resolve image-generation flag: {_de}")
            if _disable_images:
                print(f"[Universal-Generator] Image generation is DISABLED for school {school_id}")
            _context_map = get_section_context_map(class_name, subject, chapters, blueprint_dict, all_question_types, school_id=school_id)
            _paper_data, _in_tok, _out_tok = generate_paper_parallel(
                blueprint=blueprint_dict,
                pattern=pattern,
                context_map=_context_map,
                difficulty=difficulty,
                class_name=class_name,
                subject=subject,
                chapters=chapters,
                disable_images=_disable_images,
                unit_map=unit_map,
                creative_ratio=creative_ratio,
            )
            print("[Universal-Generator] ✅ Parallel pipeline succeeded — rendering")
            return _render_paper_from_data(
                _paper_data, blueprint_dict, class_name, subject, chapters,
                additional_context, pattern,
                total_input_tokens=_in_tok, total_output_tokens=_out_tok,
            )
        except Exception as _par_exc:
            print(f"[Universal-Generator] ⚠️  Parallel pipeline failed ({_par_exc}), using single-prompt fallback")

        # ── Attempt 2 (fallback): single-prompt approach ─────────────────────
        variation_offset = random.randint(0, 20)
        context_data = get_universal_context(class_name, subject, chapters, all_question_types, variation_offset=variation_offset, school_id=school_id)
        context_text = context_data['context_text']

        file_path, summary, total_cost, in_tok, out_tok = generate_with_universal_prompt(
            class_name, subject, chapters, difficulty, pattern, blueprint,
            context_text, all_question_types, summary_file, model_source, additional_context,
            creative_ratio=creative_ratio,
        )
        return file_path, summary, total_cost, in_tok, out_tok
        
    except Exception as e:
        error_msg = f"Universal generation failed: {str(e)}"
        print(f"[Universal-Generator] ❌ {error_msg}")
        
        with open(summary_file, "a", encoding="utf-8") as f:
            f.write(f"\n❌ ERROR: {error_msg}\n")
            f.write(f"Traceback: {traceback.format_exc()}\n")
        
        raise Exception(error_msg)

def _self_contained_directive() -> str:
    """The section pipeline's self-containment rule, reused by the single-prompt fallback.

    Same wording in both paths on purpose: a paper that falls back must not start asking
    "In Activity 6.2, …" just because the parallel pipeline was unavailable.
    """
    from .section_generator import SELF_CONTAINED_RULE
    return SELF_CONTAINED_RULE


def _source_mix_directive(creative_ratio) -> str:
    """The generate page's source-mix meter, as a whole-paper rule for the single-prompt path.

    The parallel pipeline spends the percentage question by question (see
    section_generator.plan_creative_allocation); this fallback builds ONE prompt for the whole
    paper, so the same setting can only be stated as a proportion. Empty at 0 — the default,
    where every question stays grounded in the reference material.
    """
    try:
        ratio = max(0, min(100, int(creative_ratio or 0)))
    except (TypeError, ValueError):
        ratio = 0
    if not ratio:
        return ""
    return (
        f"\nSOURCE MIX — MANDATORY: about {ratio}% of this paper's questions must be YOUR OWN "
        "compositions — original questions on the same chapters, built on a fresh scenario, "
        "example or set of numbers that does NOT appear in the reference material. The other "
        f"{100 - ratio}% must be grounded in the reference material. Questions that have to "
        "quote the book (extracts, prescribed passages, map work) always stay book-based, and "
        "an own question must never be a lightly reworded reference-material question.\n"
    )


def _difficulty_directive(difficulty: str) -> str:
    d = difficulty.strip().lower()
    if d == "hard":
        return """DIFFICULTY — HARD (NON-NEGOTIABLE REQUIREMENTS):
- Every question MUST demand multi-step reasoning or higher-order thinking (Bloom's: analysis, synthesis, evaluation)
- BANNED question formats: "What is", "Define", "Name", "List" — pure recall is strictly forbidden
- MCQ: all four options must be factually plausible with subtle distinctions — no obviously wrong distractors
- Assertion-Reason: pick non-obvious relationships where naive reasoning leads to the WRONG answer
- Short Answer: require "Explain why", "Justify", "Derive", "Predict", "Compare" — never "State" or "Describe"
- Long Answer: must integrate concepts from 2+ topics, require critical evaluation or synthesis, not narration
- Numericals: multi-step with unit conversions, formula derivations, or conceptual twists — single-step sums are banned
- Case-based: scenario must require inference; answers must NOT be directly lifted from the passage
- Passages (English): dense academic prose; test inference and implicit meaning — NOT literal comprehension
- A student who studied only once should answer fewer than 25% of questions correctly
- Prefer: exceptions to rules, edge cases, counter-intuitive results, cross-topic links, real-world applications"""
    elif d == "medium":
        return """DIFFICULTY — MEDIUM:
- Mix of application (50%) and analytical questions (50%)
- MCQ: 2 clearly wrong options, 2 plausible distractors
- Short Answer: require understanding + some application, not just recall
- Target standard CBSE board-exam level"""
    else:
        return """DIFFICULTY — EASY:
- Majority direct knowledge and basic application questions
- Clear, unambiguous wording; suitable for revision and below-average students
- MCQ: one obvious correct answer with straightforward distractors"""


def generate_with_universal_prompt(class_name, subject, chapters, difficulty, pattern, blueprint, context_text, question_types, summary_file, model_source, additional_context="",
                                   creative_ratio=0):
    """Generate paper using universal prompt system"""

    print(f"[Universal-Prompt] Building adaptive prompt for {subject}")
    print(f"[Universal-Prompt] Blueprint structure: {list(blueprint.keys())}")

    # Blueprint is already normalized by BlueprintManager.get_blueprint()
    blueprint_dict = blueprint

    # Detect compound subjects (Science = Biology+Chemistry+Physics;
    # Social Science = History+Geography+Political Science+Economics)
    compound_subject_spec = ""
    pattern_sections = getattr(pattern, 'sections', None) or []
    if pattern_sections and isinstance(pattern_sections[0], dict) and 'subject' in pattern_sections[0]:
        lines = [
            f"IMPORTANT: This is a COMPOUND subject paper. Generate questions separately for each component subject:",
        ]
        for sec in pattern_sections:
            sec_name = sec.get('name', '')
            sec_subj = sec.get('subject', '')
            sec_marks = sec.get('marks', 0)
            sec_qs = sec.get('questions', 0)
            sec_notes = sec.get('notes', '')
            # MI-02: include chapter routing per sub-subject in the compound spec
        # chapters list is from the user's selection; map them by sub-subject heuristic
        sec_chapter_hint = ""
        if chapters:
            sec_chapter_hint = f" (draw chapters relevant to {sec_subj} from: {', '.join(chapters)})"
        lines.append(
                f"  § {sec_name} — {sec_subj}: {sec_qs} questions, {sec_marks} marks{sec_chapter_hint}"
                + (f". {sec_notes}" if sec_notes else "")
            )
        lines.append("Each section must be labelled with the component subject name (e.g. '§ A — Biology').")
        lines.append("IMPORTANT: All questions in each § must be strictly from that section's component subject only.")
        compound_subject_spec = "\n".join(lines)
        print(f"[Universal-Prompt] Compound subject detected: {[s.get('subject') for s in pattern_sections]}")

    # ── Fallback: detect discipline from section title/name when 'subject' field is absent ──
    # Handles Social Science (History/Geography/PolSci/Economics) and Science (Phy/Chem/Bio)
    if not compound_subject_spec and pattern_sections:
        _DISC_RE = re.compile(
            r'\b(history|geography|political[\s\-]science|economics|civics|'
            r'physics|chemistry|biology)\b',
            re.IGNORECASE,
        )
        # Keywords that classify a chapter into a discipline
        _DISC_CHAPTER_KW = {
            'History':            ['nationalism', 'global world', 'globalworld', 'industriali',
                                   'print culture', 'rise of', 'making of', 'age of', 'work life',
                                   'indo-china', 'indo china'],
            'Geography':          ['resource', 'agriculture', 'water', 'forest', 'wildlife',
                                   'mineral', 'manufacturing', 'lifeline', 'land', 'soil', 'energy',
                                   'crops', 'irrigation'],
            'Political Science':  ['power sharing', 'federali', 'democracy', 'gender', 'religion',
                                   'political part', 'struggle', 'outcome', 'challenge', 'caste'],
            'Civics':             ['power sharing', 'federali', 'democracy', 'gender', 'religion',
                                   'political part', 'struggle', 'outcome', 'challenge', 'caste'],
            'Economics':          ['development', 'sector', 'money', 'credit', 'globalisa',
                                   'globaliza', 'consumer', 'income', 'poverty', 'employment'],
            'Physics':            ['motion', 'force', 'gravit', 'electricity', 'magnetis',
                                   'light', 'optic', 'nuclear', 'current', 'refraction'],
            'Chemistry':          ['chemical', 'acid', 'base', 'salt', 'metal', 'carbon',
                                   'periodic', 'reaction', 'compound', 'oxide'],
            'Biology':            ['life process', 'reproduction', 'heredit', 'evolution',
                                   'control', 'nervous', 'environment', 'ecosystem', 'organism'],
        }

        def _chapters_for_disc(disc, all_chapters):
            kws = _DISC_CHAPTER_KW.get(disc, [])
            matched = [c for c in all_chapters
                       if any(kw in c.lower() for kw in kws)]
            return matched or all_chapters  # fallback: all chapters

        routed = []
        for sec in pattern_sections:
            title_raw = sec.get('title', '') or sec.get('name', '') or ''
            m = _DISC_RE.search(title_raw)
            if m:
                routed.append((sec, m.group(1).title()))

        if routed:
            lines = [
                "IMPORTANT: This paper covers multiple sub-disciplines. "
                "Each section corresponds to exactly ONE sub-discipline. "
                "You MUST generate questions for each section ONLY from the chapters listed for that section. "
                "DO NOT cross-pollinate — Geography questions must never appear in the History section, etc.",
            ]
            for sec, disc in routed:
                sec_name = sec.get('name', '')
                sec_marks = sec.get('marks', 0)
                relevant = _chapters_for_disc(disc, chapters) if chapters else []
                ch_hint = f" | chapters: {', '.join(relevant)}" if relevant else ""
                lines.append(f"  Section {sec_name} → {disc} ({sec_marks} marks){ch_hint}")
            compound_subject_spec = "\n".join(lines)
            print(f"[Universal-Prompt] Title-based compound subject routing: "
                  f"{[(s.get('name',''), d) for s, d in routed]}")

    # Extract passage and extract instructions from pattern sections
    passage_instructions = ""
    if hasattr(pattern, 'sections') and pattern.sections:
        print(f"[Universal-Prompt] Extracting passage instructions from pattern sections")
        for section in pattern.sections:
            if isinstance(section, dict):
                section_name = section.get('name', 'Section')
                # Check for passage_instruction field
                if 'passage_instruction' in section:
                    passage_instructions += f"\n- {section_name}: {section['passage_instruction']}"
                    print(f"[Universal-Prompt]   Found passage instruction for {section_name}")
                # Check for extract_instruction field
                if 'extract_instruction' in section:
                    passage_instructions += f"\n- {section_name}: {section['extract_instruction']}"
                    print(f"[Universal-Prompt]   Found extract instruction for {section_name}")
                # Check subsections for extract instructions
                if 'subsections' in section and isinstance(section['subsections'], list):
                    for subsection in section['subsections']:
                        if isinstance(subsection, dict) and 'extract_instruction' in subsection:
                            sub_name = subsection.get('name', 'Subsection')
                            passage_instructions += f"\n- {section_name} > {sub_name}: {subsection['extract_instruction']}"
                            print(f"[Universal-Prompt]   Found extract instruction for {sub_name}")

    # Check if this is a complex blueprint (has non-empty subsections)
    is_complex_blueprint = False
    for section_key, section_data in blueprint_dict.items():
        if isinstance(section_data, dict) and section_data.get('subsections'):
            is_complex_blueprint = True
            break

    print(f"[Universal-Prompt] Blueprint type: {'Complex' if is_complex_blueprint else 'Simple'}")

    if is_complex_blueprint:
        # Use the complex blueprint directly as the schema
        blueprint_schema = json.dumps(blueprint_dict, indent=2)
        sections_spec = "Follow the exact blueprint structure provided below."
    else:
        # Build sections specification from simple blueprint
        sections_spec = ""
        # Use blueprint_dict which has been converted to proper format
        for section_name, section_data in blueprint_dict.items():
            section_title = section_data.get('title', '')
            section_marks = section_data.get('marks', 0)
            section_qtypes = section_data.get('question_types', [])

            qtype_strs = [qt.get('type', str(qt)) if isinstance(qt, dict) else str(qt) for qt in section_qtypes]
            sections_spec += f"""
Section {section_name} - {section_title} ({section_marks} marks):
Question Types: {', '.join(qtype_strs)}
"""
        # Still use the original blueprint structure for the schema
        blueprint_schema = json.dumps(blueprint, indent=2)

    # Get question type instructions
    question_type_instructions = get_question_type_instructions(question_types, subject)
    
    # Build the universal prompt
    passage_section = f"""
PASSAGE AND EXTRACT GENERATION INSTRUCTIONS (CRITICAL):
{passage_instructions if passage_instructions else "No specific passage/extract instructions found in pattern"}

PASSAGE/EXTRACT REQUIREMENTS:
1. If a section requires passages: Generate passages inline and include in JSON with "passage" key
2. If a section requires extracts: Generate extract text and include in JSON with "extract" key
3. Questions must reference the passages/extracts you generate
4. Include passage/extract text directly in the section, NOT in a separate file

JSON STRUCTURE FOR SECTIONS WITH PASSAGES/EXTRACTS:
- Add "passage" key: "passage": "Generated passage text here"
- Add "extract" key: "extract": "Generated extract text here"
- Keep "questions" array with questions based on passage/extract
- Questions must reference and quote the passage/extract

Example with passage:
{{
  "Reading Comprehension": {{
    "marks": 5,
    "questions_count": 5,
    "question_types": ["Short Answer"],
    "instructions": ["Read the passage carefully"],
    "constraints": {{}},
    "passage": "Mrs. Jones had an unusual garden...[full passage text here]",
    "questions": [
      {{"qnum": "1a", "text": "What was unusual about Mrs. Jones' garden?", "marks": 1}},
      ...
    ]
  }}
}}

Example with extract:
{{
  "Literature": {{
    "marks": 5,
    "questions_count": 3,
    "question_types": ["Extract"],
    "instructions": ["Answer with reference to the extract"],
    "constraints": {{}},
    "extract": "[Extract text from literature piece here]",
    "questions": [
      {{"qnum": "2a", "text": "Extract question based on above text", "marks": 5}},
      ...
    ]
  }}
}}
"""

    # Build the universal prompt
    # Merge any caller-provided additional context with retrieved context
    combined_context = context_text
    if additional_context:
        combined_context = f"{context_text}\n\n[ADDITIONAL CONTEXT]\n{additional_context}"

    # Get variation instructions
    variation_hints = get_variation_instructions()
    variation_instructions = "\n".join([f"- {hint}" for hint in variation_hints])

    # Language papers (Hindi/Tamil/Sanskrit) must be written in that language/script even on
    # this single-prompt fallback path.
    from .section_generator import _language_directive, _is_english_subject
    language_block = _language_directive(subject)

    # English grammar and creative writing are composed from the model's own knowledge — the
    # context material is off-limits for both. The per-section pipeline enforces this by
    # withholding retrieval from those sections outright; this single-prompt fallback shares one
    # context for the whole paper, so it can only be stated as a rule. See
    # section_generator.english_own_slot_kinds.
    english_grammar_block = ""
    if _is_english_subject(subject):
        english_grammar_block = """
ENGLISH GRAMMAR — ABSOLUTE RULE:
Every grammar question (gap filling, editing/omission, reordering, tenses, voice, narration,
articles, prepositions, determiners, modals, subject-verb agreement, clauses and phrases,
sentence transformation, punctuation, parts of speech) MUST be composed ENTIRELY from your own
knowledge of English grammar. Take NOTHING from the CONTEXT MATERIAL for a grammar question —
not a sentence, phrase, wording, name, character, place, chapter title or storyline. Write your
own example sentences about everyday situations and set their "chapter_tag" to "Grammar".

CREATIVE WRITING — ABSOLUTE RULE:
Every writing task (article, formal/informal letter, letter to the editor, notice, classified or
display advertisement, poster, speech, debate, report, story, diary entry, email, invitation,
analytical or descriptive paragraph, précis, note-making) MUST set a SELF-CONTAINED, real-world
brief composed ENTIRELY from your own knowledge. Do NOT base the task on, quote, summarise or
even MENTION a textbook chapter, story, poem, poet, author or character, and never open with
"After reading ...", "Based on your reading of ..." or "Inspired by the poem ...". The student
must be able to write the answer without having read any textbook. Use everyday situations from
school life, the neighbourhood, the environment, health, technology, sport or current affairs,
and where a question offers an internal choice BOTH options must be independent briefs of this
kind. Set their "chapter_tag" to "Writing".

The context material is for the reading and literature sections only.
"""

    prompt = f"""You are an expert question paper generator for CBSE {class_name} {subject} examinations.
{language_block}
CONTEXT MATERIAL:
{combined_context}
{english_grammar_block}

EXAMINATION SPECIFICATIONS:
- Class: {class_name}
- Subject: {subject}
- Chapters to Cover: {', '.join(chapters) if chapters else 'All chapters'}
{f"{chr(10)}{compound_subject_spec}" if compound_subject_spec else ""}
CHAPTER DISTRIBUTION — CRITICAL:
You MUST spread questions across ALL {len(chapters) if chapters else 1} chapter(s) listed above.
{chr(10).join(f'  Chapter {i+1}: {c}' for i, c in enumerate(chapters)) if chapters else ''}
Do NOT take all or most questions from a single chapter. Each chapter must contribute at least one question.
Aim for roughly equal representation: ~{max(1, round((blueprint and sum(v.get('marks',0) for v in blueprint.values() if isinstance(v,dict)) or 20) / max(1, len(chapters))))} marks per chapter.
{"STRICT SECTION-CHAPTER ROUTING: Each section above lists its allowed chapters. You MUST only use those chapters for that section. A Geography chapter (e.g. Agriculture, Water Resources) MUST NOT appear in a History section and vice versa." if compound_subject_spec else ""}

QUESTION PAPER STRUCTURE:
{sections_spec}

EXACT QUESTION DISTRIBUTION (MANDATORY — match these counts and marks precisely):
{_exact_distribution_spec(blueprint_dict)}
RULES FOR THE DISTRIBUTION ABOVE:
- Generate EXACTLY the number of each question type shown — do not merge, drop, add, or re-balance types.
- Each question's "marks" MUST equal the per-type marks shown (NEVER the section average).
- Each question MUST include a "type" field (one of: MCQ, VSA, SA, LA, CBQ) AND a "subtype" field:
  use "assertion_reason" for Assertion-Reason questions, "source_based" for case/source-based,
  "image_based" for diagram questions, otherwise "standard".
- The per-type marks across each section MUST sum to that section's total marks.

IMAGE-BASED QUESTIONS:
- When the distribution/instructions call for an image- or diagram-based question, add an
  "image_prompt" field to that question: a self-contained visual description (20-40 words)
  an AI image model can render. The image is placed ABOVE the question, so phrase the text
  like "Study the diagram above and answer:". Only add "image_prompt" to questions that need it.

BLUEPRINT SCHEMA (Follow this exact structure):
{blueprint_schema}

QUESTION TYPE REQUIREMENTS:
{question_type_instructions}

CRITICAL VARIATION REQUIREMENTS:
IMPORTANT: Generate COMPLETELY DIFFERENT questions from any previous papers. Each question must be unique and original.
{variation_instructions}

MATHEMATICAL NOTATION RULES (strictly follow for all math subjects):
- Powers/exponents: write as x² y³ Aⁿ⁻¹ using Unicode superscripts, NOT x^2 or A^(n-1)
- Multiplication: use × (U+00D7), NOT * or x  (e.g. 3 × 4, |A| × |B|)
- Determinants/absolute value: use | | notation as-is (e.g. |A|, |adj A|)
- Fractions in text: write as "p/q" or "p over q"
- Square root: write √x not sqrt(x)
- Common expressions: n² not n^2, xⁿ not x^n, (n-1) as superscript means write ⁿ⁻¹

GENERATION GUIDELINES:
1. Create questions that are appropriate for {class_name} level students
2. Ensure questions test understanding, application, and analysis
3. Use the provided context material as the primary source
4. Maintain consistency with CBSE examination patterns
5. Include a variety of question types as specified
6. Ensure proper mark distribution across sections
7. Make questions clear, unambiguous, and well-structured
8. Generate passages/extracts INLINE - do not reference external sources
9. Make passages/extracts contextually relevant to the chapters specified
10. CREATE UNIQUE QUESTIONS - Avoid repeating common question patterns or wordings
11. Use different examples, scenarios, and approaches for each question
12. Vary the complexity and angles of questions even for the same topic
13. DISTRIBUTE questions across ALL listed chapters — never cluster in one chapter

{passage_section}

CRITICAL JSON OUTPUT INSTRUCTIONS:
1. Output ONLY valid JSON - no text before or after
2. Use section names as keys (e.g., "Reading Comprehension", "Writing", "Grammar", "Literature")
3. For sections with passages: Include "passage" key with passage text BEFORE questions
4. For sections with extracts: Include "extract" key with extract text BEFORE questions
5. Each section MUST have ONLY these keys: "marks", "questions_count", "question_types", "instructions", "constraints", "passage" (if applicable), "extract" (if applicable), "questions"
6. NO other keys allowed - do not add LIT_EXTRACT_SR, LIT_EXTRACT_PM or any other nested keys
7. Each question must have: {{"qnum": (number or string), "type": "MCQ|VSA|SA|LA|CBQ", "subtype": "standard|assertion_reason|source_based|image_based", "text": "question text", "marks": marks_value}}
8. Flatten ALL questions into the "questions" array - no subsections, no nested structures
9. Do NOT include markdown, explanations, or any text - ONLY JSON

EXAMPLE OUTPUT FORMAT (WITH PASSAGES):
{{
  "Reading Comprehension": {{
    "marks": 5,
    "questions_count": 5,
    "question_types": ["Short Answer"],
    "instructions": ["Read the passage carefully"],
    "constraints": {{}},
    "passage": "Mrs. Jones had an unusual garden that caught everyone's attention. Unlike traditional gardens filled with colorful flowers and lush greenery, her garden was full of stones...[continue passage to reach word limit]",
    "questions": [
      {{"qnum": "1a", "text": "What makes Mrs. Jones' garden different from traditional gardens?", "marks": 1}},
      {{"qnum": "1b", "text": "Name two features of Mrs. Jones' garden.", "marks": 1}}
    ]
  }},
  "Literature": {{
    "marks": 5,
    "questions_count": 2,
    "question_types": ["Extract"],
    "instructions": ["Answer with reference to the extract"],
    "constraints": {{}},
    "extract": "[Full extract from literature text - should be substantial enough to base questions on]",
    "questions": [
      {{"qnum": "2a", "text": "Extract-based question 1", "marks": 2}},
      {{"qnum": "2b", "text": "Extract-based question 2", "marks": 3}}
    ]
  }}
}}

{_difficulty_directive(difficulty)}
{_source_mix_directive(creative_ratio)}
{_self_contained_directive()}
OUTPUT REQUIREMENTS:
- Generate ONLY valid JSON
- Follow the example format above
- Include all sections from the blueprint
- Total marks and questions must match the blueprint
- ALL questions must be based on provided context OR generated passages/extracts
- Passages MUST be included inline in JSON with "passage" key
- Extracts MUST be included inline in JSON with "extract" key
- Do not include any text outside JSON"""

    print(f"[Universal-Prompt] Prompt length: {len(prompt)} characters")
    
    # Log prompt to summary file
    with open(summary_file, "a", encoding="utf-8") as f:
        f.write(f"\n=== UNIVERSAL PROMPT ===\n")
        f.write(f"Blueprint type: {'Complex' if is_complex_blueprint else 'Simple'}\n")
        f.write(f"Prompt length: {len(prompt)} characters\n")
        f.write(f"Question types: {question_types}\n")
        f.write(f"Context length: {len(context_text)} characters\n")
        f.write(f"{'='*50}\n\n")
    
    # Call Bedrock with the universal prompt - use varied temperature for more variation
    total_input_tokens = 0
    total_output_tokens = 0
    try:
        # Use varied temperature for more question variation
        variation_temp = get_variation_temperature()
        print(f"[Variation] Using temperature: {variation_temp} (varied for uniqueness)")
        raw_json, input_tokens_gen, output_tokens_gen = call_bedrock(prompt, GEN_MODEL_ID, temperature=variation_temp, model_source='aws')
        total_input_tokens += input_tokens_gen
        total_output_tokens += output_tokens_gen
        print(f"[Universal-Prompt] ✅ Received response from Bedrock")
        
        # Log the complete raw response
        print(f"[Generation] Raw response received: {len(raw_json)} characters")
        print(f"[Generation] Raw response preview: {raw_json[:200]}...")
        
        # Check if the response is complete
        if not raw_json.strip().endswith('}') and not raw_json.strip().endswith('```'):
            print(f"[Generation] ⚠️  WARNING: Response appears incomplete!")
            print(f"[Generation] Last 200 chars: {raw_json[-200:]}")
            
            # Try to complete the JSON if it's clearly cut off
            if raw_json.strip().endswith('"') and raw_json.count('{') > raw_json.count('}'):
                print(f"[Generation] Attempting to complete truncated JSON...")
                # Add missing closing brackets
                missing_braces = raw_json.count('{') - raw_json.count('}')
                raw_json += '"}' * missing_braces
                print(f"[Generation] Added {missing_braces} closing brackets")
        
        # Step 2: Validate the complete response with validator model
        print(f"[Validation] Starting JSON validation...")
        print(f"[Validation] Using model: {VAL_MODEL_ID}")
        
        validator_prompt = f"""
You are a JSON validator and fixer for CBSE question papers.

TASK: Fix the following JSON to ensure it's valid and complete.

REQUIREMENTS:
- Must be valid JSON syntax
- Must follow the exact schema structure
- All required fields must be present
- No missing brackets, quotes, or commas
- All question numbers must be sequential
- All marks must add up correctly

INPUT JSON TO FIX:
{raw_json}

OUTPUT: Return ONLY the corrected JSON, no explanations.
"""
        
        validated, input_tokens_val, output_tokens_val = call_bedrock(validator_prompt, VAL_MODEL_ID, temperature=0.3, model_source='aws')
        total_input_tokens += input_tokens_val
        total_output_tokens += output_tokens_val
        
        print(f"[Validation] Validation response received: {len(validated)} characters")
        print(f"[Validation] Validation response preview: {validated[:200]}...")
        
        # Step 3: Final JSON enforcement and cleanup
        print(f"[Final-Validation] Starting final JSON enforcement...")
        paper_data = enforce_json(validated)
        
        print(f"[Final-Validation] JSON enforcement completed successfully!")
        print(f"[Final-Validation] Final data structure: {list(paper_data.keys()) if isinstance(paper_data, dict) else 'Not a dict'}")

        # Calculate cost
        cost_per_input_1k_tokens = 0.49  # INR
        cost_per_output_1k_tokens = 1.47 # INR
        total_cost = (total_input_tokens / 1000) * cost_per_input_1k_tokens + \
                     (total_output_tokens / 1000) * cost_per_output_1k_tokens
        print(f"[Cost] Total Input Tokens: {total_input_tokens}")
        print(f"[Cost] Total Output Tokens: {total_output_tokens}")
        print(f"[Cost] Calculated Cost: {total_cost:.4f} INR")
        
        # Log generation results to summary file
        if summary_file:
            with open(summary_file, "a", encoding="utf-8") as f:
                f.write(f"=== GENERATION RESULTS ===\n")
                f.write(f"Generation Model: {GEN_MODEL_ID}\n")
                f.write(f"Validation Model: {VAL_MODEL_ID}\n")
                f.write(f"Raw JSON Length: {len(raw_json)} chars\n")
                f.write(f"Validated JSON Length: {len(validated)} chars\n")
                f.write(f"Final JSON Structure: {list(paper_data.keys()) if isinstance(paper_data, dict) else 'Not a dict'}\n")
                f.write(f"Questions Generated: {len(paper_data.get('sections', {})) if isinstance(paper_data, dict) else len(paper_data)} sections\n")
                f.write(f"Total Input Tokens: {total_input_tokens}\n")
                f.write(f"Total Output Tokens: {total_output_tokens}\n")
                f.write(f"Calculated Cost: {total_cost:.4f} INR\n")
                f.write(f"Generation Process: Complete\n")
                f.write(f"{'='*50}\n\n")

        print("=" * 50)
        print("USING DIRECT JSON APPROACH")
        print("=" * 50)

        # Use in-memory paper data stored by enforce_json() — thread-safe and avoids
        # the shared temp_clean.json file that caused cross-user data contamination.
        direct_data = getattr(_request_state, 'paper_data', None) or paper_data
        print(f"[DIRECT] Using paper data with keys: {list(direct_data.keys())}")

        # The fallback skips the parallel pipeline's validation chain. Run at least V1 structural
        # validation per section so broken MCQs / wrong counts / missing fields are caught and
        # recorded (attached as _errors/_partial → visible in paper_data & temp_questions.json),
        # instead of an unvalidated paper silently reaching a teacher.
        try:
            from .section_generator import build_work_orders, validate_section_output, reconcile_uniform_marks
            _work_orders = build_work_orders(blueprint, pattern, {}, difficulty, class_name, subject, chapters, unit_map=unit_map)
            for _wo in _work_orders:
                _sec = direct_data.get(_wo.section_name)
                if isinstance(_sec, dict) and isinstance(_sec.get("questions"), list) and _sec["questions"]:
                    _errs = validate_section_output({"questions": _sec["questions"]}, _wo)
                    if _errs:
                        _sec["_errors"] = _errs
                        _sec["_partial"] = True
                        print(f"[Fallback-Validate] '{_wo.section_name}': {len(_errs)} issue(s) — {_errs[:3]}")
            # Deterministically fix per-question marks (drift + inconsistent-pattern distribution)
            # AFTER validation, so the marks audit (run later in tasks.py) doesn't fire on a slip.
            reconcile_uniform_marks(direct_data, _work_orders)
        except Exception as _ve:
            print(f"[Fallback-Validate] skipped ({_ve})")

        all_questions = render_section_questions([], direct_data, blueprint, class_name, subject, chapters, None)
        # Attach AI-generated images if prompts/markers present
        allow_images = True
        try:
            if additional_context:
                ctx = json.loads(additional_context) if isinstance(additional_context, str) else {}
                if isinstance(ctx, dict):
                    allow_images = ctx.get('allow_ai_images', True) is not False
        except Exception:
            pass
        all_questions = materialize_images(all_questions, allow=allow_images)
        try:
            num_imgs = sum(1 for t, _ in all_questions if t == 'image')
            print(f"[ImageGen] Images attached: {num_imgs}")
        except Exception:
            pass
        print(f"[DIRECT] Generated {len(all_questions)} questions")
        
        summary = {sec: {"title": sec, "marks": paper_data.get(sec, {}).get('marks', 0)}
                   for sec in paper_data.keys()}

        # Prepare header meta for renderer
        header_meta = {}
        try:
            header_meta = {
                "class_name": class_name,
                "subject": subject,
                "pattern_name": getattr(pattern, 'name', ''),
                "marks": _paper_total_marks(pattern, blueprint),
            }
            # Parse additional_context for duration/marks/class_name if JSON
            if additional_context:
                import json as _json
                try:
                    ctx_obj = _json.loads(additional_context)
                    print(f"[Header-Meta] Parsed additional_context: {ctx_obj}")
                    if isinstance(ctx_obj, dict):
                        # Override with values from generate page if provided
                        if 'class_name' in ctx_obj and ctx_obj['class_name']:
                            header_meta['class_name'] = ctx_obj['class_name']
                            print(f"[Header-Meta] Using class_name from form: {ctx_obj['class_name']}")
                        if 'duration' in ctx_obj and ctx_obj['duration']:
                            header_meta['duration'] = ctx_obj['duration']
                            print(f"[Header-Meta] Using duration from form: {ctx_obj['duration']}")
                        if 'marks' in ctx_obj and ctx_obj['marks']:
                            header_meta['marks'] = ctx_obj['marks']
                            print(f"[Header-Meta] Using marks from form: {ctx_obj['marks']}")
                        if 'test_type' in ctx_obj and ctx_obj['test_type']:
                            header_meta['test_type'] = ctx_obj['test_type']
                            print(f"[Header-Meta] Using test_type from form: {ctx_obj['test_type']}")
                        if 'school_name' in ctx_obj and ctx_obj['school_name']:
                            header_meta['school_name'] = ctx_obj['school_name']
                except Exception as e:
                    print(f"[Header-Meta] WARN: failed to parse additional_context: {e}")
                    print(f"[Header-Meta] Traceback: {traceback.format_exc()}")
            print(f"[Header-Meta] ✅ Built header metadata: {header_meta}")
            print(f"[Header-Meta]   - class_name: '{header_meta.get('class_name')}'")
            print(f"[Header-Meta]   - duration: '{header_meta.get('duration')}'")
            print(f"[Header-Meta]   - marks: '{header_meta.get('marks')}'")
            print(f"[Header-Meta]   - test_type: '{header_meta.get('test_type')}'")
            print(f"[Header-Meta]   - subject: '{header_meta.get('subject')}'")
        except Exception as _e:
            print(f"[Header-Meta] ❌ WARN: failed to set header meta: {_e}")
            print(f"[Header-Meta] Traceback: {traceback.format_exc()}")

        # Render DOCX (download should be Word) and return (file_path, summary) tuple
        file_path, summary = render_docx(class_name, subject, chapters, all_questions, summary, header_meta=header_meta)
        return file_path, summary, total_cost, total_input_tokens, total_output_tokens
        
    except Exception as e:
        error_msg = f"Universal prompt generation failed: {str(e)}"
        print(f"[Universal-Prompt] ❌ {error_msg}")
        
        with open(summary_file, "a", encoding="utf-8") as f:
            f.write(f"\n❌ ERROR: {error_msg}\n")
            f.write(f"Traceback: {traceback.format_exc()}\n")
        
        raise Exception(error_msg)

def generate_paper(class_name, subject, chapters, difficulty, pattern, section=None, model_source='aws', additional_context="", school_id=None, unit_map=None,
                   creative_ratio=0):
    """
    Main entry point for question paper generation.
    Now uses the universal generator by default, with fallback to legacy system.
    """
    print(f"[Generator] Starting paper generation for {class_name} {subject}")
    
    try:
        # Try universal generator first
        print(f"[Generator] Attempting universal generation...")
        file_path, summary, total_cost, in_tok, out_tok = generate_universal_paper(class_name, subject, chapters, difficulty, pattern, section, model_source, additional_context, school_id=school_id, unit_map=unit_map,
                                                                                  creative_ratio=creative_ratio)
        return file_path, summary, total_cost, in_tok, out_tok
        
    except Exception as e:
        print(f"[Generator] ⚠️ Universal generation failed: {e}")
        print(f"[Generator] Falling back to legacy system...")
        
        # Fallback to legacy system
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        summary_file = f"temp_generation_summary_{timestamp}.txt"
        
        with open(summary_file, "w", encoding="utf-8") as f:
            f.write(f"=== LEGACY QUESTION PAPER GENERATION ===\n")
            f.write(f"Timestamp: {datetime.now()}\n")
            f.write(f"Class: {class_name}\n")
            f.write(f"Subject: {subject}\n")
            f.write(f"Chapters: {', '.join(chapters) if chapters else 'All'}\n")
            f.write(f"Difficulty: {difficulty}\n")
            f.write(f"Pattern: {pattern.name}\n")
            f.write(f"Fallback reason: Universal generation failed - {e}\n")
            f.write(f"Model Source: {model_source}\n") # Log model source
            f.write(f"{'='*50}\n\n")
        
        print(f"[Generator] Using legacy system - Summary saved to: {summary_file}")

        # Use the pattern's structure converted to blueprint format
        blueprint = pattern_sections_to_blueprint_dict(pattern)
        if not blueprint:
            blueprint = get_blueprint(class_name, subject, section)

        if subject.lower() in ["english", "english core"]:
            file_path, summary = generate_english_paper(class_name, subject, chapters, difficulty, pattern, blueprint, summary_file, model_source)
            return file_path, summary, 0.0, 0, 0
        else:
            file_path, summary = generate_science_paper(class_name, subject, chapters, difficulty, pattern, blueprint, summary_file, model_source)
            return file_path, summary, 0.0, 0, 0
