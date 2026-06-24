import os
import docx


def extract_text_from_pdf(pdf_path):
    """Extract text content from a PDF file."""
    if not os.path.exists(pdf_path):
        return f"PDF file not found: {pdf_path}"

    try:
        import PyPDF2
        from PyPDF2.errors import PdfReadError
        with open(pdf_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text = ""
            for page in reader.pages:
                try:
                    text += (page.extract_text() or "") + "\n"
                except Exception:
                    continue
            if text.strip():
                return text[:5000]
    except Exception as e:
        print(f"[PDF Extract] PyPDF2 error: {str(e)[:100]}, trying pdfplumber...")

    try:
        import pdfplumber
        text = ""
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                try:
                    text += (page.extract_text() or "") + "\n"
                except Exception:
                    continue
        return text[:5000] if text.strip() else "Could not extract text from PDF (may be image-based)"
    except ImportError:
        return "pdfplumber not installed — run: pip install pdfplumber"
    except Exception as e:
        return f"Could not extract PDF content: {e}"


def extract_text_from_docx(docx_file):
    """Extract text content from a DOCX file."""
    try:
        document = docx.Document(docx_file)
        return "\n".join(p.text for p in document.paragraphs)
    except Exception as e:
        return f"Could not extract DOCX content: {e}"


def extract_docx_text_with_images(docx_path, paper_id):
    """Extract DOCX text for editing, emitting an ``[IMG_FILE: <rel>]`` marker for every
    embedded image in document order.

    The plain-text editor/AI round-trip otherwise loses images: text extraction drops them
    and the re-render rebuilds from text only. By saving each embedded image's bytes under
    media/generated_images/ and inserting a marker line, images survive the round-trip —
    _parse_edited_text turns the marker back into an ('image', path) tuple at render time.
    """
    from django.conf import settings
    from docx.oxml.ns import qn
    try:
        document = docx.Document(docx_path)
    except Exception as e:
        return f"Could not extract DOCX content: {e}"

    media_root = settings.MEDIA_ROOT
    out_dir = os.path.join(media_root, 'generated_images')
    os.makedirs(out_dir, exist_ok=True)

    lines, idx = [], 0
    for para in document.paragraphs:
        # Images render in their own paragraph ABOVE the question, so emit markers first.
        for blip in para._p.iter(qn('a:blip')):
            embed = blip.get(qn('r:embed'))
            if not embed:
                continue
            try:
                part = document.part.related_parts.get(embed)
                if part is None:
                    continue
                ext = os.path.splitext(str(part.partname))[1] or '.png'
                rel = f"generated_images/edit_{paper_id}_{idx}{ext}"
                with open(os.path.join(media_root, rel), 'wb') as fh:
                    fh.write(part.blob)
                lines.append(f"[IMG_FILE: {rel}]")
                idx += 1
            except Exception:
                continue
        txt = para.text
        if txt and txt.strip():
            lines.append(txt)
    return "\n".join(lines)
